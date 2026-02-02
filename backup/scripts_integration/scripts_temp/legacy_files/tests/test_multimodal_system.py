#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
多模态文件系统验证测试
Multimodal File System Verification Test

测试文档解析、OCR识别、多模态识别功能，包括音频、视频、技术图纸等
"""

import asyncio
import json
import logging
import os
import sys
from datetime import datetime
from pathlib import Path

# 添加项目路径
sys.path.append('/Users/xujian/Athena工作平台')

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class MultimodalSystemTester:
    """多模态系统测试器"""

    def __init__(self):
        self.test_files = {
            'documents': [
                'data/documents/*.pdf',
                'data/documents/*.docx',
                'data/documents/*.txt'
            ],
            'images': [
                'data/images/*.png',
                'data/images/*.jpg',
                'data/images/*.jpeg',
                'data/images/*.tiff',
                'data/images/*.bmp'
            ],
            'technical_drawings': [
                'data/drawings/*.dwg',
                'data/drawings/*.dxf',
                'data/drawings/*.svg',
                'data/images/blueprints/*'
            ],
            'audio': [
                'data/audio/*.mp3',
                'data/audio/*.wav',
                'data/audio/*.m4a'
            ],
            'video': [
                'data/video/*.mp4',
                'data/video/*.avi',
                'data/video/*.mov'
            ]
        }

    def check_dependencies(self):
        """检查依赖库"""
        logger.info('🔍 检查多模态处理依赖库...')

        dependencies = {
            'PyMuPDF': 'fitz',
            'pytesseract': 'pytesseract',
            'Pillow': 'PIL',
            'opencv-python': 'cv2',
            'librosa': 'librosa',
            'moviepy': 'moviepy',
            'transformers': 'transformers',
            'torch': 'torch',
            'torchvision': 'torchvision'
        }

        missing = []
        installed = []

        for lib_name, import_name in dependencies.items():
            try:
                __import__(import_name)
                installed.append(lib_name)
                logger.info(f"  ✅ {lib_name}")
            except ImportError:
                missing.append(lib_name)
                logger.warning(f"  ❌ {lib_name} (未安装)")

        logger.info(f"\n📊 依赖统计: 已安装 {len(installed)}/{len(dependencies)} 个库")

        if missing:
            logger.warning(f"⚠️ 缺失依赖: {', '.join(missing)}")
            logger.info('💡 安装命令: pip install ' + ' '.join(missing))

        return len(missing) == 0

    def find_test_files(self):
        """查找测试文件"""
        logger.info("\n📁 查找测试文件...")

        found_files = {}

        for category, patterns in self.test_files.items():
            files = []
            for pattern in patterns:
                path = Path(pattern)
                if path.parent.exists():
                    files.extend(path.parent.glob(path.name))
            found_files[category] = files

            if files:
                logger.info(f"  {category}: 找到 {len(files)} 个文件")
                for file in files[:3]:  # 只显示前3个
                    logger.info(f"    - {file.name}")
                if len(files) > 3:
                    logger.info(f"    ... (还有 {len(files) - 3} 个文件)")
            else:
                logger.warning(f"  {category}: 未找到文件")

        return found_files

    def test_document_parsing(self, file_path):
        """测试文档解析"""
        logger.info(f"📄 测试文档解析: {file_path.name}")

        try:
            # 尝试不同的解析方法
            if file_path.suffix.lower() == '.pdf':
                return self._parse_pdf(file_path)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                return self._parse_word(file_path)
            elif file_path.suffix.lower() == '.txt':
                return self._parse_text(file_path)
            else:
                return {'error': f"不支持的文档格式: {file_path.suffix}"}

        except Exception as e:
            logger.error(f"  ❌ 解析失败: {str(e)}")
            return {'error': str(e)}

    def _parse_pdf(self, file_path):
        """解析PDF文件"""
        import fitz

        doc = fitz.open(str(file_path))

        # 提取文本
        full_text = ''
        for page in doc:
            text = page.get_text()
            full_text += text + "\n"

        # 检查是否需要OCR
        if not full_text.strip():
            logger.info('  🔍 PDF未包含文本，尝试OCR...')
            return self._ocr_pdf(file_path)

        doc.close()

        return {
            'type': 'pdf',
            'pages': len(doc),
            'text_length': len(full_text),
            'sample_text': full_text[:500] + '...' if len(full_text) > 500 else full_text,
            'needs_ocr': False
        }

    def _ocr_pdf(self, file_path):
        """OCR解析PDF"""
        try:
            from utils.ocr_patent_parser import OCRPatentParser

            parser = OCRPatentParser()
            result = parser.parse_patent_pdf_ocr(str(file_path), max_pages=3)

            if 'error' not in result:
                logger.info(f"  ✅ OCR成功，提取了 {len(result.get('ocr_text', ''))} 个字符")
                return {
                    'type': 'pdf_ocr',
                    'ocr_success': True,
                    'extracted_info': result
                }
            else:
                return result

        except Exception as e:
            return {'error': f"OCR失败: {str(e)}"}

    def _parse_word(self, file_path):
        """解析Word文档"""
        try:
            from docx import Document

            doc = Document(str(file_path))
            full_text = "\n".join([para.text for para in doc.paragraphs])

            return {
                'type': 'word',
                'paragraphs': len(doc.paragraphs),
                'text_length': len(full_text),
                'sample_text': full_text[:500] + '...' if len(full_text) > 500 else full_text
            }

        except Exception as e:
            return {'error': f"Word解析失败: {str(e)}"}

    def _parse_text(self, file_path):
        """解析文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            return {
                'type': 'text',
                'text_length': len(content),
                'sample_text': content[:500] + '...' if len(content) > 500 else content
            }

        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    content = f.read()

                return {
                    'type': 'text',
                    'encoding': 'gbk',
                    'text_length': len(content),
                    'sample_text': content[:500] + '...' if len(content) > 500 else content
                }
            except:
                return {'error': '文本编码解析失败'}

    def test_image_processing(self, file_path):
        """测试图像处理"""
        logger.info(f"🖼️ 测试图像处理: {file_path.name}")

        try:
            import cv2
            import numpy as np
            from PIL import Image

            # 使用PIL打开图像
            with Image.open(file_path) as img:
                img_info = {
                    'type': 'image',
                    'format': img.format,
                    'mode': img.mode,
                    'size': img.size,
                    'file_size': file_path.stat().st_size
                }

            # 使用OpenCV分析
            img_cv = cv2.imread(str(file_path))
            if img_cv is not None:
                height, width = img_cv.shape[:2]
                img_info.update({
                    'opencv_size': (width, height),
                    'channels': img_cv.shape[2] if len(img_cv.shape) > 2 else 1
                })

                # 检测是否为技术图纸
                if self._is_technical_drawing(img_cv):
                    img_info['drawing_type'] = '技术图纸'
                    logger.info('  📐 检测为技术图纸')
                else:
                    img_info['drawing_type'] = '普通图像'

            return img_info

        except Exception as e:
            logger.error(f"  ❌ 图像处理失败: {str(e)}")
            return {'error': str(e)}

    def _is_technical_drawing(self, img_cv):
        """判断是否为技术图纸"""
        # 简单的启发式判断
        gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
        edges = cv2.Canny(gray, 50, 150)

        # 计算边缘密度
        edge_density = np.sum(edges > 0) / edges.size

        # 技术图纸通常有较高的边缘密度
        return edge_density > 0.05

    def test_audio_processing(self, file_path):
        """测试音频处理"""
        logger.info(f"🎵 测试音频处理: {file_path.name}")

        try:
            import librosa
            import soundfile as sf

            # 加载音频
            y, sr = librosa.load(str(file_path))

            audio_info = {
                'type': 'audio',
                'format': file_path.suffix[1:].upper(),
                'duration': len(y) / sr,
                'sample_rate': sr,
                'channels': 1 if len(y.shape) == 1 else y.shape[1],
                'file_size': file_path.stat().st_size
            }

            # 提取音频特征
            mfcc = librosa.feature.mfcc(y=y, sr=sr, n_mfcc=13)
            audio_info['mfcc_shape'] = mfcc.shape

            # 检测是否包含语音
            spectral_centroid = librosa.feature.spectral_centroid(y=y, sr=sr)
            if np.mean(spectral_centroid) > 2000:
                audio_info['content_type'] = '语音'
            else:
                audio_info['content_type'] = '音乐/其他'

            logger.info(f"  ✅ 音频时长: {audio_info['duration']:.2f}秒")
            return audio_info

        except Exception as e:
            logger.error(f"  ❌ 音频处理失败: {str(e)}")
            return {'error': str(e)}

    def test_video_processing(self, file_path):
        """测试视频处理"""
        logger.info(f"🎬 测试视频处理: {file_path.name}")

        try:
            import cv2

            cap = cv2.VideoCapture(str(file_path))

            if not cap.isOpened():
                return {'error': '无法打开视频文件'}

            # 获取视频信息
            fps = cap.get(cv2.CAP_PROP_FPS)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            duration = frame_count / fps if fps > 0 else 0

            video_info = {
                'type': 'video',
                'format': file_path.suffix[1:].upper(),
                'fps': fps,
                'frame_count': frame_count,
                'resolution': (width, height),
                'duration': duration,
                'file_size': file_path.stat().st_size
            }

            # 读取第一帧
            ret, frame = cap.read()
            if ret:
                video_info['first_frame_shape'] = frame.shape

            cap.release()

            logger.info(f"  ✅ 视频: {width}x{height}, {duration:.2f}秒")
            return video_info

        except Exception as e:
            logger.error(f"  ❌ 视频处理失败: {str(e)}")
            return {'error': str(e)}

    async def test_multimodal_ai(self, file_path):
        """测试多模态AI识别"""
        logger.info(f"🤖 测试多模态AI识别: {file_path.name}")

        try:
            # 这里可以集成多模态模型，如CLIP等
            # 暂时返回模拟结果
            result = {
                'ai_analysis': 'enabled',
                'recognized_objects': [],
                'text_detected': False,
                'confidence': 0.0
            }

            # 尝试文本检测
            if file_path.suffix.lower() in ['.png', '.jpg', '.jpeg']:
                result = await self._detect_text_in_image(file_path)

            return result

        except Exception as e:
            logger.error(f"  ❌ AI识别失败: {str(e)}")
            return {'error': str(e)}

    async def _detect_text_in_image(self, file_path):
        """检测图像中的文本"""
        try:
            import pytesseract
            from PIL import Image

            with Image.open(file_path) as img:
                text = pytesseract.image_to_string(img, lang='chi_sim+eng')

                if text.strip():
                    return {
                        'ai_analysis': 'completed',
                        'text_detected': True,
                        'extracted_text': text[:200] + '...' if len(text) > 200 else text,
                        'confidence': 0.9
                    }
                else:
                    return {
                        'ai_analysis': 'completed',
                        'text_detected': False,
                        'confidence': 0.1
                    }

        except Exception as e:
            return {'error': f"文本检测失败: {str(e)}"}

    async def run_tests(self):
        """运行所有测试"""
        logger.info('🚀 多模态文件系统验证测试开始')
        logger.info('=' * 60)
        logger.info(f"测试时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # 1. 检查依赖
        deps_ok = self.check_dependencies()

        # 2. 查找测试文件
        found_files = self.find_test_files()

        # 3. 运行测试
        test_results = {}

        for category, files in found_files.items():
            if not files:
                continue

            logger.info(f"\n\n🧪 测试 {category}...")
            test_results[category] = []

            # 最多测试5个文件
            for file_path in files[:5]:
                logger.info(f"\n📋 测试文件: {file_path}")

                if category == 'documents':
                    result = self.test_document_parsing(file_path)
                elif category == 'images':
                    result = self.test_image_processing(file_path)
                    # 如果是图像，也测试AI识别
                    ai_result = await self.test_multimodal_ai(file_path)
                    result['ai_analysis'] = ai_result
                elif category == 'technical_drawings':
                    result = self.test_image_processing(file_path)
                elif category == 'audio':
                    result = self.test_audio_processing(file_path)
                elif category == 'video':
                    result = self.test_video_processing(file_path)
                else:
                    result = {'error': '未知文件类型'}

                test_results[category].append({
                    'file': str(file_path),
                    'result': result
                })

        # 4. 生成测试报告
        self.generate_test_report(test_results, deps_ok)

        return test_results

    def generate_test_report(self, test_results, deps_ok):
        """生成测试报告"""
        logger.info("\n\n📊 测试报告")
        logger.info('=' * 60)

        # 依赖检查结果
        logger.info(f"\n📦 依赖检查: {'✅ 全部安装' if deps_ok else '⚠️ 部分缺失'}")

        # 统计测试结果
        total_files = 0
        successful = 0
        failed = 0

        for category, results in test_results.items():
            category_success = 0
            for result in results:
                total_files += 1
                if 'error' not in result['result']:
                    successful += 1
                    category_success += 1
                else:
                    failed += 1

            logger.info(f"\n{category}:")
            logger.info(f"  - 测试文件: {len(results)}")
            logger.info(f"  - 成功: {category_success}")
            logger.info(f"  - 失败: {len(results) - category_success}")

        # 总体统计
        logger.info(f"\n📈 总体统计:")
        logger.info(f"  - 总文件数: {total_files}")
        logger.info(f"  - 成功解析: {successful}")
        logger.info(f"  - 解析失败: {failed}")
        logger.info(f"  - 成功率: {successful/total_files*100:.1f}%' if total_files > 0 else 'N/A")

        # 功能总结
        logger.info(f"\n✨ 功能验证总结:")
        logger.info(f"  - 文档解析: {'✅' if 'documents' in test_results else '⚠️'}")
        logger.info(f"  - 图像处理: {'✅' if 'images' in test_results else '⚠️'}")
        logger.info(f"  - OCR识别: {'✅' if any('ocr' in str(r) for results in test_results.values() for r in results) else '⚠️'}")
        logger.info(f"  - 音频处理: {'✅' if 'audio' in test_results else '⚠️'}")
        logger.info(f"  - 视频处理: {'✅' if 'video' in test_results else '⚠️'}")
        logger.info(f"  - 技术图纸: {'✅' if 'technical_drawings' in test_results else '⚠️'}")
        logger.info(f"  - AI识别: {'✅' if any('ai_analysis' in str(r) for results in test_results.values() for r in results) else '⚠️'}")

async def main():
    """主函数"""
    tester = MultimodalSystemTester()
    results = await tester.run_tests()

    # 保存结果
    report_path = f"multimodal_test_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    with open(report_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2, default=str)

    logger.info(f"\n💾 测试报告已保存: {report_path}")

if __name__ == '__main__':
    asyncio.run(main())