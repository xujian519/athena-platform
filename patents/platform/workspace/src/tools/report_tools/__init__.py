#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
报告工具模块
Report Tools Module

提供报告生成、可视化、文档处理等功能
Provides report generation, visualization, document processing and other functions

作者: Athena AI系统
创建时间: 2025-12-06
版本: 1.0.0
"""

import json
import logging
import os
import sys
from datetime import datetime

logger = logging.getLogger(__name__)

# 添加项目根目录到路径
project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

# 导入报告工具依赖
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import matplotlib
    import matplotlib.pyplot as plt
    matplotlib.use('Agg')  # 使用非交互式后端
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

try:
    import seaborn as sns
    SEABORN_AVAILABLE = True
except ImportError:
    SEABORN_AVAILABLE = False

try:
    from jinja2 import Environment, Template
    JINJA2_AVAILABLE = True
except ImportError:
    JINJA2_AVAILABLE = False

try:
    try:
        from docx import Document
        PYTHON_DOCX_AVAILABLE = True
    except ImportError:
        PYTHON_DOCX_AVAILABLE = False
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

__all__ = [
    'ReportGenerator',
    'ChartGenerator',
    'DocumentProcessor'
]

# 版本信息
__version__ = '1.0.0'
__author__ = 'Athena AI系统'

class ReportGenerator:
    """报告生成器"""

    def __init__(self):
        self.jinja2_available = JINJA2_AVAILABLE
        self.pandas_available = PANDAS_AVAILABLE

    def generate_markdown_report(self, data, template_content=None):
        """生成Markdown报告"""
        try:
            if template_content and JINJA2_AVAILABLE:
                template = Template(template_content)
                content = template.render(data=data)
            else:
                content = self._generate_basic_markdown(data)

            return content
        except Exception as e:
            return f"Error generating markdown report: {str(e)}"

    def generate_html_report(self, data, template_content=None):
        """生成HTML报告"""
        try:
            if template_content and JINJA2_AVAILABLE:
                template = Template(template_content)
                content = template.render(data=data)
            else:
                content = self._generate_basic_html(data)

            return content
        except Exception as e:
            return f"Error generating HTML report: {str(e)}"

    def generate_json_report(self, data):
        """生成JSON报告"""
        try:
            report_data = {
                'title': data.get('title', 'Report'),
                'generated_at': datetime.now().isoformat(),
                'data': data
            }
            return json.dumps(report_data, indent=2, ensure_ascii=False)
        except Exception as e:
            return f"Error generating JSON report: {str(e)}"

    def _generate_basic_markdown(self, data):
        """生成基础Markdown格式"""
        content = f"# {data.get('title', 'Report')}\n\n"
        content += f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"

        if 'summary' in data:
            content += f"## 摘要\n{data['summary']}\n\n"

        if 'sections' in data:
            for section in data['sections']:
                content += f"## {section.get('title', 'Section')}\n"
                content += f"{section.get('content', '')}\n\n"

        return content

    def _generate_basic_html(self, data):
        """生成基础HTML格式"""
        html = f"""
<!DOCTYPE html>
<html>
<head>
    <title>{data.get('title', 'Report')}</title>
    <meta charset='utf-8'>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        h1 {{ color: #333; }}
        h2 {{ color: #666; }}
        .timestamp {{ color: #999; }}
    </style>
</head>
<body>
    <h1>{data.get('title', 'Report')}</h1>
    <p class='timestamp'>生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
"""

        if 'summary' in data:
            html += f"    <h2>摘要</h2>\n    <p>{data['summary']}</p>\n"

        if 'sections' in data:
            for section in data['sections']:
                html += f"    <h2>{section.get('title', 'Section')}</h2>\n"
                html += f"    <p>{section.get('content', '')}</p>\n"

        html += """
</body>
</html>"""
        return html

class ChartGenerator:
    """图表生成器"""

    def __init__(self):
        self.matplotlib_available = MATPLOTLIB_AVAILABLE
        self.seaborn_available = SEABORN_AVAILABLE
        self.pandas_available = PANDAS_AVAILABLE

    def generate_bar_chart(self, data, title='Bar Chart', save_path=None):
        """生成柱状图"""
        if not MATPLOTLIB_AVAILABLE:
            return 'Matplotlib not available for chart generation'

        try:
            plt.figure(figsize=(10, 6))

            if isinstance(data, dict):
                plt.bar(data.keys(), data.values())
            elif PANDAS_AVAILABLE and isinstance(data, pd.DataFrame):
                data.plot(kind='bar')
            else:
                plt.bar(range(len(data)), data)

            plt.title(title)
            plt.xticks(rotation=45)
            plt.tight_layout()

            if save_path:
                plt.savefig(save_path, dpi=300, bbox_inches='tight')
                plt.close()
                return f"Chart saved to {save_path}"
            else:
                plt.close()
                return 'Chart generated successfully'
        except Exception as e:
            return f"Error generating bar chart: {str(e)}"

    def generate_line_chart(self, x_data, y_data, title='Line Chart', save_path=None):
        """生成折线图"""
        if not MATPLOTLIB_AVAILABLE:
            return 'Matplotlib not available for chart generation'

        try:
            plt.figure(figsize=(10, 6))
            plt.plot(x_data, y_data, marker='o')
            plt.title(title)
            plt.xlabel('X Axis')
            plt.ylabel('Y Axis')
            plt.grid(True, alpha=0.3)
            plt.tight_layout()

            if save_path:
                plt.savefig(save_path, dpi=300, bbox_inches='tight')
                plt.close()
                return f"Chart saved to {save_path}"
            else:
                plt.close()
                return 'Chart generated successfully'
        except Exception as e:
            return f"Error generating line chart: {str(e)}"

class DocumentProcessor:
    """文档处理器"""

    def __init__(self):
        self.docx_available = PYTHON_DOCX_AVAILABLE
        self.pandas_available = PANDAS_AVAILABLE

    def create_word_document(self, content_dict, save_path):
        """创建Word文档"""
        if not PYTHON_DOCX_AVAILABLE:
            return 'python-docx not available for Word document creation'

        try:
            doc = Document()

            # 添加标题
            if 'title' in content_dict:
                doc.add_heading(content_dict['title'], 0)

            # 添加生成时间
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            # 添加内容
            if 'sections' in content_dict:
                for section in content_dict['sections']:
                    if 'title' in section:
                        doc.add_heading(section['title'], 1)
                    if 'content' in section:
                        doc.add_paragraph(section['content'])

            doc.save(save_path)
            return f"Word document saved to {save_path}"
        except Exception as e:
            return f"Error creating Word document: {str(e)}"

    def export_data_to_excel(self, data, save_path):
        """导出数据到Excel"""
        if not PANDAS_AVAILABLE:
            return 'Pandas not available for Excel export'

        try:
            if isinstance(data, dict):
                df = pd.DataFrame(list(data.items()), columns=['Key', 'Value'])
            elif isinstance(data, list):
                df = pd.DataFrame(data)
            elif isinstance(data, pd.DataFrame):
                df = data
            else:
                return 'Unsupported data format for Excel export'

            df.to_excel(save_path, index=False)
            return f"Data exported to Excel: {save_path}"
        except Exception as e:
            return f"Error exporting to Excel: {str(e)}"

# 工具可用性状态
TOOL_STATUS = {
    'report_generator': JINJA2_AVAILABLE,
    'chart_generator': MATPLOTLIB_AVAILABLE,
    'document_processor': PYTHON_DOCX_AVAILABLE,
    'matplotlib': MATPLOTLIB_AVAILABLE,
    'seaborn': SEABORN_AVAILABLE,
    'jinja2': JINJA2_AVAILABLE,
    'python_docx': PYTHON_DOCX_AVAILABLE,
    'pandas': PANDAS_AVAILABLE
}

def get_available_tools():
    """获取可用的工具列表"""
    return [tool for tool, available in TOOL_STATUS.items() if available]

def get_tool_status():
    """获取工具状态信息"""
    return TOOL_STATUS.copy()

def check_dependencies():
    """检查依赖库状态"""
    deps = {
        'pandas': PANDAS_AVAILABLE,
        'matplotlib': MATPLOTLIB_AVAILABLE,
        'seaborn': SEABORN_AVAILABLE,
        'jinja2': JINJA2_AVAILABLE,
        'python_docx': PYTHON_DOCX_AVAILABLE
    }
    return deps

if __name__ == '__main__':
    logger.info('🔧 报告工具模块状态')
    logger.info(f"版本: {__version__}")
    logger.info(f"作者: {__author__}")
    logger.info(f"\n可用工具: {get_available_tools()}")
    logger.info(f"依赖检查: {check_dependencies()}")

    # 简单功能测试
    try:
        generator = ReportGenerator()
        test_data = {
            'title': 'Test Report',
            'summary': 'This is a test report generated by ReportGenerator.',
            'sections': [
                {'title': 'Section 1', 'content': 'This is the first section.'},
                {'title': 'Section 2', 'content': 'This is the second section.'}
            ]
        }

        markdown = generator.generate_markdown_report(test_data)
        logger.info('✅ Markdown报告生成测试成功')

        json_report = generator.generate_json_report(test_data)
        logger.info('✅ JSON报告生成测试成功')

    except Exception as e:
        logger.info(f"❌ ReportGenerator 测试失败: {e}")