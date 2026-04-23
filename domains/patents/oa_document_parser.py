#!/usr/bin/env python3
"""
审查意见文档解析器
Office Action Document Parser

从审查意见文档中提取结构化信息，支持:
- PDF文档解析
- Word文档解析
- 图片OCR识别
- 结构化JSON输出
- Markdown确认模板生成

作者: 小诺·双鱼公主
创建时间: 2025-12-24
版本: v0.1.2 "晨星初现"
"""
from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any

# 导入平台已有的文档处理工具
try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    logger = logging.getLogger(__name__)
    logger.warning("pdfplumber未安装，PDF解析功能不可用")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger = logging.getLogger(__name__)
    logger.warning("python-docx未安装，Word解析功能不可用")

# 导入平台已有的OCR工具
try:
    from core.ai.perception.chinese_ocr_optimizer import ChineseOCROptimizer
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    ChineseOCROptimizer = None

from core.logging_config import setup_logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = setup_logging()


class DocumentType(Enum):
    """文档类型"""

    PDF = "pdf"
    WORD = "docx"
    IMAGE = "image"
    TEXT = "text"


class RejectionType(Enum):
    """驳回类型"""

    NOVELTY = "novelty"  # 新颖性
    INVENTIVENESS = "inventiveness"  # 创造性
    UTILITY = "utility"  # 实用性
    CLARITY = "clarity"  # 清晰度
    SUPPORT = "support"  # 说明书支持
    UNITY = "unity"  # 单一性


@dataclass
class PriorArtReference:
    """对比文件详细信息"""

    reference_id: str = ""  # 对比文件标识 (D1, D2等)
    publication_number: str = ""  # 公开号
    publication_date: str = ""  # 公开日
    title: str = ""  # 标题
    applicant: str = ""  # 申请人
    relevant_claims: list[str] = field(default_factory=list)  # 相关权利要求
    description: str = ""  # 审查员引用描述

    def to_dict(self) -> dict[str, Any]:
        """转换为字典"""
        return {
            "reference_id": self.reference_id,
            "publication_number": self.publication_number,
            "publication_date": self.publication_date,
            "title": self.title,
            "applicant": self.applicant,
            "relevant_claims": self.relevant_claims,
            "description": self.description,
        }


@dataclass
class ParsedOfficeAction:
    """解析后的审查意见"""

    # 基本信息
    oa_id: str = ""
    document_source: str = ""  # 原文档路径

    # 目标专利信息（新增）
    target_application_number: str = ""  # 目标专利申请号
    target_patent_title: str = ""  # 目标专利标题
    target_applicant: str = ""  # 目标专利申请人

    # 驳回信息
    rejection_type: str = ""  # 驳回类型
    rejection_reason: str = ""  # 驳回理由

    # 对比文件详细信息（增强）
    prior_art_references: list[str] = field(default_factory=list)  # 简单列表（保留兼容性）
    prior_art_details: list[PriorArtReference] = field(default_factory=list)  # 详细信息

    # 权利要求
    cited_claims: list[int] = field(default_factory=list)

    # 审查员观点
    examiner_arguments: list[str] = field(default_factory=list)
    missing_features: list[str] = field(default_factory=list)

    # 时间信息
    received_date: str = ""
    response_deadline: str = ""

    # 元数据
    parsed_at: str = field(default_factory=lambda: datetime.now().isoformat())
    confidence: float = 0.0  # 解析置信度

    # 原始文本
    raw_text: str = ""

    # 结构化表格数据
    tables: list[dict[str, Any] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        """转换为字典"""
        return {
            "oa_id": self.oa_id,
            "document_source": self.document_source,
            "target_application_number": self.target_application_number,
            "target_patent_title": self.target_patent_title,
            "target_applicant": self.target_applicant,
            "rejection_type": self.rejection_type,
            "rejection_reason": self.rejection_reason,
            "prior_art_references": self.prior_art_references,
            "prior_art_details": [ref.to_dict() for ref in self.prior_art_details],
            "cited_claims": self.cited_claims,
            "examiner_arguments": self.examiner_arguments,
            "missing_features": self.missing_features,
            "received_date": self.received_date,
            "response_deadline": self.response_deadline,
            "parsed_at": self.parsed_at,
            "confidence": self.confidence,
            "tables": self.tables,
        }

    def to_json(self, indent: int = 2) -> str:
        """转换为JSON字符串"""
        return json.dumps(self.to_dict(), ensure_ascii=False, indent=indent)

    def to_markdown(self) -> str:
        """转换为Markdown格式（用于用户确认）"""
        md = []
        md.append("# 📋 审查意见解析结果\n")
        md.append("---\n")

        # 基本信息
        md.append("## 📝 基本信息")
        md.append(f"- **审查意见ID**: `{self.oa_id or '待生成'}`")
        md.append(f"- **文档来源**: `{self.document_source}`")
        md.append(f"- **解析时间**: {self.parsed_at}")
        md.append(f"- **置信度**: {'⭐' * int(self.confidence * 5)} ({self.confidence:.1%})\n")

        # 目标专利信息（新增）
        md.append("## 🎯 目标专利信息")
        md.append(f"- **申请号**: `{self.target_application_number or '❌ 未提取'}`")
        md.append(f"- **专利标题**: {self.target_patent_title or '未提取'}")
        md.append(f"- **申请人**: {self.target_applicant or '未提取'}\n")

        # 驳回信息
        md.append("## ❌ 驳回信息")
        md.append(f"- **驳回类型**: {self._format_rejection_type()}")
        md.append(f"- **驳回理由**: {self.rejection_reason or '未提取到'}\n")

        # 时间信息
        md.append("## ⏰ 时间信息")
        md.append(f"- **收到日期**: {self.received_date or '未提取'}")
        md.append(f"- **答复期限**: {self.response_deadline or '未提取'}\n")

        # 对比文件（增强显示）
        md.append("## 📄 对比文件")
        if self.prior_art_details:
            for ref in self.prior_art_details:
                md.append(f"### {ref.reference_id or '?'}: `{ref.publication_number}`")
                if ref.title:
                    md.append(f"- **标题**: {ref.title}")
                if ref.publication_date:
                    md.append(f"- **公开日**: {ref.publication_date}")
                if ref.applicant:
                    md.append(f"- **申请人**: {ref.applicant}")
                if ref.relevant_claims:
                    md.append(f"- **相关权利要求**: {', '.join(ref.relevant_claims)}")
                if ref.description:
                    md.append(f"- **引用描述**: {ref.description[:100]}...")
                md.append("")
        elif self.prior_art_references:
            for i, ref in enumerate(self.prior_art_references, 1):
                md.append(f"{i}. `{ref}`")
        else:
            md.append("*未提取到对比文件*")
        md.append("")

        # 被引用权利要求
        md.append("## 📜 被引用权利要求")
        if self.cited_claims:
            claims_str = ", ".join([f"权利要求{c}" for c in self.cited_claims])
            md.append(f"- {claims_str}")
        else:
            md.append("*未提取到权利要求*")
        md.append("")

        # 审查员论点
        md.append("## 💬 审查员论点")
        if self.examiner_arguments:
            for i, arg in enumerate(self.examiner_arguments, 1):
                md.append(f"{i}. {arg}")
        else:
            md.append("*未提取到论点*")
        md.append("")

        # 缺失技术特征
        md.append("## 🔍 缺失技术特征")
        if self.missing_features:
            for i, feature in enumerate(self.missing_features, 1):
                md.append(f"{i}. {feature}")
        else:
            md.append("*未提取到技术特征*")
        md.append("")

        # 结构化表格
        md.append("## 📊 提取的表格")
        if self.tables:
            for table in self.tables:
                idx = table.get("table_index", "?")
                md.append(f"### 表格 {idx}")
                md.append(
                    f"- **来源**: {table.get('source', 'unknown')}, "
                    f"**行数**: {table.get('row_count', 0)}, "
                    f"**列数**: {table.get('col_count', 0)}"
                )
                if table.get("page_number", 0) > 0:
                    md.append(f"- **所在页**: {table['page_number']}")
                rows = table.get("rows", [])
                if rows:
                    md.append("")
                    md.append(self._render_table_as_markdown(rows))
                md.append("")
        else:
            md.append("*未提取到表格*")
        md.append("")

        # 原始文本预览
        md.append("## 📄 原始文本预览")
        preview = self.raw_text[:500] if self.raw_text else "无"
        if len(self.raw_text or "") > 500:
            preview += "..."
        md.append(f"```\n{preview}\n```\n")

        # 确认提示
        md.append("---\n")
        md.append("### ✅ 请确认以上信息是否准确")
        md.append("- **重点检查**:")
        md.append(f"  - 目标专利申请号是否正确: `{self.target_application_number or '❌ 未提取'}`")
        md.append(f"  - 对比文件信息是否完整: {len(self.prior_art_details)} 个对比文件")
        md.append("- 如有错误，请指出需要修改的部分")
        md.append("- 确认无误后，回复 `确认` 或 `confirm` 继续")
        md.append("")

        return "\n".join(md)

    @staticmethod
    def _render_table_as_markdown(rows: list[list[str]) -> str:
        """将二维行数据渲染为Markdown表格"""
        if not rows:
            return ""
        col_count = max(len(r) for r in rows)
        lines = []
        for i, row in enumerate(rows):
            padded = [(cell or "").replace("|", "\\|") for cell in row]
            while len(padded) < col_count:
                padded.append("")
            lines.append("| " + " | ".join(padded) + " |")
            if i == 0:
                lines.append("| " + " | ".join(["---"] * col_count) + " |")
        return "\n".join(lines)

    def _format_rejection_type(self) -> str:
        """格式化驳回类型"""
        type_mapping = {
            "novelty": "🔴 新颖性驳回",
            "inventiveness": "🟠 创造性驳回",
            "utility": "🟡 实用性驳回",
            "clarity": "🟢 清晰度驳回",
            "support": "🔵 说明书支持驳回",
            "unity": "🟣 单一性驳回",
        }
        return type_mapping.get(self.rejection_type, f"⚪ {self.rejection_type}")


class OfficeActionParser:
    """
    审查意见文档解析器

    核心功能:
    1. 多格式文档解析 (PDF/Word/图片/文本)
    2. 智能信息提取
    3. 结构化数据生成
    4. Markdown确认模板
    """

    def __init__(self, parser_factory=None):
        """初始化解析器

        Args:
            parser_factory: 可选的DocumentParserFactory实例，
                           传入后对扫描型PDF自动使用MinerU解析。
                           不传时行为与原来完全一致。
        """
        self.name = "审查意见文档解析器"
        self.version = "v0.1.2"
        self._parser_factory = parser_factory

        # 检查依赖
        self.capabilities = {
            "pdf": HAS_PDF,
            "docx": HAS_DOCX,
            "ocr": HAS_OCR,
        }

        # 初始化OCR优化器（如果可用）
        self.ocr_optimizer = None
        if HAS_OCR and ChineseOCROptimizer:
            self.ocr_optimizer = ChineseOCROptimizer(config={"enable_preprocessing": True, "enable_postprocessing": True})

        # 使用平台已有的正则模式（复用生产环境经验）
        self.patterns = self._init_patterns()

        logger.info(f"📄 {self.name} ({self.version}) 初始化完成")
        logger.info(f"   支持格式: PDF={HAS_PDF}, Word={HAS_DOCX}, OCR={HAS_OCR}")

    def detect_document_type(self, file_path: str) -> DocumentType:
        """
        检测文档类型

        Args:
            file_path: 文件路径

        Returns:
            文档类型
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return DocumentType.PDF
        elif suffix in [".docx", ".doc"]:
            return DocumentType.WORD
        elif suffix in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
            return DocumentType.IMAGE
        elif suffix in [".txt", ".md"]:
            return DocumentType.TEXT
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")

    def parse_document(self, file_path: str) -> ParsedOfficeAction:
        """
        解析文档

        Args:
            file_path: 文件路径

        Returns:
            解析后的审查意见
        """
        doc_type = self.detect_document_type(file_path)

        logger.info(f"📄 开始解析: {file_path} (类型: {doc_type.value})")

        # 根据类型选择解析方法
        tables: list[dict[str, Any] = []

        if doc_type == DocumentType.PDF:
            text, tables = self._parse_pdf(file_path)
        elif doc_type == DocumentType.WORD:
            text, tables = self._parse_word(file_path)
        elif doc_type == DocumentType.IMAGE:
            text = self._parse_image(file_path)
        elif doc_type == DocumentType.TEXT:
            text = self._parse_text(file_path)
        else:
            raise ValueError(f"不支持的文档类型: {doc_type}")

        # 提取结构化信息
        parsed = self._extract_structured_info(text, file_path)
        parsed.tables = tables

        logger.info(
            f"✅ 解析完成 (置信度: {parsed.confidence:.1%}, 表格数: {len(tables)})"
        )

        return parsed

    def _init_patterns(self) -> dict[str, re.Pattern]:
        """
        初始化正则表达式模式（复用平台已有的生产环境经验）
        参考: core/intelligence/dspy/production_docx_extractor.py
        """
        return {
            # 专利号模式
            "patent_number": re.compile(r"CN\d{9}\.\d|CN\d+[A-Z]?"),
            # 申请号模式
            "application_number": re.compile(r"[0-9]{11,13}[.X]?"),
            # 日期模式
            "date": re.compile(r"(\d{4})[年\-](\d{1,2})[月\-](\d{1,2})"),
            # 证据/对比文件模式
            "evidence": re.compile(r"对比文件\s*([0-9A-D])[::]"),
            # 审查意见中的常见模式
            "rejection_reason": re.compile(r"驳回理由|审查意见|结论"),
            "claims": re.compile(r"权利要求(\d+)"),
            "technical_feature": re.compile(r"区别技术特征|区别在于|未公开|未披露"),
        }

    def _parse_pdf(self, file_path: str) -> tuple[str, list[dict[str, Any]]:
        """
        解析PDF文档，提取文本和结构化表格
        参考: core/judgment_vector_db/data_processing/pdf_extractor.py

        当文本层不足（扫描型PDF）且有parser_factory时，自动调用文档解析服务。
        """
        if not HAS_PDF:
            raise RuntimeError("未安装pdfplumber库，请运行: pip install pdfplumber")

        text = ""
        tables: list[dict[str, Any] = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    # 提取结构化表格
                    page_tables = page.extract_tables()
                    for table_data in page_tables:
                        if not table_data:
                            continue
                        # 过滤空表格
                        has_content = any(
                            any(cell and cell.strip() for cell in row)
                            for row in table_data
                        )
                        if not has_content:
                            continue
                        tables.append({
                            "table_index": len(tables) + 1,
                            "page_number": page_num,
                            "rows": [
                                [cell.strip() if cell else "" for cell in row]
                                for row in table_data
                            ],
                            "header": (
                                [cell.strip() if cell else "" for cell in table_data[0]
                                if table_data
                                else None
                            ),
                            "row_count": len(table_data),
                            "col_count": max(len(row) for row in table_data) if table_data else 0,
                            "source": "pdf",
                        })
        except Exception as e:
            logger.error(f"PDF解析失败: {e}")
            raise

        # 文本层不足 + 有parser_factory → 尝试MinerU/Tesseract
        if len(text.strip()) < 100 and self._parser_factory:
            logger.info("PDF文本层不足(%d字符)，尝试文档解析服务", len(text.strip()))
            try:
                from core.document_parser.base import ParseRequest
                from core.document_parser.parser_factory import run_async_safe

                factory = self._parser_factory
                request = ParseRequest(file_path=file_path)
                result = run_async_safe(factory.parse(request))
                if result.content and not result.error:
                    logger.info("文档解析服务成功: backend=%s", result.backend_used)
                    text = result.content
                    # 如果pdfplumber没提取到表格，使用解析服务的表格
                    if not tables and result.tables:
                        tables = result.tables
                else:
                    logger.warning("文档解析服务失败: %s", result.error)
            except Exception as e:
                logger.warning("文档解析服务异常: %s", e)

        return text, tables

    def _parse_word(self, file_path: str) -> tuple[str, list[dict[str, Any]]:
        """
        解析Word文档，提取文本和结构化表格
        参考: core/intelligence/dspy/production_docx_extractor.py
        """
        if not HAS_DOCX:
            raise RuntimeError("未安装python-docx库，请运行: pip install python-docx")

        text = ""
        tables: list[dict[str, Any] = []
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # 提取结构化表格（同时保留文本兼容性）
            for table in doc.tables:
                table_rows: list[list[str] = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_rows.append(row_data)
                    # 向后兼容：表格文本仍拼入text供正则提取
                    text += " ".join(row_data) + " "
                text += "\n"
                if table_rows:
                    tables.append({
                        "table_index": len(tables) + 1,
                        "page_number": 0,
                        "rows": table_rows,
                        "header": table_rows[0] if table_rows else None,
                        "row_count": len(table_rows),
                        "col_count": max(len(r) for r in table_rows),
                        "source": "word",
                    })
        except Exception as e:
            logger.error(f"Word解析失败: {e}")
            raise

        return text, tables

    def _parse_image(self, file_path: str) -> str:
        """
        解析图片（使用平台已有的OCR优化器或文档解析服务）
        参考: core/perception/chinese_ocr_optimizer.py

        优先使用parser_factory（MinerU），不可用时降级到本地OCR。
        """
        # 优先尝试parser_factory（MinerU）
        if self._parser_factory:
            try:
                from core.document_parser.base import ParseRequest
                from core.document_parser.parser_factory import run_async_safe

                factory = self._parser_factory
                request = ParseRequest(file_path=file_path)
                result = run_async_safe(factory.parse(request))
                if result.content and not result.error:
                    logger.info("图片解析服务成功: backend=%s", result.backend_used)
                    return result.content
            except Exception as e:
                logger.warning("图片解析服务异常，降级到本地OCR: %s", e)

        # 降级到本地OCR
        if not HAS_OCR or not self.ocr_optimizer:
            raise RuntimeError("OCR功能不可用，请检查ChineseOCROptimizer配置")

        try:
            # 使用平台已有的OCR优化器
            import cv2
            image = cv2.imread(file_path)

            # 调用OCR优化器（这里简化处理，实际可能需要更多配置）
            result = self.ocr_optimizer.process_image(image)

            if result and hasattr(result, 'text'):
                return result.text
            else:
                return ""
        except Exception as e:
            logger.error(f"OCR解析失败: {e}")
            raise

    def _parse_text(self, file_path: str) -> str:
        """解析纯文本"""
        try:
            with open(file_path, encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试GBK编码
            with open(file_path, encoding="gbk") as f:
                return f.read()

    def _extract_structured_info(self, text: str, source: str) -> ParsedOfficeAction:
        """
        从文本中提取结构化信息

        Args:
            text: 原始文本
            source: 文档来源

        Returns:
            解析后的审查意见
        """
        parsed = ParsedOfficeAction(
            oa_id=self._generate_oa_id(),
            document_source=source,
            raw_text=text,
        )

        # 提取目标专利信息（新增）
        parsed.target_application_number = self._extract_target_application_number(text)
        parsed.target_patent_title = self._extract_target_patent_title(text)
        parsed.target_applicant = self._extract_target_applicant(text)

        # 提取驳回类型
        parsed.rejection_type = self._extract_rejection_type(text)

        # 提取驳回理由
        parsed.rejection_reason = self._extract_rejection_reason(text)

        # 提取对比文件（增强）
        parsed.prior_art_references = self._extract_prior_art(text)
        parsed.prior_art_details = self._extract_prior_art_details(text)

        # 提取权利要求
        parsed.cited_claims = self._extract_cited_claims(text)

        # 提取审查员论点
        parsed.examiner_arguments = self._extract_examiner_arguments(text)

        # 提取缺失特征
        parsed.missing_features = self._extract_missing_features(text)

        # 提取时间信息
        parsed.received_date, parsed.response_deadline = self._extract_dates(text)

        # 计算置信度
        parsed.confidence = self._calculate_confidence(parsed)

        return parsed

    def _generate_oa_id(self) -> str:
        """生成审查意见ID"""
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        return f"OA_{timestamp}"

    def _extract_target_application_number(self, text: str) -> str:
        """
        提取目标专利申请号

        审查意见中通常包含目标专利的申请号
        """
        # 尝试多种模式
        patterns = [
            r"申请号\s*[:：]?\s*([0-9]{11,13}[.X]?)",  # 申请号: 202310000001
            r"申请号\s*[:：]?\s*(CN\d{9}[.X]?)",  # 申请号: CN202310000001
            r"针对\s*(CN\d{9}[.X]?)"  # 针对CN202310000001
        ]

        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1).strip()

        return ""

    def _extract_target_patent_title(self, text: str) -> str:
        """提取目标专利标题"""
        patterns = [
            r"发明名称\s*[:：]?\s*([^\n]{5,100})",
            r"专利名称\s*[:：]?\s*([^\n]{5,100})",
            r"申请.*?名称\s*[:：]?\s*([^\n]{5,100})",
        ]

        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                title = match.group(1).strip()
                # 清理多余字符
                title = re.sub(r"\s+", "", title)
                return title

        return ""

    def _extract_target_applicant(self, text: str) -> str:
        """提取目标专利申请人"""
        patterns = [
            r"申请人\s*[:：]?\s*([^\n]{2,50})",
            r"发明人\s*[:：]?\s*([^\n]{2,50})",
        ]

        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                applicant = match.group(1).strip()
                # 清理多余字符
                applicant = re.sub(r"\s+", "", applicant)
                return applicant

        return ""

    def _extract_prior_art_details(self, text: str) -> list[PriorArtReference]:
        """
        提取对比文件详细信息

        审查意见中通常包含对比文件的详细信息，如:
        - D1: CN112345678A
        - 公开日: 2023-01-01
        - 标题: xxx
        - 申请人: xxx
        """
        references = []

        # 分割成块，每个对比文件一个块
        # 查找对比文件标识 (D1, D2等)
        ref_blocks = re.split(r"(对比文件\s*[0-9A-D]+|对比文件[0-9A-D]|D[0-9]+)", text)

        current_ref_id = ""

        for i, block in enumerate(ref_blocks):
            # 如果是对比文件标识
            if re.match(r"对比文件\s*[0-9A-D]+|对比文件[0-9A-D]|D[0-9]+", block):
                current_ref_id = block.strip()
                continue

            # 如果不是块，跳过
            if i == 0:
                continue

            # 提取详细信息
            ref = PriorArtReference()
            ref.reference_id = current_ref_id

            # 公开号
            pub_match = re.search(r"(CN\d+[A-Z]?|US\d+|EP\d+[A-Z]?|WO\d{4}/\d+)", block)
            if pub_match:
                ref.publication_number = pub_match.group(1)

            # 公开日
            date_match = re.search(r"公开日\s*[:：]?\s*(\d{4}[-年]\d{1,2}[-月]\d{1,2})", block)
            if date_match:
                ref.publication_date = date_match.group(1)

            # 标题
            title_match = re.search(r"标题\s*[:：]?\s*([^\n]{5,100})", block)
            if title_match:
                ref.title = title_match.group(1).strip()
            else:
                # 尝试提取第一句话作为标题
                sentences = re.split(r"[。；;]", block)
                if sentences:
                    ref.title = sentences[0].strip()[:100]

            # 申请人
            applicant_match = re.search(r"申请人\s*[:：]?\s*([^\n]{2,50})", block)
            if applicant_match:
                ref.applicant = applicant_match.group(1).strip()

            # 相关权利要求
            claims_match = re.findall(r"权利要求(\d+)", block)
            if claims_match:
                ref.relevant_claims = [f"权利要求{c}" for c in claims_match]

            # 引用描述（提取该块的前200个字符）
            ref.description = block.strip()[:200]

            # 只有当有公开号或引用标识时才添加
            if ref.publication_number or ref.reference_id:
                references.append(ref)

                # 更新简单列表（兼容性）
                if ref.publication_number and ref.publication_number not in self.prior_art_references:
                    pass  # 这个会在主方法中处理

        return references

    def _extract_rejection_type(self, text: str) -> str:
        """提取驳回类型"""
        text_lower = text.lower()

        # 关键词匹配
        keywords = {
            "novelty": ["新颖性", "缺乏新颖性", "不具备新颖性", "novelty"],
            "inventiveness": ["创造性", "突出的实质性特点", "显而易见", "inventive", "inventive step"],
            "utility": ["实用性", "无法实现", "utility"],
            "clarity": ["不清楚", "模糊", "清晰度", "clarity"],
            "support": ["不支持", "说明书不支持", "support"],
            "unity": ["单一性", "不属于一个发明构思", "unity"],
        }

        for reject_type, kw_list in keywords.items():
            for kw in kw_list:
                if kw in text_lower:
                    return reject_type

        return ""

    def _extract_rejection_reason(self, text: str) -> str:
        """提取驳回理由"""
        # 查找包含驳回原因的段落
        patterns = [
            r"驳回理由[：:]\s*(.{10,500})",
            r"审查意见[：:]\s*(.{10,500})",
            r"结论[：:]\s*(.{10,500})",
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.DOTALL)
            if match:
                reason = match.group(1).strip()
                # 清理多余空白
                reason = re.sub(r"\s+", " ", reason)
                return reason[:500]  # 限制长度

        return ""

    def _extract_prior_art(self, text: str) -> list[str]:
        """
        提取对比文件（使用平台已有的正则模式）
        参考: core/intelligence/dspy/production_docx_extractor.py
        """
        references = []

        # 使用平台已有的专利号模式
        if "patent_number" in self.patterns:
            matches = self.patterns["patent_number"].findall(text)
            references.extend(matches)

        # 额外支持其他专利号格式
        additional_patterns = [
            r"[A-Z]{2}\d+[A-Z]?",  # CN123456A, US1234567B2
            r"US\d+",  # 美国专利
            r"EP\d+[A-Z]?",  # 欧洲专利
            r"WO\d{4}/\d+",  # PCT专利
            r"对比文件\s*([0-9A-D])",  # D1, D2等
        ]

        for pattern in additional_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            references.extend(matches)

        # 去重并返回
        return list(set(references))

    def _extract_cited_claims(self, text: str) -> list[int]:
        """提取被引用的权利要求"""
        claims = []

        # 匹配 "权利要求1", "claim 1" 等
        patterns = [
            r"权利要求(\d+)",
            r"权利要求[第](\d+)",
            r"claim\s*(\d+)",
            r"权\s*(\d+)",
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            claims.extend([int(m) for m in matches])

        # 去重并排序
        return sorted(set(claims))

    def _extract_examiner_arguments(self, text: str) -> list[str]:
        """提取审查员论点"""
        arguments = []

        # 查找包含论点的句子
        # 通常以 "审查员认为", "对比文件公开了" 等开头
        patterns = [
            r"(审查员认为[^。]*。)",
            r"(对比文件\d*[公开了|披露了][^。]*。)",
            r"(区别在于[^。]*。)",
            r"([^。]*不具备[^。]*。)",
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text)
            arguments.extend(matches)

        # 去重并限制数量
        unique_args = list(set(arguments))[:10]
        return unique_args

    def _extract_missing_features(self, text: str) -> list[str]:
        """提取缺失的技术特征"""
        features = []

        # 查找包含"未公开"、"未披露"等关键词的句子
        patterns = [
            r"([^。]*未公开[^。]*。)",
            r"([^。]*未披露[^。]*。)",
            r"([^。]*不存在[^。]*。)",
            r"区别技术特征[：:]\s*([^。]*。)",
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text)
            features.extend(matches)

        # 去重并限制数量
        unique_features = list(set(features))[:10]
        return unique_features

    def _extract_dates(self, text: str) -> tuple[str, str]:
        """提取日期信息"""
        received_date = ""
        deadline = ""

        # 匹配日期格式 (2024-01-01, 2024年1月1日等)
        date_patterns = [
            r"(\d{4}[-年]\d{1,2}[-月]\d{1,2}[日]?)",
            r"(\d{4}/\d{1,2}/\d{1,2})",
        ]

        all_dates = []
        for pattern in date_patterns:
            matches = re.findall(pattern, text)
            all_dates.extend(matches)

        # 第一个日期通常是收到日期
        if all_dates:
            received_date = all_dates[0]

        # 查找答复期限关键词
        deadline_patterns = [
            r"答复期限[：:]\s*(\d{4}[-年]\d{1,2}[-月]\d{1,2}[日]?)",
            r"请于[^。]*?(\d{4}[-年]\d{1,2}[-月]\d{1,2}[日]?)",
        ]

        for pattern in deadline_patterns:
            match = re.search(pattern, text)
            if match:
                deadline = match.group(1)
                break

        return received_date, deadline

    def _calculate_confidence(self, parsed: ParsedOfficeAction) -> float:
        """计算解析置信度"""
        score = 0.0
        total = 6.0

        # 每个字段都有得1分
        if parsed.rejection_type:
            score += 1.0
        if parsed.rejection_reason:
            score += 1.0
        if parsed.prior_art_references:
            score += 1.0
        if parsed.cited_claims:
            score += 1.0
        if parsed.examiner_arguments:
            score += 1.0
        if parsed.received_date or parsed.response_deadline:
            score += 1.0

        return score / total


# 全局单例
_oa_parser_instance = None


def get_oa_parser() -> OfficeActionParser:
    """获取审查意见解析器单例"""
    global _oa_parser_instance
    if _oa_parser_instance is None:
        _oa_parser_instance = OfficeActionParser()
    return _oa_parser_instance


# 测试代码
async def main():
    """测试审查意见文档解析器"""

    print("\n" + "=" * 60)
    print("📄 审查意见文档解析器测试")
    print("=" * 60 + "\n")

    parser = get_oa_parser()

    # 测试文本解析
    print("📝 测试: 文本解析")

    test_text = """
    审查意见通知书

    申请号: 202310000001.X
    收到日期: 2024-01-15
    答复期限: 2024-04-15

    驳回理由:
    权利要求1不具备新颖性。对比文件D1(CN112345678A)公开了权利要求1的全部技术特征，
    包括使用深度学习进行图像识别的核心步骤。

    审查员认为:
    1. D1公开了基于卷积神经网络的图像识别方法
    2. D1也使用了注意力机制
    3. 权利要求1与D1的区别仅在于网络结构参数

    对比文件:
    - D1: CN112345678A
    - D2: US2023001234A1

    结论: 权利要求1-3不具备新颖性。
    """

    # 模拟解析
    parsed = parser._extract_structured_info(test_text, "test_input.txt")

    print("✅ 解析完成\n")

    # 输出JSON
    print("📋 JSON格式:")
    print(parsed.to_json())
    print()

    # 输出Markdown
    print("📄 Markdown确认模板:")
    print(parsed.to_markdown())

    print("\n✅ 测试完成!")


# 入口点: @async_main装饰器已添加到main函数
