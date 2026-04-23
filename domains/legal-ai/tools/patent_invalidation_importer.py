#!/usr/bin/env python3
import json
import logging
import os
import re
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


def load_neo4j_config() -> dict[str, str]:
    base_dir = Path(__file__).resolve().parents[2]
    cfg_path = (
        base_dir
        / 'infrastructure'
        / 'config'
        / 'database'
        / 'knowledge_graph_config.json'
    )
    uri = os.getenv('KNOWLEDGE_GRAPH_NEO4J_URI', 'bolt://localhost:7687')
    user = os.getenv('KNOWLEDGE_GRAPH_NEO4J_USER', 'neo4j')
    password = os.getenv('KNOWLEDGE_GRAPH_NEO4J_PASSWORD', '')
    database = os.getenv('KNOWLEDGE_GRAPH_NEO4J_DATABASE', 'neo4j')
    if cfg_path.exists():
        try:
            with open(cfg_path, encoding='utf-8') as f:
                cfg = json.load(f)
                neo = cfg.get('knowledge_graph', {}).get('neo4j', {})
                uri = neo.get('uri', uri)
                user = neo.get('user', user)
                password = neo.get('password', password)
                database = neo.get('database', database)
        except Exception as e:
            # 记录异常但不中断流程
            logger.debug(f"[patent_invalidation_importer] Exception: {e}")
    return {'uri': uri, 'user': user, 'password': password, 'database': database}


CHN_NUMS = '一二三四五六七八九十零百千万1234567890'
DATE_CN_RE = re.compile(r"(\d{4})年(\d{1,2})月(\d{1,2})日")
CLAIM_RE = re.compile(rf"权利要求第([\d{CHN_NUMS}]+)项")
INVALID_RE = re.compile(r"(宣告|决定).*无效")
MAINTAIN_RE = re.compile(r"(维持|有效)")
PATENT_NO_RE = re.compile(
    r"(?:专利(?:号|申请号)|申请公布号|申请号|专利申请公布号)[：:：]?\s*([A-Za-z0-9\.\-]+)"
)
APP_NO_RE = re.compile(r"申请号[：:：]?\s*([A-Za-z0-9\.\-]+)")
PUB_NO_RE = re.compile(r"(?:申请公布号|公开号)[：:：]?\s*([A-Za-z0-9\.\-]+)")
GRANT_NO_RE = re.compile(r"(?:授权号|授权公告号|专利号)[：:：]?\s*([A-Za-z0-9\.\-]+)")
CASE_NO_RE = re.compile(r"(案件编号|案号)[：:]?\s*([A-Za-z0-9\-]+)")
PETITIONER_RE = re.compile(r"(请求人|无效宣告请求人|申请人)[：:]?\s*(.+)")
RESPONDENT_RE = re.compile(r"(专利权人|被请求人)[：:]?\s*(.+)")
LAW_CITE_RE = re.compile(
    rf"(专利法|专利审查指南|实施细则)[（(]?(?:第([{CHN_NUMS}]+)条)?[）)]?"
)


def read_docx_text(path: Path) -> list[str]:
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError('python-docx 未安装') from e
    doc = Document(str(path))
    lines: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    try:
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t = p.text.strip()
                        if t:
                            lines.append(t)
    except Exception as e:
        # 记录异常但不中断流程
        logger.debug(f"[patent_invalidation_importer] Exception: {e}")
    try:
        for sec in doc.sections:
            for p in sec.header.paragraphs:
                t = p.text.strip()
                if t:
                    lines.append(t)
            for p in sec.footer.paragraphs:
                t = p.text.strip()
                if t:
                    lines.append(t)
    except Exception as e:
        # 记录异常但不中断流程
        logger.debug(f"[patent_invalidation_importer] Exception: {e}")
    return lines


def convert_doc_to_docx(src: Path, out_dir: Path) -> Path | None:
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / (src.stem + '.docx')
    try:
        subprocess.run(
            ['textutil', '-convert', 'docx', str(src), '-output', str(out_path)],
            check=True,
        )
        return out_path if out_path.exists() else None
    except Exception:
        return None


def cn_num_to_int(s: str) -> int | None:
    try:
        if s.isdigit():
            return int(s)
        mapping = {
            '一': 1,
            '二': 2,
            '三': 3,
            '四': 4,
            '五': 5,
            '六': 6,
            '七': 7,
            '八': 8,
            '九': 9,
            '十': 10,
        }
        if s == '十':
            return 10
        if s.startswith('十'):
            return 10 + mapping.get(s[1], 0)
        if s.endswith('十'):
            return mapping.get(s[0], 0) * 10
        if len(s) == 2 and s[1] in mapping and s[0] in mapping:
            return mapping[s[0] * 10 + mapping[s[1]
    except Exception as e:
        # 记录异常但不中断流程
        logger.debug(f"[patent_invalidation_importer] Exception: {e}")
    return None


def clean_number(s: str) -> str:
    s = s.replace('FORMTEXT', '')
    s = re.sub(r"[_\s]+", '', s)
    s = re.sub(r"[^A-Za-z0-9\.\-]", '', s)
    return s


def extract_sections(lines: list[str]) -> dict[str, str]:
    heads = {
        'request': re.compile(
            r"^(?:[（(]?[一二三四五六七八九十]+[、\.．)]\s*)?请求事项[:：]?$"
        ),
        'review': re.compile(
            r"^(?:[（(]?[一二三四五六七八九十]+[、\.．)]\s*)?(审查要点|审查意见|审查判断)[:：]?$"
        ),
        'evidence': re.compile(
            r"^(?:[（(]?[一二三四五六七八九十]+[、\.．)]\s*)?证据(?:清单)?[:：]?$"
        ),
        'conclusion': re.compile(
            r"^(?:[（(]?[一二三四五六七八九十]+[、\.．)]\s*)?(结论|决定|审查结论|处理意见)[:：]?$"
        ),
    }
    idxs: dict[str, int] = {}
    for i, ln in enumerate(lines):
        t = ln.strip()
        for k, r in heads.items():
            if k not in idxs and (t.startswith('请求事项') or r.search(t)):
                if (k == 'request' and (t.startswith('请求事项') or r.search(t))) or (
                    k != 'request' and r.search(t)
                ):
                    idxs[k] = i
    order = [k for k in ['request', 'review', 'evidence', 'conclusion'] if k in idxs]
    ret: dict[str, str] = {}
    for n, k in enumerate(order):
        start = idxs[k] + 1
        end = len(lines) if n == len(order) - 1 else idxs[order[n + 1]
        content = "\n".join([x.strip() for x in lines[start:end]).strip()
        if content:
            ret[k] = content
    return ret


def extract_fields(lines: list[str], filename: str | None = None) -> dict[str, any]:
    fields = {
        'case_number': None,
        'decision_date': None,
        'patent_number': None,
        'petitioner': None,
        'respondent': None,
        'outcome': None,
        'invalidated_claims': [],
        'law_citations': [],
        'application_number': None,
        'publication_number': None,
        'grant_number': None,
        'sections': {},
    }
    for ln in lines[:300]:
        m = CASE_NO_RE.search(ln)
        if m:
            fields['case_number'] = m.group(2).strip()
        m = PATENT_NO_RE.search(ln)
        if m:
            fields['patent_number'] = clean_number(m.group(1).strip())
        m = APP_NO_RE.search(ln)
        if m:
            fields['application_number'] = clean_number(m.group(1).strip())
        m = PUB_NO_RE.search(ln)
        if m:
            fields['publication_number'] = clean_number(m.group(1).strip())
        m = GRANT_NO_RE.search(ln)
        if m:
            fields['grant_number'] = clean_number(m.group(1).strip())
        m = PETITIONER_RE.search(ln)
        if m:
            fields['petitioner'] = m.group(2).strip()
        m = RESPONDENT_RE.search(ln)
        if m:
            fields['respondent'] = m.group(2).strip()
        m = DATE_CN_RE.search(ln)
        if m and ('决定' in ln or '宣告' in ln or '审查' in ln):
            fields['decision_date'] = (
                f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
            )
    text = "\n".join(lines)
    if INVALID_RE.search(text):
        fields['outcome'] = 'invalidated'
    elif MAINTAIN_RE.search(text):
        fields['outcome'] = 'maintained'
    for m in CLAIM_RE.finditer(text):
        num = m.group(1)
        n = cn_num_to_int(num) or num
        if n not in fields['invalidated_claims']:
            fields['invalidated_claims'].append(n)
    for m in LAW_CITE_RE.finditer(text):
        fields['law_citations'].append({'law': m.group(1), 'article': m.group(2)})
    primary = (
        fields.get('grant_number')
        or fields.get('publication_number')
        or fields.get('application_number')
        or fields.get('patent_number')
    )
    if primary:
        fields['patent_number'] = primary
    fields['sections'] = extract_sections(lines)
    if not fields.get('patent_number') and filename:
        m = re.search(r"(CN)?[A-Za-z0-9]{8,}", filename)
        if m:
            fields['patent_number'] = clean_number(m.group(0))
    if not fields.get('case_number') and filename:
        m = re.search(r"\d{4,4}[\-]?[A-Za-z]?[A-Za-z0-9]+", filename)
        if m:
            fields['case_number'] = m.group(0)
    return fields


def upsert_case(
    session, case_id: str, title: str, fields: dict[str, any], source_path: str
):
    props = {
        'id': case_id,
        'entity_type': 'invalidation_case',
        'title': title,
        'case_number': fields.get('case_number'),
        'decision_date': fields.get('decision_date'),
        'outcome': fields.get('outcome'),
        'invalidated_claims': json.dumps(
            fields.get('invalidated_claims', []), ensure_ascii=False
        ),
        'kg_domain': 'professional',
        'source': 'patent_invalidation_docs',
        'path': source_path,
        'application_number': fields.get('application_number'),
        'publication_number': fields.get('publication_number'),
        'grant_number': fields.get('grant_number'),
        'created_at': datetime.now().isoformat(),
    }
    session.run(
        'MERGE (n:Entity {id:$id}) SET n.entity_type=$entity_type, n.title=$title, n.case_number=$case_number, n.decision_date=$decision_date, n.outcome=$outcome, n.invalidated_claims=$invalidated_claims, n.kg_domain=$kg_domain, n.source=$source, n.path=$path, n.application_number=$application_number, n.publication_number=$publication_number, n.grant_number=$grant_number, n.created_at=$created_at',
        **props,
    )


def link_entities(session, case_id: str, fields: dict[str, any]) -> None:
    primary = fields.get('patent_number')
    if primary:
        aliases = []
        for k in ('grant_number', 'publication_number', 'application_number'):
            v = fields.get(k)
            if v and v != primary:
                aliases.append(v)
        pprops = {
            'id': f"PAT::{primary}",
            'entity_type': 'patent',
            'title': primary,
            'kg_domain': 'professional',
            'application_number': fields.get('application_number'),
            'publication_number': fields.get('publication_number'),
            'grant_number': fields.get('grant_number'),
            'aliases': json.dumps(aliases, ensure_ascii=False),
            'created_at': datetime.now().isoformat(),
        }
        session.run(
            'MERGE (p:Entity {id:$id}) SET p.entity_type=$entity_type, p.title=$title, p.kg_domain=$kg_domain, p.application_number=$application_number, p.publication_number=$publication_number, p.grant_number=$grant_number, p.aliases=$aliases, p.created_at=$created_at',
            **pprops,
        )
        session.run(
            'MATCH (c:Entity {id:$cid}), (p:Entity {id:$pid}) MERGE (c)-[:INVOLVES]->(p)',
            cid=case_id,
            pid=pprops['id'],
        )
        for claim in fields.get('invalidated_claims', []):
            clid = f"CLAIM::{primary}::{claim}"
            session.run(
                "MERGE (cl:Entity {id:$id}) SET cl.entity_type='claim', cl.title=$title, cl.number=$num, cl.kg_domain='professional', cl.created_at=$ts",
                id=clid,
                title=f"权利要求第{claim}项",
                num=str(claim),
                ts=datetime.now().isoformat(),
            )
            session.run(
                'MATCH (c:Entity {id:$cid}), (cl:Entity {id:$clid}) MERGE (c)-[r:INVALIDATES]->(cl) SET r.created_at=$ts',
                cid=case_id,
                clid=clid,
                ts=datetime.now().isoformat(),
            )
    for role, pattern in (('PETITIONER', 'petitioner'), ('RESPONDENT', 'respondent')):
        name = fields.get(pattern)
        if name:
            pid = f"PARTY::{name}"
            session.run(
                "MERGE (o:Entity {id:$id}) SET o.entity_type='party', o.title=$title, o.kg_domain='professional', o.created_at=$ts",
                id=pid,
                title=name,
                ts=datetime.now().isoformat(),
            )
            session.run(
                f"MATCH (c:Entity {{id:$cid}}), (o:Entity {{id:$pid}}) MERGE (c)-[:{role}]->(o)",
                cid=case_id,
                pid=pid,
            )
    for cite in fields.get('law_citations', []):
        law = cite.get('law')
        art = cite.get('article')
        if law:
            res = session.run(
                "MATCH (l:Entity{entity_type:'law'}) WHERE l.title CONTAINS $law RETURN l.id as id, l.title as title LIMIT 1",
                law=law,
            )
            rec = res.single()
            if rec and art:
                ares = session.run(
                    "MATCH (a:Entity{entity_type:'legal_article'})-[:BELONGS_TO]->(l:Entity{id:$lid}) WHERE a.title CONTAINS $atitle RETURN a.id as id LIMIT 1",
                    lid=rec['id'],
                    atitle=f"第{art}条",
                )
                arec = ares.single()
                if arec:
                    session.run(
                        'MATCH (c:Entity {id:$cid}), (a:Entity {id:$aid}) MERGE (c)-[:CITES]->(a)',
                        cid=case_id,
                        aid=arec['id'],
                    )
    secs = fields.get('sections', {})
    for stype, content in secs.items():
        sid = f"SEC::{case_id}::{stype}"
        session.run(
            "MERGE (s:Entity {id:$id}) SET s.entity_type='section', s.title=$title, s.section_type=$stype, s.content=$content, s.kg_domain='professional', s.created_at=$ts",
            id=sid,
            title={
                'request': '请求事项',
                'review': '审查要点',
                'evidence': '证据',
                'conclusion': '结论',
            }.get(stype, stype),
            stype=stype,
            content=content,
            ts=datetime.now().isoformat(),
        )
        session.run(
            'MATCH (c:Entity {id:$cid}), (s:Entity {id:$sid}) MERGE (c)-[:HAS_SECTION]->(s)',
            cid=case_id,
            sid=sid,
        )


def import_invalidation_kg(base_dir: str, limit: int = 500) -> Any:
    cfg = load_neo4j_config()
    from neo4j import GraphDatabase

    driver = GraphDatabase.driver(cfg['uri'], auth=(cfg['user'], cfg['password']))
    src = Path(base_dir)
    converted_dir = Path(
        Path(__file__).resolve().parents[2] / 'data' / 'patent_invalidation_docx'
    )
    files: list[Path] = []
    for p in src.rglob('*'):
        if p.is_file() and (p.suffix.lower() in ('.doc', '.docx')):
            files.append(p)
            if len(files) >= limit:
                break
    count = 0
    stats = {
        'total': 0,
        'processed': 0,
        'failures': {'conversion_failed': 0, 'read_failed': 0, 'no_structure': 0},
    }
    failed_samples: list[dict[str, str] = []
    with driver.session(database=cfg['database']) as session:
        try:
            session.run(
                'CREATE CONSTRAINT entity_id_unique IF NOT EXISTS FOR (n:Entity) REQUIRE n.id IS UNIQUE'
            )
        except Exception as e:
            # 记录异常但不中断流程
            logger.debug(f"[patent_invalidation_importer] Exception: {e}")
        try:
            session.run(
                'CREATE INDEX entity_type_index IF NOT EXISTS FOR (n:Entity) ON (n.entity_type)'
            )
        except Exception as e:
            # 记录异常但不中断流程
            logger.debug(f"[patent_invalidation_importer] Exception: {e}")
        try:
            session.run(
                'CREATE INDEX entity_title_index IF NOT EXISTS FOR (n:Entity) ON (n.title)'
            )
        except Exception as e:
            # 记录异常但不中断流程
            logger.debug(f"[patent_invalidation_importer] Exception: {e}")
        stats['total'] = len(files)
        for f in files:
            docx_path = (
                f
                if f.suffix.lower() == '.docx'
                else convert_doc_to_docx(f, converted_dir)
            )
            if not docx_path:
                stats['failures']['conversion_failed'] += 1
                if len(failed_samples) < 20:
                    failed_samples.append(
                        {'path': str(f), 'reason': 'conversion_failed'}
                    )
                continue
            try:
                lines = read_docx_text(docx_path)
            except Exception:
                stats['failures']['read_failed'] += 1
                if len(failed_samples) < 20:
                    failed_samples.append({'path': str(f), 'reason': 'read_failed'})
                continue
            title = None
            for ln in lines[:10]:
                if not title and ln:
                    title = ln.strip()
            title = title or docx_path.stem
            fields = extract_fields(lines, filename=str(f.name))
            case_id = f"INV::{fields.get('case_number') or docx_path.stem}"[:128]
            if not fields.get('case_number') and not (
                fields.get('grant_number')
                or fields.get('publication_number')
                or fields.get('application_number')
                or fields.get('patent_number')
            ):
                stats['failures']['no_structure'] += 1
                if len(failed_samples) < 20:
                    failed_samples.append({'path': str(f), 'reason': 'no_structure'})
            upsert_case(session, case_id, title, fields, str(f))
            link_entities(session, case_id, fields)
            count += 1
            stats['processed'] = count
    driver.close()
    total_fail = sum(stats['failures'].values())
    fail_rate = (total_fail / stats['total']) if stats['total'] else 0.0
    return {
        'processed': count,
        'total': stats['total'],
        'failures': stats['failures'],
        'failure_rate': fail_rate,
        'failed_samples': failed_samples,
    }


def merge_patent_aliases(limit: int = 10000) -> Any:
    cfg = load_neo4j_config()
    from neo4j import GraphDatabase

    driver = GraphDatabase.driver(cfg['uri'], auth=(cfg['user'], cfg['password']))
    stats = {
        'patent_alias_nodes_deleted': 0,
        'involves_migrated': 0,
        'claims_merged': 0,
    }
    with driver.session(database=cfg['database']) as session:
        res = session.run(
            """
            MATCH (p:Entity {entity_type:'patent'})
            WHERE p.aliases IS NOT NULL OR p.application_number IS NOT NULL OR p.publication_number IS NOT NULL OR p.grant_number IS NOT NULL
            RETURN p.id as main_id, p.aliases as aliases, p.application_number as app, p.publication_number as pub, p.grant_number as grant
            LIMIT $lim
            """,
            lim=limit,
        )
        for rec in res:
            main_id = rec['main_id']
            aliases_raw = rec['aliases']
            app = rec['app']
            pub = rec['pub']
            grant = rec['grant']
            alias_set = set()
            try:
                if aliases_raw:
                    alias_set.update(json.loads(aliases_raw))
            except Exception as e:
                # 记录异常但不中断流程
                logger.debug(f"[patent_invalidation_importer] Exception: {e}")
            for v in (app, pub, grant):
                if v:
                    alias_set.add(v)
            alias_set = {a for a in alias_set if f"PAT::{a}" != main_id}
            for a in alias_set:
                alias_pat_id = f"PAT::{a}"
                session.run(
                    """
                    MATCH (c:Entity)-[r:INVOLVES]->(pa:Entity {id:$alias_pat_id})
                    MATCH (p:Entity {id:$main_id})
                    MERGE (c)-[r2:INVOLVES]->(p)
                    SET r2.created_at = coalesce(r.created_at, r2.created_at)
                    DELETE r
                    """,
                    alias_pat_id=alias_pat_id,
                    main_id=main_id,
                )
                cres = session.run(
                    'MATCH (c:Entity)-[:INVOLVES]->(p:Entity {id:$main_id}) RETURN count(c) as cnt',
                    main_id=main_id,
                ).single()
                if cres:
                    stats['involves_migrated'] += int(cres['cnt']) if cres['cnt'] else 0
                clres = session.run(
                    """
                    MATCH (cl:Entity {entity_type:'claim'})
                    WHERE cl.id STARTS WITH $prefix
                    RETURN cl.id as id, cl.number as num
                    """,
                    prefix=f"CLAIM::{a}::",
                )
                for clrec in clres:
                    alias_claim_id = clrec['id']
                    num = str(clrec['num']) if clrec['num'] is not None else ''
                    main_claim_id = f"CLAIM::{main_id.split('PAT::',1)[1]}::{num}"
                    session.run(
                        """
                        MATCH (c:Entity)-[r:INVALIDATES]->(cl:Entity {id:$alias_claim_id})
                        MERGE (tcl:Entity {id:$main_claim_id})
                        SET tcl.entity_type='claim', tcl.title=$title, tcl.number=$num, tcl.kg_domain='professional', tcl.created_at=coalesce(tcl.created_at, $now)
                        MERGE (c)-[r2:INVALIDATES]->(tcl)
                        SET r2.created_at = coalesce(r.created_at, r2.created_at)
                        DELETE r
                        """,
                        alias_claim_id=alias_claim_id,
                        main_claim_id=main_claim_id,
                        title=f"权利要求第{num}项" if num else "权利要求",
                        num=num,
                        now=datetime.now().isoformat(),
                    )
                    session.run(
                        'MATCH (cl:Entity {id:$alias_claim_id}) DETACH DELETE cl',
                        alias_claim_id=alias_claim_id,
                    )
                    stats['claims_merged'] += 1
                session.run(
                    'MATCH (pa:Entity {id:$alias_pat_id}) DETACH DELETE pa',
                    alias_pat_id=alias_pat_id,
                )
                stats['patent_alias_nodes_deleted'] += 1
    driver.close()
    return stats


def normalize_patent_number(s: str) -> str:
    s = (s or '').strip()
    s = s.replace('FORMTEXT', '')
    s = re.sub(r"[_\s]", '', s)
    s = s.upper()
    s = s.replace('ZL', '')
    s = re.sub(r"[()（）]", '', s)
    s = re.sub(r"\.(\d)$", r'\1', s)
    s = re.sub(r"[^A-Z0-9]", '', s)
    return s


def heuristic_merge_old_nodes(limit: int = 50000) -> Any:
    cfg = load_neo4j_config()
    from neo4j import GraphDatabase

    driver = GraphDatabase.driver(cfg['uri'], auth=(cfg['user'], cfg['password']))
    stats = {
        'duplicate_groups': 0,
        'patent_nodes_deleted': 0,
        'involves_migrated': 0,
        'claims_merged': 0,
    }
    with driver.session(database=cfg['database']) as session:
        res = session.run(
            """
            MATCH (p:Entity {entity_type:'patent'})
            RETURN p.id as id, p.title as title, p.application_number as app, p.publication_number as pub, p.grant_number as grant
            LIMIT $lim
            """,
            lim=limit,
        )
        by_norm = {}
        nodes = []
        for rec in res:
            pid = rec['id']
            title = rec['title'] or ''
            app = rec['app']
            pub = rec['pub']
            grant = rec['grant']
            norm = normalize_patent_number(title or pid.split('PAT::', 1)[1])
            nodes.append((pid, title, app, pub, grant, norm))
            by_norm.setdefault(norm, []).append(pid)
        for norm, ids in by_norm.items():
            if len(ids) <= 1:
                continue
            main_id = sorted(ids, key=lambda x: len(x))[-1]
            stats['duplicate_groups'] += 1
            for alias_id in ids:
                if alias_id == main_id:
                    continue
                session.run(
                    """
                    MATCH (c:Entity)-[r:INVOLVES]->(pa:Entity {id:$alias_id})
                    MATCH (p:Entity {id:$main_id})
                    MERGE (c)-[r2:INVOLVES]->(p)
                    SET r2.created_at = coalesce(r.created_at, r2.created_at)
                    DELETE r
                    """,
                    alias_id=alias_id,
                    main_id=main_id,
                )
                cres = session.run(
                    'MATCH (c:Entity)-[:INVOLVES]->(p:Entity {id:$main_id}) RETURN count(c) as cnt',
                    main_id=main_id,
                ).single()
                if cres and cres['cnt']:
                    stats['involves_migrated'] += int(cres['cnt']) or 0
                alias_num = alias_id.split('PAT::', 1)[1]
                clres = session.run(
                    """
                    MATCH (cl:Entity {entity_type:'claim'})
                    WHERE cl.id STARTS WITH $prefix
                    RETURN cl.id as id, cl.number as num
                    """,
                    prefix=f"CLAIM::{alias_num}::",
                )
                for clrec in clres:
                    alias_claim_id = clrec['id']
                    num = str(clrec['num']) if clrec['num'] is not None else ''
                    main_num = main_id.split('PAT::', 1)[1]
                    main_claim_id = f"CLAIM::{main_num}::{num}"
                    session.run(
                        """
                        MATCH (c:Entity)-[r:INVALIDATES]->(cl:Entity {id:$alias_claim_id})
                        MERGE (tcl:Entity {id:$main_claim_id})
                        SET tcl.entity_type='claim', tcl.title=$title, tcl.number=$num, tcl.kg_domain='professional', tcl.created_at=coalesce(tcl.created_at, $now)
                        MERGE (c)-[r2:INVALIDATES]->(tcl)
                        SET r2.created_at = coalesce(r.created_at, r2.created_at)
                        DELETE r
                        """,
                        alias_claim_id=alias_claim_id,
                        main_claim_id=main_claim_id,
                        title=f"权利要求第{num}项" if num else "权利要求",
                        num=num,
                        now=datetime.now().isoformat(),
                    )
                    session.run(
                        'MATCH (cl:Entity {id:$alias_claim_id}) DETACH DELETE cl',
                        alias_claim_id=alias_claim_id,
                    )
                    stats['claims_merged'] += 1
                session.run(
                    'MATCH (pa:Entity {id:$alias_id}) DETACH DELETE pa',
                    alias_id=alias_id,
                )
                stats['patent_nodes_deleted'] += 1
    driver.close()
    return stats


def graph_statistics() -> Any:
    cfg = load_neo4j_config()
    from neo4j import GraphDatabase

    driver = GraphDatabase.driver(cfg['uri'], auth=(cfg['user'], cfg['password']))
    with driver.session(database=cfg['database']) as session:
        cases = session.run(
            "MATCH (n:Entity {entity_type:'invalidation_case'}) RETURN count(n) as c"
        ).single()['c']
        patents = session.run(
            "MATCH (n:Entity {entity_type:'patent'}) RETURN count(n) as c"
        ).single()['c']
        claims = session.run(
            "MATCH (n:Entity {entity_type:'claim'}) RETURN count(n) as c"
        ).single()['c']
        sections = session.run(
            "MATCH (n:Entity {entity_type:'section'}) RETURN count(n) as c"
        ).single()['c']
        involves = session.run(
            'MATCH ()-[r:INVOLVES]->() RETURN count(r) as c'
        ).single()['c']
        invalidates = session.run(
            'MATCH ()-[r:INVALIDATES]->() RETURN count(r) as c'
        ).single()['c']
        cites = session.run('MATCH ()-[r:CITES]->() RETURN count(r) as c').single()['c']
        has_section = session.run(
            'MATCH ()-[r:HAS_SECTION]->() RETURN count(r) as c'
        ).single()['c']
        covered_cases = session.run(
            "MATCH (c:Entity {entity_type:'invalidation_case'})-[:INVOLVES]->() RETURN count(distinct c) as c"
        ).single()['c']
        covered_ratio = float(covered_cases) / float(cases) if cases else 0.0
    driver.close()
    return {
        'nodes': {
            'cases': cases,
            'patents': patents,
            'claims': claims,
            'sections': sections,
        },
        'relationships': {
            'INVOLVES': involves,
            'INVALIDATES': invalidates,
            'CITES': cites,
            'HAS_SECTION': has_section,
        },
        'coverage': {'cases_with_patent': covered_cases, 'ratio': covered_ratio},
    }


def main() -> None:
    if os.environ.get('INVALIDATION_MERGE_ALIASES') == '1':
        res = merge_patent_aliases()
        print(json.dumps(res, ensure_ascii=False))
        return
    if os.environ.get('INVALIDATION_HEURISTIC_MERGE') == '1':
        res = heuristic_merge_old_nodes()
        print(json.dumps(res, ensure_ascii=False))
        return
    if os.environ.get('INVALIDATION_STATS') == '1':
        res = graph_statistics()
        print(json.dumps(res, ensure_ascii=False))
        return
    base_dir = os.environ.get(
        'INVALIDATION_BASE_DIR', '/Users/xujian/学习资料/语料/专利无效宣告原文'
    )
    max_files = int(os.environ.get('INVALIDATION_MAX_FILES', '500'))
    res = import_invalidation_kg(base_dir, limit=max_files)
    print(json.dumps(res, ensure_ascii=False))


if __name__ == '__main__':
    main()
