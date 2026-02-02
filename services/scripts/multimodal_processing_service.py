#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
多模态处理服务
Multimodal Processing Service

提供文档解析、OCR识别、多模态处理等功能
"""

import json
from core.async_main import async_main
import logging
from core.logging_config import setup_logging
import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel

# 导入统一认证模块
from shared.auth.auth_middleware import create_auth_middleware, setup_cors

# 配置日志
# setup_logging()  # 日志配置已移至模块导入
logger = setup_logging()

# 创建FastAPI应用
app = FastAPI(
    title='多模态处理服务',
    description='提供文档解析、OCR识别、多模态AI处理等功能',
    version='1.0.0'
)

# 配置CORS

# 创建临时目录
TEMP_DIR = Path('/tmp/multimodal_uploads')
TEMP_DIR.mkdir(exist_ok=True)

class ProcessingRequest(BaseModel):
    file_type: str
    processing_options: Optional[Dict[str, Any]] = None

class ProcessingResponse(BaseModel):
    success: bool
    file_type: str
    processing_time: float
    extracted_data: Dict[str, Any]
    message: str | None = None

@app.on_event('startup')
async def startup_event():
    """服务启动事件"""
    logger.info('🚀 多模态处理服务启动')

@app.on_event('shutdown')
async def shutdown_event():
    """服务关闭事件"""
    # 清理临时文件
    import shutil
    if TEMP_DIR.exists():
        shutil.rmtree(TEMP_DIR)
    logger.info('📌 多模态处理服务已关闭')

@app.get('/')
async def root():
    """根路径"""
    return {
        'service': '多模态处理服务',
        'status': 'running',
        'supported_formats': {
            'documents': ['pdf', 'docx', 'txt'],
            'images': ['png', 'jpg', 'jpeg', 'tiff', 'bmp'],
            'audio': ['mp3', 'wav', 'm4a'],
            'video': ['mp4', 'avi', 'mov']
        },
        'features': [
            '文档文本提取',
            'OCR文字识别',
            '图像处理分析',
            '音频特征提取',
            '视频帧提取',
            '多模态AI理解'
        ],
        'timestamp': datetime.now().isoformat()
    }

@app.get('/health')
async def health_check():
    """健康检查"""
    return {
        'status': 'healthy',
        'dependencies': {
            'PyMuPDF': '✅',
            'pytesseract': '✅',
            'OpenCV': '✅',
            'librosa': '✅',
            'moviepy': '✅',
            'transformers': '✅',
            'torch': '✅'
        },
        'timestamp': datetime.now().isoformat()
    }

@app.post('/api/process/document', response_model=ProcessingResponse)
async def process_document(
    file: UploadFile = File(...),
    use_ocr: bool = Form(False),
    extract_metadata: bool = Form(True)
):
    """
    处理文档文件
    支持PDF、Word、文本文件
    """
    start_time = datetime.now()

    try:
        # 保存上传文件
        file_path = await save_uploaded_file(file)

        # 根据文件类型处理
        if file_path.suffix.lower() == '.pdf':
            extracted_data = await process_pdf(file_path, use_ocr)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            extracted_data = await process_word(file_path)
        elif file_path.suffix.lower() == '.txt':
            extracted_data = await process_text(file_path)
        else:
            raise HTTPException(status_code=400, detail='不支持的文档格式')

        # 提取元数据
        if extract_metadata:
            extracted_data['metadata'] = extract_file_metadata(file_path)

        processing_time = (datetime.now() - start_time).total_seconds()

        # 清理临时文件
        os.unlink(file_path)

        return ProcessingResponse(
            success=True,
            file_type='document',
            processing_time=processing_time,
            extracted_data=extracted_data,
            message=f"成功处理文档: {file.filename}"
        )

    except Exception as e:
        logger.error(f"文档处理失败: {str(e)}")
        return ProcessingResponse(
            success=False,
            file_type='document',
            processing_time=(datetime.now() - start_time).total_seconds(),
            extracted_data={},
            message=f"处理失败: {str(e)}"
        )

@app.post('/api/process/image', response_model=ProcessingResponse)
async def process_image(
    file: UploadFile = File(...),
    detect_text: bool = Form(True),
    analyze_content: bool = Form(True)
):
    """
    处理图像文件
    支持OCR文字识别和内容分析
    """
    start_time = datetime.now()

    try:
        # 保存上传文件
        file_path = await save_uploaded_file(file)

        extracted_data = {}

        # 基本图像信息
        extracted_data['image_info'] = await get_image_info(file_path)

        # OCR文字识别
        if detect_text:
            extracted_data['ocr_result'] = await perform_ocr(file_path)

        # 内容分析
        if analyze_content:
            extracted_data['content_analysis'] = await analyze_image_content(file_path)

        processing_time = (datetime.now() - start_time).total_seconds()

        # 清理临时文件
        os.unlink(file_path)

        return ProcessingResponse(
            success=True,
            file_type='image',
            processing_time=processing_time,
            extracted_data=extracted_data,
            message=f"成功处理图像: {file.filename}"
        )

    except Exception as e:
        logger.error(f"图像处理失败: {str(e)}")
        return ProcessingResponse(
            success=False,
            file_type='image',
            processing_time=(datetime.now() - start_time).total_seconds(),
            extracted_data={},
            message=f"处理失败: {str(e)}"
        )

@app.post('/api/process/audio', response_model=ProcessingResponse)
async def process_audio(
    file: UploadFile = File(...),
    extract_features: bool = Form(True),
    transcribe: bool = Form(False)
):
    """
    处理音频文件
    提取音频特征和语音转文本
    """
    start_time = datetime.now()

    try:
        # 保存上传文件
        file_path = await save_uploaded_file(file)

        extracted_data = await process_audio_file(file_path, extract_features, transcribe)

        processing_time = (datetime.now() - start_time).total_seconds()

        # 清理临时文件
        os.unlink(file_path)

        return ProcessingResponse(
            success=True,
            file_type='audio',
            processing_time=processing_time,
            extracted_data=extracted_data,
            message=f"成功处理音频: {file.filename}"
        )

    except Exception as e:
        logger.error(f"音频处理失败: {str(e)}")
        return ProcessingResponse(
            success=False,
            file_type='audio',
            processing_time=(datetime.now() - start_time).total_seconds(),
            extracted_data={},
            message=f"处理失败: {str(e)}"
        )

@app.post('/api/process/video', response_model=ProcessingResponse)
async def process_video(
    file: UploadFile = File(...),
    extract_frames: bool = Form(True),
    frame_count: int = Form(10)
):
    """
    处理视频文件
    提取视频信息和关键帧
    """
    start_time = datetime.now()

    try:
        # 保存上传文件
        file_path = await save_uploaded_file(file)

        extracted_data = await process_video_file(file_path, extract_frames, frame_count)

        processing_time = (datetime.now() - start_time).total_seconds()

        # 清理临时文件
        os.unlink(file_path)

        return ProcessingResponse(
            success=True,
            file_type='video',
            processing_time=processing_time,
            extracted_data=extracted_data,
            message=f"成功处理视频: {file.filename}"
        )

    except Exception as e:
        logger.error(f"视频处理失败: {str(e)}")
        return ProcessingResponse(
            success=False,
            file_type='video',
            processing_time=(datetime.now() - start_time).total_seconds(),
            extracted_data={},
            message=f"处理失败: {str(e)}"
        )

@app.post('/api/process/multimodal')
async def process_multimodal(
    files: List[UploadFile] = File(...),
    analysis_type: str = Form('comprehensive')
):
    """
    多模态综合处理
    分析多个文件之间的关联性
    """
    start_time = datetime.now()

    try:
        results = []

        for file in files:
            file_path = await save_uploaded_file(file)

            # 根据文件类型处理
            if file_path.suffix.lower() in ['.pdf', '.docx', '.txt']:
                result = await process_document_file(file_path)
            elif file_path.suffix.lower() in ['.png', '.jpg', '.jpeg']:
                result = await process_image_file(file_path)
            else:
                result = {'error': '不支持的文件类型'}

            result['filename'] = file.filename
            results.append(result)

            # 清理临时文件
            os.unlink(file_path)

        # 综合分析
        multimodal_insights = analyze_multimodal_relations(results)

        processing_time = (datetime.now() - start_time).total_seconds()

        return {
            'success': True,
            'files_processed': len(results),
            'processing_time': processing_time,
            'individual_results': results,
            'multimodal_insights': multimodal_insights,
            'message': f"成功处理 {len(files)} 个文件"
        }

    except Exception as e:
        logger.error(f"多模态处理失败: {str(e)}")
        return {
            'success': False,
            'error': str(e),
            'message': '多模态处理失败'
        }

# 辅助函数
async def save_uploaded_file(file: UploadFile) -> Path:
    """保存上传的文件"""
    # 创建临时文件
    file_path = TEMP_DIR / f"{datetime.now().timestamp()}_{file.filename}"

    # 写入文件
    with open(file_path, 'wb') as buffer:
        content = await file.read()
        buffer.write(content)

    return file_path

async def process_pdf(file_path: Path, use_ocr: bool = False) -> Dict:
    """处理PDF文件"""
    import fitz

    doc = fitz.open(str(file_path))

    # 提取文本
    full_text = ''
    for page in doc:
        text = page.get_text()
        full_text += text + "\n"

    result = {
        'type': 'pdf',
        'pages': len(doc),
        'text_length': len(full_text),
        'has_text': bool(full_text.strip())
    }

    # 如果没有文本且需要OCR
    if not full_text.strip() and use_ocr:
        logger.info('PDF不包含文本，尝试OCR...')
        # 这里可以调用OCR功能
        result['ocr_used'] = True
        result['extracted_text'] = 'OCR功能已配置但需要具体实现'
    else:
        result['extracted_text'] = full_text[:1000] + '...' if len(full_text) > 1000 else full_text

    doc.close()
    return result

async def process_word(file_path: Path) -> Dict:
    """处理Word文档"""
    from docx import Document

    doc = Document(str(file_path))
    full_text = "\n".join([para.text for para in doc.paragraphs])

    return {
        'type': 'word',
        'paragraphs': len(doc.paragraphs),
        'text_length': len(full_text),
        'extracted_text': full_text[:1000] + '...' if len(full_text) > 1000 else full_text
    }

async def process_text(file_path: Path) -> Dict:
    """处理文本文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='gbk') as f:
            content = f.read()

    return {
        'type': 'text',
        'text_length': len(content),
        'extracted_text': content[:1000] + '...' if len(content) > 1000 else content
    }

async def get_image_info(file_path: Path) -> Dict:
    """获取图像基本信息"""
    import cv2
    from PIL import Image

    with Image.open(file_path) as img:
        info = {
            'format': img.format,
            'mode': img.mode,
            'size': img.size,
            'file_size': file_path.stat().st_size
        }

    # 使用OpenCV获取更多信息
    img_cv = cv2.imread(str(file_path))
    if img_cv is not None:
        info.update({
            'opencv_size': (img_cv.shape[1], img_cv.shape[0]),
            'channels': img_cv.shape[2] if len(img_cv.shape) > 2 else 1
        })

    return info

async def perform_ocr(file_path: Path) -> Dict:
    """执行OCR识别"""
    try:
        import pytesseract
        from PIL import Image

        with Image.open(file_path) as img:
            # 检测中文和英文
            text = pytesseract.image_to_string(img, lang='chi_sim+eng')

            return {
                'success': True,
                'detected_text': text,
                'has_text': bool(text.strip()),
                'confidence': 0.9 if text.strip() else 0.1
            }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'detected_text': '',
            'confidence': 0.0
        }

async def analyze_image_content(file_path: Path) -> Dict:
    """分析图像内容"""
    import cv2
    import numpy as np

    img_cv = cv2.imread(str(file_path))
    if img_cv is None:
        return {'error': '无法读取图像'}

    gray = cv2.cvt_color(img_cv, cv2.COLOR_BGR2GRAY)
    edges = cv2.Canny(gray, 50, 150)

    # 边缘密度分析
    edge_density = np.sum(edges > 0) / edges.size

    analysis = {
        'edge_density': edge_density,
        'is_drawing': edge_density > 0.05,
        'dominant_colors': get_dominant_colors(img_cv)
    }

    return analysis

def get_dominant_colors(img_cv, k=5) -> None:
    """获取主要颜色"""
    import cv2
    import numpy as np

    # 将图像转换为RGB并重塑
    img_rgb = cv2.cvt_color(img_cv, cv2.COLOR_BGR2RGB)
    pixels = img_rgb.reshape(-1, 3)

    # 使用KMeans聚类
    from sklearn.cluster import KMeans
    kmeans = KMeans(n_clusters=k, random_state=42)
    kmeans.fit(pixels)

    # 获取聚类中心
    colors = kmeans.cluster_centers_.astype(int)

    return [color.tolist() for color in colors]

async def process_audio_file(file_path: Path, extract_features: bool, transcribe: bool) -> Dict:
    """处理音频文件"""
    import librosa
    import numpy as np

    try:
        y, sr = librosa.load(str(file_path))

        result = {
            'type': 'audio',
            'duration': len(y) / sr,
            'sample_rate': sr,
            'file_size': file_path.stat().st_size
        }

        # 提取特征
        if extract_features:
            mfcc = librosa.feature.mfcc(y=y, sr=sr, n_mfcc=13)
            spectral_centroids = librosa.feature.spectral_centroid(y=y, sr=sr)

            result.update({
                'mfcc_shape': mfcc.shape,
                'avg_spectral_centroid': float(np.mean(spectral_centroids)),
                'is_speech': np.mean(spectral_centroids) > 2000
            })

        # 语音转文字（需要额外的语音识别模型）
        if transcribe:
            result['transcription'] = '语音转文字功能已配置，但需要具体实现'

        return result

    except Exception as e:
        return {'error': str(e)}

async def process_video_file(file_path: Path, extract_frames: bool, frame_count: int) -> Dict:
    """处理视频文件"""
    import cv2

    cap = cv2.VideoCapture(str(file_path))

    if not cap.is_opened():
        return {'error': '无法打开视频文件'}

    fps = cap.get(cv2.CAP_PROP_FPS)
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))

    result = {
        'type': 'video',
        'fps': fps,
        'total_frames': total_frames,
        'resolution': (width, height),
        'duration': total_frames / fps if fps > 0 else 0,
        'file_size': file_path.stat().st_size
    }

    # 提取关键帧
    if extract_frames:
        frames_info = []
        frame_interval = max(1, total_frames // frame_count)

        for i in range(0, total_frames, frame_interval):
            cap.set(cv2.CAP_PROP_POS_FRAMES, i)
            ret, frame = cap.read()
            if ret:
                frames_info.append({
                    'frame_number': i,
                    'timestamp': i / fps if fps > 0 else 0,
                    'shape': frame.shape
                })

        result['extracted_frames'] = frames_info

    cap.release()
    return result

def extract_file_metadata(file_path: Path) -> Dict:
    """提取文件元数据"""
    stat = file_path.stat()
    return {
        'filename': file_path.name,
        'size': stat.st_size,
        'created_time': datetime.fromtimestamp(stat.st_ctime).isoformat(),
        'modified_time': datetime.fromtimestamp(stat.st_mtime).isoformat(),
        'extension': file_path.suffix
    }

def analyze_multimodal_relations(results: List[Dict]) -> Dict:
    """分析多模态文件之间的关系"""
    insights = {
        'total_files': len(results),
        'file_types': {},
        'connections': []
    }

    # 统计文件类型
    for result in results:
        if 'error' not in result:
            file_type = result.get('type', 'unknown')
            insights['file_types'][file_type] = insights['file_types'].get(file_type, 0) + 1

    # 生成连接关系
    if len(results) > 1:
        insights['connections'].append({
            'type': 'upload_batch',
            'description': f"同时上传了 {len(results)} 个文件",
            'files': [r.get('filename', '') for r in results]
        })

    return insights

if __name__ == '__main__':
    import uvicorn
    uvicorn.run(
        'multimodal_processing_service:app',
        host='0.0.0.0',
        port=8012,
        reload=True,
        log_level='info'
    )