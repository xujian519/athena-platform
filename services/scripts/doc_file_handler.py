#!/usr/bin/env python3
"""
DOC文件处理模块
DOC File Handler for Athena Multimodal System
支持传统Microsoft Word DOC格式的文本提取和转换
"""

import logging
import os
import shutil
import tempfile
from pathlib import Path
from typing import Any

from core.logging_config import setup_logging

# 导入DOC处理相关库
try:
    import doc2docx
    DOC2DOCX_AVAILABLE = True
except ImportError:
    DOC2DOCX_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import olefile
    OLEFILE_AVAILABLE = True
except ImportError:
    OLEFILE_AVAILABLE = False

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = setup_logging()

class DOCFileHandler:
    """DOC文件处理器"""

    def __init__(self):
        self.doc2docx_available = DOC2DOCX_AVAILABLE
        self.docx_available = DOCX_AVAILABLE
        self.olefile_available = OLEFILE_AVAILABLE

        logger.info("DOC处理器初始化:")
        logger.info(f"  - doc2docx: {'✓' if self.doc2docx_available else '✗'}")
        logger.info(f"  - python-docx: {'✓' if self.docx_available else '✗'}")
        logger.info(f"  - olefile: {'✓' if self.olefile_available else '✗'}")

    def is_supported_format(self, file_path: str) -> bool:
        """检查文件格式是否支持"""
        ext = Path(file_path).suffix.lower()
        return ext == '.doc'

    async def extract_text_from_doc(self, doc_path: str, options: dict[str, Any] = None) -> dict[str, Any]:
        """
        从DOC文件中提取文本

        Args:
            doc_path: DOC文件路径
            options: 处理选项

        Returns:
            包含提取结果的字典
        """
        if not os.path.exists(doc_path):
            return {
                'success': False,
                'error': '文件不存在',
                'file_path': doc_path
            }

        if not self.is_supported_format(doc_path):
            return {
                'success': False,
                'error': f'不支持的文件格式: {Path(doc_path).suffix}',
                'file_path': doc_path
            }

        try:
            # 方法1: 使用doc2docx转换后处理
            if self.doc2docx_available and self.docx_available:
                return await self._extract_via_doc2docx(doc_path, options)

            # 方法2: 使用olefile直接提取
            elif self.olefile_available:
                return await self._extract_via_olefile(doc_path, options)

            else:
                return {
                    'success': False,
                    'error': '缺少必要的处理库，请安装doc2docx或olefile',
                    'file_path': doc_path,
                    'available_libraries': {
                        'doc2docx': self.doc2docx_available,
                        'python-docx': self.docx_available,
                        'olefile': self.olefile_available
                    }
                }

        except Exception as e:
            logger.error(f"DOC文件处理失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_path': doc_path,
                'exception_type': type(e).__name__
            }

    async def _extract_via_doc2docx(self, doc_path: str, options: dict[str, Any] = None) -> dict[str, Any]:
        """使用doc2docx转换并提取文本"""
        temp_docx_path = None
        try:
            # 创建临时目录
            temp_dir = tempfile.mkdtemp(prefix='athena_doc_')

            # 转换DOC到DOCX
            logger.info(f"正在转换DOC文件: {doc_path}")
            doc2docx.convert(doc_path, temp_dir)

            # 查找转换后的DOCX文件
            temp_docx_path = None
            for file in os.listdir(temp_dir):
                if file.endswith('.docx'):
                    temp_docx_path = os.path.join(temp_dir, file)
                    break

            if not temp_docx_path:
                raise Exception('DOCX转换失败，未找到转换后的文件')

            # 使用python-docx提取文本
            logger.info(f"正在提取DOCX文本: {temp_docx_path}")
            doc = docx.Document(temp_docx_path)

            # 提取段落文本
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)

            # 提取表格文本
            tables_text = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))
                if table_text:
                    tables_text.append('\n'.join(table_text))

            # 合并所有文本
            full_text = '\n\n'.join(paragraphs)
            if tables_text:
                full_text += '\n\n[表格内容]\n' + '\n\n'.join(tables_text)

            # 提取文档元数据
            metadata = {
                'paragraphs_count': len(paragraphs),
                'tables_count': len(doc.tables),
                'file_size': os.path.getsize(doc_path),
                'conversion_method': 'doc2docx'
            }

            # 检查文档属性
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                metadata['title'] = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                metadata['author'] = doc.core_properties.author

            return {
                'success': True,
                'text_content': full_text,
                'paragraphs': paragraphs,
                'tables': tables_text,
                'metadata': metadata,
                'file_path': doc_path,
                'extraction_method': 'doc2docx_conversion'
            }

        except Exception as e:
            logger.error(f"doc2docx转换失败: {e}")
            raise e
        finally:
            # 清理临时文件
            if temp_docx_path and os.path.exists(temp_docx_path):
                try:
                    os.unlink(temp_docx_path)
                except Exception as e:
                    logger.error(f"Error: {e}", exc_info=True)
            if 'temp_dir' in locals() and os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir)
                except Exception as e:
                    logger.error(f"Error: {e}", exc_info=True)

    async def _extract_via_olefile(self, doc_path: str, options: dict[str, Any] = None) -> dict[str, Any]:
        """使用olefile直接提取DOC文件文本"""
        try:
            if not olefile.is_ole_file(doc_path):
                raise Exception('文件不是有效的OLE格式')

            logger.info(f"使用olefile解析DOC文件: {doc_path}")

            with olefile.OleFileIO(doc_path) as ole:
                # 尝试提取WordDocument流
                if ole.exists('WordDocument'):
                    # 这是一个简化的实现
                    # 实际的DOC文件解析非常复杂，需要理解Microsoft Word文件格式

                    # 获取文件基本信息
                    metadata = {
                        'file_size': os.path.getsize(doc_path),
                        'ole_streams': ole.listdir(),
                        'extraction_method': 'olefile_basic'
                    }

                    # 由于DOC格式的复杂性，这里只返回基本信息
                    # 建议使用doc2docx进行完整转换
                    return {
                        'success': True,
                        'text_content': '',
                        'paragraphs': [],
                        'tables': [],
                        'metadata': metadata,
                        'file_path': doc_path,
                        'extraction_method': 'olefile_basic',
                        'warning': 'olefile只能提供基本文件信息，建议安装doc2docx库进行完整文本提取'
                    }
                else:
                    raise Exception('文件中没有找到WordDocument流')

        except Exception as e:
            logger.error(f"olefile提取失败: {e}")
            raise e

    async def convert_doc_to_docx(self, doc_path: str, output_dir: str | None = None) -> dict[str, Any]:
        """
        将DOC文件转换为DOCX格式

        Args:
            doc_path: DOC文件路径
            output_dir: 输出目录，如果为None则使用临时目录

        Returns:
            包含转换结果的字典
        """
        if not self.doc2docx_available:
            return {
                'success': False,
                'error': 'doc2docx库不可用，无法进行转换'
            }

        try:
            # 确定输出目录
            if output_dir is None:
                output_dir = tempfile.mkdtemp(prefix='athena_doc_convert_')
            elif not os.path.exists(output_dir):
                os.makedirs(output_dir)

            # 转换文件
            logger.info(f"正在转换DOC到DOCX: {doc_path} -> {output_dir}")
            doc2docx.convert(doc_path, output_dir)

            # 查找转换后的文件
            converted_files = []
            for file in os.listdir(output_dir):
                if file.endswith('.docx'):
                    converted_files.append(os.path.join(output_dir, file))

            if not converted_files:
                raise Exception('转换完成但未找到输出文件')

            return {
                'success': True,
                'converted_files': converted_files,
                'output_directory': output_dir,
                'original_file': doc_path,
                'conversion_method': 'doc2docx'
            }

        except Exception as e:
            logger.error(f"DOC转换失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'original_file': doc_path
            }

    def get_supported_info(self) -> dict[str, Any]:
        """获取支持的信息"""
        return {
            'supported_formats': ['.doc'],
            'primary_method': 'doc2docx_conversion' if self.doc2docx_available else 'olefile_basic',
            'available_libraries': {
                'doc2docx': self.doc2docx_available,
                'python-docx': self.docx_available,
                'olefile': self.olefile_available
            },
            'capabilities': {
                'text_extraction': self.doc2docx_available and self.docx_available,
                'table_extraction': self.doc2docx_available and self.docx_available,
                'metadata_extraction': self.doc2docx_available and self.docx_available,
                'format_conversion': self.doc2docx_available
            },
            'recommendations': [
                '安装doc2docx库以获得最佳的DOC文件支持',
                '对于复杂的DOC文件，建议先转换为DOCX格式',
                'olefile方法只能提供基本的文件信息'
            ]
        }

# 全局实例
doc_handler = None

def get_doc_handler() -> DOCFileHandler:
    """获取DOC处理器实例"""
    global doc_handler
    if doc_handler is None:
        doc_handler = DOCFileHandler()
    return doc_handler

# 测试函数
async def test_doc_handler():
    """测试DOC处理器"""
    handler = get_doc_handler()

    # 显示支持信息
    info = handler.get_supported_info()
    logger.info('DOC处理器支持信息:')
    print(json.dumps(info, indent=2, ensure_ascii=False))

    return handler

if __name__ == '__main__':
    import asyncio
    import json

    # 运行测试
    handler = asyncio.run(test_doc_handler())
