#!/usr/bin/env python3
"""
Dolphin文档解析服务
ByteDance Dolphin Document Parser Service for Athena Platform
集成ByteDance Dolphin文档图像解析能力，提供专业文档版面分析和OCR功能
"""

import json
import logging
import os
import tempfile
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

import cv2
import fitz  # PyMuPDF

# 文档处理库
import layoutparser as lp
import numpy as np
import yaml
from paddleocr import PaddleOCR

# 配置日志 - 必须在导入检查之前
logger = logging.getLogger('DolphinParser')

# DOC文件处理
try:
    from doc_file_handler import get_doc_handler
    DOC_HANDLER_AVAILABLE = True
except ImportError:
    DOC_HANDLER_AVAILABLE = False
    logger.info('警告: DOC处理器导入失败')

import aiofiles

# FastAPI相关
import uvicorn
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from pydantic import BaseModel, Field

# 导入统一认证模块

# 配置日志
# setup_logging()  # 日志配置已移至模块导入
logger = logging.getLogger('DolphinParser')

# 创建FastAPI应用
app = FastAPI(
    title='Dolphin文档解析服务',
    description='基于ByteDance Dolphin的专业文档图像解析服务',
    version='1.0.0',
    docs_url='/docs',
    redoc_url='/redoc'
)

# CORS配置

# 配置
DOLPHIN_PORT = 8013
UPLOAD_DIR = '/tmp/dolphin_uploads'
PROCESSED_DIR = '/tmp/dolphin_processed'
CACHE_DIR = '/tmp/dolphin_cache'

# 创建必要的目录
for directory in [UPLOAD_DIR, PROCESSED_DIR, CACHE_DIR]:
    os.makedirs(directory, exist_ok=True)

# 全局变量
ocr_model = None
layout_model = None

def load_config():
    """加载配置文件"""
    config_path = os.path.join(os.path.dirname(__file__), 'dolphin_config.yaml')
    if os.path.exists(config_path):
        with open(config_path, encoding='utf-8') as f:
            return yaml.safe_load(f)
    return {}

CONFIG = load_config()

class DolphinDocumentParser:
    """Dolphin文档解析器"""

    def __init__(self):
        self.ocr = None
        self.layout_model = None
        self.supported_formats = {'.pdf', '.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.docx', '.doc'}
        self.doc_handler = None
        self._initialize_models()

    def _initialize_models(self):
        """初始化OCR和布局检测模型"""
        try:
            # 初始化PaddleOCR
            self.ocr = PaddleOCR(
                use_angle_cls=True,
                lang='ch',  # 支持中英文
                show_log=False,
                use_gpu=False
            )
            logger.info('PaddleOCR初始化成功')

            # 初始化Layout Parser模型
            # 使用PubLayNet预训练模型进行版面分析
            model_config = lp.Detectron2LayoutModel(
                config_path='lp://PubLayNet/mask_rcnn_R_50_FPN_3x/config',
                model_path='lp://PubLayNet/mask_rcnn_R_50_FPN_3x/model',
                extra_config=[
                    'MODEL.ROI_HEADS.SCORE_THRESH_TEST', 0.8,
                    'MODEL.ROI_HEADS.NMS_THRESH_TEST', 0.5,
                ]
            )
            self.layout_model = model_config
            logger.info('Layout Parser模型初始化成功')

        except Exception as e:
            logger.error(f"模型初始化失败: {e}")
            self.ocr = None
            self.layout_model = None

        # 初始化DOC处理器
        if DOC_HANDLER_AVAILABLE:
            try:
                self.doc_handler = get_doc_handler()
                logger.info('DOC处理器初始化成功')
            except Exception as e:
                logger.error(f"DOC处理器初始化失败: {e}")
                self.doc_handler = None

    def is_supported_format(self, file_path: str) -> bool:
        """检查文件格式是否支持"""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_formats

    async def parse_document(self, file_path: str, options: dict[str, Any] = None) -> dict[str, Any]:
        """
        解析文档，提取版面信息和文本内容

        Args:
            file_path: 文档路径
            options: 解析选项

        Returns:
            解析结果字典
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        if not self.is_supported_format(file_path):
            raise ValueError(f"不支持的文件格式: {Path(file_path).suffix}")

        start_time = time.time()
        file_ext = Path(file_path).suffix.lower()

        try:
            if file_ext == '.pdf':
                return await self._parse_pdf(file_path, options)
            elif file_ext in {'.jpg', '.jpeg', '.png', '.tiff', '.bmp'}:
                return await self._parse_image(file_path, options)
            elif file_ext == '.docx':
                return await self._parse_docx(file_path, options)
            elif file_ext == '.doc':
                return await self._parse_doc(file_path, options)
            else:
                raise ValueError(f"暂不支持的格式: {file_ext}")

        except Exception as e:
            logger.error(f"文档解析失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_path': file_path,
                'processing_time': time.time() - start_time
            }

    async def _parse_pdf(self, pdf_path: str, options: dict[str, Any] = None) -> dict[str, Any]:
        """解析PDF文档"""
        results = []
        try:
            # 使用PyMuPDF打开PDF
            pdf_document = fitz.open(pdf_path)
            total_pages = len(pdf_document)

            for page_num in range(min(total_pages, options.get('max_pages', 10) if options else 10)):
                page = pdf_document[page_num]

                # 提取页面为图像
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2倍分辨率
                img_data = pix.tobytes('png')

                # 保存临时图像
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    tmp_file.write(img_data)
                    tmp_image_path = tmp_file.name

                # 解析页面图像
                page_result = await self._parse_image(tmp_image_path, {
                    'page_number': page_num + 1,
                    'total_pages': total_pages
                })

                results.append(page_result)

                # 清理临时文件
                os.unlink(tmp_image_path)

            pdf_document.close()

            return {
                'success': True,
                'file_type': 'PDF',
                'total_pages': total_pages,
                'pages': results,
                'processing_time': time.time() - time.time()
            }

        except Exception as e:
            logger.error(f"PDF解析失败: {e}")
            return {'success': False, 'error': str(e)}

    async def _parse_image(self, image_path: str, options: dict[str, Any] = None) -> dict[str, Any]:
        """解析图像文档"""
        try:
            # 读取图像
            image = cv2.imread(image_path)
            if image is None:
                raise ValueError('无法读取图像文件')

            # 布局分析
            layout_result = None
            if self.layout_model:
                layout_result = await self._analyze_layout(image)

            # OCR文本提取
            ocr_result = None
            if self.ocr:
                ocr_result = await self._extract_text_with_ocr(image_path)

            # 合并结果
            merged_result = self._merge_layout_and_ocr(layout_result, ocr_result, options)

            return {
                'success': True,
                'file_type': 'Image',
                'image_size': image.shape[:2],
                'layout_analysis': layout_result,
                'ocr_result': ocr_result,
                'merged_content': merged_result,
                'page_info': options or {},
                'processing_time': time.time() - time.time()
            }

        except Exception as e:
            logger.error(f"图像解析失败: {e}")
            return {'success': False, 'error': str(e)}

    async def _parse_docx(self, docx_path: str, options: dict[str, Any] = None) -> dict[str, Any]:
        """解析Word文档"""
        try:
            import docx

            doc = docx.Document(docx_path)
            paragraphs = []
            tables = []

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'text': para.text,
                        'style': para.style.name,
                        'level': para.style.next_paragraph_style or 0
                    })

            # 提取表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            return {
                'success': True,
                'file_type': 'DOCX',
                'paragraphs': paragraphs,
                'tables': tables,
                'processing_time': time.time() - time.time()
            }

        except Exception as e:
            logger.error(f"DOCX解析失败: {e}")
            return {'success': False, 'error': str(e)}

    async def _analyze_layout(self, image: np.ndarray) -> dict[str, Any]:
        """分析文档版面布局"""
        try:
            # 使用Layout Parser进行版面分析
            layout_result = self.layout_model.detect(image)

            # 整理版面分析结果
            layout_blocks = []
            for block in layout_result:
                layout_blocks.append({
                    'type': block.type,  # text, title, figure, table, list
                    'bbox': block.block.coordinates(),  # 边界框
                    'score': block.score  # 置信度
                })

            return {
                'blocks': layout_blocks,
                'total_blocks': len(layout_blocks),
                'block_types': list({block['type'] for block in layout_blocks})
            }

        except Exception as e:
            logger.error(f"版面分析失败: {e}")
            return None

    async def _extract_text_with_ocr(self, image_path: str) -> dict[str, Any]:
        """使用OCR提取文本"""
        try:
            # 使用PaddleOCR进行文本识别
            ocr_results = self.ocr.ocr(image_path, cls=True)

            text_blocks = []
            full_text = ''

            for page in ocr_results:
                for line in page:
                    if line:
                        # 提取文本框信息
                        box = line[0]  # 文本框坐标
                        text = line[1][0]  # 识别文本
                        confidence = line[1][1]  # 置信度

                        text_blocks.append({
                            'text': text,
                            'bbox': box,
                            'confidence': confidence
                        })

                        full_text += text + "\n"

            return {
                'text_blocks': text_blocks,
                'full_text': full_text.strip(),
                'total_blocks': len(text_blocks)
            }

        except Exception as e:
            logger.error(f"OCR文本提取失败: {e}")
            return None

    def _merge_layout_and_ocr(self, layout_result: dict, ocr_result: dict, options: dict) -> dict[str, Any]:
        """合并版面分析和OCR结果"""
        if not layout_result or not ocr_result:
            return {}

        # 根据版面分析结果对OCR文本进行分类
        categorized_content = {
            'titles': [],
            'paragraphs': [],
            'figures': [],
            'tables': [],
            'lists': []
        }

        # 简单的文本分类逻辑（可以根据需要优化）
        for block in layout_result['blocks']:
            block_type = block['type']
            block_bbox = block['bbox']

            # 查找对应的OCR文本块
            matching_texts = []
            for text_block in ocr_result['text_blocks']:
                if self._is_bbox_overlap(block_bbox, text_block['bbox']):
                    matching_texts.append(text_block['text'])

            if matching_texts:
                categorized_content[f"{block_type}s"].append({
                    'bbox': block_bbox,
                    'texts': matching_texts,
                    'confidence': block['score']
                })

        return {
            'categorized_content': categorized_content,
            'structured_text': self._generate_structured_text(categorized_content)
        }

    def _is_bbox_overlap(self, bbox1: list, bbox2: list, threshold: float = 0.5) -> bool:
        """判断两个边界框是否重叠"""
        # 计算交集面积
        x_left = max(bbox1[0], bbox2[0])
        y_top = max(bbox1[1], bbox2[1])
        x_right = min(bbox1[2], bbox2[2])
        y_bottom = min(bbox1[3], bbox2[3])

        if x_right < x_left or y_bottom < y_top:
            return False

        intersection_area = (x_right - x_left) * (y_bottom - y_top)

        # 计算并集面积
        area1 = (bbox1[2] - bbox1[0]) * (bbox1[3] - bbox1[1])
        area2 = (bbox2[2] - bbox2[0]) * (bbox2[3] - bbox2[1])
        union_area = area1 + area2 - intersection_area

        if union_area == 0:
            return False

        return (intersection_area / union_area) >= threshold

    def _generate_structured_text(self, categorized_content: dict) -> str:
        """生成结构化文本"""
        structured_text = ''

        # 按优先级组织内容
        content_order = ['titles', 'paragraphs', 'tables', 'lists', 'figures']

        for content_type in content_order:
            items = categorized_content.get(content_type, [])
            if items:
                if content_type == 'titles':
                    for item in items:
                        structured_text += f"\n# {' '.join(item['texts'])}\n\n"
                elif content_type == 'paragraphs':
                    for item in items:
                        structured_text += f"{' '.join(item['texts'])}\n\n"
                elif content_type == 'tables':
                    for item in items:
                        structured_text += f"[表格内容]\n{' '.join(item['texts'])}\n\n"
                elif content_type == 'lists':
                    for item in items:
                        structured_text += f"[列表项]\n{' '.join(item['texts'])}\n"
                elif content_type == 'figures':
                    for item in items:
                        structured_text += f"[图表]\n{' '.join(item['texts'])}\n\n"

        return structured_text.strip()

    async def _parse_doc(self, doc_path: str, options: dict[str, Any] = None) -> dict[str, Any]:
        """解析DOC文件"""
        try:
            if not self.doc_handler:
                # 如果DOC处理器不可用，尝试转换为DOCX
                logger.warning('DOC处理器不可用，尝试格式转换')
                return await self._fallback_doc_processing(doc_path, options)

            # 使用DOC处理器提取文本
            logger.info(f"使用DOC处理器解析: {doc_path}")
            doc_result = await self.doc_handler.extract_text_from_doc(doc_path, options)

            if not doc_result.get('success', False):
                return {
                    'success': False,
                    'error': f"DOC解析失败: {doc_result.get('error', '未知错误')}",
                    'file_type': 'DOC',
                    'file_path': doc_path
                }

            # 处理提取到的文本
            text_content = doc_result.get('text_content', '')
            paragraphs = doc_result.get('paragraphs', [])
            tables = doc_result.get('tables', [])
            metadata = doc_result.get('metadata', {})

            # 构建结果
            result = {
                'success': True,
                'file_type': 'DOC',
                'text_content': text_content,
                'paragraphs_count': len(paragraphs),
                'tables_count': len(tables),
                'metadata': metadata,
                'extraction_method': doc_result.get('extraction_method', 'unknown'),
                'processing_time': 0.0,
                'structured_analysis': self._analyze_doc_content(paragraphs, tables)
            }

            # 如果有文本内容，尝试进行版面分析模拟
            if text_content.strip():
                result['layout_simulation'] = self._simulate_layout_analysis(paragraphs, tables)

            return result

        except Exception as e:
            logger.error(f"DOC解析失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_type': 'DOC',
                'file_path': doc_path
            }

    async def _fallback_doc_processing(self, doc_path: str, options: dict[str, Any] = None) -> dict[str, Any]:
        """备用DOC处理方法"""
        try:
            # 尝试将DOC转换为DOCX
            logger.info(f"尝试将DOC转换为DOCX: {doc_path}")
            conversion_result = await self.doc_handler.convert_doc_to_docx(doc_path)

            if conversion_result.get('success', False):
                # 获取转换后的DOCX文件
                converted_files = conversion_result.get('converted_files', [])
                if converted_files:
                    docx_path = converted_files[0]

                    # 解析转换后的DOCX文件
                    logger.info(f"解析转换后的DOCX文件: {docx_path}")
                    docx_result = await self._parse_docx(docx_path, options)

                    # 添加转换信息
                    if docx_result.get('success', False):
                        docx_result['conversion_info'] = {
                            'original_format': 'DOC',
                            'converted_from': doc_path,
                            'conversion_method': 'doc2docx'
                        }

                    return docx_result

            raise Exception('DOC转换失败')

        except Exception as e:
            logger.error(f"备用DOC处理失败: {e}")
            return {
                'success': False,
                'error': f"DOC文件处理失败: {str(e)}。建议安装doc2docx库或手动将DOC转换为DOCX格式。",
                'file_type': 'DOC',
                'file_path': doc_path,
                'recommendations': [
                    '安装doc2docx库: pip install doc2docx',
                    '使用Microsoft Word将DOC另存为DOCX格式',
                    '使用在线转换工具进行格式转换'
                ]
            }

    def _analyze_doc_content(self, paragraphs: list[str], tables: list[str]) -> dict[str, Any]:
        """分析DOC内容"""
        analysis = {
            'content_summary': {
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables),
                'has_content': len(paragraphs) > 0 or len(tables) > 0
            },
            'content_structure': {
                'paragraph_lengths': [len(p) for p in paragraphs],
                'table_contents_count': len(tables)
            },
            'quality_indicators': {
                'empty_paragraphs': sum(1 for p in paragraphs if not p.strip()),
                'short_paragraphs': sum(1 for p in paragraphs if len(p.strip()) < 20),
                'long_paragraphs': sum(1 for p in paragraphs if len(p.strip()) > 500)
            }
        }

        return analysis

    def _simulate_layout_analysis(self, paragraphs: list[str], tables: list[str]) -> dict[str, Any]:
        """模拟版面分析（基于文本内容）"""
        # 简单的版面分析模拟
        layout_blocks = []

        # 为段落创建版面块
        for i, paragraph in enumerate(paragraphs):
            if paragraph.strip():
                # 简单的标题识别
                if len(paragraph.strip()) < 50 and paragraph.strip().istitle():
                    block_type = 'title'
                elif paragraph.strip().startswith(('•', '-', '*', '1.', '2.')):
                    block_type = 'list'
                else:
                    block_type = 'text'

                layout_blocks.append({
                    'type': block_type,
                    'content': paragraph.strip(),
                    'position': i,
                    'simulated': True
                })

        # 为表格创建版面块
        for i, table in enumerate(tables):
            if table.strip():
                layout_blocks.append({
                    'type': 'table',
                    'content': table.strip(),
                    'position': len(paragraphs) + i,
                    'simulated': True
                })

        return {
            'blocks': layout_blocks,
            'total_blocks': len(layout_blocks),
            'block_types': list({block['type'] for block in layout_blocks}),
            'note': '这是基于文本内容的模拟版面分析，不是真正的版面检测'
        }

# 全局解析器实例
dolphin_parser = DolphinDocumentParser()

# API模型
class ParseRequest(BaseModel):
    """解析请求模型"""
    file_path: str | None = None
    options: dict[str, Any] = Field(default_factory=dict)

class ParseResponse(BaseModel):
    """解析响应模型"""
    success: bool
    message: str
    result: dict[str, Any] | None = None
    timestamp: datetime

# API端点
@app.get('/')
async def root():
    """根端点"""
    return {
        'service': 'Dolphin文档解析服务',
        'version': '1.0.0',
        'status': 'running',
        'models': {
            'ocr': dolphin_parser.ocr is not None,
            'layout_parser': dolphin_parser.layout_model is not None
        },
        'supported_formats': list(dolphin_parser.supported_formats),
        'timestamp': datetime.now().isoformat()
    }

@app.get('/health')
async def health_check():
    """健康检查"""
    return {
        'status': 'healthy',
        'models': {
            'ocr': 'initialized' if dolphin_parser.ocr else 'not_initialized',
            'layout_parser': 'initialized' if dolphin_parser.layout_model else 'not_initialized'
        },
        'system_resources': {
            'upload_directory': os.path.exists(UPLOAD_DIR),
            'processed_directory': os.path.exists(PROCESSED_DIR),
            'cache_directory': os.path.exists(CACHE_DIR)
        },
        'timestamp': datetime.now().isoformat()
    }

@app.post('/parse', response_model=ParseResponse)
async def parse_document(
    file: UploadFile = File(...),
    include_layout: bool = Form(default=True),
    include_ocr: bool = Form(default=True),
    max_pages: int = Form(default=10)
):
    """
    解析上传的文档

    Args:
        file: 上传的文档文件
        include_layout: 是否包含版面分析
        include_ocr: 是否包含OCR文本提取
        max_pages: 最大处理页数（仅PDF）
    """
    if not file:
        raise HTTPException(status_code=400, detail='没有上传文件')

    # 生成临时文件路径
    file_id = str(uuid.uuid4())
    safe_filename = f"{file_id}_{file.filename}"
    temp_file_path = os.path.join(UPLOAD_DIR, safe_filename)

    try:
        # 保存上传的文件
        async with aiofiles.open(temp_file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        # 准备解析选项
        options = {
            'include_layout': include_layout,
            'include_ocr': include_ocr,
            'max_pages': max_pages
        }

        # 解析文档
        result = await dolphin_parser.parse_document(temp_file_path, options)

        return ParseResponse(
            success=result.get('success', False),
            message='文档解析成功' if result.get('success') else f"文档解析失败: {result.get('error', '未知错误')}",
            result=result,
            timestamp=datetime.now()
        )

    except Exception as e:
        logger.error(f"文档解析异常: {e}")
        raise HTTPException(status_code=500, detail=f"文档解析失败: {str(e)}") from e

    finally:
        # 清理临时文件
        if os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except Exception as e:
                logger.error(f"Error: {e}", exc_info=True)

@app.post('/parse-url')
async def parse_document_from_url(
    url: str = Form(...),
    options: str = Form(default='{}')
):
    """从URL解析文档"""
    try:
        import requests

        # 下载文件
        response = requests.get(url, timeout=30)
        response.raise_for_status()

        # 保存临时文件
        file_id = str(uuid.uuid4())
        temp_file_path = os.path.join(UPLOAD_DIR, f"{file_id}_downloaded")

        with open(temp_file_path, 'wb') as f:
            f.write(response.content)

        # 解析选项
        parse_options = json.loads(options) if options else {}

        # 解析文档
        result = await dolphin_parser.parse_document(temp_file_path, parse_options)

        # 清理临时文件
        os.unlink(temp_file_path)

        return ParseResponse(
            success=result.get('success', False),
            message='URL文档解析成功' if result.get('success') else f"URL文档解析失败: {result.get('error', '未知错误')}",
            result=result,
            timestamp=datetime.now()
        )

    except Exception as e:
        logger.error(f"URL文档解析异常: {e}")
        raise HTTPException(status_code=500, detail=f"URL文档解析失败: {str(e)}") from e

@app.get('/formats')
async def get_supported_formats():
    """获取支持的文件格式"""
    return {
        'supported_formats': list(dolphin_parser.supported_formats),
        'format_details': {
            '.pdf': 'PDF文档，支持多页处理',
            '.jpg/.jpeg': 'JPEG图像文件',
            '.png': 'PNG图像文件',
            '.tiff/.bmp': 'TIFF/BMP图像文件',
            '.docx': 'Word文档（文本提取）'
        },
        'features': {
            'layout_analysis': '版面分析 - 识别标题、段落、表格、图表、列表',
            'ocr_text_extraction': 'OCR文字提取 - 中英文文字识别',
            'structured_output': '结构化输出 - 按内容类型组织文本',
            'multi_page_support': '多页支持 - PDF多页文档处理'
        }
    }

@app.get('/stats')
async def get_service_stats():
    """获取服务统计信息"""
    return {
        'service': 'Dolphin文档解析服务',
        'version': '1.0.0',
        'uptime': '统计功能开发中',
        'models_status': {
            'ocr': dolphin_parser.ocr is not None,
            'layout_parser': dolphin_parser.layout_model is not None
        },
        'supported_formats_count': len(dolphin_parser.supported_formats),
        'temp_directories': {
            'upload_dir': UPLOAD_DIR,
            'processed_dir': PROCESSED_DIR,
            'cache_dir': CACHE_DIR
        }
    }

# 启动事件
@app.on_event('startup')
async def startup_event():
    """应用启动事件"""
    logger.info('Dolphin文档解析服务启动完成')

@app.on_event('shutdown')
async def shutdown_event():
    """应用关闭事件"""
    logger.info('Dolphin文档解析服务正在关闭')

# 启动服务
if __name__ == '__main__':
    uvicorn.run(
        'dolphin_parser_service:app',
        host='0.0.0.0',
        port=DOLPHIN_PORT,
        reload=True,
        log_level='info'
    )
