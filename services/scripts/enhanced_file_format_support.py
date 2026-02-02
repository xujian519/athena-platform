#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
增强文件格式支持模块
Enhanced File Format Support Module for Athena Multimodal System
"""

import csv
from core.async_main import async_main
import gzip
import json
import logging
from core.logging_config import setup_logging
import mimetypes
import os
import subprocess
import sys
import tarfile
import tempfile
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

# 专业库导入
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import xlrd
    XLDR_AVAILABLE = True
except ImportError:
    XLDR_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import librosa
    LIBROSA_AVAILABLE = True
except ImportError:
    LIBROSA_AVAILABLE = False

try:
    import cv2
    OPENCV_AVAILABLE = True
except ImportError:
    OPENCV_AVAILABLE = False

try:
    import pdfplumber
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    import h5py
    H5PY_AVAILABLE = True
except ImportError:
    H5PY_AVAILABLE = False

try:
    import pyarrow as pa
    PYARROW_AVAILABLE = True
except ImportError:
    PYARROW_AVAILABLE = False

try:
    import net_cdf4
    NETCDF4_AVAILABLE = True
except ImportError:
    NETCDF4_AVAILABLE = False

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = setup_logging()

class EnhancedFileFormatSupport:
    """增强文件格式支持类"""

    def __init__(self):
        self.supported_formats = self._get_comprehensive_format_list()
        self.format_handlers = self._initialize_format_handlers()
        self._setup_mime_types()

    def _setup_mime_types(self) -> Any:
        """设置MIME类型映射"""
        # 扩展MIME类型映射
        custom_mimes = {
            '.md': 'text/markdown',
            '.tex': 'application/x-latex',
            '.dxf': 'application/dxf',
            '.dwg': 'application/acad',
            '.stl': 'application/sla',
            '.obj': 'model/obj',
            '.ply': 'application/ply',
            '.fbx': 'application/x-fbx',
            '.3ds': 'application/x-3ds',
            '.csv': 'text/csv',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel',
            '.parquet': 'application/vnd.apache.parquet',
            '.arrow': 'application/vnd.apache.arrow',
            '.hdf5': 'application/x-hdf',
            '.h5': 'application/x-hdf',
            '.nc': 'application/x-netcdf',
            '.cdf': 'application/x-netcdf',
            '.svg': 'image/svg+xml',
            '.ai': 'application/illustrator',
            '.eps': 'application/postscript',
            '.psd': 'application/x-photoshop',
            '.kra': 'application/x-krita',
            '.afphoto': 'application/x-affinity-photo',
            '.epub': 'application/epub+zip',
            '.mobi': 'application/x-mobipocket-ebook',
            '.azw': 'application/vnd.amazon.ebook',
            '.azw3': 'application/vnd.amazon.ebook',
            '.pages': 'application/x-iwork-pages-sffpages',
            '.numbers': 'application/x-iwork-numbers-sffnumbers',
            '.key': 'application/x-iwork-keynote-sffkey',
            '.mat': 'application/x-matlab-data',
            '.fig': 'application/x-matlab-figure',
            '.m': 'text/x-matlab',
            '.py': 'text/x-python',
            '.ipynb': 'application/x-jupyter-notebook',
            '.r': 'text/x-r-source',
            '.rdata': 'application/x-r-data',
            '.sas7bdat': 'application/x-sas-data',
            '.sav': 'application/x-spss-sav',
            '.dta': 'application/x-stata',
            '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
            '.ppt': 'application/vnd.ms-powerpoint',
            '.odt': 'application/vnd.oasis.opendocument.text',
            '.ods': 'application/vnd.oasis.opendocument.spreadsheet',
            '.odp': 'application/vnd.oasis.opendocument.presentation',
            '.pages': 'application/x-iwork-pages-sffpages',
            '.abw': 'application/x-abiword',
            '.zim': 'application/x-zim',
            '.chm': 'application/x-chm',
            '.hlp': 'application/x-winhelp',
            '.man': 'text/x-troff-man',
            '.1': 'text/x-troff-man',
            '.2': 'text/x-troff-man',
            '.3': 'text/x-troff-man',
            '.4': 'text/x-troff-man',
            '.5': 'text/x-troff-man',
            '.6': 'text/x-troff-man',
            '.7': 'text/x-troff-man',
            '.8': 'text/x-troff-man',
            '.9': 'text/x-troff-man',
            '.nfo': 'text/x-nfo',
            '.log': 'text/x-log',
            '.ini': 'text/x-ini',
            '.cfg': 'text/x-ini',
            '.conf': 'text/x-ini',
            '.toml': 'text/x-toml',
            '.yaml': 'text/x-yaml',
            '.yml': 'text/x-yaml',
            '.json': 'application/json',
            '.xml': 'application/xml',
            '.svgz': 'image/svg+xml-compressed',
            '.ttf': 'font/ttf',
            '.otf': 'font/otf',
            '.woff': 'font/woff',
            '.woff2': 'font/woff2',
            '.eot': 'application/vnd.ms-fontobject',
            '.sql': 'application/sql',
            '.db': 'application/x-sqlite3',
            '.sqlite': 'application/x-sqlite3',
            '.sqlite3': 'application/x-sqlite3',
            '.7z': 'application/x-7z-compressed',
            '.rar': 'application/x-rar-compressed',
            '.tar': 'application/x-tar',
            '.gz': 'application/gzip',
            '.bz2': 'application/x-bzip2',
            '.xz': 'application/x-xz',
            '.lzma': 'application/x-lzma',
            '.Z': 'application/x-compress',
            '.rpm': 'application/x-rpm',
            '.deb': 'application/x-debian-package',
            '.dmg': 'application/x-apple-diskimage',
            '.img': 'application/x-apple-diskimage',
            '.iso': 'application/x-iso9660-image',
            '.vhd': 'application/x-vhd',
            '.vmdk': 'application/x-vmdk',
            '.ova': 'application/x-ova',
            '.ovf': 'application/x-ovf',
        }

        for ext, mime in custom_mimes.items():
            mimetypes.add_type(mime, ext)

    def _get_comprehensive_format_list(self) -> Dict[str, Dict[str, Any]]:
        """获取全面的文件格式支持列表"""
        return {
            # 文档类
            'documents': {
                'text_formats': {
                    'txt': {'name': '纯文本', 'handler': 'text_handler', 'mime': 'text/plain'},
                    'md': {'name': 'Markdown', 'handler': 'markdown_handler', 'mime': 'text/markdown'},
                    'tex': {'name': 'LaTeX', 'handler': 'latex_handler', 'mime': 'application/x-latex'},
                    'rst': {'name': 're_structured_text', 'handler': 'text_handler', 'mime': 'text/x-rst'},
                    'org': {'name': 'Org-mode', 'handler': 'text_handler', 'mime': 'text/x-org'},
                    'adoc': {'name': 'AsciiDoc', 'handler': 'text_handler', 'mime': 'text/x-asciidoc'},
                    'log': {'name': '日志文件', 'handler': 'text_handler', 'mime': 'text/x-log'},
                    'csv': {'name': 'CSV表格', 'handler': 'csv_handler', 'mime': 'text/csv'},
                    'tsv': {'name': 'TSV表格', 'handler': 'csv_handler', 'mime': 'text/tab-separated-values'},
                    'json': {'name': 'JSON数据', 'handler': 'json_handler', 'mime': 'application/json'},
                    'xml': {'name': 'XML文档', 'handler': 'xml_handler', 'mime': 'application/xml'},
                    'yaml': {'name': 'YAML配置', 'handler': 'yaml_handler', 'mime': 'text/x-yaml'},
                    'toml': {'name': 'TOML配置', 'handler': 'toml_handler', 'mime': 'text/x-toml'},
                    'ini': {'name': 'INI配置', 'handler': 'text_handler', 'mime': 'text/x-ini'},
                    'cfg': {'name': '配置文件', 'handler': 'text_handler', 'mime': 'text/x-ini'},
                    'sql': {'name': 'SQL脚本', 'handler': 'text_handler', 'mime': 'application/sql'},
                    'py': {'name': 'Python脚本', 'handler': 'code_handler', 'mime': 'text/x-python'},
                    'js': {'name': 'JavaScript', 'handler': 'code_handler', 'mime': 'text/javascript'},
                    'java': {'name': 'Java源码', 'handler': 'code_handler', 'mime': 'text/x-java-source'},
                    'cpp': {'name': 'C++源码', 'handler': 'code_handler', 'mime': 'text/x-c++src'},
                    'c': {'name': 'C源码', 'handler': 'code_handler', 'mime': 'text/x-csrc'},
                    'h': {'name': 'C头文件', 'handler': 'code_handler', 'mime': 'text/x-chdr'},
                    'sh': {'name': 'Shell脚本', 'handler': 'code_handler', 'mime': 'text/x-shellscript'},
                    'bat': {'name': '批处理文件', 'handler': 'code_handler', 'mime': 'text/x-bat'},
                    'ps1': {'name': 'PowerShell', 'handler': 'code_handler', 'mime': 'text/x-powershell'},
                    'r': {'name': 'R脚本', 'handler': 'code_handler', 'mime': 'text/x-r-source'},
                    'm': {'name': 'MATLAB脚本', 'handler': 'code_handler', 'mime': 'text/x-matlab'},
                    'ipynb': {'name': 'Jupyter Notebook', 'handler': 'jupyter_handler', 'mime': 'application/x-jupyter-notebook'},
                },
                'office_documents': {
                    'pdf': {'name': 'PDF文档', 'handler': 'pdf_handler', 'mime': 'application/pdf'},
                    'docx': {'name': 'Word文档', 'handler': 'docx_handler', 'mime': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'},
                    'doc': {'name': 'Word 97-2003', 'handler': 'doc_handler', 'mime': 'application/msword'},
                    'dotx': {'name': 'Word模板', 'handler': 'docx_handler', 'mime': 'application/vnd.openxmlformats-officedocument.wordprocessingml.template'},
                    'xlsx': {'name': 'Excel工作簿', 'handler': 'xlsx_handler', 'mime': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'},
                    'xls': {'name': 'Excel 97-2003', 'handler': 'xls_handler', 'mime': 'application/vnd.ms-excel'},
                    'xltx': {'name': 'Excel模板', 'handler': 'xlsx_handler', 'mime': 'application/vnd.openxmlformats-officedocument.spreadsheetml.template'},
                    'pptx': {'name': 'PowerPoint演示', 'handler': 'pptx_handler', 'mime': 'application/vnd.openxmlformats-officedocument.presentationml.presentation'},
                    'ppt': {'name': 'PowerPoint 97-2003', 'handler': 'ppt_handler', 'mime': 'application/vnd.ms-powerpoint'},
                    'potx': {'name': 'PowerPoint模板', 'handler': 'pptx_handler', 'mime': 'application/vnd.openxmlformats-officedocument.presentationml.template'},
                    'pages': {'name': 'Pages文档', 'handler': 'pages_handler', 'mime': 'application/x-iwork-pages-sffpages'},
                    'numbers': {'name': 'Numbers表格', 'handler': 'numbers_handler', 'mime': 'application/x-iwork-numbers-sffnumbers'},
                    'key': {'name': 'Keynote演示', 'handler': 'keynote_handler', 'mime': 'application/x-iwork-keynote-sffkey'},
                },
                'technical_documents': {
                    'odt': {'name': 'OpenDocument文本', 'handler': 'odt_handler', 'mime': 'application/vnd.oasis.opendocument.text'},
                    'ods': {'name': 'OpenDocument表格', 'handler': 'ods_handler', 'mime': 'application/vnd.oasis.opendocument.spreadsheet'},
                    'odp': {'name': 'OpenDocument演示', 'handler': 'odp_handler', 'mime': 'application/vnd.oasis.opendocument.presentation'},
                    'rtf': {'name': '富文本格式', 'handler': 'rtf_handler', 'mime': 'application/rtf'},
                    'wpd': {'name': 'WordPerfect', 'handler': 'wpd_handler', 'mime': 'application/vnd.wordperfect'},
                    'wps': {'name': 'WPS文档', 'handler': 'wps_handler', 'mime': 'application/vnd.ms-works'},
                    'abw': {'name': 'AbiWord', 'handler': 'abw_handler', 'mime': 'application/x-abiword'},
                    'chm': {'name': 'CHM帮助文件', 'handler': 'chm_handler', 'mime': 'application/x-chm'},
                    'hlp': {'name': 'HLP帮助文件', 'handler': 'hlp_handler', 'mime': 'application/x-winhelp'},
                    'epub': {'name': 'EPUB电子书', 'handler': 'epub_handler', 'mime': 'application/epub+zip'},
                    'mobi': {'name': 'Mobipocket电子书', 'handler': 'mobi_handler', 'mime': 'application/x-mobipocket-ebook'},
                    'azw': {'name': 'Kindle电子书', 'handler': 'azw_handler', 'mime': 'application/vnd.amazon.ebook'},
                    'azw3': {'name': 'Kindle电子书', 'handler': 'azw3_handler', 'mime': 'application/vnd.amazon.ebook'},
                    'fb2': {'name': 'FictionBook', 'handler': 'fb2_handler', 'mime': 'application/x-fictionbook+xml'},
                    'lit': {'name': 'Microsoft Reader', 'handler': 'lit_handler', 'mime': 'application/x-ms-reader'},
                    'pdb': {'name': 'Palm文档', 'handler': 'pdb_handler', 'mime': 'application/vnd.palm'},
                    'pdb': {'name': 'Protein数据库', 'handler': 'pdb_handler', 'mime': 'chemical/x-pdb'},
                }
            },

            # 图像类
            'images': {
                'raster_images': {
                    'png': {'name': 'PNG图像', 'handler': 'image_handler', 'mime': 'image/png'},
                    'jpg': {'name': 'JPEG图像', 'handler': 'image_handler', 'mime': 'image/jpeg'},
                    'jpeg': {'name': 'JPEG图像', 'handler': 'image_handler', 'mime': 'image/jpeg'},
                    'gif': {'name': 'GIF图像', 'handler': 'image_handler', 'mime': 'image/gif'},
                    'bmp': {'name': 'BMP图像', 'handler': 'image_handler', 'mime': 'image/bmp'},
                    'tiff': {'name': 'TIFF图像', 'handler': 'image_handler', 'mime': 'image/tiff'},
                    'tif': {'name': 'TIFF图像', 'handler': 'image_handler', 'mime': 'image/tiff'},
                    'webp': {'name': 'WebP图像', 'handler': 'image_handler', 'mime': 'image/webp'},
                    'jp2': {'name': 'JPEG 2000', 'handler': 'image_handler', 'mime': 'image/jp2'},
                    'jpx': {'name': 'JPEG 2000', 'handler': 'image_handler', 'mime': 'image/jpx'},
                    'ico': {'name': 'ICO图标', 'handler': 'image_handler', 'mime': 'image/x-icon'},
                    'icns': {'name': 'macOS图标', 'handler': 'image_handler', 'mime': 'image/x-icns'},
                    'tga': {'name': 'TGA图像', 'handler': 'image_handler', 'mime': 'image/x-targa'},
                    'pcx': {'name': 'PCX图像', 'handler': 'image_handler', 'mime': 'image/x-pcx'},
                    'pic': {'name': 'PIC图像', 'handler': 'image_handler', 'mime': 'image/x-pict'},
                    'pict': {'name': 'PICT图像', 'handler': 'image_handler', 'mime': 'image/x-pict'},
                    'sgi': {'name': 'SGI图像', 'handler': 'image_handler', 'mime': 'image/sgi'},
                    'rgb': {'name': 'RGB图像', 'handler': 'image_handler', 'mime': 'image/x-rgb'},
                    'rgba': {'name': 'RGBA图像', 'handler': 'image_handler', 'mime': 'image/x-rgba'},
                    'exr': {'name': 'OpenEXR图像', 'handler': 'image_handler', 'mime': 'image/x-exr'},
                    'hdr': {'name': 'HDR图像', 'handler': 'image_handler', 'mime': 'image/vnd.radiance'},
                    'heic': {'name': 'HEIC图像', 'handler': 'image_handler', 'mime': 'image/heic'},
                    'heif': {'name': 'HEIF图像', 'handler': 'image_handler', 'mime': 'image/heif'},
                    'avif': {'name': 'AVIF图像', 'handler': 'image_handler', 'mime': 'image/avif'},
                    'jxl': {'name': 'JPEG XL图像', 'handler': 'image_handler', 'mime': 'image/jxl'},
                },
                'vector_images': {
                    'svg': {'name': 'SVG矢量图', 'handler': 'svg_handler', 'mime': 'image/svg+xml'},
                    'svgz': {'name': 'SVG压缩', 'handler': 'svg_handler', 'mime': 'image/svg+xml-compressed'},
                    'ai': {'name': 'Illustrator文件', 'handler': 'ai_handler', 'mime': 'application/illustrator'},
                    'eps': {'name': 'PostScript', 'handler': 'eps_handler', 'mime': 'application/postscript'},
                    'ps': {'name': 'PostScript', 'handler': 'ps_handler', 'mime': 'application/postscript'},
                    'pdf': {'name': 'PDF矢量', 'handler': 'pdf_handler', 'mime': 'application/pdf'},
                    'wmf': {'name': 'Windows图元', 'handler': 'wmf_handler', 'mime': 'image/x-wmf'},
                    'emf': {'name': '增强图元', 'handler': 'emf_handler', 'mime': 'image/x-emf'},
                    'cgm': {'name': 'CGM图形', 'handler': 'cgm_handler', 'mime': 'image/cgm'},
                    'drw': {'name': 'MicroGrafx', 'handler': 'drw_handler', 'mime': 'application/x-draw'},
                    'dxf': {'name': 'AutoCAD DXF', 'handler': 'dxf_handler', 'mime': 'application/dxf'},
                    'dwg': {'name': 'AutoCAD DWG', 'handler': 'dwg_handler', 'mime': 'application/acad'},
                    'svg': {'name': 'SVG动画', 'handler': 'svg_handler', 'mime': 'image/svg+xml'},
                },
                'technical_images': {
                    'psd': {'name': 'Photoshop文件', 'handler': 'psd_handler', 'mime': 'application/x-photoshop'},
                    'psb': {'name': 'Photoshop大图', 'handler': 'psb_handler', 'mime': 'application/x-photoshop'},
                    'kra': {'name': 'Krita文件', 'handler': 'kra_handler', 'mime': 'application/x-krita'},
                    'afphoto': {'name': 'Affinity Photo', 'handler': 'afphoto_handler', 'mime': 'application/x-affinity-photo'},
                    'afdesign': {'name': 'Affinity Designer', 'handler': 'afdesign_handler', 'mime': 'application/x-affinity-designer'},
                    'xcf': {'name': 'GIMP文件', 'handler': 'xcf_handler', 'mime': 'image/x-xcf'},
                    'ora': {'name': 'OpenRaster', 'handler': 'ora_handler', 'mime': 'image/openraster'},
                    'pdn': {'name': 'Paint.NET', 'handler': 'pdn_handler', 'mime': 'image/x-paintnet'},
                    'cpt': {'name': 'CorelDRAW模板', 'handler': 'cpt_handler', 'mime': 'application/coreldraw'},
                    'tiff': {'name': '多页TIFF', 'handler': 'tiff_handler', 'mime': 'image/tiff'},
                    'dcx': {'name': 'PCX多页', 'handler': 'dcx_handler', 'mime': 'image/x-dcx'},
                    'jng': {'name': 'JPEG网络图', 'handler': 'jng_handler', 'mime': 'image/x-jng'},
                    'mng': {'name': '多层网络图', 'handler': 'mng_handler', 'mime': 'video/x-mng'},
                    'miff': {'name': 'Magick图像', 'handler': 'miff_handler', 'mime': 'application/x-magick-image'},
                    'xbm': {'name': 'X位图', 'handler': 'xbm_handler', 'mime': 'image/x-xbitmap'},
                    'xpm': {'name': 'X像素图', 'handler': 'xpm_handler', 'mime': 'image/x-xpixmap'},
                    'pcd': {'name': 'PhotoCD', 'handler': 'pcd_handler', 'mime': 'image/x-photo-cd'},
                }
            },

            # 音频类
            'audio': {
                'lossless_audio': {
                    'wav': {'name': 'WAV音频', 'handler': 'audio_handler', 'mime': 'audio/wav'},
                    'flac': {'name': 'FLAC音频', 'handler': 'audio_handler', 'mime': 'audio/flac'},
                    'alac': {'name': 'ALAC音频', 'handler': 'audio_handler', 'mime': 'audio/x-alac'},
                    'ape': {'name': 'Monkey音频', 'handler': 'audio_handler', 'mime': 'audio/x-ape'},
                    'wv': {'name': 'WavPack', 'handler': 'audio_handler', 'mime': 'audio/x-wavpack'},
                    'tta': {'name': 'True Audio', 'handler': 'audio_handler', 'mime': 'audio/x-tta'},
                    'shn': {'name': 'Shorten', 'handler': 'audio_handler', 'mime': 'audio/x-shn'},
                    'dsd': {'name': 'DSD音频', 'handler': 'audio_handler', 'mime': 'audio/x-dsd'},
                    'dsf': {'name': 'DSD文件', 'handler': 'audio_handler', 'mime': 'audio/x-dsf'},
                    'dff': {'name': 'DSDFF', 'handler': 'audio_handler', 'mime': 'audio/x-dff'},
                },
                'lossy_audio': {
                    'mp3': {'name': 'MP3音频', 'handler': 'audio_handler', 'mime': 'audio/mpeg'},
                    'm4a': {'name': 'M4A音频', 'handler': 'audio_handler', 'mime': 'audio/mp4'},
                    'aac': {'name': 'AAC音频', 'handler': 'audio_handler', 'mime': 'audio/aac'},
                    'ogg': {'name': 'OGG音频', 'handler': 'audio_handler', 'mime': 'audio/ogg'},
                    'opus': {'name': 'Opus音频', 'handler': 'audio_handler', 'mime': 'audio/opus'},
                    'wma': {'name': 'WMA音频', 'handler': 'audio_handler', 'mime': 'audio/x-ms-wma'},
                    'ra': {'name': 'RealAudio', 'handler': 'audio_handler', 'mime': 'audio/x-pn-realaudio'},
                    'rm': {'name': 'RealMedia', 'handler': 'audio_handler', 'mime': 'audio/x-pn-realaudio'},
                    '3gp': {'name': '3GPP音频', 'handler': 'audio_handler', 'mime': 'audio/3gpp'},
                    'amr': {'name': 'AMR音频', 'handler': 'audio_handler', 'mime': 'audio/amr'},
                    'gsm': {'name': 'GSM音频', 'handler': 'audio_handler', 'mime': 'audio/x-gsm'},
                    'au': {'name': 'AU音频', 'handler': 'audio_handler', 'mime': 'audio/basic'},
                    'snd': {'name': 'SND音频', 'handler': 'audio_handler', 'mime': 'audio/basic'},
                    'mid': {'name': 'MIDI文件', 'handler': 'midi_handler', 'mime': 'audio/midi'},
                    'kar': {'name': 'KAR文件', 'handler': 'midi_handler', 'mime': 'audio/midi'},
                }
            },

            # 视频类
            'video': {
                'container_formats': {
                    'mp4': {'name': 'MP4视频', 'handler': 'video_handler', 'mime': 'video/mp4'},
                    'avi': {'name': 'AVI视频', 'handler': 'video_handler', 'mime': 'video/x-msvideo'},
                    'mov': {'name': 'QuickTime', 'handler': 'video_handler', 'mime': 'video/quicktime'},
                    'mkv': {'name': 'Matroska', 'handler': 'video_handler', 'mime': 'video/x-matroska'},
                    'webm': {'name': 'WebM视频', 'handler': 'video_handler', 'mime': 'video/webm'},
                    'flv': {'name': 'Flash视频', 'handler': 'video_handler', 'mime': 'video/x-flv'},
                    'wmv': {'name': 'WMV视频', 'handler': 'video_handler', 'mime': 'video/x-ms-wmv'},
                    '3gp': {'name': '3GPP视频', 'handler': 'video_handler', 'mime': 'video/3gpp'},
                    'ogv': {'name': 'OGG视频', 'handler': 'video_handler', 'mime': 'video/ogg'},
                    'm4v': {'name': 'M4V视频', 'handler': 'video_handler', 'mime': 'video/x-m4v'},
                    'mxf': {'name': 'MXF容器', 'handler': 'video_handler', 'mime': 'application/mxf'},
                    'ts': {'name': 'MPEG传输流', 'handler': 'video_handler', 'mime': 'video/mp2t'},
                    'mts': {'name': 'AVCHD视频', 'handler': 'video_handler', 'mime': 'video/mp2t'},
                    'vob': {'name': 'DVD视频', 'handler': 'video_handler', 'mime': 'video/x-ms-vob'},
                    'divx': {'name': 'DivX视频', 'handler': 'video_handler', 'mime': 'video/x-divx'},
                    'xvid': {'name': 'Xvid视频', 'handler': 'video_handler', 'mime': 'video/x-xvid'},
                },
                'professional_video': {
                    'prores': {'name': 'ProRes视频', 'handler': 'video_handler', 'mime': 'video/x-apple-prores'},
                    'dnxhd': {'name': 'DNxHD视频', 'handler': 'video_handler', 'mime': 'video/x-dnxhd'},
                    'avchd': {'name': 'AVCHD视频', 'handler': 'video_handler', 'mime': 'video/x-avchd'},
                    'xavc': {'name': 'XAVC视频', 'handler': 'video_handler', 'mime': 'video/x-xavc'},
                    'hevc': {'name': 'H.265视频', 'handler': 'video_handler', 'mime': 'video/x-h265'},
                    'h264': {'name': 'H.264视频', 'handler': 'video_handler', 'mime': 'video/x-h264'},
                    'vp8': {'name': 'VP8视频', 'handler': 'video_handler', 'mime': 'video/x-vp8'},
                    'vp9': {'name': 'VP9视频', 'handler': 'video_handler', 'mime': 'video/x-vp9'},
                    'av1': {'name': 'AV1视频', 'handler': 'video_handler', 'mime': 'video/x-av1'},
                }
            },

            # 3D模型类
            '3d_models': {
                'cad_models': {
                    'stl': {'name': 'STL 3D模型', 'handler': 'stl_handler', 'mime': 'application/sla'},
                    'obj': {'name': 'Wavefront OBJ', 'handler': 'obj_handler', 'mime': 'model/obj'},
                    'ply': {'name': 'Stanford PLY', 'handler': 'ply_handler', 'mime': 'application/ply'},
                    'fbx': {'name': 'Autodesk FBX', 'handler': 'fbx_handler', 'mime': 'application/x-fbx'},
                    '3ds': {'name': '3D Studio', 'handler': '3ds_handler', 'mime': 'application/x-3ds'},
                    'dae': {'name': 'Collada DAE', 'handler': 'dae_handler', 'mime': 'model/vnd.collada+xml'},
                    'iges': {'name': 'IGES文件', 'handler': 'iges_handler', 'mime': 'model/iges'},
                    'igs': {'name': 'IGES文件', 'handler': 'iges_handler', 'mime': 'model/iges'},
                    'step': {'name': 'STEP文件', 'handler': 'step_handler', 'mime': 'application/step'},
                    'stp': {'name': 'STEP文件', 'handler': 'step_handler', 'mime': 'application/step'},
                    'x3d': {'name': 'X3D模型', 'handler': 'x3d_handler', 'mime': 'model/x3d+xml'},
                    'wrl': {'name': 'VRML模型', 'handler': 'wrl_handler', 'mime': 'model/vrml'},
                    'vrml': {'name': 'VRML模型', 'handler': 'wrl_handler', 'mime': 'model/vrml'},
                    'ctm': {'name': 'OpenCTM', 'handler': 'ctm_handler', 'mime': 'model/x-openctm'},
                },
                'rendering_models': {
                    'blend': {'name': 'Blender文件', 'handler': 'blend_handler', 'mime': 'application/x-blender'},
                    'max': {'name': '3ds Max', 'handler': 'max_handler', 'mime': 'application/x-3dsmax'},
                    'ma': {'name': 'Maya文件', 'handler': 'maya_handler', 'mime': 'application/x-maya'},
                    'mb': {'name': 'Maya Binary', 'handler': 'maya_handler', 'mime': 'application/x-maya'},
                    'c4d': {'name': 'Cinema4D', 'handler': 'c4d_handler', 'mime': 'application/x-cinema4d'},
                    'lwo': {'name': 'LightWave', 'handler': 'lwo_handler', 'mime': 'application/x-lightwave'},
                    'lws': {'name': 'LightWave Scene', 'handler': 'lws_handler', 'mime': 'application/x-lightwave'},
                    'geo': {'name': 'Houdini', 'handler': 'geo_handler', 'mime': 'application/x-houdini'},
                    'abc': {'name': 'Alembic', 'handler': 'abc_handler', 'mime': 'application/x-alembic'},
                    'usd': {'name': 'Universal Scene', 'handler': 'usd_handler', 'mime': 'model/x-usd'},
                    'usda': {'name': 'ASCII USD', 'handler': 'usd_handler', 'mime': 'model/x-usda'},
                    'usdc': {'name': 'Crate USD', 'handler': 'usd_handler', 'mime': 'model/x-usdc'},
                    'usdz': {'name': 'Universal Scene AR', 'handler': 'usdz_handler', 'mime': 'model/x-usdz-ar'},
                }
            },

            # 数据格式类
            'data_formats': {
                'scientific_data': {
                    'hdf5': {'name': 'HDF5文件', 'handler': 'hdf5_handler', 'mime': 'application/x-hdf'},
                    'h5': {'name': 'HDF5文件', 'handler': 'hdf5_handler', 'mime': 'application/x-hdf'},
                    'nc': {'name': 'NetCDF文件', 'handler': 'netcdf_handler', 'mime': 'application/x-netcdf'},
                    'cdf': {'name': 'NetCDF文件', 'handler': 'netcdf_handler', 'mime': 'application/x-netcdf'},
                    'grib': {'name': 'GRIB气象', 'handler': 'grib_handler', 'mime': 'application/x-grib'},
                    'mat': {'name': 'MATLAB数据', 'handler': 'mat_handler', 'mime': 'application/x-matlab-data'},
                    'fig': {'name': 'MATLAB图形', 'handler': 'fig_handler', 'mime': 'application/x-matlab-figure'},
                    'arff': {'name': 'WEKA数据', 'handler': 'arff_handler', 'mime': 'text/x-arff'},
                    'sav': {'name': 'SPSS数据', 'handler': 'sav_handler', 'mime': 'application/x-spss-sav'},
                    'zsav': {'name': 'SPSS压缩', 'handler': 'sav_handler', 'mime': 'application/x-spss-sav'},
                    'dta': {'name': 'Stata数据', 'handler': 'dta_handler', 'mime': 'application/x-stata'},
                    'sas7bdat': {'name': 'SAS数据', 'handler': 'sas_handler', 'mime': 'application/x-sas-data'},
                    'xpt': {'name': 'SAS传输', 'handler': 'sas_handler', 'mime': 'application/x-sas-transport'},
                    'por': {'name': 'SPSS Portable', 'handler': 'por_handler', 'mime': 'application/x-spss-por'},
                },
                'big_data': {
                    'parquet': {'name': 'Apache Parquet', 'handler': 'parquet_handler', 'mime': 'application/vnd.apache.parquet'},
                    'arrow': {'name': 'Apache Arrow', 'handler': 'arrow_handler', 'mime': 'application/vnd.apache.arrow'},
                    'orc': {'name': 'Apache ORC', 'handler': 'orc_handler', 'mime': 'application/vnd.apache.orc'},
                    'avro': {'name': 'Apache Avro', 'handler': 'avro_handler', 'mime': 'application/avro'},
                    'feather': {'name': 'Feather格式', 'handler': 'feather_handler', 'mime': 'application/vnd.apache.feather'},
                    'pkl': {'name': 'Pickle文件', 'handler': 'pickle_handler', 'mime': 'application/x-pickle'},
                    'pickle': {'name': 'Pickle文件', 'handler': 'pickle_handler', 'mime': 'application/x-pickle'},
                    'joblib': {'name': 'Joblib文件', 'handler': 'joblib_handler', 'mime': 'application/x-joblib'},
                    'dill': {'name': 'Dill文件', 'handler': 'dill_handler', 'mime': 'application/x-dill'},
                    'bson': {'name': 'BSON文件', 'handler': 'bson_handler', 'mime': 'application/x-bson'},
                    'msgpack': {'name': 'MessagePack', 'handler': 'msgpack_handler', 'mime': 'application/x-msgpack'},
                }
            },

            # 压缩格式类
            'archives': {
                'compression_formats': {
                    'zip': {'name': 'ZIP压缩', 'handler': 'zip_handler', 'mime': 'application/zip'},
                    'rar': {'name': 'RAR压缩', 'handler': 'rar_handler', 'mime': 'application/x-rar-compressed'},
                    '7z': {'name': '7-Zip压缩', 'handler': 'sevenz_handler', 'mime': 'application/x-7z-compressed'},
                    'tar': {'name': 'TAR归档', 'handler': 'tar_handler', 'mime': 'application/x-tar'},
                    'gz': {'name': 'GZIP压缩', 'handler': 'gzip_handler', 'mime': 'application/gzip'},
                    'bz2': {'name': 'BZIP2压缩', 'handler': 'bzip2_handler', 'mime': 'application/x-bzip2'},
                    'xz': {'name': 'XZ压缩', 'handler': 'xz_handler', 'mime': 'application/x-xz'},
                    'lzma': {'name': 'LZMA压缩', 'handler': 'lzma_handler', 'mime': 'application/x-lzma'},
                    'Z': {'name': 'compress压缩', 'handler': 'compress_handler', 'mime': 'application/x-compress'},
                    'arj': {'name': 'ARJ压缩', 'handler': 'arj_handler', 'mime': 'application/x-arj'},
                    'lha': {'name': 'LHA压缩', 'handler': 'lha_handler', 'mime': 'application/x-lha'},
                    'lzh': {'name': 'LZH压缩', 'handler': 'lzh_handler', 'mime': 'application/x-lzh'},
                    'zoo': {'name': 'ZOO压缩', 'handler': 'zoo_handler', 'mime': 'application/x-zoo'},
                    'arc': {'name': 'ARC压缩', 'handler': 'arc_handler', 'mime': 'application/x-arc'},
                    'pak': {'name': 'PAK压缩', 'handler': 'pak_handler', 'mime': 'application/x-pak'},
                    'ace': {'name': 'ACE压缩', 'handler': 'ace_handler', 'mime': 'application/x-ace'},
                    'sit': {'name': 'StuffIt', 'handler': 'sit_handler', 'mime': 'application/x-stuffit'},
                    'sitx': {'name': 'StuffIt X', 'handler': 'sitx_handler', 'mime': 'application/x-stuffit'},
                    'hqx': {'name': 'BinHex', 'handler': 'hqx_handler', 'mime': 'application/x-stuffit'},
                    'binhex': {'name': 'BinHex', 'handler': 'hqx_handler', 'mime': 'application/x-stuffit'},
                    'uue': {'name': 'UU编码', 'handler': 'uue_handler', 'mime': 'text/x-uuencode'},
                    'xxe': {'name': 'XX编码', 'handler': 'xxe_handler', 'mime': 'text/x-xxencode'},
                    'b64': {'name': 'Base64', 'handler': 'base64_handler', 'mime': 'text/x-base64'},
                },
                'disk_images': {
                    'iso': {'name': 'ISO镜像', 'handler': 'iso_handler', 'mime': 'application/x-iso9660-image'},
                    'dmg': {'name': 'macOS镜像', 'handler': 'dmg_handler', 'mime': 'application/x-apple-diskimage'},
                    'img': {'name': '磁盘镜像', 'handler': 'img_handler', 'mime': 'application/x-disk-image'},
                    'vhd': {'name': 'Virtual HD', 'handler': 'vhd_handler', 'mime': 'application/x-vhd'},
                    'vmdk': {'name': 'VMware', 'handler': 'vmdk_handler', 'mime': 'application/x-vmdk'},
                    'ova': {'name': 'Virtual Appliance', 'handler': 'ova_handler', 'mime': 'application/x-ova'},
                    'ovf': {'name': 'Virtual Format', 'handler': 'ovf_handler', 'mime': 'application/x-ovf'},
                    'vdi': {'name': 'VirtualBox', 'handler': 'vdi_handler', 'mime': 'application/x-virtualbox-vdi'},
                    'qcow': {'name': 'QCOW', 'handler': 'qcow_handler', 'mime': 'application/x-qcow'},
                    'qcow2': {'name': 'QCOW2', 'handler': 'qcow_handler', 'mime': 'application/x-qcow2'},
                    'qed': {'name': 'QED', 'handler': 'qed_handler', 'mime': 'application/x-qed'},
                    'raw': {'name': 'Raw镜像', 'handler': 'raw_handler', 'mime': 'application/x-raw-disk-image'},
                    'img': {'name': 'Raw镜像', 'handler': 'raw_handler', 'mime': 'application/x-raw-disk-image'},
                    'bin': {'name': 'Binary镜像', 'handler': 'raw_handler', 'mime': 'application/x-raw-disk-image'},
                    'cue': {'name': 'CUE文件', 'handler': 'cue_handler', 'mime': 'application/x-cue'},
                    'toc': {'name': 'TOC文件', 'handler': 'toc_handler', 'mime': 'application/x-toc'},
                }
            },

            # 系统和可执行类
            'system': {
                'executables': {
                    'exe': {'name': 'Windows可执行', 'handler': 'executable_handler', 'mime': 'application/x-msdownload'},
                    'msi': {'name': 'Windows安装', 'handler': 'msi_handler', 'mime': 'application/x-msi'},
                    'com': {'name': 'macOS应用', 'handler': 'executable_handler', 'mime': 'application/x-mac-app'},
                    'app': {'name': 'macOS应用', 'handler': 'executable_handler', 'mime': 'application/x-mac-app'},
                    'deb': {'name': 'Debian包', 'handler': 'deb_handler', 'mime': 'application/x-debian-package'},
                    'rpm': {'name': 'RPM包', 'handler': 'rpm_handler', 'mime': 'application/x-rpm'},
                    'pkg': {'name': 'macOS包', 'handler': 'pkg_handler', 'mime': 'application/x-newton-compatible-pkg'},
                    'dmg': {'name': 'macOS镜像', 'handler': 'dmg_handler', 'mime': 'application/x-apple-diskimage'},
                    'snap': {'name': 'Snap包', 'handler': 'snap_handler', 'mime': 'application/x-snap'},
                    'flatpak': {'name': 'Flatpak', 'handler': 'flatpak_handler', 'mime': 'application/x-flatpak'},
                    'appimage': {'name': 'AppImage', 'handler': 'appimage_handler', 'mime': 'application/x-appimage'},
                    'run': {'name': '可执行文件', 'handler': 'executable_handler', 'mime': 'application/x-executable'},
                    'bin': {'name': '二进制文件', 'handler': 'binary_handler', 'mime': 'application/octet-stream'},
                    'dll': {'name': 'Windows DLL', 'handler': 'dll_handler', 'mime': 'application/x-msdownload'},
                    'so': {'name': 'Linux共享库', 'handler': 'dll_handler', 'mime': 'application/x-sharedlib'},
                    'dylib': {'name': 'macOS动态库', 'handler': 'dll_handler', 'mime': 'application/x-mach-binary'},
                    'a': {'name': '静态库', 'handler': 'library_handler', 'mime': 'application/x-archive'},
                    'lib': {'name': '静态库', 'handler': 'library_handler', 'mime': 'application/x-archive'},
                    'o': {'name': '目标文件', 'handler': 'object_handler', 'mime': 'application/x-object'},
                    'obj': {'name': '目标文件', 'handler': 'object_handler', 'mime': 'application/x-object'},
                }
            }
        }

    def _initialize_format_handlers(self) -> Dict[str, Any]:
        """初始化格式处理器"""
        handlers = {
            # 文本处理器
            'text_handler': self.handle_text,
            'markdown_handler': self.handle_markdown,
            'latex_handler': self.handle_latex,
            'csv_handler': self.handle_csv,
            'json_handler': self.handle_json,
            'xml_handler': self.handle_xml,
            'yaml_handler': self.handle_yaml,
            'toml_handler': self.handle_toml,
            'code_handler': self.handle_code,
            'jupyter_handler': self.handle_jupyter,

            # 办公文档处理器
            'pdf_handler': self.handle_pdf,
            'docx_handler': self.handle_docx,
            'doc_handler': self.handle_doc,
            'xlsx_handler': self.handle_xlsx,
            'xls_handler': self.handle_xls,
            'pptx_handler': self.handle_pptx,
            'ppt_handler': self.handle_ppt,

            # 图像处理器
            'image_handler': self.handle_image,
            'svg_handler': self.handle_svg,
            'ai_handler': self.handle_ai,
            'eps_handler': self.handle_eps,
            'psd_handler': self.handle_psd,
            'dxf_handler': self.handle_dxf,
            'dwg_handler': self.handle_dwg,

            # 音频处理器
            'audio_handler': self.handle_audio,
            'midi_handler': self.handle_midi,

            # 视频处理器
            'video_handler': self.handle_video,

            # 3D模型处理器
            'stl_handler': self.handle_stl,
            'obj_handler': self.handle_obj,
            'ply_handler': self.handle_ply,
            'fbx_handler': self.handle_fbx,

            # 数据格式处理器
            'hdf5_handler': self.handle_hdf5,
            'netcdf_handler': self.handle_netcdf,
            'mat_handler': self.handle_mat,
            'parquet_handler': self.handle_parquet,
            'arrow_handler': self.handle_arrow,

            # 压缩文件处理器
            'zip_handler': self.handle_zip,
            'tar_handler': self.handle_tar,
            'gzip_handler': self.handle_gzip,
            'rar_handler': self.handle_rar,
            'sevenz_handler': self.handle_sevenz,

            # 可执行文件处理器
            'executable_handler': self.handle_executable,
            'library_handler': self.handle_library,
            'object_handler': self.handle_object,

            # 默认处理器
            'default_handler': self.handle_default,
        }

        return handlers

    def detect_format(self, file_path: str) -> Dict[str, Any | None]:
        """检测文件格式"""
        path = Path(file_path)

        if not path.exists():
            return None

        # 获取文件扩展名
        ext = path.suffix.lower()

        # 查找支持的格式
        for category, formats in self.supported_formats.items():
            for subcategory, format_dict in formats.items():
                if ext in format_dict:
                    format_info = format_dict[ext].copy()
                    format_info['category'] = category
                    format_info['subcategory'] = subcategory
                    format_info['extension'] = ext
                    return format_info

        return None

    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """获取文件信息"""
        path = Path(file_path)

        if not path.exists():
            return {'error': '文件不存在'}

        # 检测格式
        format_info = self.detect_format(file_path)

        # 获取文件基本信息
        stat = path.stat()

        # 检测MIME类型
        mime_type, _ = mimetypes.guess_type(file_path)

        # 使用file命令获取更详细的类型信息
        file_type = self._get_file_type(file_path)

        return {
            'file_path': str(path),
            'file_name': path.name,
            'file_size': stat.st_size,
            'extension': path.suffix.lower(),
            'mime_type': mime_type,
            'file_type': file_type,
            'format_info': format_info,
            'created_time': stat.st_ctime,
            'modified_time': stat.st_mtime,
            'is_binary': self._is_binary_file(file_path),
            'is_text': self._is_text_file(file_path),
            'is_image': self._is_image_file(file_path),
            'is_audio': self._is_audio_file(file_path),
            'is_video': self._is_video_file(file_path),
            'is_archive': self._is_archive_file(file_path),
            'is_executable': path.stat().st_mode & 0o111 != 0,
        }

    def _get_file_type(self, file_path: str) -> str:
        """使用file命令获取文件类型"""
        try:
            result = subprocess.run(['file', '-b', file_path],
                                  capture_output=True, text=True, timeout=5)
            return result.stdout.strip()
        except (FileNotFoundError, PermissionError, OSError):
            return 'Unknown'

    def _is_binary_file(self, file_path: str) -> bool:
        """检查是否为二进制文件"""
        try:
            with open(file_path, 'rb') as f:
                chunk = f.read(1024)
                return b'\0' in chunk
        except (FileNotFoundError, PermissionError, OSError):
            return False

    def _is_text_file(self, file_path: str) -> bool:
        """检查是否为文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                f.read(1024)
                return True
        except UnicodeDecodeError:
            return False
        except (FileNotFoundError, PermissionError, OSError):
            return False

    def _is_image_file(self, file_path: str) -> bool:
        """检查是否为图像文件"""
        try:
            ext = Path(file_path).suffix.lower()
            image_extensions = {
                '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif',
                '.webp', '.svg', '.ico', '.psd', '.ai', '.eps', '.raw'
            }
            return ext in image_extensions
        except (FileNotFoundError, PermissionError, OSError):
            return False

    def _is_audio_file(self, file_path: str) -> bool:
        """检查是否为音频文件"""
        try:
            ext = Path(file_path).suffix.lower()
            audio_extensions = {
                '.mp3', '.wav', '.flac', '.aac', '.ogg', '.wma', '.m4a',
                '.alac', '.ape', '.wv', '.tta', '.shn', '.dsd', '.opus'
            }
            return ext in audio_extensions
        except (FileNotFoundError, PermissionError, OSError):
            return False

    def _is_video_file(self, file_path: str) -> bool:
        """检查是否为视频文件"""
        try:
            ext = Path(file_path).suffix.lower()
            video_extensions = {
                '.mp4', '.avi', '.mov', '.mkv', '.webm', '.flv', '.wmv',
                '.3gp', '.ogv', '.m4v', '.mpg', '.mpeg', '.ts', '.mts'
            }
            return ext in video_extensions
        except (FileNotFoundError, PermissionError, OSError):
            return False

    def _is_archive_file(self, file_path: str) -> bool:
        """检查是否为压缩文件"""
        try:
            ext = Path(file_path).suffix.lower()
            archive_extensions = {
                '.zip', '.rar', '.7z', '.tar', '.gz', '.bz2', '.xz',
                '.lzma', '.Z', '.arj', '.lha', '.sit', '.ace'
            }
            return ext in archive_extensions
        except (FileNotFoundError, PermissionError, OSError):
            return False

    # 格式处理器方法
    async def handle_text(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            return {
                'status': 'success',
                'content_type': 'text',
                'content': content,
                'encoding': 'utf-8',
                'lines': len(content.splitlines()),
                'words': len(content.split()),
                'characters': len(content),
                'handler': 'text_handler'
            }
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'text_handler'}

    async def handle_markdown(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理Markdown文件"""
        try:
            if not MARKDOWN_AVAILABLE:
                return await self.handle_text(file_path, options)

            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # 解析Markdown
            md = markdown.Markdown(extensions=['tables', 'toc', 'codehilite'])
            html = md.convert(content)

            return {
                'status': 'success',
                'content_type': 'markdown',
                'content': content,
                'html': html,
                'handler': 'markdown_handler'
            }
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'markdown_handler'}

    async def handle_csv(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理CSV文件"""
        try:
            if PANDAS_AVAILABLE:
                df = pd.read_csv(file_path)

                return {
                    'status': 'success',
                    'content_type': 'csv_dataframe',
                    'shape': df.shape,
                    'columns': df.columns.tolist(),
                    'dtypes': df.dtypes.to_dict(),
                    'head': df.head(10).to_dict('records'),
                    'tail': df.tail(10).to_dict('records'),
                    'info': df.info().to_string(),
                    'describe': df.describe().to_dict(),
                    'handler': 'csv_handler'
                }
            else:
                # 纯Python CSV处理
                with open(file_path, 'r', encoding='utf-8') as f:
                    reader = csv.DictReader(f)
                    rows = list(reader)

                return {
                    'status': 'success',
                    'content_type': 'csv_text',
                    'rows': len(rows),
                    'columns': rows[0].keys() if rows else [],
                    'sample': rows[:10],
                    'handler': 'csv_handler'
                }
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'csv_handler'}

    async def handle_json(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理JSON文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            return {
                'status': 'success',
                'content_type': 'json',
                'data': data,
                'keys': list(data.keys()) if isinstance(data, dict) else [],
                'handler': 'json_handler'
            }
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'json_handler'}

    async def handle_xml(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理XML文件"""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()

            def xml_to_dict(element) -> None:
                result = {}
                if element.text and element.text.strip():
                    result['text'] = element.text.strip()
                if element.attrib:
                    result['attributes'] = element.attrib
                for child in element:
                    if child.tag in result:
                        if not isinstance(result[child.tag], list):
                            result[child.tag] = [result[child.tag]]
                        result[child.tag].append(xml_to_dict(child))
                    else:
                        result[child.tag] = xml_to_dict(child)
                return result

            return {
                'status': 'success',
                'content_type': 'xml',
                'root_tag': root.tag,
                'data': xml_to_dict(root),
                'handler': 'xml_handler'
            }
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'xml_handler'}

    async def handle_image(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理图像文件"""
        try:
            if PIL_AVAILABLE:
                with Image.open(file_path) as img:
                    return {
                        'status': 'success',
                        'content_type': 'image',
                        'size': img.size,
                        'mode': img.mode,
                        'format': img.format,
                        'info': img.info,
                        'handler': 'image_handler'
                    }
            else:
                # 使用OpenCV
                if OPENCV_AVAILABLE:
                    img = cv2.imread(file_path)
                    height, width = img.shape[:2]
                    return {
                        'status': 'success',
                        'content_type': 'image',
                        'size': (width, height),
                        'channels': img.shape[2] if len(img.shape) > 2 else 1,
                        'handler': 'image_handler'
                    }
                else:
                    # 基础文件信息
                    return await self.handle_default(file_path, options)
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'image_handler'}

    async def handle_pdf(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理PDF文件"""
        try:
            if PDF_AVAILABLE and pdfplumber:
                with pdfplumber.open(file_path) as pdf:
                    text_content = ''
                    pages_info = []

                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        text_content += page_text + "\n"

                        pages_info.append({
                            'page_number': page_num + 1,
                            'text_length': len(page_text),
                            'bbox': page.bbox,
                            'width': page.width,
                            'height': page.height
                        })

                    return {
                        'status': 'success',
                        'content_type': 'pdf',
                        'text': text_content,
                        'pages': len(pdf.pages),
                        'pages_info': pages_info,
                        'handler': 'pdf_handler'
                    }
            else:
                # 使用PyPDF2
                if PDF_AVAILABLE and PyPDF2:
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text_content = ''
                        pages_info = []

                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text()
                            text_content += page_text + "\n"

                            pages_info.append({
                                'page_number': page_num + 1,
                                'text_length': len(page_text),
                                'width': page.mediabox.width,
                                'height': page.mediabox.height
                            })

                        return {
                            'status': 'success',
                            'content_type': 'pdf',
                            'text': text_content,
                            'pages': len(pdf_reader.pages),
                            'pages_info': pages_info,
                            'handler': 'pdf_handler'
                        }
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'pdf_handler'}

    async def handle_xlsx(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理Excel XLSX文件"""
        try:
            if PANDAS_AVAILABLE and OPENPYXL_AVAILABLE:
                df = pd.read_excel(file_path)
                return {
                    'status': 'success',
                    'content_type': 'excel_xlsx',
                    'shape': df.shape,
                    'columns': df.columns.tolist(),
                    'dtypes': df.dtypes.to_dict(),
                    'head': df.head(10).to_dict('records'),
                    'info': df.info().to_string(),
                    'describe': df.describe().to_dict(),
                    'handler': 'xlsx_handler'
                }
            else:
                return await self.handle_default(file_path, options)
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'xlsx_handler'}

    async def handle_audio(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理音频文件"""
        try:
            if LIBROSA_AVAILABLE:
                y, sr = librosa.load(file_path)
                duration = librosa.get_duration(y=y, sr=sr)

                return {
                    'status': 'success',
                    'content_type': 'audio',
                    'duration': duration,
                    'sample_rate': sr,
                    'channels': 1 if y.ndim == 1 else y.shape[1],
                    'samples': len(y),
                    'handler': 'audio_handler'
                }
            else:
                return await self.handle_default(file_path, options)
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'audio_handler'}

    async def handle_video(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理视频文件"""
        try:
            # 使用ffprobe获取视频信息
            cmd = [
                'ffprobe', '-v', 'quiet', '-print_format', 'json',
                '-show_format', '-show_streams', file_path
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)

            if result.returncode == 0:
                data = json.loads(result.stdout)
                format_info = data.get('format', {})
                streams = data.get('streams', [])

                video_streams = [s for s in streams if s.get('codec_type') == 'video']
                audio_streams = [s for s in streams if s.get('codec_type') == 'audio']

                return {
                    'status': 'success',
                    'content_type': 'video',
                    'format': format_info,
                    'video_streams': len(video_streams),
                    'audio_streams': len(audio_streams),
                    'total_streams': len(streams),
                    'handler': 'video_handler'
                }
            else:
                return await self.handle_default(file_path, options)
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'video_handler'}

    async def handle_zip(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理ZIP压缩文件"""
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                file_list = zip_file.namelist()
                file_info = []

                for file_name in file_list:
                    info = zip_file.getinfo(file_name)
                    file_info.append({
                        'name': file_name,
                        'size': info.file_size,
                        'compressed_size': info.compress_size,
                        'date_time': info.date_time,
                        'is_dir': info.is_dir()
                    })

                return {
                    'status': 'success',
                    'content_type': 'archive_zip',
                    'files': file_list,
                    'file_count': len(file_list),
                    'file_info': file_info,
                    'total_size': info.file_size,
                    'handler': 'zip_handler'
                }
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'zip_handler'}

    async def handle_default(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """默认处理器"""
        try:
            path = Path(file_path)
            stat = path.stat()

            return {
                'status': 'success',
                'content_type': 'unknown',
                'file_size': stat.st_size,
                'file_name': path.name,
                'extension': path.suffix,
                'handler': 'default_handler'
            }
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'default_handler'}

    # 其他处理器的占位符方法
    async def handle_latex(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理LaTeX文件"""
        return await self.handle_text(file_path, options)

    async def handle_yaml(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理YAML文件"""
        return await self.handle_text(file_path, options)

    async def handle_toml(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理TOML文件"""
        return await self.handle_text(file_path, options)

    async def handle_code(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理代码文件"""
        return await self.handle_text(file_path, options)

    async def handle_jupyter(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理Jupyter Notebook"""
        try:
            import nbformat
            with open(file_path, 'r', encoding='utf-8') as f:
                notebook = nbformat.read(f, as_version=4)

            return {
                'status': 'success',
                'content_type': 'jupyter_notebook',
                'cells': len(notebook.cells),
                'notebook': notebook,
                'metadata': notebook.metadata,
                'handler': 'jupyter_handler'
            }
        except Exception as e:
            return await self.handle_text(file_path, options)

    async def handle_docx(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理Word DOCX文件"""
        try:
            if DOCX_AVAILABLE:
                doc = docx.Document(file_path)
                text_content = '\n'.join([para.text for para in doc.paragraphs])

                return {
                    'status': 'success',
                    'content_type': 'word_docx',
                    'text': text_content,
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables),
                    'handler': 'docx_handler'
                }
            else:
                return await self.handle_default(file_path, options)
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'docx_handler'}

    async def handle_doc(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理Word DOC文件"""
        return await self.handle_default(file_path, options)

    async def handle_xls(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理Excel XLS文件"""
        try:
            if PANDAS_AVAILABLE and XLDR_AVAILABLE:
                df = pd.read_excel(file_path)
                return {
                    'status': 'success',
                    'content_type': 'excel_xls',
                    'shape': df.shape,
                    'columns': df.columns.tolist(),
                    'head': df.head(10).to_dict('records'),
                    'handler': 'xls_handler'
                }
            else:
                return await self.handle_default(file_path, options)
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'xls_handler'}

    async def handle_pptx(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理PowerPoint PPTX文件"""
        return await self.handle_default(file_path, options)

    async def handle_ppt(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理PowerPoint PPT文件"""
        return await self.handle_default(file_path, options)

    async def handle_svg(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理SVG矢量图"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            return {
                'status': 'success',
                'content_type': 'svg_vector',
                'content': content,
                'handler': 'svg_handler'
            }
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'svg_handler'}

    async def handle_ai(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理Adobe Illustrator文件"""
        return await self.handle_default(file_path, options)

    async def handle_eps(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理PostScript文件"""
        return await self.handle_default(file_path, options)

    async def handle_ps(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理PostScript文件"""
        return await self.handle_default(file_path, options)

    async def handle_psd(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理Photoshop文件"""
        return await self.handle_default(file_path, options)

    async def handle_dxf(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理DXF文件"""
        return await self.handle_default(file_path, options)

    async def handle_dwg(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理DWG文件"""
        return await self.handle_default(file_path, options)

    async def handle_midi(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理MIDI文件"""
        return await self.handle_default(file_path, options)

    async def handle_stl(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理STL 3D模型"""
        return await self.handle_default(file_path, options)

    async def handle_obj(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理OBJ 3D模型"""
        return await self.handle_default(file_path, options)

    async def handle_ply(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理PLY 3D模型"""
        return await self.handle_default(file_path, options)

    async def handle_fbx(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理FBX 3D模型"""
        return await self.handle_default(file_path, options)

    async def handle_hdf5(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理HDF5文件"""
        try:
            if H5PY_AVAILABLE:
                with h5py.File(file_path, 'r') as f:
                    return {
                        'status': 'success',
                        'content_type': 'hdf5',
                        'keys': list(f.keys()),
                        'handler': 'hdf5_handler'
                    }
            else:
                return await self.handle_default(file_path, options)
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'hdf5_handler'}

    async def handle_netcdf(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理NetCDF文件"""
        try:
            if NETCDF4_AVAILABLE:
                with net_cdf4.Dataset(file_path) as ds:
                    return {
                        'status': 'success',
                        'content_type': 'netcdf',
                        'dimensions': dict(ds.dimensions),
                        'variables': list(ds.variables.keys()),
                        'attributes': dict(ds.ncattrs()),
                        'handler': 'netcdf_handler'
                    }
            else:
                return await self.handle_default(file_path, options)
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'netcdf_handler'}

    async def handle_mat(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理MATLAB文件"""
        return await self.handle_default(file_path, options)

    async def handle_parquet(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理Parquet文件"""
        try:
            if PYARROW_AVAILABLE:
                table = pa.read_table(file_path)
                return {
                    'status': 'success',
                    'content_type': 'parquet',
                    'shape': table.shape,
                    'columns': table.column_names,
                    'schema': table.schema,
                    'handler': 'parquet_handler'
                }
            else:
                return await self.handle_default(file_path, options)
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'parquet_handler'}

    async def handle_arrow(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理Arrow文件"""
        try:
            if PYARROW_AVAILABLE:
                    with pa.ipc.open_stream(file_path) as reader:
                        table = reader.read_all()
                        return {
                            'status': 'success',
                            'content_type': 'arrow',
                            'shape': table.shape,
                            'columns': table.column_names,
                            'schema': table.schema,
                            'handler': 'arrow_handler'
                        }
            else:
                return await self.handle_default(file_path, options)
        except Exception as e:
            return {'status': 'error', 'error': str(e), 'handler': 'arrow_handler'}

    async def handle_tar(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理TAR文件"""
        return await self.handle_default(file_path, options)

    async def handle_gzip(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理GZIP文件"""
        return await self.handle_default(file_path, options)

    async def handle_rar(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理RAR文件"""
        return await self.handle_default(file_path, options)

    async def handle_sevenz(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理7Z文件"""
        return await self.handle_default(file_path, options)

    async def handle_executable(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理可执行文件"""
        return await self.handle_default(file_path, options)

    async def handle_library(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理库文件"""
        return await self.handle_default(file_path, options)

    async def handle_object(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理目标文件"""
        return await self.handle_default(file_path, options)

    def get_supported_extensions(self) -> Dict[str, List[str]]:
        """获取所有支持的文件扩展名"""
        extensions = {}

        for category, formats in self.supported_formats.items():
            category_extensions = []
            for subcategory, format_dict in formats.items():
                category_extensions.extend(format_dict.keys())
            extensions[category] = category_extensions

        return extensions

    def get_handler(self, format_info: Dict[str, Any]) -> callable:
        """获取格式处理器"""
        handler_name = format_info.get('handler', 'default_handler')
        return self.format_handlers.get(handler_name, self.format_handlers['default_handler'])

    async def process_file(self, file_path: str, options: Dict = None) -> Dict[str, Any]:
        """处理文件"""
        # 获取文件格式信息
        format_info = self.detect_format(file_path)

        if not format_info:
            # 无法识别格式，使用默认处理器
            return await self.handle_default(file_path, options)

        # 获取处理器
        handler = self.get_handler(format_info)

        # 处理文件
        result = await handler(file_path, options)

        # 添加格式信息
        result['format_info'] = format_info

        return result

def demo_enhanced_support() -> Any:
    """演示增强的文件格式支持"""
    logger.info('🚀 增强文件格式支持演示')
    logger.info(str('=' * 60))

    support = EnhancedFileFormatSupport()

    logger.info(f"📊 支持的文件类别数量: {len(support.supported_formats)}")

    # 统计支持格式数量
    total_formats = 0
    for category, formats in support.supported_formats.items():
        category_count = sum(len(format_dict) for format_dict in formats.values())
        total_formats += category_count
        logger.info(f"   {category}: {category_count} 种格式")

    logger.info(f"\n📈 总共支持: {total_formats} 种文件格式")

    # 显示各类别详细格式
    for category, formats in support.supported_formats.items():
        logger.info(f"\n📂 {category.upper()}")
        logger.info(str('-' * 40))
        for subcategory, format_dict in formats.items():
            logger.info(f"  📁 {subcategory}:")
            for ext, info in format_dict.items():
                logger.info(f"    .{ext} - {info['name']} ({info.get('mime', 'unknown')})")

    logger.info(f"\n💡 可用处理器:")
    available_handlers = []
    for name, handler in support.format_handlers.items():
        if handler != support.format_handlers['default_handler']:
            available_handlers.append(name)

    if available_handlers:
        logger.info(f"   {', '.join(available_handlers)}")

    logger.info(f"\n📋 库依赖状态:")
    dependencies = {
        'Pandas': PANDAS_AVAILABLE,
        'PIL': PIL_AVAILABLE,
        'OpenCV': OPENCV_AVAILABLE,
        'Librosa': LIBROSA_AVAILABLE,
        'PDF Libraries': PDF_AVAILABLE,
        'DocX': DOCX_AVAILABLE,
        'Markdown': MARKDOWN_AVAILABLE,
        'HDF5': H5PY_AVAILABLE,
        'PyArrow': PYARROW_AVAILABLE,
        'NetCDF4': NETCDF4_AVAILABLE,
        'OpenPyXL': OPENPYXL_AVAILABLE,
        'xlrd': XLDR_AVAILABLE,
    }

    for lib, available in dependencies.items():
        status = '✅' if available else '❌'
        logger.info(f"   {status} {lib}")

if __name__ == '__main__':
    demo_enhanced_support()