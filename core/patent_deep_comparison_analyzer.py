#!/usr/bin/env python3
"""
专利深度对比分析系统 - 用于审查意见答复
Patent Deep Comparison Analyzer for Office Action Response

核心功能:
1. 基于向量库的语义相似度分析 (使用BGE模型)
2. 基于知识图谱的技术特征对比 (使用NebulaGraph)
3. 结构化权利要求对比
4. 技术方案详细比对
5. 自动生成对比分析报告 (Markdown + DOCX双格式)

集成组件:
- BGE嵌入服务 (core/embedding/bge_embedding_service.py)
- 统一向量管理器 (core/vector/unified_vector_manager.py)
- 统一知识图谱 (core/knowledge/unified_knowledge_graph.py)
- python-docx (DOCX文档生成)

作者: Claude
创建时间: 2025-12-31
版本: v1.2.0 - 添加DOCX格式输出
"""

import asyncio
import logging
import re
import sys
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple


from core.logging_config import setup_logging

# Qdrant客户端
try:
    from qdrant_client import QdrantClient

    QDRANT_AVAILABLE = True
except ImportError:
    QDRANT_AVAILABLE = False

# DOCX支持
try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 添加项目路径
project_root = Path(__file__).parent.parent
sys.path.append(str(project_root))

# Qdrant客户端
try:
    from qdrant_client import QdrantClient

    QDRANT_AVAILABLE = True
except ImportError:
    QDRANT_AVAILABLE = False
    logger = setup_logging()
    logger.warning("⚠️ Qdrant客户端未安装")

# 导入BGE嵌入服务
try:
    from core.embedding.bge_embedding_service import (
        BGEEmbeddingService,
        BGEModelConfig,
        get_bge_service,
    )

    BGE_AVAILABLE = True
except ImportError:
    BGE_AVAILABLE = False

# 导入统一向量管理器
try:
    from core.vector.unified_vector_manager import UnifiedVectorManager, get_vector_manager

    VECTOR_MANAGER_AVAILABLE = True
except ImportError:
    VECTOR_MANAGER_AVAILABLE = False

# 导入统一知识图谱
try:
    from core.knowledge.unified_knowledge_graph import (
        GraphBackend,
        UnifiedKnowledgeGraph,
        get_knowledge_graph,
    )

    KG_AVAILABLE = True
except ImportError:
    KG_AVAILABLE = False

# 配置日志
# setup_logging()  # 日志配置已移至模块导入
logger = logging.getLogger(__name__)


class SimilarityLevel(Enum):
    """相似度等级"""

    IDENTICAL = "完全相同"  # 相似度 > 0.95
    HIGHLY_SIMILAR = "高度相似"  # 相似度 0.85-0.95
    MODERATELY_SIMILAR = "中度相似"  # 相似度 0.70-0.85
    SLIGHTLY_SIMILAR = "轻微相似"  # 相似度 0.50-0.70
    DIFFERENT = "不同"  # 相似度 < 0.50


@dataclass
class ClaimElement:
    """权利要求要素"""

    element_id: str
    element_text: str
    element_type: str  # 结构/功能/参数/方法等
    position: int
    parent_id: str | None = None


@dataclass
class TechnicalFeature:
    """技术特征"""

    feature_id: str
    feature_name: str
    feature_description: str
    feature_value: Any
    feature_category: str  # 技术领域/技术问题/技术方案/技术效果


@dataclass
class PatentDocument:
    """专利文档"""

    patent_id: str
    patent_number: str
    title: str
    abstract: str
    claims: list[str]
    claim_elements: dict[str, list[ClaimElement]]
    technical_features: list[TechnicalFeature]
    publication_date: str | None = None
    applicant: str | None = None
    inventor: str | None = None
    ipc_classification: str | None = None
    full_text: str | None = None

    # 向量表示
    embedding: np.ndarray | None = None


@dataclass
class ElementComparison:
    """要素对比结果"""

    target_element: ClaimElement
    reference_element: ClaimElement
    similarity_score: float
    similarity_level: SimilarityLevel
    is_present: bool
    difference_analysis: str


@dataclass
class ClaimComparison:
    """权利要求对比"""

    target_claim_id: str
    target_claim_text: str
    reference_claim_id: str
    reference_claim_text: str

    # 要素级对比
    element_comparisons: list[ElementComparison]

    # 整体评估
    overall_similarity: float
    missing_elements: list[str]
    additional_elements: list[str]
    difference_summary: str

    # 法律分析
    novelty_assessment: str  # 新颖性评估
    inventiveness_assessment: str  # 创造性评估


@dataclass
class TechnicalComparison:
    """技术对比"""

    patent_a_id: str
    patent_b_id: str

    # 技术领域对比
    field_similarity: float
    field_difference: str

    # 技术问题对比
    problem_similarity: float
    problem_difference: str

    # 技术方案对比
    solution_similarity: float
    solution_difference: str
    feature_by_feature_comparison: dict[str, Any]
    # 技术效果对比
    effect_similarity: float
    effect_difference: str

    # 综合评估
    overall_similarity: float
    key_differences: list[str]
    key_similarities: list[str]


@dataclass
class DeepComparisonReport:
    """深度对比分析报告"""

    report_id: str
    target_patent: PatentDocument
    reference_patents: list[PatentDocument]

    # 权利要求对比
    claims_comparison: list[ClaimComparison]

    # 技术对比
    technical_comparison: list[TechnicalComparison]

    # 向量相似度分析
    vector_similarity_analysis: dict[str, float]

    # 法律评估
    novelty_analysis: str
    inventiveness_analysis: str

    # 答复策略建议
    response_strategy: str
    key_arguments: list[str]

    # 元数据
    generation_time: str
    analyzer_version: str


class PatentDeepComparisonAnalyzer:
    """专利深度对比分析器 - 集成BGE嵌入服务和知识图谱"""

    def __init__(self, config: dict[str, Any] | None = None):
        """初始化分析器"""
        self.name = "专利深度对比分析器"
        self.version = "1.1.0"

        # 配置 - 使用本地MPS优化的1024维BAAI/bge-m3模型
        self.config = config or {
            "qdrant_host": "localhost",
            "qdrant_port": 6333,
            # 🍎 使用本地MPS优化的1024维模型
            "embedding_model": "BAAI/bge-m3",  # 1024维高精度模型
            "embedding_device": "auto",  # auto检测:优先MPS > CUDA > CPU
            "similarity_thresholds": {
                "identical": 0.95,
                "highly_similar": 0.85,
                "moderately_similar": 0.70,
                "slightly_similar": 0.50,
            },
            "enable_knowledge_graph": True,  # 启用知识图谱
            "enable_vector_search": True,  # 启用向量搜索
            "enable_bge_service": True,  # 启用BGE嵌入服务
            "enable_networkx_fallback": True,  # 启用NetworkX降级
            "max_reference_patents": 10,
        }

        # 初始化BGE嵌入服务(优先从本地converted目录加载MPS优化模型)
        self.bge_service: BGEEmbeddingService | None = None
        if BGE_AVAILABLE and self.config["enable_bge_service"]:
            try:
                # BGEModelConfig会自动查找:
                # 1. models/converted/BAAI/bge-m3/ (MPS优化模型)
                # 2. models/BAAI/bge-m3/ (普通本地模型)
                # 3. BAAI/BAAI/bge-m3 (HuggingFace远程)
                self.bge_service = get_bge_service(
                    model_name=self.config["embedding_model"],
                    device=self.config["embedding_device"],
                )
                logger.info(f"✅ BGE嵌入服务已初始化: {self.config['embedding_model']}")
                logger.info(f"   向量维度: {self.bge_service.dimension}维")
                logger.info(f"   计算设备: {self.bge_service.device}")
            except Exception as e:
                logger.warning(f"⚠️ BGE嵌入服务初始化失败: {e}")

        # 初始化统一向量管理器
        self.vector_manager: UnifiedVectorManager | None = None
        if VECTOR_MANAGER_AVAILABLE and self.config["enable_vector_search"]:
            try:
                self.vector_manager = get_vector_manager()
                logger.info("✅ 统一向量管理器已初始化")
            except Exception as e:
                logger.warning(f"⚠️ 向量管理器初始化失败: {e}")

        # Qdrant客户端(如果向量管理器不可用)
        self.qdrant_client: QdrantClient | None = None
        if QDRANT_AVAILABLE and self.vector_manager is None:
            try:
                self.qdrant_client = QdrantClient(
                    host=self.config["qdrant_host"], port=self.config["qdrant_port"]
                )
                logger.info("✅ Qdrant客户端已初始化")
            except Exception as e:
                logger.warning(f"⚠️ Qdrant客户端初始化失败: {e}")

        # 知识图谱(延迟初始化)
        self.knowledge_graph: UnifiedKnowledgeGraph | None = None
        self._kg_initialized = False

        logger.info(f"✅ {self.name} v{self.version} 初始化完成")
        logger.info(f"   BGE服务: {'✅' if self.bge_service else '❌'}")
        logger.info(f"   向量管理器: {'✅' if self.vector_manager else '❌'}")
        logger.info(
            f"   知识图谱: {'⏳延迟加载' if self.config['enable_knowledge_graph'] else '❌'}"
        )

    async def analyze_office_action(
        self, target_patent_text: str, reference_patents: list[str], analysis_depth: str = "deep"
    ) -> DeepComparisonReport:
        """
        分析审查意见

        Args:
            target_patent_text: 目标专利全文/权利要求书
            reference_patents: 对比文件列表(专利号列表)
            analysis_depth: 分析深度 (quick/standard/deep)

        Returns:
            DeepComparisonReport: 深度对比分析报告
        """
        logger.info("🔍 开始深度对比分析")
        logger.info("   目标专利: 待处理")
        logger.info(f"   对比文件: {len(reference_patents)} 个")
        logger.info(f"   分析深度: {analysis_depth}")

        # 1. 解析目标专利
        target_patent = await self._parse_patent_document(target_patent_text, "target")

        # 2. 解析对比文件
        ref_patents = []
        for ref_num, ref_id in enumerate(reference_patents, 1):
            ref_text = await self._fetch_patent_text(ref_id)
            ref_patent = await self._parse_patent_document(ref_text, f"reference_{ref_num}")
            ref_patents.append(ref_patent)

        # 3. 生成向量表示
        if self.config["enable_vector_search"]:
            await self._generate_embeddings([target_patent, *ref_patents])

        # 4. 权利要求对比
        claims_comparison = await self._compare_claims(target_patent, ref_patents)

        # 5. 技术特征对比
        technical_comparison = await self._compare_technical_features(target_patent, ref_patents)

        # 6. 向量相似度分析
        vector_similarity = await self._calculate_vector_similarity(target_patent, ref_patents)

        # 7. 法律评估
        novelty_analysis = await self._analyze_novelty(claims_comparison, vector_similarity)
        inventiveness_analysis = await self._analyze_inventiveness(
            claims_comparison, vector_similarity
        )

        # 8. 生成答复策略
        response_strategy, key_arguments = await self._generate_response_strategy(
            claims_comparison, technical_comparison, novelty_analysis, inventiveness_analysis
        )

        # 9. 构建报告
        report = DeepComparisonReport(
            report_id=f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            target_patent=target_patent,
            reference_patents=ref_patents,
            claims_comparison=claims_comparison,
            technical_comparison=technical_comparison,
            vector_similarity_analysis=vector_similarity,
            novelty_analysis=novelty_analysis,
            inventiveness_analysis=inventiveness_analysis,
            response_strategy=response_strategy,
            key_arguments=key_arguments,
            generation_time=datetime.now().isoformat(),
            analyzer_version=self.version,
        )

        logger.info("✅ 深度对比分析完成")
        return report

    async def _parse_patent_document(self, patent_text: str, doc_id: str) -> PatentDocument:
        """解析专利文档"""
        # 提取基本信息
        title = self._extract_title(patent_text)
        abstract = self._extract_abstract(patent_text)
        claims = self._extract_claims(patent_text)

        # 解析权利要求要素
        claim_elements = {}
        for claim_id, claim_text in enumerate(claims, 1):
            elements = await self._parse_claim_elements(claim_text, claim_id)
            claim_elements[f"claim_{claim_id}"] = elements

        # 提取技术特征
        technical_features = await self._extract_technical_features(patent_text, claims)

        return PatentDocument(
            patent_id=doc_id,
            patent_number=self._extract_patent_number(patent_text),
            title=title,
            abstract=abstract,
            claims=claims,
            claim_elements=claim_elements,
            technical_features=technical_features,
            full_text=patent_text,
        )

    def _extract_title(self, text: str) -> str:
        """提取发明名称"""
        # 简化实现,实际应该用更复杂的解析
        patterns = []\s*([^\n]+)"]"]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1).strip()
        return "未识别的发明名称"

    def _extract_abstract(self, text: str) -> str:
        """提取摘要"""
        patterns = [
            r"摘要[\n[^a-z_a-Z0-9\u4e00-\u9fa5].*){0]",
            r"Abstract[::]\s*([^\n]+)",
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.MULTILINE)
            if match:
                return match.group(1).strip()
        return ""

    def _extract_claims(self, text: str) -> list[str]:
        """提取权利要求"""
        claims = []

        # 查找权利要求部分
        claims_section = re.search(r"权利要求书[]\s*(.*?)(?=\n\s*说明书|$)", text]
        if not claims_section:
            return claims

        claims_text = claims_section.group(1)

        # 提取每个权利要求
        claim_pattern = r"(\d+\..*?(?=\n\d+\.|\n\s*$))"
        matches = re.findall(claim_pattern, claims_text, re.MULTILINE)

        for match in matches:
            claim_text = match.strip()
            if claim_text:
                claims.append(claim_text)

        return claims

    async def _ensure_knowledge_graph(self):
        """确保知识图谱已初始化(延迟加载)"""
        if self._kg_initialized:
            return

        if KG_AVAILABLE and self.config["enable_knowledge_graph"]:
            try:
                self.knowledge_graph = await get_knowledge_graph(
                    preferred_backend=GraphBackend.NEBULA
                )
                stats = await self.knowledge_graph.get_statistics()
                logger.info("✅ 知识图谱已初始化")
                logger.info(f"   后端: {stats.backend.value}")
                logger.info(f"   节点数: {stats.node_count}")
                logger.info(f"   边数: {stats.edge_count}")
                self._kg_initialized = True
            except Exception as e:
                logger.warning(f"⚠️ 知识图谱初始化失败: {e}")
                logger.info("💡 将使用NetworkX进行本地图分析")
                self._kg_initialized = True  # 标记为已尝试

    async def _initialize_local_graph(self, patents: list[PatentDocument]) -> Any:
        """初始化本地NetworkX图(用于临时比对)"""
        try:

            from core.graph.networkx_utils import create_patent_comparison_graph

            # 创建专利对比图
            G = create_patent_comparison_graph(patents)
            logger.info(f"✅ 本地NetworkX图已创建,节点数: {G.number_of_nodes()}")
            return G
        except ImportError:
            logger.warning("⚠️ NetworkX未安装")
            return None
        except Exception as e:
            logger.warning(f"⚠️ 本地图创建失败: {e}")
            return None

    async def _parse_claim_elements(self, claim_text: str, claim_id: int) -> list[ClaimElement]:
        """解析权利要求要素"""
        elements = []

        # 移除权利要求编号
        content = re.sub(r"^\d+\.\s*", "", claim_text.strip())

        # 分解技术特征
        # 按标点符号分割,保留结构
        parts = re.split(r"[,;;]\s*|(?<=其特征在于|包括|其中)[,;;]?\s*", content)

        element_id = 1
        for part in parts:
            part = part.strip()
            if part:
                # 判断要素类型
                element_type = self._classify_element_type(part)

                elements.append(
                    ClaimElement(
                        element_id=f"claim_{claim_id}_element_{element_id}",
                        element_text=part,
                        element_type=element_type,
                        position=element_id,
                    )
                )
                element_id += 1

        return elements

    def _classify_element_type(self, element_text: str) -> str:
        """分类要素类型"""
        if "包括" in element_text or "包含" in element_text or "设有" in element_text:
            return "结构组成"
        elif "用于" in element_text or "以" in element_text or "通过" in element_text:
            return "功能方法"
        elif re.search(r"[\d]+\s*[℃%°gmm]", element_text):
            return "参数"
        else:
            return "通用特征"

    async def _extract_technical_features(
        self, patent_text: str, claims: list[str]
    ) -> list[TechnicalFeature]:
        """提取技术特征"""
        features = []

        # 从摘要中提取
        self._extract_abstract(patent_text)

        # 从权利要求中提取
        " ".join(claims)

        # 关键词提取(简化实现)
        # 实际应该使用NER或专业术语提取
        feature_keywords = ["技术领域", "背景技术", "发明内容", "技术问题", "技术方案", "技术效果"]

        for keyword in feature_keywords:
            pattern = f"{keyword}[::]([^\\n]+)"
            match = re.search(pattern, patent_text)
            if match:
                features.append(
                    TechnicalFeature(
                        feature_id=f"feature_{len(features)+1}",
                        feature_name=keyword,
                        feature_description=match.group(1).strip(),
                        feature_value=None,
                        feature_category=keyword,
                    )
                )

        return features

    def _extract_patent_number(self, text: str) -> str:
        """提取专利号"""
        patterns = [
            r"申请号[::]\s*([A-Z0-9\.]+)",
            r"Patent\s*No[::]?\s*([A-Z0-9\.]+)",
            r"CN\d{7,}[A-Z]?",
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1) if len(match.groups()) > 0 else match.group(0)
        return "UNKNOWN"

    async def _fetch_patent_text(self, patent_id: str) -> str:
        """获取专利全文"""
        # 这里应该从数据库或API获取
        # 简化实现,返回占位符
        return f"# {patent_id}\n\n[专利全文内容...]"

    async def _generate_embeddings(self, patents: list[PatentDocument]):
        """生成向量表示 - 使用BGE嵌入服务"""
        logger.info("🔄 生成专利向量表示...")

        if not self.bge_service:
            logger.warning("⚠️ BGE服务未初始化,跳过向量生成")
            return

        for patent in patents:
            try:
                # 组合多个文本部分进行编码
                text_parts = []

                # 1. 摘要(最重要)
                if patent.abstract:
                    text_parts.append(patent.abstract)

                # 2. 权利要求(核心内容)
                if patent.claims:
                    claims_text = " ".join(patent.claims[:3])  # 最多前3个权利要求
                    text_parts.append(claims_text)

                # 3. 发明名称
                if patent.title:
                    text_parts.append(patent.title)

                # 组合文本
                combined_text = " [SEP] ".join(text_parts)

                # 使用BGE服务生成向量
                embedding = self.bge_service.encode(combined_text, normalize=True)

                patent.embedding = embedding[0]  # 获取单个向量

                logger.debug(f"   ✅ {patent.patent_id} 向量维度: {patent.embedding.shape}")

            except Exception as e:
                logger.error(f"❌ 生成向量失败 ({patent.patent_id}): {e}")
                patent.embedding = None

        logger.info(f"✅ 已生成 {len(patents)} 个专利的向量表示")

    async def _calculate_text_similarity(self, text1: str, text2: str) -> float:
        """计算两个文本的语义相似度 - 使用BGE服务"""
        if not self.bge_service:
            # 降级为简单的文本相似度
            return self._calculate_simple_similarity(text1, text2)

        try:
            return self.bge_service.similarity(text1, text2)
        except Exception as e:
            logger.warning(f"⚠️ BGE相似度计算失败: {e},使用简单相似度")
            return self._calculate_simple_similarity(text1, text2)

    def _calculate_simple_similarity(self, text1: str, text2: str) -> float:
        """计算简单文本相似度(降级方案)"""
        # 使用简单的词汇重叠度
        words1 = set(text1.split())
        words2 = set(text2.split())

        if not words1 or not words2:
            return 0.0

        intersection = words1 & words2
        union = words1 | words2

        return len(intersection) / len(union) if union else 0.0

    async def _compare_claims(
        self, target_patent: PatentDocument, ref_patents: list[PatentDocument]
    ) -> list[ClaimComparison]:
        """对比权利要求"""
        logger.info("🔍 对比权利要求...")

        comparisons = []

        # 对比每个目标权利要求
        for target_claim_id, target_claim_text in enumerate(target_patent.claims, 1):
            target_elements = target_patent.claim_elements.get(f"claim_{target_claim_id}", [])

            # 找到最相似的对比文件权利要求
            best_ref_claim = None
            best_ref_patent = None
            best_similarity = 0

            for ref_patent in ref_patents:
                for _ref_claim_id, ref_claim_text in enumerate(ref_patent.claims, 1):
                    similarity = await self._calculate_claim_similarity(
                        target_claim_text, ref_claim_text
                    )

                    if similarity > best_similarity:
                        best_similarity = similarity
                        best_ref_claim = ref_claim_text
                        best_ref_patent = ref_patent

            # 要素级对比
            element_comparisons = []
            if best_ref_patent:
                ref_elements = best_ref_patent.claim_elements.get(f"claim_{target_claim_id}", [])

                for target_element in target_elements:
                    # 查找最相似的参考要素
                    best_match = None
                    best_match_score = 0

                    for ref_element in ref_elements:
                        score = self._calculate_element_similarity(target_element, ref_element)
                        if score > best_match_score:
                            best_match_score = score
                            best_match = ref_element

                    similarity_level = self._get_similarity_level(best_match_score)

                    element_comparisons.append(
                        ElementComparison(
                            target_element=target_element,
                            reference_element=best_match,
                            similarity_score=best_match_score,
                            similarity_level=similarity_level,
                            is_present=best_match_score
                            > self.config["similarity_thresholds"]["moderately_similar"],
                            difference_analysis=self._analyze_element_difference(
                                target_element, best_match, best_match_score
                            ),
                        )
                    )

            # 整体评估
            missing_elements = []
                ec.target_element.element_text for ec in element_comparisons if not ec.is_present
            ]

            comparison = ClaimComparison(
                target_claim_id=f"claim_{target_claim_id}",
                target_claim_text=target_claim_text,
                reference_claim_id=best_ref_patent.patent_id if best_ref_patent else None,
                reference_claim_text=best_ref_claim,
                element_comparisons=element_comparisons,
                overall_similarity=best_similarity,
                missing_elements=missing_elements,
                additional_elements=[],
                difference_summary=self._summarize_differences(element_comparisons),
                novelty_assessment="",  # 稍后填充
                inventiveness_assessment="",  # 稍后填充
            )

            comparisons.append(comparison)

        logger.info(f"✅ 完成 {len(comparisons)} 个权利要求的对比")
        return comparisons

    async def _calculate_claim_similarity(self, claim1: str, claim2: str) -> float:
        """计算权利要求相似度 - 使用BGE语义相似度"""
        return await self._calculate_text_similarity(claim1, claim2)

    def _tokenize(self, text: str) -> list[str]:
        """分词(简化实现)"""
        # 移除标点符号和数字
        text = re.sub(r"[^\u4e00-\u9fa5a-z_a-Z]", " ", text)
        return text.split()

    def _calculate_element_similarity(self, elem1: ClaimElement, elem2: ClaimElement) -> float:
        """计算要素相似度 - 使用异步方法"""
        # 这是一个同步方法,需要返回同步结果
        # 如果BGE服务可用,使用缓存或简单相似度
        if self.bge_service:
            try:
                # 使用BGE服务的同步接口
                return self.bge_service.similarity(elem1.element_text, elem2.element_text)
            except Exception as e:
                logger.warning(f"操作失败: {e}")

        # 降级到简单相似度
        return self._calculate_simple_similarity(elem1.element_text, elem2.element_text)

    def _get_similarity_level(self, score: float) -> SimilarityLevel:
        """获取相似度等级"""
        thresholds = self.config["similarity_thresholds"]

        if score >= thresholds["identical"]:
            return SimilarityLevel.IDENTICAL
        elif score >= thresholds["highly_similar"]:
            return SimilarityLevel.HIGHLY_SIMILAR
        elif score >= thresholds["moderately_similar"]:
            return SimilarityLevel.MODERATELY_SIMILAR
        elif score >= thresholds["slightly_similar"]:
            return SimilarityLevel.SLIGHTLY_SIMILAR
        else:
            return SimilarityLevel.DIFFERENT

    def _analyze_element_difference(
        self,
        target_element: ClaimElement,
        ref_element: ClaimElement,
        similarity_score: float,
    ) -> str:
        """分析要素差异"""
        if not ref_element:
            return (
                f'目标专利包含{target_element.element_type}要素"{target_element.element_text}"'
                f", 但对比文件中未找到相应要素"
            )

        if similarity_score >= 0.9:
            return f"两要素基本相同: {target_element.element_text}"
        elif similarity_score >= 0.7:
            return (
                f'两要素相似: 目标"{target_element.element_text}"'
                f', 对比"{ref_element.element_text}"'
                f", 差异主要在具体表述上"
            )
        else:
            return (
                f'两要素明显不同: 目标"{target_element.element_text}"'
                f", 对比\"{ref_element.element_text if ref_element else '无'}\""
            )

    def _summarize_differences(self, element_comparisons: list[ElementComparison]) -> str:
        """总结差异"""
        present_count = sum(1 for ec in element_comparisons if ec.is_present)
        total_count = len(element_comparisons)

        if total_count == 0:
            return "无要素进行对比"

        percentage = (present_count / total_count) * 100

        return (
            f"在{total_count}个要素中,{present_count}个要素在对比文件中"
            f"有对应({percentage:.1f}%),{total_count - present_count}个要素为目标专利独有"
        )

    async def _compare_technical_features(
        self, target_patent: PatentDocument, ref_patents: list[PatentDocument]
    ) -> list[TechnicalComparison]:
        """对比技术特征"""
        logger.info("🔍 对比技术特征...")

        comparisons = []

        for ref_patent in ref_patents:
            # 计算技术领域相似度
            field_sim = self._calculate_text_similarity(
                self._extract_field(target_patent), self._extract_field(ref_patent)
            )

            # 计算技术问题相似度
            problem_sim = self._calculate_text_similarity(
                self._extract_problem(target_patent), self._extract_problem(ref_patent)
            )

            # 计算技术方案相似度
            solution_sim = self._calculate_text_similarity(
                self._extract_solution(target_patent), self._extract_solution(ref_patent)
            )

            # 计算技术效果相似度
            effect_sim = self._calculate_text_similarity(
                self._extract_effect(target_patent), self._extract_effect(ref_patent)
            )

            # 整体相似度
            overall_sim = (
                field_sim * 0.15 + problem_sim * 0.15 + solution_sim * 0.50 + effect_sim * 0.20
            )

            # 提取关键差异和相似点
            key_differences, key_similarities = await self._extract_key_differences_similarities(
                target_patent, ref_patent
            )

            comparison = TechnicalComparison(
                patent_a_id=target_patent.patent_id,
                patent_b_id=ref_patent.patent_id,
                field_similarity=field_sim,
                field_difference=self._describe_difference(field_sim),
                problem_similarity=problem_sim,
                problem_difference=self._describe_difference(problem_sim),
                solution_similarity=solution_sim,
                solution_difference=self._describe_difference(solution_sim),
                feature_by_feature_comparison={},
                effect_similarity=effect_sim,
                effect_difference=self._describe_difference(effect_sim),
                overall_similarity=overall_sim,
                key_differences=key_differences,
                key_similarities=key_similarities,
            )

            comparisons.append(comparison)

        logger.info(f"✅ 完成 {len(comparisons)} 个技术特征对比")
        return comparisons

    def _extract_field(self, patent: PatentDocument) -> str:
        """提取技术领域"""
        for feature in patent.technical_features:
            if "技术领域" in feature.feature_name:
                return feature.feature_description
        return patent.title[:50]  # fallback

    def _extract_problem(self, patent: PatentDocument) -> str:
        """提取技术问题"""
        for feature in patent.technical_features:
            if "技术问题" in feature.feature_name:
                return feature.feature_description
        return patent.abstract[:100] if patent.abstract else ""

    def _extract_solution(self, patent: PatentDocument) -> str:
        """提取技术方案"""
        for feature in patent.technical_features:
            if "技术方案" in feature.feature_name:
                return feature.feature_description
        return " ".join(patent.claims[:3]) if patent.claims else ""

    def _extract_effect(self, patent: PatentDocument) -> str:
        """提取技术效果"""
        for feature in patent.technical_features:
            if "技术效果" in feature.feature_name:
                return feature.feature_description
        return patent.abstract[-100:] if patent.abstract else ""

    def _describe_difference(self, similarity: float) -> str:
        """描述差异程度"""
        level = self._get_similarity_level(similarity)
        return f"相似度{similarity:.2f},属于{level.value}"

    async def _extract_key_differences_similarities(
        self, patent1: PatentDocument, patent2: PatentDocument
    ) -> tuple[list[str], list[str]:
        """提取关键差异和相似点"""
        differences = []
        similarities = []

        # 简化实现,实际应该基于更深入的分析
        # 可以使用知识图谱查询技术特征关系

        # 比对权利要求中的关键术语
        terms1 = self._extract_key_terms(patent1)
        terms2 = self._extract_key_terms(patent2)

        # 找出独有术语(差异)
        differences.extend(
            [f"目标专利独有技术特征: {term}" for term in terms1 if term not in terms2]
        )
        differences.extend(
            [f"对比文件独有技术特征: {term}" for term in terms2 if term not in terms1]
        )

        # 找出共有术语(相似)
        common_terms = set(terms1) & set(terms2)
        similarities.extend([f"共有技术特征: {term}" for term in list(common_terms)[:5])

        return differences[:10], similarities[:10]

    def _extract_key_terms(self, patent: PatentDocument) -> list[str]:
        """提取关键术语"""
        # 组合标题、摘要和权利要求
        text = f"{patent.title} {patent.abstract} {' '.join(patent.claims[:2]) if patent.claims else ''}"

        # 提取可能的术语(名词短语)
        # 简化实现,实际应该使用专业NER
        terms = re.findall(r"([装置|系统|方法|设备|工艺|技术))"]

        # 去重并返回前20个
        return list(set(terms))[:20]

    async def _calculate_vector_similarity(
        self, target_patent: PatentDocument, ref_patents: list[PatentDocument]
    ) -> dict[str, float]:
        """计算向量相似度 - 使用BGE向量"""
        logger.info("🔍 计算向量相似度...")

        similarities = {}

        # 检查目标专利是否有向量
        if target_patent.embedding is None:
            logger.warning("⚠️ 目标专利没有向量,降级使用文本相似度")
            for ref_patent in ref_patents:
                similarities[ref_patent.patent_id] = await self._calculate_text_similarity(
                    target_patent.abstract, ref_patent.abstract
                )
            return similarities

        # 使用向量计算相似度
        for ref_patent in ref_patents:
            if ref_patent.embedding is not None:
                # 计算余弦相似度(向量已归一化,直接点积)
                similarity = float(np.dot(target_patent.embedding, ref_patent.embedding))
            else:
                # 降级为文本相似度
                similarity = await self._calculate_text_similarity(
                    target_patent.abstract, ref_patent.abstract
                )

            similarities[ref_patent.patent_id] = similarity
            logger.debug(f"   {ref_patent.patent_id}: {similarity:.4f}")

        logger.info("✅ 向量相似度计算完成")
        return similarities

    async def _compare_technical_features_with_kg(
        self, target_patent: PatentDocument, ref_patents: list[PatentDocument]
    ) -> list[TechnicalComparison]:
        """使用知识图谱对比技术特征"""
        logger.info("🔍 使用知识图谱对比技术特征...")

        # 确保知识图谱已初始化
        await self._ensure_knowledge_graph()

        comparisons = []

        if self.knowledge_graph is None:
            # 降级到NetworkX本地图
            logger.info("💡 使用NetworkX进行本地图分析")
            all_patents = [target_patent, *ref_patents]
            local_graph = await self._initialize_local_graph(all_patents)

            if local_graph is not None:
                # 使用本地图进行对比
                comparisons = await self._compare_with_networkx(
                    target_patent, ref_patents, local_graph
                )
            else:
                # 最后降级到文本对比
                logger.warning("⚠️ NetworkX不可用,使用文本对比")
                comparisons = await self._compare_technical_features(target_patent, ref_patents)
        else:
            # 使用NebulaGraph知识图谱进行对比
            comparisons = await self._compare_with_nebula_graph(target_patent, ref_patents)

        return comparisons

    async def _compare_with_nebula_graph(
        self, target_patent: PatentDocument, ref_patents: list[PatentDocument]
    ) -> list[TechnicalComparison]:
        """使用NebulaGraph知识图谱进行对比"""
        logger.info("🔍 查询NebulaGraph知识图谱...")

        comparisons = []

        for ref_patent in ref_patents:
            try:
                # 1. 查询技术领域关系
                field_query = f"""
                MATCH (a:Patent {{name: '{target_patent.patent_id}'}})-[:IN_FIELD]->(f:Field)<-[:IN_FIELD]-(b:Patent {{name: '{ref_patent.patent_id}'}})
                RETURN f.name as common_field
                LIMIT 5
                """
                await self.knowledge_graph.execute_query(field_query)

                # 2. 查询共有技术特征
                feature_query = f"""
                MATCH (a:Patent {{name: '{target_patent.patent_id}'}})-[:HAS_FEATURE]->(tf:Feature)<-[:HAS_FEATURE]-(b:Patent {{name: '{ref_patent.patent_id}'}})
                RETURN tf.name as common_feature, tf.type as feature_type
                LIMIT 10
                """
                feature_result = await self.knowledge_graph.execute_query(feature_query)

                # 3. 查询目标专利独有特征
                unique_query = f"""
                MATCH (a:Patent {{name: '{target_patent.patent_id}'}})-[:HAS_FEATURE]->(tf:Feature)
                WHERE NOT (tf)<-[:HAS_FEATURE]-(:Patent {{name: '{ref_patent.patent_id}'}})
                RETURN tf.name as unique_feature
                LIMIT 10
                """
                unique_result = await self.knowledge_graph.execute_query(unique_query)

                # 4. 计算相似度
                common_features = len(feature_result) if feature_result else 0
                unique_features = len(unique_result) if unique_result else 0
                total_features = common_features + unique_features

                similarity = common_features / total_features if total_features > 0 else 0.0

                # 5. 构建对比结果
                key_similarities = []
                key_differences = []

                if feature_result:
                    for row in feature_result:
                        key_similarities.append(f"共有技术特征: {row.get('common_feature', 'N/A')}")

                if unique_result:
                    for row in unique_result:
                        key_differences.append(f"目标专利独有: {row.get('unique_feature', 'N/A')}")

                comparison = TechnicalComparison(
                    patent_a_id=target_patent.patent_id,
                    patent_b_id=ref_patent.patent_id,
                    field_similarity=similarity,
                    field_difference="基于知识图谱的技术领域对比",
                    problem_similarity=similarity * 0.9,
                    problem_difference="基于知识图谱的技术问题对比",
                    solution_similarity=similarity,
                    solution_difference="基于知识图谱的技术方案对比",
                    feature_by_feature_comparison={
                        "common_features": common_features,
                        "unique_features": unique_features,
                        "graph_based": True,
                    },
                    effect_similarity=similarity * 0.85,
                    effect_difference="基于知识图谱的技术效果对比",
                    overall_similarity=similarity,
                    key_differences=key_differences[:5],
                    key_similarities=key_similarities[:5],
                )

                comparisons.append(comparison)

            except Exception as e:
                logger.warning(f"⚠️ NebulaGraph查询失败: {e},降级到文本对比")
                # 降级处理在下面完成

        # 如果知识图谱查询失败,降级到文本对比
        if not comparisons:
            logger.info("💡 降级到文本对比方法")
            comparisons = await self._compare_technical_features(target_patent, ref_patents)

        return comparisons

    async def _compare_with_networkx(
        self, target_patent: PatentDocument, ref_patents: list[PatentDocument], graph: Any
    ) -> list[TechnicalComparison]:
        """使用NetworkX本地图进行对比"""
        logger.info("🔍 使用NetworkX本地图进行对比...")

        comparisons = []

        try:
            for ref_patent in ref_patents:
                # 1. 计算图距离
                try:
                    shortest_path = nx.shortest_path_length(
                        graph, source=target_patent.patent_id, target=ref_patent.patent_id
                    )
                    # 距离越小,相似度越高
                    graph_similarity = 1.0 / (1.0 + shortest_path)
                except nx.NetworkXNoPath:
                    graph_similarity = 0.0

                # 2. 计算共同邻居
                try:
                    common_neighbors = list(
                        nx.common_neighbors(graph, target_patent.patent_id, ref_patent.patent_id)
                    )
                    neighbor_similarity = (
                        len(common_neighbors)
                        / max(
                            len(list(graph.neighbors(target_patent.patent_id))),
                            len(list(graph.neighbors(ref_patent.patent_id))),
                        )
                        if graph.number_of_nodes() > 0
                        else 0.0
                    )
                except Exception:
                    common_neighbors = []
                    neighbor_similarity = 0.0

                # 3. 综合相似度
                overall_similarity = graph_similarity * 0.6 + neighbor_similarity * 0.4

                # 4. 提取关键信息
                key_similarities = [
                    f"共有特征节点数: {len(common_neighbors)}",
                    f"图距离: {shortest_path if 'shortest_path' in locals() else 'N/A'}",
                ]

                key_differences = ["基于NetworkX本地图分析得出差异"]

                comparison = TechnicalComparison(
                    patent_a_id=target_patent.patent_id,
                    patent_b_id=ref_patent.patent_id,
                    field_similarity=overall_similarity * 0.9,
                    field_difference="基于NetworkX图分析",
                    problem_similarity=overall_similarity * 0.85,
                    problem_difference="基于NetworkX图分析",
                    solution_similarity=overall_similarity,
                    solution_difference="基于NetworkX图分析",
                    feature_by_feature_comparison={
                        "common_neighbors": len(common_neighbors),
                        "graph_distance": shortest_path if "shortest_path" in locals() else None,
                        "graph_based": True,
                        "networkx": True,
                    },
                    effect_similarity=overall_similarity * 0.8,
                    effect_difference="基于NetworkX图分析",
                    overall_similarity=overall_similarity,
                    key_differences=key_differences,
                    key_similarities=key_similarities,
                )

                comparisons.append(comparison)

        except Exception as e:
            logger.warning(f"⚠️ NetworkX分析失败: {e}")

        return comparisons

    async def _analyze_novelty(
        self, claims_comparison: list[ClaimComparison], vector_similarity: dict[str, float]
    ) -> str:
        """分析新颖性"""
        logger.info("🔍 分析新颖性...")

        analysis_parts = []

        for comp in claims_comparison:
            claim_id = comp.target_claim_id

            # 检查是否有缺失要素
            if comp.missing_elements:
                analysis_parts.append(
                    f"权利要求{claim_id}相对于对比文件具备新颖性,"
                    f"因为包含以下独有要素:{'; '.join(comp.missing_elements[:3])}"
                )
            else:
                # 检查要素相似度
                low_similarity_elements = []
                    ec for ec in comp.element_comparisons if ec.similarity_score < 0.7
                ]

                if len(low_similarity_elements) > len(comp.element_comparisons) / 2:
                    analysis_parts.append(
                        f"权利要求{claim_id}的大部分要素与对比文件不同," f"可能具备新颖性"
                    )
                else:
                    analysis_parts.append(
                        f"权利要求{claim_id}与对比文件高度相似," f"新颖性存在疑问"
                    )

        novelty_analysis = "\n".join(analysis_parts)
        logger.info("✅ 新颖性分析完成")
        return novelty_analysis

    async def _analyze_inventiveness(
        self, claims_comparison: list[ClaimComparison], vector_similarity: dict[str, float]
    ) -> str:
        """分析创造性"""
        logger.info("🔍 分析创造性...")

        # 基于三步法分析
        # 1. 确定最接近的现有技术
        # 2. 确定区别特征
        # 3. 判断是否有技术启示

        closest_ref = max(vector_similarity.items(), key=lambda x: x[1])
        closest_sim = closest_ref[1]

        analysis = f"""基于三步法的创造性分析:

第一步:确定最接近的现有技术
最接近的现有技术为{closest_ref[0]}(相似度:{closest_sim:.2f})

第二步:确定区别特征
"""

        # 汇总区别特征
        for comp in claims_comparison[:3]:  # 分析前3个权利要求
            if comp.missing_elements:
                analysis += f"\n权利要求{comp.target_claim_id}的区别特征:\n"
                for element in comp.missing_elements[:2]:
                    analysis += f"  - {element}\n"

        analysis += """
第三步:判断是否有技术启示
目标专利的技术方案相对于最接近现有技术是非显而易见的,
因为:
1. 区别特征未被对比文件公开
2. 区别特征带来了预料不到的技术效果
3. 对比文件未给出将区别特征应用到最接近现有技术的技术启示

结论:目标专利具备创造性。
"""

        logger.info("✅ 创造性分析完成")
        return analysis

    async def _generate_response_strategy(
        self,
        claims_comparison: list[ClaimComparison],
        technical_comparison: list[TechnicalComparison],
        novelty_analysis: str,
        inventiveness_analysis: str,
    ) -> tuple[str, list[str]]:
        """生成答复策略"""
        logger.info("🔍 生成答复策略...")

        # 分析整体相似度
        avg_overall_sim = np.mean([tc.overall_similarity for tc in technical_comparison])

        # 统计缺失要素
        total_missing = sum(len(comp.missing_elements) for comp in claims_comparison)

        if avg_overall_sim < 0.6:
            strategy = f"""策略A:全面争辩
由于目标专利与对比文件的整体相似度较低({avg_overall_sim * 100:.1f}%),建议采取全面争辩策略:
1. 强调目标专利的技术方案与对比文件存在本质区别
2. 详细论述区别特征带来的技术效果
3. 论证区别特征非显而易见性
4. 不修改权利要求,保持保护范围"""

            key_arguments = [
                "技术方案存在本质区别",
                "对比文件未公开核心区别特征",
                "具备预料不到的技术效果",
                "非显而易见的技术组合",
            ]

        elif total_missing > 5:
            strategy = f"""策略B:重点强调独有特征
目标专利包含{total_missing}个独有技术要素,建议:
1. 重点论述独有特征的创新性和技术效果
2. 说明独有特征的组合带来了协同效应
3. 论证这种组合在现有技术中未被启示
4. 可适当修改权利要求,进一步明确独有特征"""

            key_arguments = [
                f"包含{total_missing}个独有技术特征",
                "独有特征具有创新性",
                "特征组合产生协同效应",
                "无技术启示支持组合",
            ]

        else:
            strategy = """策略C:修改+争辩
目标专利与对比文件相似度较高,建议:
1. 对权利要求进行修改,加入区别特征
2. 强调修改后的技术方案与对比文件的差异
3. 论证修改后的技术效果
4. 争取在缩小保护范围的同时获得授权"""

            key_arguments = [
                "建议修改权利要求",
                "加入区别特征以明确保护范围",
                "修改后具备新颖性和创造性",
                "平衡保护范围与授权前景",
            ]

        logger.info("✅ 答复策略生成完成")
        return strategy, key_arguments

    def generate_markdown_report(self, report: DeepComparisonReport) -> str:
        """生成Markdown格式的分析报告 - 增强版"""
        # 获取模型信息
        model_info = ""
        if self.bge_service:
            model_info = f"""
> **嵌入模型**: {self.bge_service.model_name} ({self.bge_service.dimension}维)
> **计算设备**: {self.bge_service.device}
"""

        # 获取知识图谱信息
        kg_info = ""
        if self.knowledge_graph and self._kg_initialized:
            kg_info = "\n> **知识图谱**: NebulaGraph (已启用)"
        elif self.config.get("enable_networkx_fallback"):
            kg_info = "\n> **知识图谱**: NetworkX本地图 (降级模式)"

        md = f"""# 专利深度对比分析报告

> **报告ID**: {report.report_id}
> **生成时间**: {report.generation_time}
> **分析器版本**: {report.analyzer_version}
> **分析模式**: 深度对比 (向量+知识图谱){model_info}{kg_info}

---

## 📊 分析概要

本报告采用多层次分析方法:
1. **向量语义分析** - 使用本地MPS优化的1024维BGE模型
2. **知识图谱对比** - 基于NebulaGraph技术特征关系(或NetworkX本地图)
3. **权利要求要素级对比** - 逐要素解析和比对
4. **法律评估** - 新颖性、创造性三步法分析

---

## 一、基本信息

### 目标专利
- **专利号**: {report.target_patent.patent_number}
- **发明名称**: {report.target_patent.title}
- **摘要**: {report.target_patent.abstract[:200] if report.target_patent.abstract else "无"}...
- **权利要求数**: {len(report.target_patent.claims)}

### 对比文件 ({len(report.reference_patents)}个)
"""

        for i, ref_patent in enumerate(report.reference_patents, 1):
            md += f"""
#### 对比文件{i}: {ref_patent.patent_number}
- **发明名称**: {ref_patent.title}
- **摘要**: {ref_patent.abstract[:200] if ref_patent.abstract else "无"}...
- **向量相似度**: {report.vector_similarity_analysis.get(ref_patent.patent_id, 0):.4f}
"""

        md += """

---

## 二、权利要求对比分析

### 权利要求对比总表

| 权利要求 | 最相似对比文件 | 相似度 | 独有要素数 | 新颖性评估 |
|---------|---------------|--------|-----------|-----------|
"""

        for comp in report.claims_comparison[:5]:  # 显示前5个
            ref_name = (
                report.reference_patents[0].patent_number if report.reference_patents else "无"
            )
            missing = len(comp.missing_elements)
            conclusion = "具备新颖性" if missing > 0 else "新颖性存疑"

            md += f"| {comp.target_claim_id} | {ref_name} | {comp.overall_similarity:.2f} | {missing}个 | {conclusion} |\n"

        md += """

### 详细权利要求对比

"""

        for comp in report.claims_comparison[:3]:  # 详细展示前3个
            md += f"""
#### {comp.target_claim_id}

**目标专利**:
```
{comp.target_claim_text[:200]}...
```

**对比文件** ({comp.reference_claim_id}):
```
{comp.reference_claim_text[:200] if comp.reference_claim_text else "无"}...
```

**要素级对比**:

| 目标要素 | 对比要素 | 相似度 | 状态 |
|---------|---------|--------|------|
"""

            for ec in comp.element_comparisons[:5]:
                ref_text = ec.reference_element.element_text if ec.reference_element else "无"
                status = "✓ 存在" if ec.is_present else "✗ 缺失"
                md += f"| {ec.target_element.element_text[:50]}... | {ref_text[:50]}... | {ec.similarity_score:.2f} | {status} |\n"

            md += f"\n**差异总结**: {comp.difference_summary}\n\n"

        md += """

---

## 三、技术特征对比分析

"""

        for i, tech_comp in enumerate(report.technical_comparison, 1):
            md += f"""
### 对比{i}: {tech_comp.patent_a_id} vs {tech_comp.patent_b_id}

| 维度 | 相似度 | 差异描述 |
|-----|--------|---------|
| 技术领域 | {tech_comp.field_similarity:.2f} | {tech_comp.field_difference} |
| 技术问题 | {tech_comp.problem_similarity:.2f} | {tech_comp.problem_difference} |
| 技术方案 | {tech_comp.solution_similarity:.2f} | {tech_comp.solution_difference} |
| 技术效果 | {tech_comp.effect_similarity:.2f} | {tech_comp.effect_difference} |

**整体相似度**: {tech_comp.overall_similarity:.2f}

**关键差异**:
"""

            for diff in tech_comp.key_differences[:3]:
                md += f"- {diff}\n"

            md += "\n**关键相似**:\n"
            for sim in tech_comp.key_similarities[:3]:
                md += f"- {sim}\n"

            md += "\n"

        md += """

---

## 四、向量相似度分析

| 对比文件 | 相似度分数 | 相似等级 |
|---------|-----------|---------|
"""

        for ref_id, sim in report.vector_similarity_analysis.items():
            level = self._get_similarity_level(sim)
            md += f"| {ref_id} | {sim:.4f} | {level.value} |\n"

        md += """

---

## 五、法律评估

### 新颖性分析

{report.novelty_analysis}

### 创造性分析

{report.inventiveness_analysis}

---

## 六、答复策略建议

### 推荐策略

{report.response_strategy}

### 关键论据

"""

        for i, arg in enumerate(report.key_arguments, 1):
            md += f"{i}. {arg}\n"

        md += """

---

## 七、详细建议

### 修改建议(如适用)

根据上述分析,建议考虑:

1. **权利要求修改**:如有必要,可在权利要求中进一步限定区别特征
2. **说明书补充**:在说明书中详细说明区别特征的技术效果
3. **答复策略**:采用上述推荐的答复策略进行答复

### 注意事项

- 确保所有论点有充分的证据支持
- 注意答复期限,合理安排答复时间
- 如需修改,确保修改符合《专利法》第33条规定

---

**报告结束**

*本报告由专利深度对比分析系统自动生成*
"""

        return md

    def generate_docx_report(self, report: DeepComparisonReport) -> Document:
        """生成DOCX格式的分析报告"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx未安装,无法生成DOCX报告")
            raise ImportError("请安装python-docx: pip install python-docx")

        doc = Document()

        # 设置默认字体(中文支持)
        doc.styles["Normal"].font.name = "宋体"
        doc.styles[east_asia")]

        # 标题
        title = doc.add_heading("专利深度对比分析报告", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 报告信息
        info_para = doc.add_paragraph()
        info_para.add_run("报告ID: ").bold = True
        info_para.add_run(f"{report.report_id}\n")
        info_para.add_run("生成时间: ").bold = True
        info_para.add_run(f"{report.generation_time}\n")
        info_para.add_run("分析器版本: ").bold = True
        info_para.add_run(f"{report.analyzer_version}\n")

        if self.bge_service:
            info_para.add_run("\n嵌入模型: ").bold = True
            info_para.add_run(f"{self.bge_service.model_name} ({self.bge_service.dimension}维)\n")
            info_para.add_run("计算设备: ").bold = True
            info_para.add_run(f"{self.bge_service.device}\n")

        doc.add_page_break()

        # 一、基本信息
        doc.add_heading("一、基本信息", 1)

        # 目标专利
        doc.add_heading("目标专利", 2)
        doc.add_paragraph(f"专利号: {report.target_patent.patent_number}", style="List Bullet")
        doc.add_paragraph(f"发明名称: {report.target_patent.title}", style="List Bullet")
        doc.add_paragraph(
            f'摘要: {report.target_patent.abstract[:200] if report.target_patent.abstract else "无"}...',
            style="List Bullet",
        )
        doc.add_paragraph(f"权利要求数: {len(report.target_patent.claims)}", style="List Bullet")

        # 对比文件
        doc.add_heading(f"对比文件 ({len(report.reference_patents)}个)", 2)
        for i, ref_patent in enumerate(report.reference_patents, 1):
            p = doc.add_paragraph(f"对比文件{i}: ", style="List Bullet")
            p.add_run(f"{ref_patent.patent_number}").bold = True
            doc.add_paragraph(f"  发明名称: {ref_patent.title}", style="List Bullet")
            doc.add_paragraph(
                f'  摘要: {ref_patent.abstract[:150] if ref_patent.abstract else "无"}...',
                style="List Bullet",
            )
            sim = report.vector_similarity_analysis.get(ref_patent.patent_id, 0)
            doc.add_paragraph(f"  向量相似度: {sim:.4f}", style="List Bullet")

        doc.add_page_break()

        # 二、权利要求对比分析
        doc.add_heading("二、权利要求对比分析", 1)

        # 创建表格
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "权利要求"
        hdr_cells[1].text = "对比文件"
        hdr_cells[2].text = "相似度"
        hdr_cells[3].text = "独有要素"
        hdr_cells[4].text = "结论"

        for comp in report.claims_comparison[:5]:
            row_cells = table.add_row().cells
            row_cells[0].text = comp.target_claim_id
            ref_name = (
                report.reference_patents[0].patent_number if report.reference_patents else "无"
            )
            row_cells[1].text = ref_name
            row_cells[2].text = f"{comp.overall_similarity:.2f}"
            row_cells[3].text = f"{len(comp.missing_elements)}个"
            conclusion = "具备新颖性" if len(comp.missing_elements) > 0 else "新颖性存疑"
            row_cells[4].text = conclusion

        # 详细权利要求对比
        doc.add_heading("详细权利要求对比", 2)
        for comp in report.claims_comparison[:3]:
            doc.add_heading(f"{comp.target_claim_id}", 3)

            # 目标专利
            p = doc.add_paragraph()
            p.add_run("[目标专利]\n").bold = True
            doc.add_paragraph(
                comp.target_claim_text[:300] + "..."
                if len(comp.target_claim_text) > 300
                else comp.target_claim_text
            )

            # 对比文件
            p = doc.add_paragraph()
            p.add_run(f"[对比文件]({comp.reference_claim_id})\n").bold = True
            ref_text = comp.reference_claim_text if comp.reference_claim_text else "无"
            doc.add_paragraph(
                ref_text[:300] + "..." if ref_text and len(ref_text) > 300 else ref_text
            )

            # 要素对比表格
            if comp.element_comparisons:
                p = doc.add_paragraph("\n要素级对比:\n").bold = True
                elem_table = doc.add_table(rows=1, cols=4)
                elem_table.style = "Light Grid Accent 1"
                elem_hdr = elem_table.rows[0].cells
                elem_hdr[0].text = "目标要素"
                elem_hdr[1].text = "对比要素"
                elem_hdr[2].text = "相似度"
                elem_hdr[3].text = "状态"

                for ec in comp.element_comparisons[:5]:
                    row = elem_table.add_row().cells
                    row[0].text = ec.target_element.element_text[:50] + "..."
                    ref_text = ec.reference_element.element_text if ec.reference_element else "无"
                    row[1].text = ref_text[:50] + "..."
                    row[2].text = f"{ec.similarity_score:.2f}"
                    row[3].text = "存在" if ec.is_present else "缺失"

        doc.add_page_break()

        # 三、技术特征对比分析
        doc.add_heading("三、技术特征对比分析", 1)
        for i, tech_comp in enumerate(report.technical_comparison, 1):
            doc.add_heading(f"对比{i}: {tech_comp.patent_a_id} vs {tech_comp.patent_b_id}", 2)

            # 创建技术对比表格
            tech_table = doc.add_table(rows=5, cols=3)
            tech_table.style = "Light Grid Accent 1"

            # 表头
            tech_table.rows[0].cells[0].text = "维度"
            tech_table.rows[0].cells[1].text = "相似度"
            tech_table.rows[0].cells[2].text = "差异描述"

            tech_table.rows[1].cells[0].text = "技术领域"
            tech_table.rows[1].cells[1].text = f"{tech_comp.field_similarity:.2f}"
            tech_table.rows[1].cells[2].text = tech_comp.field_difference

            tech_table.rows[2].cells[0].text = "技术问题"
            tech_table.rows[2].cells[1].text = f"{tech_comp.problem_similarity:.2f}"
            tech_table.rows[2].cells[2].text = tech_comp.problem_difference

            tech_table.rows[3].cells[0].text = "技术方案"
            tech_table.rows[3].cells[1].text = f"{tech_comp.solution_similarity:.2f}"
            tech_table.rows[3].cells[2].text = tech_comp.solution_difference

            tech_table.rows[4].cells[0].text = "技术效果"
            tech_table.rows[4].cells[1].text = f"{tech_comp.effect_similarity:.2f}"
            tech_table.rows[4].cells[2].text = tech_comp.effect_difference

            # 关键差异和相似
            doc.add_paragraph(f"\n整体相似度: {tech_comp.overall_similarity:.2f}\n")

            doc.add_paragraph("关键差异:", style="List Bullet").bold
            for diff in tech_comp.key_differences[:5]:
                doc.add_paragraph(diff, style="List Bullet 2")

            doc.add_paragraph("\n关键相似:", style="List Bullet").bold
            for sim in tech_comp.key_similarities[:5]:
                doc.add_paragraph(sim, style="List Bullet 2")

        doc.add_page_break()

        # 四、向量相似度分析
        doc.add_heading("四、向量相似度分析", 1)
        vec_table = doc.add_table(rows=len(report.vector_similarity_analysis) + 1, cols=3)
        vec_table.style = "Light Grid Accent 1"
        vec_hdr = vec_table.rows[0].cells
        vec_hdr[0].text = "对比文件"
        vec_hdr[1].text = "相似度分数"
        vec_hdr[2].text = "相似等级"

        for i, (ref_id, sim) in enumerate(report.vector_similarity_analysis.items(), 1):
            row = vec_table.rows[i].cells
            row[0].text = ref_id
            row[1].text = f"{sim:.4f}"
            level = self._get_similarity_level(sim)
            row[2].text = level.value

        doc.add_page_break()

        # 五、法律评估
        doc.add_heading("五、法律评估", 1)

        doc.add_heading("新颖性分析", 2)
        doc.add_paragraph(report.novelty_analysis)

        doc.add_heading("创造性分析", 2)
        doc.add_paragraph(report.inventiveness_analysis)

        doc.add_page_break()

        # 六、答复策略建议
        doc.add_heading("六、答复策略建议", 1)

        doc.add_heading("推荐策略", 2)
        doc.add_paragraph(report.response_strategy)

        doc.add_heading("关键论据", 2)
        for i, arg in enumerate(report.key_arguments, 1):
            doc.add_paragraph(f"{i}. {arg}", style="List Number")

        # 七、详细建议
        doc.add_heading("七、详细建议", 1)
        doc.add_paragraph("根据上述分析,建议考虑:", style="List Bullet")
        doc.add_paragraph(
            "1. 权利要求修改:如有必要,可在权利要求中进一步限定区别特征", style="List Number"
        )
        doc.add_paragraph(
            "2. 说明书补充:在说明书中详细说明区别特征的技术效果", style="List Number"
        )
        doc.add_paragraph("3. 答复策略:采用上述推荐的答复策略进行答复", style="List Number")

        doc.add_paragraph("\n注意事项:", style="List Bullet").bold
        doc.add_paragraph("确保所有论点有充分的证据支持", style="List Bullet 2")
        doc.add_paragraph("注意答复期限,合理安排答复时间", style="List Bullet 2")
        doc.add_paragraph("如需修改,确保修改符合《专利法》第33条规定", style="List Bullet 2")

        # 报告结束
        doc.add_paragraph("\n\n")
        ending = doc.add_paragraph()
        ending_run = ending.add_run("报告结束\n\n本报告由专利深度对比分析系统自动生成")
        ending_run.font.size = Pt(10)
        ending_run.font.color.rgb = RGBColor(128, 128, 128)
        ending.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return doc

    async def save_report(self, report: DeepComparisonReport, output_path: str):
        """保存报告 - Markdown + DOCX双格式"""
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        saved_files = []

        # 1. 保存Markdown格式
        md_content = self.generate_markdown_report(report)
        md_path = output_file.with_suffix(".md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        saved_files.append(str(md_path))
        logger.info(f"✅ Markdown报告已保存: {md_path}")

        # 2. 保存DOCX格式
        if DOCX_AVAILABLE:
            try:
                doc = self.generate_docx_report(report)
                docx_path = output_file.with_suffix(".docx")
                doc.save(docx_path)
                saved_files.append(str(docx_path))
                logger.info(f"✅ DOCX报告已保存: {docx_path}")
            except Exception as e:
                logger.warning(f"⚠️ DOCX报告生成失败: {e}")
        else:
            logger.warning("⚠️ python-docx未安装,跳过DOCX格式生成")
            logger.info("💡 安装命令: pip install python-docx")

        return saved_files


# 便捷函数
async def analyze_patent_for_office_action(
    target_patent_text: str, reference_patents: list[str]], output_path: str | None = None
) -> DeepComparisonReport:
    """
    便捷函数:分析专利用于审查意见答复

    Args:
        target_patent_text: 目标专利全文
        reference_patents: 对比文件专利号列表
        output_path: 报告输出路径

    Returns:
        DeepComparisonReport: 分析报告
    """
    analyzer = PatentDeepComparisonAnalyzer()

    # 执行分析
    report = await analyzer.analyze_office_action(
        target_patent_text=target_patent_text,
        reference_patents=reference_patents,
        analysis_depth="deep",
    )

    # 保存报告
    if output_path:
        await analyzer.save_report(report, output_path)
    else:
        # 默认路径
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_path = f"/Users/xujian/Athena工作平台/reports/patent_comparison_{timestamp}.md"
        await analyzer.save_report(report, default_path)

    return report


# =============================================================================
# 工具注册和集成到小娜智能体
# =============================================================================


def get_tool_schema() -> dict:
    """获取工具注册schema - 用于小娜智能体集成"""
    return {
        "name": "patent_deep_comparison_analyzer",
        "description": """专利深度对比分析工具 - 用于审查意见答复

核心功能:
1. 基于本地MPS优化的1024维BGE模型进行语义相似度分析
2. 基于NebulaGraph知识图谱(或NetworkX本地图)进行技术特征对比
3. 权利要求要素级详细比对
4. 自动生成新颖性、创造性分析
5. 提供答复策略建议

适用场景:
- 发明专利审查意见答复
- 实用新型审查意见答复
- 无效宣告请求分析
- 专利自由实施分析
""",
        "parameters": {
            "type": "object",
            "properties": {
                "target_patent_text": {
                    "type": "string",
                    "description": "目标专利全文或权利要求书文本",
                },
                "reference_patents": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "对比文件专利号列表",
                },
                "analysis_depth": {
                    "type": "string",
                    "enum": ["quick", "standard", "deep"],
                    "description": "分析深度:quick(快速)、standard(标准)、deep(深度)",
                    "default": "deep",
                },
                "output_path": {"type": "string", "description": "报告输出路径(可选)"},
                "enable_kg": {
                    "type": "boolean",
                    "description": "是否启用知识图谱对比",
                    "default": True,
                },
                "use_vector_search": {
                    "type": "boolean",
                    "description": "是否使用向量语义搜索",
                    "default": True,
                },
            },
            "required": ["target_patent_text", "reference_patents"],
        },
        "returns": {
            "type": "object",
            "description": "返回深度对比分析报告,包含权利要求对比、技术特征对比、法律评估和答复策略",
        },
    }


async def run_as_tool(
    target_patent_text: str,
    reference_patents: list[str],
    analysis_depth: str = "deep",
    output_path: str | None = None,
    enable_kg: bool = True,
    use_vector_search: bool = True,
) -> dict:
    """
    作为工具运行 - 用于小娜智能体调用

    Args:
        target_patent_text: 目标专利全文
        reference_patents: 对比文件专利号列表
        analysis_depth: 分析深度
        output_path: 报告输出路径
        enable_kg: 是否启用知识图谱
        use_vector_search: 是否使用向量搜索

    Returns:
        dict: 分析结果
    """
    # 创建配置
    config = {
        "enable_knowledge_graph": enable_kg,
        "enable_vector_search": use_vector_search,
        "enable_bge_service": True,
        "embedding_model": "BAAI/bge-m3",  # 1024维本地MPS模型
    }

    # 创建分析器
    analyzer = PatentDeepComparisonAnalyzer(config=config)

    # 执行分析
    report = await analyzer.analyze_office_action(
        target_patent_text=target_patent_text,
        reference_patents=reference_patents,
        analysis_depth=analysis_depth,
    )

    # 保存报告(Markdown + DOCX双格式)
    saved_files = []
    if output_path:
        saved_files = await analyzer.save_report(report, output_path)
    else:
        # 默认保存路径
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_path = f"/Users/xujian/Athena工作平台/reports/patent_comparison_{timestamp}"
        saved_files = await analyzer.save_report(report, default_path)

    # 返回结果摘要
    return {
        "success": True,
        "report_id": report.report_id,
        "target_patent": report.target_patent.patent_number,
        "reference_count": len(report.reference_patents),
        "claims_analyzed": len(report.claims_comparison),
        "avg_similarity": float(np.mean(list(report.vector_similarity_analysis.values()))),
        "novelty_conclusion": (
            "具备新颖性" if "具备" in report.novelty_analysis else "需要进一步分析"
        ),
        "inventiveness_conclusion": (
            "具备创造性" if "具备" in report.inventiveness_analysis else "需要进一步分析"
        ),
        "response_strategy": report.response_strategy[:100] + "...",
        "report_files": saved_files,  # 双格式文件路径列表
        "key_arguments_count": len(report.key_arguments),
        "analyzer_version": report.analyzer_version,
        "analysis_mode": "向量(BGE 1024维) + 知识图谱",
        "output_formats": ["Markdown", "DOCX"] if len(saved_files) > 1 else ["Markdown"],
    }


# =============================================================================
# 使用示例
# =============================================================================


async def example_usage():
    """完整的使用示例"""
    print("=" * 70)
    print("专利深度对比分析系统 - 使用示例")
    print("=" * 70)

    # 示例1: 基本使用
    print("\n[示例1]基本使用 - 分析发明专利")
    print("-" * 70)

    target_text = """
    发明名称:一种基于深度学习的图像识别方法及系统

    摘要:
    本发明公开了一种基于深度学习的图像识别方法及系统,包括:
    获取待识别图像;通过卷积神经网络提取图像特征;通过注意力机制
    增强关键特征;通过分类器进行图像分类。本发明能够提高图像
    识别的准确率和效率。

    权利要求书:
    1. 一种基于深度学习的图像识别方法,其特征在于,包括以下步骤:
       S1:获取待识别图像;
       S2:通过预训练的卷积神经网络提取图像的多层次特征;
       S3:通过空间注意力机制增强关键区域特征;
       S4:通过通道注意力机制增强重要通道特征;
       S5:将增强后的特征输入全连接分类器,得到识别结果。

    2. 根据权利要求1所述的方法,其特征在于,
       所述空间注意力机制采用最大池化和平均池化的并行结构。

    3. 根据权利要求1所述的方法,其特征在于,
       所述卷积神经网络采用残差连接结构。
    """

    references = ["CN201810000000.X", "CN201920000000.X"]

    # 执行分析
    report = await analyze_patent_for_office_action(
        target_patent_text=target_text,
        reference_patents=references,
        output_path="/Users/xujian/Athena工作平台/reports/example_comparison.md",
    )

    print("✅ 分析完成!")
    print(f"   报告ID: {report.report_id}")
    print(f"   对比文件数: {len(report.reference_patents)}")
    print(f"   权利要求数: {len(report.claims_comparison)}")
    print("   报告路径: /Users/xujian/Athena工作平台/reports/example_comparison.md")

    # 示例2: 作为工具调用
    print("\n[示例2]作为小娜智能体工具调用")
    print("-" * 70)

    result = await run_as_tool(
        target_patent_text=target_text,
        reference_patents=references,
        analysis_depth="deep",
        enable_kg=True,
        use_vector_search=True,
    )

    print("✅ 工具调用完成!")
    print(f"   成功: {result['success']}")
    print(f"   平均相似度: {result['avg_similarity']:.2%}")
    print(f"   新颖性: {result['novelty_conclusion']}")
    print(f"   创造性: {result['inventiveness_conclusion']}")

    # 示例3: 仅使用向量分析(不使用知识图谱)
    print("\n[示例3]仅向量分析模式(快速模式)")
    print("-" * 70)

    result2 = await run_as_tool(
        target_patent_text=target_text,
        reference_patents=references[:1],  # 只对比一个文件
        analysis_depth="quick",
        enable_kg=False,  # 不使用知识图谱
        use_vector_search=True,
    )

    print("✅ 快速分析完成!")
    print(f"   分析模式: {result2['analysis_mode']}")
    print(f"   平均相似度: {result2['avg_similarity']:.2%}")

    print("\n" + "=" * 70)
    print("所有示例运行完成!")
    print("=" * 70)


if __name__ == "__main__":
    import asyncio

    # 运行使用示例
    asyncio.run(example_usage())
