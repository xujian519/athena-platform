#!/usr/bin/env python3

"""
法律写作素材库管理器 - 增强版
Legal Writing Materials Manager - Enhanced Version

功能:
1. 利用RAG管理器进行语义检索
2. 从语料文件夹获取全文内容(支持.md/.txt/.json/.docx/.pdf)
3. 内置缓存机制,提升性能
4. 支持多种文档格式
"""

import asyncio
import json
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)


@dataclass
class CachedDocument:
    """缓存的文档"""

    content: str
    cached_at: datetime
    access_count: int = 0
    file_path: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)

    def is_expired(self, ttl_hours: int = 24) -> bool:
        """检查缓存是否过期"""
        return datetime.now() - self.cached_at > timedelta(hours=ttl_hours)


class WritingMaterialsManagerEnhanced:
    """
    法律写作素材库管理器 - 增强版

    新增功能:
    1. 内置LRU缓存机制
    2. 支持更多文档格式
    3. 更智能的文件路径推断
    4. 批量预加载热门文档
    """

    def __init__(
        self,
        rag_manager=None,
        materials_path: Optional[Path] = None,
        cache_size: int = 100,
        cache_ttl_hours: int = 24,
    ):
        """
        初始化素材库管理器

        Args:
            rag_manager: RAG管理器(用于检索)
            materials_path: 素材路径(默认为/Users/xujian/语料)
            cache_size: 缓存大小(文档数量)
            cache_ttl_hours: 缓存过期时间(小时)
        """
        self.rag_manager = rag_manager

        if materials_path is None:
            materials_path = Path("/Users/xujian/语料")

        self.materials_path = materials_path
        self.cache_size = cache_size
        self.cache_ttl_hours = cache_ttl_hours

        # 文档缓存 {file_key: CachedDocument}
        self._document_cache: dict[str, CachedDocument] = {}

        # 访问顺序(用于LRU)
        self._access_order: list[str] = []

        # 文件路径映射配置
        self._init_path_mappings()

        logger.info("✅ 写作素材库管理器初始化完成(增强版)")
        logger.info(f"   素材路径: {self.materials_path}")
        logger.info(f"   RAG管理器: {'已配置' if rag_manager else '未配置'}")
        logger.info(f"   缓存配置: 大小={cache_size}, TTL={cache_ttl_hours}小时")

    def _init_path_mappings(self):
        """初始化文件路径映射"""
        self.collections = {
            "patent_judgments": {
                "path": self.materials_path / "专利判决",
                "json_path": self.materials_path / "专利判决-json",
                "id_field": "judgment_id",
                "title_field": "title",
                "file_field": "file_name",
            },
            "case_data_1024": {
                "path": self.materials_path / "判决文书",
                "json_path": None,  # 可选
                "id_field": "id",
                "title_field": "title",
                "file_field": "file_name",
            },
            "patent_rules_1024": {
                "path": self.materials_path / "司法解释",
                "json_path": None,
                "id_field": "id",
                "title_field": "title",
                "file_field": "file_name",
            },
            "patent_legal": {
                "path": self.materials_path / "Laws-1.0.0",
                "json_path": None,
                "id_field": "id",
                "title_field": "title",
                "file_field": "file_name",
            },
            # 无效复审决定可能有不同的结构
            "invalidation_decisions": {
                "path": self.materials_path / "无效复审决定",
                "json_path": None,
                "id_field": "decision_id",
                "title_field": "title",
                "file_field": "file_name",
            },
        }

    # ========== 缓存管理 ==========

    def _get_cache_key(self, file_path: Path) -> str:
        """生成缓存键"""
        return str(file_path)

    def _get_from_cache(self, file_path: Path) -> Optional[str]:
        """从缓存获取文档"""
        key = self._get_cache_key(file_path)

        if key in self._document_cache:
            cached_doc = self._document_cache[key]

            # 检查是否过期
            if cached_doc.is_expired(self.cache_ttl_hours):
                del self._document_cache[key]
                if key in self._access_order:
                    self._access_order.remove(key)
                return None

            # 更新访问记录
            cached_doc.access_count += 1
            if key in self._access_order:
                self._access_order.remove(key)
            self._access_order.append(key)

            return cached_doc.content

        return None

    def _add_to_cache(self, file_path: Path, content: str, metadata: Optional[dict[str, Any]] = None):
        """添加到缓存"""
        key = self._get_cache_key(file_path)

        # 如果缓存已满,删除最久未使用的项
        if len(self._document_cache) >= self.cache_size and key not in self._document_cache:
            if self._access_order:
                oldest_key = self._access_order.pop(0)
                if oldest_key in self._document_cache:
                    del self._document_cache[oldest_key]

        # 添加新缓存项
        self._document_cache[key] = CachedDocument(
            content=content,
            cached_at=datetime.now(),
            file_path=str(file_path),
            metadata=metadata or {},
        )

        if key in self._access_order:
            self._access_order.remove(key)
        self._access_order.append(key)

    def clear_cache(self):
        """清空缓存"""
        self._document_cache.clear()
        self._access_order.clear()
        logger.info("   文档缓存已清空")

    def get_cache_stats(self) -> dict[str, Any]:
        """获取缓存统计"""
        total_access = sum(doc.access_count for doc in self._document_cache.values())

        return {
            "cached_documents": len(self._document_cache),
            "cache_size_limit": self.cache_size,
            "total_access_count": total_access,
            "avg_access_per_doc": (
                total_access / len(self._document_cache) if self._document_cache else 0
            ),
        }

    # ========== 文档检索 ==========

    async def search_relevant_examples(
        self, topic: str, section: str, role: str = "scholar", top_k: int = 3
    ) -> list[dict[str, str]]:
        """
        搜索相关的写作示例

        Args:
            topic: 主题
            section: 章节名称
            role: 写作角色
            top_k: 返回数量

        Returns:
            相关示例列表
        """
        examples = []

        if not self.rag_manager:
            logger.warning("RAG管理器未配置,无法检索素材")
            return examples

        try:
            # 1. 根据角色选择任务类型
            task_type = self._role_to_task_type(role)

            # 2. 使用RAG进行语义检索
            search_query = f"{topic} {section}"
            rag_context = await self.rag_manager.retrieve(query=search_query, task_type=task_type)

            # 3. 获取检索结果
            retrieval_results = rag_context.get_top_results(top_k * 2)

            # 4. 根据检索结果获取全文内容
            for result in retrieval_results[:top_k]:
                full_content = self._get_full_content_enhanced(result)

                if full_content and len(full_content) > 200:
                    examples.append(
                        {
                            "title": result.metadata.get("title", result.source),
                            "source": result.source,
                            "content": (
                                full_content[:1500] + "..."
                                if len(full_content) > 1500
                                else full_content
                            ),
                            "relevance_score": result.score,
                        }
                    )

                if len(examples) >= top_k:
                    break

            logger.info(f"   找到 {len(examples)} 个相关示例")

        except Exception as e:
            logger.warning(f"   检索写作示例失败: {e}")

        return examples

    def _role_to_task_type(self, role: str) -> str:
        """将写作角色映射到RAG任务类型"""
        role_task_map = {
            "patent_attorney": "oa_response",
            "lawyer": "tech_analysis",
            "judge": "patent_search",
            "scholar": "creativity_analysis",
        }
        return role_task_map.get(role, "default")

    # ========== 全文获取(增强版)==========

    def _get_full_content_enhanced(self, retrieval_result: Any) -> Optional[str]:
        """
        获取全文内容(增强版)

        支持多种检索方式:
        1. 直接从元数据获取文件路径
        2. 从JSON元数据文件获取路径
        3. 根据命名规则推断路径
        4. 直接使用metadata中的content
        """
        metadata = getattr(retrieval_result, "metadata", {})
        source = getattr(retrieval_result, "source", "")

        # 方法1: 检查缓存
        cached_content = self._get_from_cache(Path(source))
        if cached_content:
            return cached_content

        content = None

        # 方法2: 从元数据获取文件路径
        file_path = metadata.get("file_path") or metadata.get("source_file") or metadata.get("path")
        if file_path:
            content = self._read_and_cache_file(Path(file_path), metadata)
            if content:
                return content

        # 方法3: 从关联的JSON文件获取路径
        json_path = metadata.get("json_path") or metadata.get("metadata_file")
        if json_path:
            content = self._read_from_json_metadata(Path(json_path))
            if content:
                return content

        # 方法4: 根据source和命名规则推断
        content = self._infer_and_load_file(source, metadata)
        if content:
            return content

        # 方法5: 直接使用metadata中的内容
        content = metadata.get("content") or metadata.get("text") or metadata.get("full_text")
        if content:
            return content

        # 方法6: 尝试从检索结果的content字段获取
        if hasattr(retrieval_result, "content"):
            return retrieval_result.content

        return None

    def _read_from_json_metadata(self, json_path: Path) -> Optional[str]:
        """从JSON元数据文件读取内容"""
        try:
            if not json_path.exists():
                # 尝试在materials_path下查找
                json_path = self.materials_path / json_path.name

            if json_path.exists() and json_path.suffix == ".json":
                with open(json_path, encoding="utf-8") as f:
                    data = json.load(f)

                # 提取文件路径和内容
                file_path_str = data.get("metadata", {}).get("file_path")
                if file_path_str:
                    content = self._read_and_cache_file(Path(file_path_str), data)
                    if content:
                        return content

                # 如果没有文件路径,尝试从JSON中提取内容字段
                for field in ["content", "text", "full_text", "reasoning", "facts"]:
                    if data.get(field):
                        if isinstance(data[field], list):
                            return "\n".join(data[field])
                        return str(data[field])

        except Exception as e:
            logger.debug(f"从JSON读取失败 {json_path}: {e}")

        return None

    def _infer_and_load_file(self, source: str, metadata: dict[str, Any]) -> Optional[str]:
        """根据命名规则推断并加载文件"""
        # 提取可能的文件名
        filename = None

        # 从metadata获取
        filename = metadata.get("file_name") or metadata.get("filename")

        # 从source提取
        if not filename:
            # 尝试从source中提取案号/决定号
            # 匹配案号格式:(年份)法院代字第XX号
            match = re.search(r"(?\d{4})?[^)第]+第\d+号", source)
            if match:
                filename = match.group(0) + ".md"

        if not filename:
            # 尝试使用ID作为文件名
            doc_id = (
                metadata.get("id") or metadata.get("judgment_id") or metadata.get("decision_id")
            )
            if doc_id:
                filename = f"{doc_id}.md"

        if filename:
            # 在各个可能的目录下搜索
            search_paths = []
            for _collection_name, config in self.collections.items():
                if config["path"].exists():
                    search_paths.append(config["path"])

            for search_path in search_paths:
                file_path = search_path / filename
                if file_path.exists():
                    content = self._read_and_cache_file(file_path, metadata)
                    if content:
                        return content

                # 尝试其他扩展名
                for ext in [".txt", ".json", ".docx"]:
                    alt_path = search_path / (filename.rsplit(".", 1)[0] + ext)
                    if alt_path.exists():
                        content = self._read_and_cache_file(alt_path, metadata)
                        if content:
                            return content

        return None

    def _read_and_cache_file(
        self, file_path: Path, metadata: Optional[dict[str, Any]] = None
    ) -> Optional[str]:
        """读取文件并缓存"""
        # 先检查缓存
        cached = self._get_from_cache(file_path)
        if cached:
            return cached

        # 读取文件
        content = self._read_file_content(file_path)

        # 添加到缓存
        if content:
            self._add_to_cache(file_path, content, metadata)

        return content

    def _read_file_content(self, file_path: Path) -> Optional[str]:
        """读取文件内容(支持多种格式)"""
        try:
            suffix = file_path.suffix.lower()

            if suffix in [".md", ".txt"]:
                with open(file_path, encoding="utf-8") as f:
                    return f.read()

            elif suffix == ".json":
                with open(file_path, encoding="utf-8") as f:
                    data = json.load(f)

                # 尝试提取内容字段
                if isinstance(data, dict):
                    # 优先查找content相关字段
                    for field in ["content", "text", "full_text", "reasoning"]:
                        if field in data:
                            value = data[field]
                            if isinstance(value, list):
                                return "\n".join(str(v) for v in value)
                            return str(value)

                    # 如果没有content字段,返回整个JSON的格式化字符串
                    return json.dumps(data, ensure_ascii=False, indent=2)

                return str(data)

            elif suffix == ".docx":
                # 尝试使用python-docx读取
                try:
                    from docx import Document

                    doc = Document(file_path)
                    return "\n".join([para.text for para in doc.paragraphs])
                except ImportError:
                    logger.debug(f"python-docx未安装,无法读取{file_path}")
                    return f"[Word文档: {file_path.name}]"

            elif suffix == ".pdf":
                # 尝试使用PyPDF2读取
                try:
                    import PyPDF2

                    with open(file_path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text()
                        return text
                except ImportError:
                    logger.debug(f"PyPDF2未安装,无法读取{file_path}")
                    return f"[PDF文档: {file_path.name}]"
                except Exception as e:
                    logger.debug(f"PDF读取失败 {file_path}: {e}")
                    return f"[PDF文档: {file_path.name}]"

            else:
                return f"[文件: {file_path.name}, 类型: {suffix}]"

        except Exception as e:
            logger.debug(f"读取文件失败 {file_path}: {e}")
            return None

    # ========== 批量操作 ==========

    async def preload_popular_documents(
        self, document_ids: list[str], collection: str = "patent_judgments"
    ):
        """
        批量预加载热门文档到缓存

        Args:
            document_ids: 文档ID列表
            collection: 集合名称
        """
        if collection not in self.collections:
            logger.warning(f"未知的集合: {collection}")
            return

        config = self.collections[collection]
        loaded = 0

        for doc_id in document_ids:
            # 构建可能的文件路径
            possible_files = [
                config["path"] / f"{doc_id}.md",
                config["path"] / f"{doc_id}.txt",
                config["path"] / f"{doc_id}.json",
            ]

            # 如果有json_path,也检查那里
            if config["json_path"]:
                possible_files.append(config["json_path"] / f"{doc_id}.json")

            for file_path in possible_files:
                if file_path.exists():
                    content = self._read_and_cache_file(file_path, {"id": doc_id})
                    if content:
                        loaded += 1
                        break

        logger.info(f"   预加载完成: {loaded}/{len(document_ids)} 个文档")

    # ========== 兼容性接口 ==========

    def build_writing_prompt_with_examples(
        self, topic: str, section: str, role: str, base_style_prompt: str
    ) -> str:
        """构建包含示例的写作提示词"""
        examples = asyncio.run(
            self.search_relevant_examples(topic=topic, section=section, role=role, top_k=2)
        )

        examples_section = ""
        if examples:
            examples_section = "\n### 参考示例\n\n"
            examples_section += "以下是从平台真实法律文档中提取的相关示例:\n\n"
            for i, example in enumerate(examples, 1):
                examples_section += f"#### 示例 {i}: {example['title']}\n\n"
                examples_section += f"**来源**: {example['source']}\n\n"
                examples_section += f"```\n{example['content']}\n```\n\n"
        else:
            examples_section = "\n### 参考示例\n\n暂无相关示例。\n"

        return f"""{base_style_prompt}

## 参考素材

{examples_section}

请参考上述示例的风格和结构,但不要抄袭具体内容。
"""

    def search_materials(
        self, query: str, category: Optional[str] = None, top_k: int = 5
    ) -> list[dict[str, Any]]:
        """搜索素材(兼容旧版本接口)"""
        if not self.rag_manager:
            return []

        try:
            task_type = category or "default"
            rag_context = asyncio.run(self.rag_manager.retrieve(query=query, task_type=task_type))

            results = []
            for result in rag_context.get_top_results(top_k):
                results.append(
                    {
                        "id": result.source,
                        "title": result.metadata.get("title", result.source),
                        "source": result.source,
                        "content": (
                            result.content[:500] + "..."
                            if len(result.content) > 500
                            else result.content
                        ),
                        "relevance_score": result.score,
                        "metadata": result.metadata,
                    }
                )

            return results

        except Exception as e:
            logger.warning(f"检索失败: {e}")
            return []


# 全局单例
_materials_manager_enhanced_instance = None


def get_materials_manager_enhanced(
    rag_manager=None,
    materials_path: Optional[Path] = None,
    cache_size: int = 100,
    cache_ttl_hours: int = 24,
) -> WritingMaterialsManagerEnhanced:
    """获取素材库管理器单例(增强版)"""
    global _materials_manager_enhanced_instance

    if _materials_manager_enhanced_instance is None:
        _materials_manager_enhanced_instance = WritingMaterialsManagerEnhanced(
            rag_manager=rag_manager,
            materials_path=materials_path,
            cache_size=cache_size,
            cache_ttl_hours=cache_ttl_hours,
        )

    return _materials_manager_enhanced_instance


if __name__ == "__main__":
    # 测试代码
    import sys

    sys.path.insert(0, str(Path(__file__).parent.parent.parent))



    async def test():
        print("法律写作素材库管理器 - 增强版测试")
        print("=" * 70)

        # 创建素材管理器(不配置RAG,测试基本功能)
        manager = get_materials_manager_enhanced(rag_manager=None)

        # 测试文件读取
        print("\n📄 测试文件读取:")
        test_file = Path("/Users/xujian/语料/专利判决/(2000)高知初字第37号.md")
        if test_file.exists():
            content = manager._read_and_cache_file(test_file, {"test": True})
            print(f"   文件: {test_file.name}")
            print(f"   长度: {len(content)} 字符")
            print(f"   预览: {content[:100]}...")

        # 测试缓存
        print("\n💾 测试缓存:")
        stats = manager.get_cache_stats()
        print(f"   缓存文档数: {stats['cached_documents']}")
        print(f"   总访问次数: {stats['total_access_count']}")

        # 再次读取应该命中缓存
        manager._read_and_cache_file(test_file)
        stats2 = manager.get_cache_stats()
        print(f"   二次读取后访问次数: {stats2['total_access_count']}")

        print("\n✅ 测试完成!")

    asyncio.run(test())

