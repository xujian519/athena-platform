#!/usr/bin/env python3
"""
专利复审决定深度处理器
Patent Re-examination Decision Deep Processor

支持DOC和DOCX格式的深度解析、向量化、知识图谱构建

作者: Athena平台团队
创建时间: 2026-01-11
版本: v1.0.0
"""

from __future__ import annotations
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

# 文档处理
from docx import Document

from core.logging_config import setup_logging

# DOC处理 (需要antiword或python-docx2txt)
try:
    from docx2python import docx2python

    DOCX2PYTHON_AVAILABLE = True
except ImportError:
    DOCX2PYTHON_AVAILABLE = False

# 向量化

# PostgreSQL

# Neo4j

# Qdrant

# 配置日志
# setup_logging()  # 日志配置已移至模块导入
logger = setup_logging()


@dataclass
class PatentDecisionData:
    """专利复审决定数据结构"""

    # 基础信息
    file_path: str
    file_format: str  # 'doc' or 'docx'
    file_size: int

    # 决定书信息
    decision_number: str = ""
    decision_type: str = ""  # 无效宣告请求审查决定
    decision_date: str = ""

    # 当事人信息
    petitioner: str = ""  # 请求人
    patent_owner: str = ""  # 专利权人
    patent_number: str = ""  # 专利号
    patent_title: str = ""  # 专利名称

    # 争议焦点
    dispute_foci: list[str] = field(default_factory=list)

    # 决定结果
    decision_result: str = ""  # 维持/宣告无效/部分无效
    decision_reason: str = ""  # 决定理由

    # 法律依据
    legal_basis: list[str] = field(default_factory=list)

    # 完整内容
    full_content: str = ""

    # 分段内容
    sections: dict[str, str] = field(default_factory=dict)

    # 元数据
    metadata: dict[str, Any] = field(default_factory=dict)


class DOCXProcessor:
    """DOCX文档处理器"""

    def __init__(self):
        """初始化处理器"""
        logger.info("✅ DOCX处理器初始化完成")

    def can_process(self, file_path: str) -> bool:
        """判断是否可以处理该文件"""
        return file_path.lower().endswith(".docx")

    def extract_content(self, file_path: str) -> PatentDecisionData:
        """
        从DOCX文件提取内容

        Args:
            file_path: 文件路径

        Returns:
            PatentDecisionData对象
        """
        logger.info(f"📄 处理DOCX文件: {file_path}")

        try:
            # 使用python-docx读取
            doc = Document(file_path)

            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            full_content = "\n".join(paragraphs)

            # 提取表格内容(如果有)
            self._extract_tables(doc)

            # 创建数据对象
            data = PatentDecisionData(
                file_path=file_path,
                file_format="docx",
                file_size=os.path.getsize(file_path),
                full_content=full_content,
            )

            # 智能解析决定书内容
            self._parse_decision_content(data, paragraphs)

            logger.info(f"✅ DOCX文件处理完成: {data.decision_number or file_path}")
            return data

        except Exception as e:
            logger.error(f"❌ DOCX文件处理失败: {e}")
            # 如果python-docx失败,尝试docx2python
            if DOCX2PYTHON_AVAILABLE:
                return self._extract_with_docx2python(file_path)
            raise

    def _extract_tables(self, doc: Document) -> list[list[str]]:
        """提取表格内容"""
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables.append(table_data)
        return tables

    def _extract_with_docx2python(self, file_path: str) -> PatentDecisionData:
        """使用docx2python提取内容(备用方法)"""
        logger.info(f"🔄 使用docx2python处理: {file_path}")

        content = docx2python(file_path)

        # 提取正文内容
        body_content = content.body
        full_content = ""
        for page in body_content:
            for line in page:
                full_content += "\n".join(line) + "\n"

        data = PatentDecisionData(
            file_path=file_path,
            file_format="docx",
            file_size=os.path.getsize(file_path),
            full_content=full_content,
        )

        return data

    def _parse_decision_content(self, data: PatentDecisionData, paragraphs: list[str]) -> Any:
        """
        智能解析决定书内容

        提取关键信息:决定号、当事人、争议焦点、决定结果等
        """
        # 1. 提取决定号
        re.compile(r"无效宣告请求审查决定\s*[((]?(\d{1,4})?[))]?")

        # 2. 提取当事人信息
        petitioner_pattern = re.compile(r"请求人[::]\s*([^\n]+)")
        owner_pattern = re.compile(r"专利权人[::]\s*([^\n]+)")

        # 3. 提取专利号
        patent_number_pattern = re.compile(
            r"ZL\s*(\d{12,}[\.X]?\d*)|" r"\d{12,}[\.X]?\d*|" r"(CN)?(\d{9,})"
        )

        # 4. 提取决定结果
        result_patterns = [
            re.compile(r"维持.*有效", re.IGNORECASE),
            re.compile(r"宣告.*(?:专利权)?(?:全部|部分)?无效", re.IGNORECASE),
            re.compile(r"驳回.*请求", re.IGNORECASE),
        ]

        # 遍历段落提取信息
        for _i, para in enumerate(paragraphs):
            # 提取决定号
            if "决定号" in para or "无效宣告请求审查决定" in para:
                match = re.search(r"决定号[：:]*\s*([^\s]+)", para)
                if match:
                    data.decision_number = match.group(1)
                else:
                    # 尝试从文件名提取
                    filename = Path(data.file_path).stem
                    data.decision_number = filename

            # 提取请求人
            match = petitioner_pattern.search(para)
            if match:
                data.petitioner = match.group(1).strip()

            # 提取专利权人
            match = owner_pattern.search(para)
            if match:
                data.patent_owner = match.group(1).strip()

            # 提取专利号
            match = patent_number_pattern.search(para)
            if match:
                # 找最长匹配
                for group in match.groups():
                    if group:
                        data.patent_number = group
                        break

            # 提取决定结果
            for pattern in result_patterns:
                if pattern.search(para):
                    if "维持" in para:
                        data.decision_result = "维持有效"
                    elif "无效" in para:
                        data.decision_result = "宣告无效"
                    elif "驳回" in para:
                        data.decision_result = "驳回请求"
                    break

        # 5. 提取争议焦点(通常在"理由"部分)
        dispute_section = self._extract_dispute_foci(paragraphs)
        data.dispute_foci = dispute_section

        # 6. 提取决定理由
        reason_section = self._extract_decision_reason(paragraphs)
        data.decision_reason = reason_section

        # 7. 提取法律依据
        legal_basis = self._extract_legal_basis(paragraphs)
        data.legal_basis = legal_basis

        logger.info(f"   决定号: {data.decision_number}")
        logger.info(f"   请求人: {data.petitioner}")
        logger.info(f"   专利权人: {data.patent_owner}")
        logger.info(f"   专利号: {data.patent_number}")
        logger.info(f"   决定结果: {data.decision_result}")
        logger.info(f"   争议焦点数: {len(data.dispute_foci)}")

    def _extract_dispute_foci(self, paragraphs: list[str]) -> list[str]:
        """提取争议焦点"""
        foci = []

        # 寻找"争议焦点"或"焦点"段落
        in_focus_section = False
        focus_patterns = [
            re.compile(r"争议焦点[::]"),
            re.compile(r"焦点[::]"),
            re.compile(r"本案.*?焦点"),
        ]

        for para in paragraphs:
            # 检查是否进入焦点段落
            for pattern in focus_patterns:
                if pattern.search(para):
                    in_focus_section = True
                    focus_text = para
                    if len(focus_text) > 10:  # 避免只匹配到标题
                        foci.append(focus_text.strip())
                    break

            # 如果在焦点段落中,继续提取
            else:
                if in_focus_section:
                    if len(para) > 20 and not para.startswith(("一", "二", "三", "1.", "2.", "3.")):
                        foci.append(para.strip())
                    else:
                        in_focus_section = False

        return foci[:5]  # 最多返回5个争议焦点

    def _extract_decision_reason(self, paragraphs: list[str]) -> str:
        """提取决定理由"""
        reason_parts = []

        # 寻找"理由"或"决定理由"段落
        in_reason_section = False
        reason_patterns = [
            re.compile(r"理由[::]"),
            re.compile(r"决定理由[::]"),
            re.compile(r"审查结论[::]"),
        ]

        for para in paragraphs:
            # 检查是否进入理由段落
            for pattern in reason_patterns:
                if pattern.search(para):
                    in_reason_section = True
                    break

            # 如果在理由段落中,提取内容
            else:
                if in_reason_section:
                    reason_parts.append(para)

        return "\n".join(reason_parts[:10])  # 最多提取前10段

    def _extract_legal_basis(self, paragraphs: list[str]) -> list[str]:
        """提取法律依据"""
        basis = []

        # 匹配专利法条款引用
        patent_law_pattern = re.compile(
            r"专利法(?:实施细则)?第?([一二三四五六七八九十百千]+|\d+)条", re.IGNORECASE
        )

        for para in paragraphs:
            matches = patent_law_pattern.findall(para)
            if matches:
                basis.extend(matches)

        return list(set(basis))  # 去重


class DOCProcessor:
    """DOC文档处理器(旧格式)"""

    def __init__(self):
        """初始化处理器"""
        self.antiword_available = self._check_antiword()
        self.textract_available = self._check_textract()
        logger.info("✅ DOC处理器初始化完成")

    def _check_antiword(self) -> bool:
        """检查antiword是否可用"""
        import shutil

        return shutil.which("antiword") is not None

    def _check_textract(self) -> bool:
        """检查textract是否可用"""
        try:
            import textract

            return True
        except ImportError:
            return False

    def can_process(self, file_path: str) -> bool:
        """判断是否可以处理该文件"""
        return file_path.lower().endswith(".doc")

    def extract_content(self, file_path: str) -> PatentDecisionData:
        """
        从DOC文件提取内容

        Args:
            file_path: 文件路径

        Returns:
            PatentDecisionData对象
        """
        logger.info(f"📄 处理DOC文件: {file_path}")

        # 方法1: 尝试使用antiword
        if self.antiword_available:
            return self._extract_with_antiword(file_path)

        # 方法2: 尝试使用textract
        elif self.textract_available:
            return self._extract_with_textract(file_path)

        # 方法3: 尝试使用python-docx(可能对旧DOC有效)
        else:
            try:
                return self._extract_with_docx_compat(file_path)
            except Exception as e:
                logger.debug(f"空except块已触发: {e}")
                logger.error(f"❌ 无法处理DOC文件: {file_path}")
                logger.error("   建议安装: brew install antiword")
                raise

    def _extract_with_antiword(self, file_path: str) -> PatentDecisionData:
        """使用antiword提取DOC内容"""
        import subprocess

        try:
            # 调用antiword
            result = subprocess.run(
                ["antiword", "-m", file_path],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="ignore",
            )

            if result.returncode == 0:
                content = result.stdout

                data = PatentDecisionData(
                    file_path=file_path,
                    file_format="doc",
                    file_size=os.path.getsize(file_path),
                    full_content=content,
                )

                # 解析内容(与DOCX相同的解析逻辑)
                paragraphs = content.split("\n")
                self._parse_decision_content(data, paragraphs)

                logger.info(f"✅ DOC文件处理完成(antiword): {file_path}")
                return data
            else:
                raise Exception(f"antiword失败: {result.stderr}")

        except Exception as e:
            logger.error(f"❌ antiword提取失败: {e}")
            raise

    def _extract_with_textract(self, file_path: str) -> PatentDecisionData:
        """使用textract提取内容"""
        import textract

        try:
            content = textract.process(file_path).decode("utf-8")

            data = PatentDecisionData(
                file_path=file_path,
                file_format="doc",
                file_size=os.path.getsize(file_path),
                full_content=content,
            )

            paragraphs = content.split("\n")
            self._parse_decision_content(data, paragraphs)

            logger.info(f"✅ DOC文件处理完成(textract): {file_path}")
            return data

        except Exception as e:
            logger.error(f"❌ textract提取失败: {e}")
            raise

    def _extract_with_docx_compat(self, file_path: str) -> PatentDecisionData:
        """尝试用python-docx处理DOC(可能失败)"""
        try:
            doc = Document(file_path)

            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            full_content = "\n".join(paragraphs)

            data = PatentDecisionData(
                file_path=file_path,
                file_format="doc",
                file_size=os.path.getsize(file_path),
                full_content=full_content,
            )

            self._parse_decision_content(data, paragraphs)

            logger.info(f"✅ DOC文件处理完成(docx兼容): {file_path}")
            return data

        except Exception as e:
            logger.error(f"❌ python-docx处理DOC失败: {e}")
            raise

    def _parse_decision_content(self, data: PatentDecisionData, paragraphs: list[str]) -> Any:
        """解析决定内容(复用DOCX的解析逻辑)"""
        docx_processor = DOCXProcessor()
        docx_processor._parse_decision_content(data, paragraphs)


class PatentDecisionProcessor:
    """专利复审决定统一处理器"""

    def __init__(self):
        """初始化处理器"""
        self.docx_processor = DOCXProcessor()
        self.doc_processor = DOCProcessor()

        # 初始化向量化模型(延迟加载)
        self.model = None

        logger.info("✅ 专利复审决定处理器初始化完成")

    def _get_model(self) -> Any:
        """获取BGE-M3模型"""
        if self.model is None:
            from core.embedding.bge_embedding_service import BGEEmbeddingService

            embed_service = BGEEmbeddingService(model_name="bge-m3", device="cpu")
            self.model = embed_service.model
            logger.info("✅ BGE-M3模型加载完成")

        return self.model

    def process_file(self, file_path: str) -> PatentDecisionData:
        """
        处理单个文件(自动识别格式)

        Args:
            file_path: 文件路径

        Returns:
            PatentDecisionData对象
        """
        # 判断文件格式并选择处理器
        if file_path.lower().endswith(".docx"):
            return self.docx_processor.extract_content(file_path)
        elif file_path.lower().endswith(".doc"):
            return self.doc_processor.extract_content(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_path}")

    def process_directory(
        self, directory: str, limit: int | None = None, pattern: str | None = None
    ) -> list[PatentDecisionData]:
        """
        批量处理目录中的文件

        Args:
            directory: 目录路径
            limit: 限制处理数量(用于测试)
            pattern: 文件名匹配模式

        Returns:
            PatentDecisionData列表
        """
        logger.info(f"📁 开始处理目录: {directory}")

        # 收集文件
        all_files = []
        for root, _dirs, files in os.walk(directory):
            for file in files:
                file_path = os.path.join(root, file)

                # 过滤文件格式
                if file_path.lower().endswith((".docx", ".doc")):
                    # 应用文件名模式过滤
                    if pattern is None or re.search(pattern, file):
                        all_files.append(file_path)

        logger.info(f"   找到文件: {len(all_files)}个")

        # 限制处理数量
        if limit:
            all_files = all_files[:limit]
            logger.info(f"   限制处理: {limit}个")

        # 批量处理
        results = []
        success_count = 0
        fail_count = 0

        for i, file_path in enumerate(all_files, 1):
            try:
                logger.info(f"\n[{i}/{len(all_files)}] 处理: {Path(file_path).name}")
                data = self.process_file(file_path)
                results.append(data)
                success_count += 1

                # 每100个文件输出一次统计
                if i % 100 == 0:
                    logger.info("\n📊 进度统计:")
                    logger.info(f"   已处理: {i}/{len(all_files)}")
                    logger.info(f"   成功: {success_count}")
                    logger.info(f"   失败: {fail_count}")

            except Exception as e:
                logger.error(f"❌ 处理失败: {file_path}")
                logger.error(f"   错误: {e}")
                fail_count += 1

        logger.info("\n✅ 批量处理完成!")
        logger.info(f"   总数: {len(all_files)}")
        logger.info(f"   成功: {success_count}")
        logger.info(f"   失败: {fail_count}")

        return results

    def vectorize_decision(
        self, data: PatentDecisionData, granularity: str = "multi"
    ) -> dict[str, list[tuple[str, list[float]]]]:
        """
        将决定书向量化

        Args:
            data: 决定书数据
            granularity: 粒度 ('single', 'multi')

        Returns:
            字典,包含不同粒度的向量
        """
        model = self._get_model()

        vectors = {
            "document": [],  # 文档级
            "focus": [],  # 争议焦点级
            "paragraph": [],  # 段落级
            "sentence": [],  # 语句级
        }

        if granularity == "single":
            # 仅文档级向量
            vector = model.encode([data.full_content])[0]
            vectors["document"].append(("doc_full", vector.tolist()))

        else:
            # 多粒度向量

            # 1. 文档级向量(摘要)
            summary = self._create_summary(data)
            vectors["document"].append(("summary", model.encode([summary])[0].tolist()))

            # 2. 争议焦点级向量
            for i, focus in enumerate(data.dispute_foci, 1):
                if focus:
                    vector = model.encode([focus])[0]
                    vectors["focus"].append((f"focus_{i}", vector.tolist()))

            # 3. 段落级向量
            paragraphs = data.full_content.split("\n\n")
            for i, para in enumerate(paragraphs):
                if len(para.strip()) > 50:  # 过滤短段落
                    vector = model.encode([para])[0]
                    vectors["paragraph"].append((f"para_{i}", vector.tolist()))

            # 4. 关键语句级向量(可选)
            key_sentences = self._extract_key_sentences(data)
            for i, sent in enumerate(key_sentences):
                vector = model.encode([sent])[0]
                vectors["sentence"].append((f"sent_{i}", vector.tolist()))

        logger.info("✅ 向量化完成:")
        logger.info(f"   文档级: {len(vectors['document'])}个")
        logger.info(f"   焦点级: {len(vectors['focus'])}个")
        logger.info(f"   段落级: {len(vectors['paragraph'])}个")
        logger.info(f"   语句级: {len(vectors['sentence'])}个")

        return vectors

    def _create_summary(self, data: PatentDecisionData) -> str:
        """创建决定书摘要"""
        parts = []

        if data.decision_number:
            parts.append(f"决定号:{data.decision_number}")

        if data.patent_number:
            parts.append(f"专利号:{data.patent_number}")

        if data.petitioner:
            parts.append(f"请求人:{data.petitioner}")

        if data.patent_owner:
            parts.append(f"专利权人:{data.patent_owner}")

        if data.dispute_foci:
            parts.append(f"争议焦点:{'; '.join(data.dispute_foci[:3])}")

        if data.decision_result:
            parts.append(f"决定结果:{data.decision_result}")

        if data.decision_reason:
            parts.append(f"决定理由摘要:{data.decision_reason[:200]}")

        return "\n".join(parts)

    def _extract_key_sentences(self, data: PatentDecisionData) -> list[str]:
        """提取关键语句"""
        key_sentences = []

        # 提取包含法条引用的句子
        paragraphs = data.full_content.split("\n")
        legal_pattern = re.compile(r"专利法(?:实施细则)?第[一二三四五六七八九十百千\d]+条")

        for para in paragraphs:
            if legal_pattern.search(para):
                key_sentences.append(para)

        return key_sentences[:20]  # 最多返回20个关键句


# =============================================================================
# 使用示例
# =============================================================================

if __name__ == "__main__":
    import sys

    processor = PatentDecisionProcessor()

    # 测试单个文件
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        print(f"\n📄 处理文件: {file_path}\n")

        data = processor.process_file(file_path)

        print("\n✅ 提取结果:")
        print(f"   决定号: {data.decision_number}")
        print(f"   请求人: {data.petitioner}")
        print(f"   专利权人: {data.patent_owner}")
        print(f"   专利号: {data.patent_number}")
        print(f"   决定结果: {data.decision_result}")
        print(f"   争议焦点: {len(data.dispute_foci)}个")

        # 向量化测试
        print("\n🔄 开始向量化...")
        vectors = processor.vectorize_decision(data, granularity="multi")

        print("\n✅ 向量化结果:")
        for level, vec_list in vectors.items():
            print(f"   {level}: {len(vec_list)}个向量")

    else:
        # 批量处理测试
        print("\n📁 批量处理模式")
        print("使用方法: python patent_decision_processor.py <file_path>\n")

        # 测试处理5个文件
        print("\n🧪 测试处理5个文件...")
        results = processor.process_directory(
            "/Volumes/AthenaData/07_Corpus_Data/语料/专利/专利无效复审决定原文", limit=5
        )

        print(f"\n✅ 处理完成: {len(results)}个文件")
        for data in results:
            print(f"   - {data.decision_number or Path(data.file_path).name}")
