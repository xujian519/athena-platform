#!/usr/bin/env python3
"""
文档智能处理器
Intelligent Document Processor

针对法律文档和专利文档的结构化处理
作者: 小诺·双鱼座
创建时间: 2025-12-21
版本: v1.0.0 "智能切分"
"""

from __future__ import annotations
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any

# 支持的文档类型
try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """文档类型"""

    PATENT_LAW = "patent_law"  # 专利法
    EXAMINATION_GUIDE = "exam_guide"  # 审查指南
    LEGAL_CLAUSE = "legal_clause"  # 法律条款
    PATENT_APPLICATION = "patent_app"  # 专利申请
    LEGAL_CASE = "legal_case"  # 法律案例
    TECHNICAL_DOC = "technical_doc"  # 技术文档
    UNKNOWN = "unknown"  # 未知类型


@dataclass
class DocumentChunk:
    """文档块"""

    chunk_id: str
    content: str
    chunk_type: str  # article, clause, claim, paragraph, etc.
    metadata: dict[str, Any] = field(default_factory=dict)
    source_info: dict[str, Any] = field(default_factory=dict)
    hierarchy_level: int = 0
    parent_id: str | None = None
    child_ids: list[str] = field(default_factory=list)


@dataclass
class ProcessedDocument:
    """处理后的文档"""

    doc_id: str
    doc_type: DocumentType
    title: str
    chunks: list[DocumentChunk]
    metadata: dict[str, Any] = field(default_factory=dict)
    processing_stats: dict[str, Any] = field(default_factory=dict)


class DocumentProcessor:
    """智能文档处理器"""

    def __init__(self):
        # 法律文档切分规则
        self.legal_patterns = {
            "article": re.compile(
                r"第[一二三四五六七八九十百千万\d]+条[::\s]*(.*?)(?=\n第[一二三四五六七八九十百千万\d]+条|第[一二三四五六七八九十百千万\d]+章|\n第[一二三四五六七八九十百千万\d]+节|$)",
                re.DOTALL,
            ),
            "chapter": re.compile(
                r"第[一二三四五六七八九十百千万\d]+章[::\s]*(.*?)(?=\n第[一二三四五六七八九十百千万\d]+章|\n第[一二三四五六七八九十百千万\d]+节|$)",
                re.DOTALL,
            ),
            "section": re.compile(
                r"第[一二三四五六七八九十百千万\d]+节[::\s]*(.*?)(?=\n第[一二三四五六七八九十百千万\d]+节|\n第[一二三四五六七八九十百千万\d]+章|$)",
                re.DOTALL,
            ),
            "paragraph": re.compile(
                r"(\d+\..*?)(?=\n\d+\.|\n[一二三四五六七八九十百千万\d]+、|\n[A-Z]\.|$)", re.DOTALL
            ),
        }

        # 专利文档切分规则
        self.patent_patterns = {
            "claim": re.compile(r"(\d+\..*?权利要求.*?)(?=\n\d+\.|$)", re.DOTALL | re.IGNORECASE),
            "technical_field": re.compile(
                r"技术领域[::\s]*(.*?)(?=\n[背景技术|发明内容|附图说明|具体实施方式])", re.DOTALL
            ),
            "background_art": re.compile(
                r"背景技术[::\s]*(.*?)(?=\n[技术领域|发明内容|附图说明|具体实施方式])", re.DOTALL
            ),
            "summary": re.compile(
                r"发明内容[::\s]*(.*?)(?=\n[技术领域|背景技术|附图说明|具体实施方式])", re.DOTALL
            ),
            "brief_description": re.compile(
                r"附图说明[::\s]*(.*?)(?=\n[技术领域|背景技术|发明内容|具体实施方式])", re.DOTALL
            ),
            "detailed_description": re.compile(r"具体实施方式[::\s]*(.*?)(?=$)", re.DOTALL),
        }

        # 通用段落切分
        self.paragraph_splitter = re.compile(r"\n\s*\n|\r\n\s*\r\n")

    async def process_document(self, file_path: str | Path) -> ProcessedDocument | None:
        """
        处理文档

        Args:
            file_path: 文件路径

        Returns:
            Optional[ProcessedDocument]: 处理后的文档
        """
        file_path = Path(file_path)

        try:
            # 1. 读取文档内容
            content = await self._read_document(file_path)
            if not content:
                return None

            # 2. 识别文档类型
            doc_type = self._detect_document_type(file_path, content)

            # 3. 提取文档标题
            title = self._extract_title(file_path, content, doc_type)

            # 4. 结构化切分
            chunks = await self._split_content(content, doc_type)

            # 5. 生成处理统计
            stats = self._generate_processing_stats(content, chunks)

            # 6. 构建结果
            processed_doc = ProcessedDocument(
                doc_id=str(uuid.uuid4()),
                doc_type=doc_type,
                title=title,
                chunks=chunks,
                metadata={
                    "source_file": str(file_path),
                    "file_size": file_path.stat().st_size,
                    "processed_at": datetime.now().isoformat(),
                },
                processing_stats=stats,
            )

            logger.info(f"✅ 文档处理完成: {file_path.name} -> {len(chunks)}个块")
            return processed_doc

        except Exception as e:
            logger.error(f"❌ 文档处理失败 {file_path}: {e}")
            return None

    async def _read_document(self, file_path: Path) -> str | None:
        """读取文档内容"""
        try:
            if file_path.suffix.lower() == ".txt":
                return file_path.read_text(encoding="utf-8")

            elif file_path.suffix.lower() == ".pdf" and PDF_AVAILABLE:
                return self._read_pdf(file_path)

            elif file_path.suffix.lower() in [".docx", ".doc"] and DOCX_AVAILABLE:
                return self._read_docx(file_path)

            else:
                logger.error(f"❌ 不支持的文件格式: {file_path.suffix}")
                return None

        except Exception as e:
            logger.error(f"❌ 文件读取失败 {file_path}: {e}")
            return None

    def _read_pdf(self, file_path: Path) -> str | None:
        """读取PDF文件"""
        if not PDF_AVAILABLE:
            logger.error("❌ PyPDF2未安装,无法处理PDF文件")
            return None

        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = []

                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    content.append(f"--- 第{page_num + 1}页 ---\n{text}")

                return "\n".join(content)

        except Exception as e:
            logger.error(f"❌ PDF读取失败 {file_path}: {e}")
            return None

    def _read_docx(self, file_path: Path) -> str | None:
        """读取DOCX文件"""
        if not DOCX_AVAILABLE:
            logger.error("❌ python-docx未安装,无法处理DOCX文件")
            return None

        try:
            doc = docx.Document(file_path)
            content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)

            return "\n".join(content)

        except Exception as e:
            logger.error(f"❌ DOCX读取失败 {file_path}: {e}")
            return None

    def _detect_document_type(self, file_path: Path, content: str) -> DocumentType:
        """检测文档类型"""
        file_name = file_path.name.lower()
        content_lower = content.lower()

        # 基于文件名检测
        if "专利法" in file_name or "patent law" in file_name:
            return DocumentType.PATENT_LAW
        elif "审查指南" in file_name or "examination guide" in file_name:
            return DocumentType.EXAMINATION_GUIDE
        elif "权利要求" in file_name or "claim" in file_name:
            return DocumentType.PATENT_APPLICATION
        elif "案例" in file_name or "判决" in file_name or "case" in file_name:
            return DocumentType.LEGAL_CASE

        # 基于内容检测
        if "专利法" in content_lower or "中华人民共和国专利法" in content_lower:
            return DocumentType.PATENT_LAW
        elif "审查指南" in content_lower or "发明专利审查指南" in content_lower:
            return DocumentType.EXAMINATION_GUIDE
        elif "权利要求书" in content_lower or "技术领域" in content_lower:
            return DocumentType.PATENT_APPLICATION
        elif "法院" in content_lower or "判决书" in content_lower:
            return DocumentType.LEGAL_CASE

        return DocumentType.UNKNOWN

    def _extract_title(self, file_path: Path, content: str, doc_type: DocumentType) -> str:
        """提取文档标题"""
        # 尝试从内容第一行提取标题
        lines = content.split("\n")
        first_line = lines[0].strip() if lines else ""

        if first_line and len(first_line) < 200:  # 合理的标题长度
            return first_line

        # 基于文档类型生成标题
        type_titles = {
            DocumentType.PATENT_LAW: "中华人民共和国专利法",
            DocumentType.EXAMINATION_GUIDE: "专利审查指南",
            DocumentType.PATENT_APPLICATION: "专利申请文件",
            DocumentType.LEGAL_CASE: "法律案例文件",
            DocumentType.TECHNICAL_DOC: "技术文档",
        }

        title = type_titles.get(doc_type, "未知文档")
        return f"{title} - {file_path.name}"

    async def _split_content(self, content: str, doc_type: DocumentType) -> list[DocumentChunk]:
        """结构化切分内容"""
        chunks = []

        if doc_type in [
            DocumentType.PATENT_LAW,
            DocumentType.EXAMINATION_GUIDE,
            DocumentType.LEGAL_CLAUSE,
        ]:
            # 法律文档结构化切分
            chunks.extend(self._split_legal_document(content))

        elif doc_type == DocumentType.PATENT_APPLICATION:
            # 专利文档结构化切分
            chunks.extend(self._split_patent_document(content))

        else:
            # 通用段落切分
            chunks.extend(self._split_generic_document(content))

        # 确保每个块都有唯一的ID
        for i, chunk in enumerate(chunks):
            chunk.chunk_id = f"chunk_{i:04d}"

        return chunks

    def _split_legal_document(self, content: str) -> list[DocumentChunk]:
        """切分法律文档"""
        chunks = []

        # 尝试按条切分
        articles = self.legal_patterns["article"].findall(content)
        if articles:
            for i, article_content in enumerate(articles):
                chunk = DocumentChunk(
                    chunk_id=f"article_{i+1}",
                    content=article_content.strip(),
                    chunk_type="article",
                    metadata={"article_number": i + 1, "document_type": "legal"},
                )
                chunks.append(chunk)

        # 如果没有找到条,尝试按段落切分
        if not chunks:
            paragraphs = self.paragraph_splitter.split(content)
            for i, paragraph in enumerate(paragraphs):
                if paragraph.strip():
                    chunk = DocumentChunk(
                        chunk_id=f"para_{i+1}",
                        content=paragraph.strip(),
                        chunk_type="paragraph",
                        metadata={"paragraph_number": i + 1, "document_type": "legal"},
                    )
                    chunks.append(chunk)

        return chunks

    def _split_patent_document(self, content: str) -> list[DocumentChunk]:
        """切分专利文档"""
        chunks = []

        # 按专利文档结构切分
        structure_patterns = {
            "technical_field": self.patent_patterns["technical_field"],
            "background_art": self.patent_patterns["background_art"],
            "summary": self.patent_patterns["summary"],
            "claims": self.patent_patterns["claim"],
            "detailed_description": self.patent_patterns["detailed_description"],
        }

        for section_name, pattern in structure_patterns.items():
            matches = pattern.findall(content)
            for i, match_content in enumerate(matches):
                chunk = DocumentChunk(
                    chunk_id=f"{section_name}_{i+1}",
                    content=match_content.strip(),
                    chunk_type=section_name,
                    metadata={
                        "section": section_name,
                        "section_index": i + 1,
                        "document_type": "patent",
                    },
                )
                chunks.append(chunk)

        # 如果没有找到结构化内容,使用通用切分
        if not chunks:
            chunks = self._split_generic_document(content)

        return chunks

    def _split_generic_document(self, content: str) -> list[DocumentChunk]:
        """通用文档切分"""
        chunks = []

        # 按段落切分
        paragraphs = self.paragraph_splitter.split(content)
        for i, paragraph in enumerate(paragraphs):
            if paragraph.strip():
                chunk = DocumentChunk(
                    chunk_id=f"para_{i+1}",
                    content=paragraph.strip(),
                    chunk_type="paragraph",
                    metadata={"paragraph_number": i + 1, "document_type": "generic"},
                )
                chunks.append(chunk)

        return chunks

    def _generate_processing_stats(
        self, content: str, chunks: list[DocumentChunk]
    ) -> dict[str, Any]:
        """生成处理统计信息"""
        return {
            "total_characters": len(content),
            "total_chunks": len(chunks),
            "average_chunk_size": (
                sum(len(c.content) for c in chunks) / len(chunks) if chunks else 0
            ),
            "chunk_types": list({c.chunk_type for c in chunks}),
            "processing_time": datetime.now().isoformat(),
        }


# 导入uuid
import uuid


# 使用示例
async def process_file_example(file_path: str) -> ProcessedDocument | None:
    """处理文件示例"""
    processor = DocumentProcessor()
    return await processor.process_document(file_path)
