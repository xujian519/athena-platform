#!/usr/bin/env python3
"""
商标法律法规导入到法律向量库和知识图谱的专用工具
将商标相关法律法规导入到:
  - PostgreSQL: legal_documents 表
  - Qdrant: legal_knowledge 集合
  - NebulaGraph: legal_kg 空间

作者: 小诺·双鱼座（Athena平台AI智能体）
创建时间: 2025-12-26
"""

from __future__ import annotations
import asyncio
import logging
import os
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF

# BGE向量嵌入服务
import psycopg2


def get_env_password(key: str, default: str = "") -> str:
    """从环境变量获取密码"""
    return os.environ.get(key, default)
from docx import Document
from nebula3.Config import Config
from nebula3.gclient.net import ConnectionPool
from psycopg2.extras import Json
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, PointStruct, VectorParams
from sentence_transformers import SentenceTransformer

# 使用安全哈希函数替代不安全的MD5/SHA1
from production.utils.security_helpers import short_hash

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.file_handler(f'/Users/xujian/Athena工作平台/logs/trademark_law_import_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log'),
        logging.stream_handler()
    ]
)
logger = logging.getLogger(__name__)


@dataclass
class TrademarkLawDocument:
    """商标法律文档数据结构"""
    doc_id: str  # PostgreSQL UUID (会自动生成)
    doc_ref_id: str  # 引用ID（用于外部引用）
    doc_type: str  # law/ regulation/guideline/interpretation/court_guide
    title: str
    content: str
    source_file: str
    file_type: str  # pdf/docx/doc
    metadata: dict[str, Any] = field(default_factory=dict)

    def __post_init__(self):
        self.char_count = len(self.content)
        self.word_count = len([w for w in self.content if w.strip()])

        # 生成引用ID（用于外部标识）
        if not self.doc_ref_id or self.doc_ref_id == "":
            content_hash = short_hash(f"{self.title}{self.source_file}".encode())[:16]
            self.doc_ref_id = f"trademark_{self.doc_type}_{content_hash}"

        # doc_id保持为None，让PostgreSQL自动生成UUID


class TrademarkLawImporter:
    """商标法律法规导入器"""

    def __init__(self, use_bge: bool = True):
        self.base_dir = Path("/Users/xujian/Athena工作平台")
        self.source_dir = Path("/Volumes/AthenaData/语料/商标相关法律法规")

        # 数据库配置
        self.pg_conn = None
        self.qdrant_client = None
        self.nebula_pool = None

        # 集合名称
        self.collection_name = "legal_knowledge"
        self.space_name = "legal_kg"

        # BGE模型配置
        self.use_bge = use_bge
        self.bge_model = None
        self.bge_dimension = 1024  # bge-large-zh-v1.5

        logger.info("商标法律法规导入器初始化完成")
        logger.info(f"源目录: {self.source_dir}")
        logger.info(f"BGE向量化: {'启用' if use_bge else '禁用'}")

    async def initialize_databases(self):
        """初始化数据库连接"""
        logger.info("=" * 60)
        logger.info("🚀 初始化数据库连接")
        logger.info("=" * 60)

        # 1. PostgreSQL连接
        try:
            self.pg_conn = psycopg2.connect(
                host="127.0.0.1",
                port=5432,
                database="patent_legal_db",
                user="postgres",
                password=get_env_password('NEBULA_PASSWORD')
            )
            logger.info("✅ PostgreSQL连接成功")
        except Exception as e:
            logger.error(f"❌ PostgreSQL连接失败: {e}")
            raise

        # 2. Qdrant客户端
        try:
            self.qdrant_client = QdrantClient(
                url="http://localhost:6333",
                timeout=30
            )

            # 确保集合存在
            collections = self.qdrant_client.get_collections().collections
            collection_names = [c.name for c in collections]

            if self.collection_name not in collection_names:
                logger.info(f"📋 创建Qdrant集合: {self.collection_name}")
                self.qdrant_client.create_collection(
                    collection_name=self.collection_name,
                    vectors_config=VectorParams(size=self.bge_dimension, distance=Distance.COSINE)
                )
            else:
                logger.info(f"✅ Qdrant集合已存在: {self.collection_name}")

            logger.info("✅ Qdrant连接成功")

        except Exception as e:
            logger.error(f"❌ Qdrant连接失败: {e}")
            raise

        # 3. 加载BGE模型（使用Apple Silicon优化版本）
        if self.use_bge:
            try:
                logger.info("🔄 加载BGE模型...")
                # 使用Apple Silicon优化的本地模型
                apple_silicon_model_path = "/Users/xujian/Athena工作平台/models/converted/bge-large-zh-v1.5"
                logger.info(f"使用Apple Silicon优化模型: {apple_silicon_model_path}")

                # 验证模型目录存在
                if not Path(apple_silicon_model_path).exists():
                    raise FileNotFoundError(f"模型目录不存在: {apple_silicon_model_path}")

                self.bge_model = SentenceTransformer(apple_silicon_model_path)
                logger.info(f"✅ BGE模型加载成功，向量维度: {self.bge_dimension}")

            except Exception as e:
                logger.error(f"❌ BGE模型加载失败: {e}")
                logger.warning("⚠️ 将跳过向量生成")
                self.use_bge = False

        # 4. NebulaGraph连接
        try:
            config = Config()
            config.max_connection_pool_size = 10
            self.nebula_pool = ConnectionPool()
            self.nebula_pool.init([('127.0.0.1', 9669)], config)

            # 验证连接
            session = self.nebula_pool.get_session('root', 'nebula')

            # 确保空间存在
            result = session.execute(f"USE {self.space_name}")
            if not result.is_succeeded():
                logger.info(f"📋 创建NebulaGraph空间: {self.space_name}")
                result = session.execute(f"CREATE SPACE IF NOT EXISTS {self.space_name}")
                if not result.is_succeeded():
                    logger.error(f"创建空间失败: {result.error_msg()}")

            session.release()
            logger.info("✅ NebulaGraph连接成功")

        except Exception as e:
            logger.error(f"❌ NebulaGraph连接失败: {e}")
            logger.warning("⚠️ 将继续使用PostgreSQL和Qdrant")

    def parse_pdf_document(self, file_path: Path) -> list[TrademarkLawDocument]:
        """解析PDF文档"""
        logger.info(f"📄 解析PDF: {file_path.name}")

        documents = []

        try:
            pdf_document = fitz.open(str(file_path))
            total_pages = pdf_document.page_count

            logger.info(f"   总页数: {total_pages}")

            # 提取所有文本
            full_text = ""
            for page_num in range(total_pages):
                page = pdf_document[page_num]
                text = page.get_text()
                full_text += text + "\n\n"

            # 根据文件名确定文档类型
            filename = file_path.stem
            if "商标法" in filename and "实施条例" not in filename:
                doc_type = "law"
                title = "中华人民共和国商标法"
            elif "实施条例" in filename:
                doc_type = "regulation"
                title = "中华人民共和国商标法实施条例"
            elif "审查指南" in filename:
                doc_type = "guideline"
                title = "商标审查审理指南"
            elif "司法解释" in filename or "解释" in filename:
                doc_type = "interpretation"
                title = filename.replace("_", " ").replace(".doc", "")
            elif "高院" in filename or "审理指南" in filename:
                doc_type = "court_guide"
                title = filename.replace("_", " ").replace(".doc", "")
            else:
                doc_type = "other"
                title = filename

            # 创建文档对象
            doc = TrademarkLawDocument(
                doc_id=None,  # PostgreSQL UUID（数据库自动生成）
                doc_ref_id="",  # 引用ID（将在__post_init__中生成）
                doc_type=doc_type,
                title=title,
                content=full_text,
                source_file=str(file_path),
                file_type="pdf",
                metadata={
                    "total_pages": total_pages,
                    "effective_date": self.extract_date_from_filename(filename),
                    "authority": self.extract_authority(doc_type)
                }
            )

            documents.append(doc)
            logger.info(f"   ✅ 解析完成: {len(full_text)}字符")

            pdf_document.close()

        except Exception as e:
            logger.error(f"   ❌ PDF解析失败: {e}")

        return documents

    def parse_docx_document(self, file_path: Path) -> list[TrademarkLawDocument]:
        """解析DOCX文档"""
        logger.info(f"📝 解析DOCX: {file_path.name}")

        documents = []

        try:
            doc = Document(str(file_path))

            # 提取所有段落文本
            full_text = "\n".join([para.text for para in doc.paragraphs])

            # 确定文档类型
            filename = file_path.stem
            if "商标法" in filename:
                doc_type = "law"
                title = "中华人民共和国商标法"
            elif "实施条例" in filename:
                doc_type = "regulation"
                title = "中华人民共和国商标法实施条例"
            elif "解释" in filename:
                doc_type = "interpretation"
                title = filename.replace("_", " ")
            else:
                doc_type = "other"
                title = filename

            # 创建文档对象
            law_doc = TrademarkLawDocument(
                doc_id=None,  # PostgreSQL UUID（数据库自动生成）
                doc_ref_id="",  # 引用ID（将在__post_init__中生成）
                doc_type=doc_type,
                title=title,
                content=full_text,
                source_file=str(file_path),
                file_type="docx",
                metadata={
                    "paragraphs": len(doc.paragraphs),
                    "effective_date": self.extract_date_from_filename(filename),
                    "authority": self.extract_authority(doc_type)
                }
            )

            documents.append(law_doc)
            logger.info(f"   ✅ 解析完成: {len(full_text)}字符")

        except Exception as e:
            logger.error(f"   ❌ DOCX解析失败: {e}")

        return documents

    def extract_date_from_filename(self, filename: str) -> str:
        """从文件名提取生效日期"""
        # 格式通常是: _20190423 或 _20201229
        import re
        match = re.search(r'_(\d{8})', filename)
        if match:
            date_str = match.group(1)
            try:
                return f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"
            except Exception as e:
                logger.debug(f"空except块已触发: {e}")
                return date_str
        return ""

    def extract_authority(self, doc_type: str) -> str:
        """提取发布机关"""
        authorities = {
            "law": "全国人大常委会",
            "regulation": "国务院",
            "guideline": "国家知识产权局",
            "interpretation": "最高人民法院",
            "court_guide": "北京市高级人民法院"
        }
        return authorities.get(doc_type, "")

    def split_document_into_chunks(self, doc: TrademarkLawDocument, chunk_size: int = 1000) -> list[dict[str, Any]]:
        """将文档切分成块"""
        logger.info(f"📊 切分文档: {doc.title}")

        chunks = []
        content = doc.content
        total_chars = len(content)

        # 使用doc_ref_id而不是doc_id来生成chunk_id（因为doc_id可能是None）
        for i in range(0, total_chars, chunk_size):
            chunk_text = content[i:i + chunk_size]

            chunk_id = f"{doc.doc_ref_id}_chunk_{i // chunk_size:04d}"

            chunk = {
                "chunk_id": chunk_id,
                "doc_id": doc.doc_id,  # PostgreSQL UUID
                "doc_ref_id": doc.doc_ref_id,  # 引用ID
                "title": doc.title,
                "doc_type": doc.doc_type,
                "content": chunk_text,
                "chunk_index": i // chunk_size,
                "source_file": doc.source_file,
                "metadata": {
                    **doc.metadata,
                    "char_count": len(chunk_text)
                }
            }

            chunks.append(chunk)

        logger.info(f"   ✅ 切分成 {len(chunks)} 个块")
        return chunks

    async def save_to_postgresql(self, doc: TrademarkLawDocument):
        """保存到PostgreSQL legal_documents表"""
        logger.info(f"💾 保存到PostgreSQL: {doc.title}")

        try:
            cursor = self.pg_conn.cursor()

            # 检查是否已存在
            cursor.execute(
                "SELECT id FROM legal_documents WHERE file_path = %s",
                (doc.source_file,)
            )

            if cursor.fetchone():
                logger.info("   ⚠️ 文档已存在，跳过")
                cursor.close()
                return None

            # 插入文档（不指定id，让数据库自动生成UUID）
            cursor.execute(
                """
                INSERT INTO legal_documents (
                    file_path, file_name, document_type,
                    content, metadata, text_length, created_at
                ) VALUES (%s, %s, %s, %s, %s, %s, %s)
                RETURNING id
                """,
                (
                    doc.source_file,
                    Path(doc.source_file).name,
                    doc.doc_type,
                    doc.content,
                    Json(doc.metadata),
                    doc.char_count,
                    datetime.now()
                )
            )

            # 获取自动生成的UUID
            doc_id = cursor.fetchone()[0]
            doc.doc_id = str(doc_id)  # 保存到对象中供后续使用

            self.pg_conn.commit()
            cursor.close()

            logger.info(f"   ✅ 保存成功，ID: {doc.doc_id}")

        except Exception as e:
            logger.error(f"   ❌ 保存失败: {e}")
            self.pg_conn.rollback()
            doc.doc_id = None

    async def generate_and_save_vectors(self, chunks: list[dict[str, Any]]):
        """生成向量并保存到Qdrant"""
        logger.info(f"📊 生成并向量化 {len(chunks)} 个块")

        if not self.use_bge or not self.bge_model:
            logger.info("   ⏭️ BGE模型未加载，跳过向量化")
            return

        # 批量生成向量
        batch_size = 32
        vectors_generated = 0
        vectors_saved = 0

        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            batch_texts = [chunk["content"] for chunk in batch]
            batch_ids = [chunk["chunk_id"] for chunk in batch]

            try:
                # 生成向量
                logger.info(f"   🔄 生成向量 [{i+1}-{min(i+batch_size, len(chunks))}/{len(chunks)}]")
                embeddings = self.bge_model.encode(
                    batch_texts,
                    normalize_embeddings=True,
                    show_progress_bar=False
                )

                # 准备Qdrant点
                points = []
                for _j, (chunk_id, embedding, chunk) in enumerate(zip(batch_ids, embeddings, batch, strict=False)):
                    # 使用chunk_id的哈希作为点ID
                    point_id = hash(chunk_id) & 0xFFFFFFFFFFFFFFFF  # 转换为64位无符号整数

                    points.append(PointStruct(
                        id=point_id,
                        vector=embedding.tolist(),
                        payload={
                            "chunk_id": chunk_id,
                            "doc_id": chunk.get("doc_id", ""),
                            "doc_ref_id": chunk.get("doc_ref_id", ""),
                            "title": chunk["title"],
                            "doc_type": chunk["doc_type"],
                            "content": chunk["content"],
                            "chunk_index": chunk["chunk_index"],
                            "source_file": chunk["source_file"],
                            **chunk["metadata"]
                        }
                    ))
                    vectors_generated += 1

                # 保存到Qdrant
                logger.info(f"   💾 保存到Qdrant: {len(points)} 个向量")
                self.qdrant_client.upsert(
                    collection_name=self.collection_name,
                    points=points
                )
                vectors_saved += len(points)

            except Exception as e:
                logger.error(f"   ❌ 批次处理失败: {e}")
                continue

        logger.info(f"   ✅ 完成 {vectors_generated} 个向量生成，{vectors_saved} 个向量保存")

    def save_to_knowledge_graph(self, doc: TrademarkLawDocument) -> None:
        """保存到NebulaGraph legal_kg空间"""
        logger.info(f"🕸️ 保存到知识图谱: {doc.title}")
        logger.info("   ⏭️ 知识图谱存储暂未实现（需要先创建TAG定义）")
        # TODO: 需要先在legal_kg空间创建对应的TAG定义
        # 目前只保存到PostgreSQL和Qdrant

    async def process_file(self, file_path: Path):
        """处理单个文件"""
        logger.info("=" * 60)
        logger.info(f"📄 处理文件: {file_path.name}")
        logger.info("=" * 60)

        # 1. 解析文档
        if file_path.suffix.lower() == '.pdf':
            documents = self.parse_pdf_document(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            documents = self.parse_docx_document(file_path)
        else:
            logger.warning(f"   ⚠️ 不支持的文件类型: {file_path.suffix}")
            return

        if not documents:
            logger.error("   ❌ 文档解析失败")
            return

        # 2. 保存到PostgreSQL
        for doc in documents:
            await self.save_to_postgresql(doc)

            # 3. 切分成块
            chunks = self.split_document_into_chunks(doc)

            # 4. 生成向量（需要BGE服务）
            await self.generate_and_save_vectors(chunks)

            # 5. 保存到知识图谱
            self.save_to_knowledge_graph(doc)

    async def process_all_files(self):
        """处理所有文件"""
        logger.info("=" * 60)
        logger.info("🚀 开始处理所有商标法律法规文件")
        logger.info("=" * 60)

        # 获取所有文件
        files = list(self.source_dir.glob("*"))
        files = [f for f in files if f.is_file()]

        logger.info(f"📊 找到 {len(files)} 个文件")

        # 按类型排序（先处理小文件）
        files.sort(key=lambda f: f.stat().st_size)

        for idx, file_path in enumerate(files, 1):
            logger.info(f"\n📄 进度: [{idx}/{len(files)}] {file_path.name}")
            await self.process_file(file_path)

        logger.info("=" * 60)
        logger.info("✅ 所有文件处理完成")
        logger.info("=" * 60)

    def close(self) -> Any:
        """关闭数据库连接"""
        if self.pg_conn:
            self.pg_conn.close()
            logger.info("✅ PostgreSQL连接已关闭")

        if self.nebula_pool:
            self.nebula_pool.close()
            logger.info("✅ NebulaGraph连接已关闭")


async def main():
    """主函数"""
    importer = TrademarkLawImporter()

    try:
        # 初始化数据库连接
        await importer.initialize_databases()

        # 处理所有文件
        await importer.process_all_files()

    finally:
        importer.close()


if __name__ == "__main__":
    asyncio.run(main())
