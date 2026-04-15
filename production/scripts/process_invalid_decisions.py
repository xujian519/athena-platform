#!/usr/bin/env python3
"""
专利无效复审决定原文批量处理脚本
将DOCX文件解析后导入到PostgreSQL和Qdrant

作者: 小诺·双鱼座（Athena平台AI智能体）
创建时间: 2025-12-26
"""

from __future__ import annotations
import asyncio
import logging
import os
import re
import subprocess
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

import psycopg2


def get_env_password(key: str, default: str = "") -> str:
    """从环境变量获取密码"""
    return os.environ.get(key, default)
from docx import Document
from nebula3.Config import Config
from nebula3.gclient.net import ConnectionPool
from psycopg2.extras import Json
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, PointStruct, VectorParams
from sentence_transformers import SentenceTransformer

# 使用安全哈希函数替代不安全的MD5/SHA1
from production.utils.security_helpers import short_hash

# 尝试导入docx2python用于处理.doc文件
try:
    from docx2python import docx2python
    HAS_DOCX2PYTHON = True
except ImportError:
    HAS_DOCX2PYTHON = False

# 添加项目路径
sys.path.insert(0, str(Path(__file__).parent.parent))

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.file_handler(f'/Users/xujian/Athena工作平台/logs/invalid_decision_processor_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log'),
        logging.stream_handler()
    ]
)
logger = logging.getLogger(__name__)


@dataclass
class InvalidDecisionDocument:
    """无效决定文档数据结构"""
    doc_id: str
    decision_number: str  # 决定号
    patent_number: str  # 专利号
    title: str
    content: str
    source_file: str
    file_type: str  # docx/doc
    metadata: dict[str, Any] = field(default_factory=dict)

    def __post_init__(self):
        self.char_count = len(self.content)
        # 如果没有决定号，从文件名提取
        if not self.decision_number:
            filename = Path(self.source_file).stem
            # 尝试提取9位数字作为决定号
            match = re.search(r'(\d{9})', filename)
            if match:
                self.decision_number = match.group(1)


class InvalidDecisionImporter:
    """专利无效复审决定导入器"""

    def __init__(self, use_bge: bool = True):
        self.base_dir = Path("/Users/xujian/Athena工作平台")
        # 修正为实际的目录路径
        self.source_dir = Path("/Volumes/AthenaData/语料/专利/专利无效复审决定原文")

        # 数据库配置
        self.pg_conn = None
        self.qdrant_client = None
        self.nebula_pool = None

        # 集合名称
        self.collection_name = "patent_decisions"
        self.space_name = "patent_rules"

        # BGE模型配置
        self.use_bge = use_bge
        self.bge_model = None
        self.bge_dimension = 1024  # bge-large-zh-v1.5
        self.chunk_size = 1000  # 文本分块大小

        logger.info("专利无效复审决定导入器初始化完成")
        logger.info(f"源目录: {self.source_dir}")
        logger.info(f"BGE向量化: {'启用' if use_bge else '禁用'}")

    async def initialize_databases(self):
        """初始化数据库连接"""
        logger.info("=" * 60)
        logger.info("🚀 初始化数据库连接")
        logger.info("=" * 60)

        # 1. PostgreSQL连接
        try:
            self.pg_conn = psycopg2.connect(
                host="127.0.0.1",
                port=5432,
                database="patent_legal_db",
                user="postgres",
                password=get_env_password('NEBULA_PASSWORD')
            )
            logger.info("✅ PostgreSQL连接成功")

            # 创建表（如果不存在）
            cursor = self.pg_conn.cursor()
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS patent_invalid_decisions (
                    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                    decision_number VARCHAR(50),
                    patent_number VARCHAR(50),
                    file_path TEXT UNIQUE NOT NULL,
                    file_name VARCHAR(255),
                    document_type VARCHAR(50),
                    content TEXT,
                    metadata JSONB,
                    text_length INTEGER,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );

                CREATE INDEX IF NOT EXISTS idx_decision_number ON patent_invalid_decisions(decision_number);
                CREATE INDEX IF NOT EXISTS idx_patent_number ON patent_invalid_decisions(patent_number);
                CREATE INDEX IF NOT EXISTS idx_file_path ON patent_invalid_decisions(file_path);
            """)
            self.pg_conn.commit()
            cursor.close()

        except Exception as e:
            logger.error(f"❌ PostgreSQL连接失败: {e}")
            raise

        # 2. Qdrant客户端
        try:
            self.qdrant_client = QdrantClient(
                url="http://localhost:6333",
                timeout=30
            )

            # 确保集合存在
            collections = self.qdrant_client.get_collections().collections
            collection_names = [c.name for c in collections]

            if self.collection_name not in collection_names:
                logger.info(f"📋 创建Qdrant集合: {self.collection_name}")
                self.qdrant_client.create_collection(
                    collection_name=self.collection_name,
                    vectors_config=VectorParams(size=self.bge_dimension, distance=Distance.COSINE)
                )
            else:
                logger.info(f"✅ Qdrant集合已存在: {self.collection_name}")

            logger.info("✅ Qdrant连接成功")

        except Exception as e:
            logger.error(f"❌ Qdrant连接失败: {e}")
            raise

        # 3. 加载BGE模型（使用Apple Silicon优化版本）
        if self.use_bge:
            try:
                logger.info("🔄 加载BGE模型...")
                apple_silicon_model_path = "/Users/xujian/Athena工作平台/models/converted/bge-large-zh-v1.5"
                logger.info(f"使用Apple Silicon优化模型: {apple_silicon_model_path}")

                if not Path(apple_silicon_model_path).exists():
                    raise FileNotFoundError(f"模型目录不存在: {apple_silicon_model_path}")

                self.bge_model = SentenceTransformer(apple_silicon_model_path)
                logger.info(f"✅ BGE模型加载成功，向量维度: {self.bge_dimension}")

            except Exception as e:
                logger.error(f"❌ BGE模型加载失败: {e}")
                logger.warning("⚠️ 将跳过向量生成")
                self.use_bge = False

        # 4. NebulaGraph连接（知识图谱）
        try:
            config = Config()
            config.max_connection_pool_size = 10
            self.nebula_pool = ConnectionPool()
            self.nebula_pool.init([('127.0.0.1', 9669)], config)

            # 验证连接并创建Schema
            session = self.nebula_pool.get_session('root', 'nebula')

            # 确保空间存在
            result = session.execute(f"USE {self.space_name}")
            if not result.is_succeeded():
                logger.info(f"📋 创建NebulaGraph空间: {self.space_name}")
                result = session.execute(f"CREATE SPACE IF NOT EXISTS {self.space_name} (partition_num=10, replica_factor=1, vid_type=FIXED_STRING(256));")
                if not result.is_succeeded():
                    logger.warning(f"创建空间失败: {result.error_msg()}")
                else:
                    # 等待空间生效
                    import time
                    time.sleep(10)
                    result = session.execute(f"USE {self.space_name}")

            # 如果空间可用，创建TAG和EDGE
            if result.is_succeeded():
                # 创建TAG（顶点类型）
                tags = [
                    "patent_decision(decision_id string, decision_number string, patent_number string, title string, content_length string, source_file string)",
                    "decision_block(block_id string, content string, block_index string)",
                    "legal_ref(ref_id string, law_name string, article string, reference_type string)",
                    "legal_reference(ref_id string, patent_number string, reference_type string)"
                ]

                for tag in tags:
                    try:
                        result = session.execute(f"CREATE TAG IF NOT EXISTS {tag};")
                        if result.is_succeeded():
                            logger.debug(f"   ✅ 创建TAG: {tag.split('(')[0]}")
                    except Exception:
                        logger.debug(f"   ⚠️ TAG可能已存在: {tag.split('(')[0]}")

                # 创建EDGE（关系类型）
                edges = [
                    "contains(sequence string)",
                    "cites(citation_type string)",
                    "decides(outcome string, date string)",
                    "precedes(temporal_relation string)",
                    "refers_to(reference_type string)"
                ]

                for edge in edges:
                    try:
                        result = session.execute(f"CREATE EDGE IF NOT EXISTS {edge};")
                        if result.is_succeeded():
                            logger.debug(f"   ✅ 创建EDGE: {edge.split('(')[0]}")
                    except Exception:
                        logger.debug(f"   ⚠️ EDGE可能已存在: {edge.split('(')[0]}")

            session.release()
            logger.info("✅ NebulaGraph连接成功")

        except Exception as e:
            logger.error(f"❌ NebulaGraph连接失败: {e}")
            logger.warning("⚠️ 将继续使用PostgreSQL和Qdrant，跳过知识图谱功能")

    def parse_docx_document(self, file_path: Path) -> InvalidDecisionDocument | None:
        """解析DOCX文档"""
        logger.debug(f"📝 解析DOCX: {file_path.name}")

        try:
            doc = Document(str(file_path))

            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            full_text = "\n".join(paragraphs)

            if not full_text:
                logger.warning(f"   ⚠️ 文档为空: {file_path.name}")
                return None

            # 从文本提取元数据
            decision_number = ""
            patent_number = ""

            # 提取决定号 (第XXXXX号)
            match = re.search(r'第(\d+)号', full_text)
            if match:
                decision_number = match.group(1)

            # 提取专利号 (CNXXXXXXXXX.X 或 10位数字)
            match = re.search(r'(CN\d{9}\.\d|[0-9X]{10})', full_text)
            if match:
                patent_number = match.group(1)

            # 生成标题
            title = f"专利无效复审决定 {decision_number or file_path.stem}"

            # 创建文档对象
            decision_doc = InvalidDecisionDocument(
                doc_id="",  # 将在保存到数据库后生成
                decision_number=decision_number,
                patent_number=patent_number,
                title=title,
                content=full_text,
                source_file=str(file_path),
                file_type="docx",
                metadata={
                    "paragraphs": len(paragraphs),
                    "filename": file_path.name
                }
            )

            logger.debug(f"   ✅ 解析完成: {len(full_text)}字符")
            return decision_doc

        except Exception as e:
            logger.error(f"   ❌ DOCX解析失败 {file_path.name}: {e}")
            return None

    def parse_doc_document(self, file_path: Path) -> InvalidDecisionDocument | None:
        """解析旧版DOC文档（使用docx2python或textutil）"""
        logger.info(f"📝 解析DOC: {file_path.name}")

        try:
            # 方法1: 使用docx2python（如果可用）
            if HAS_DOCX2PYTHON:
                try:
                    content_iter = docx2python(str(file_path))
                    full_text = content_iter.text
                except Exception as e:
                    logger.warning(f"   docx2python失败: {e}")
                    full_text = None
            else:
                full_text = None

            # 方法2: 使用macOS textutil命令
            if not full_text or len(full_text.strip()) < 100:
                try:
                    result = subprocess.run(
                        ['textutil', '-convert', 'txt', '-stdout', str(file_path)],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    if result.returncode == 0:
                        full_text = result.stdout
                except Exception as e:
                    logger.warning(f"   textutil失败: {e}")

            if not full_text or len(full_text.strip()) < 50:
                logger.warning("   ⚠️ 无法提取文本内容")
                return None

            # 从文本提取元数据
            decision_number = ""
            patent_number = ""

            # 提取决定号 (第XXXXX号)
            match = re.search(r'第(\d+)号', full_text)
            if match:
                decision_number = match.group(1)

            # 提取专利号 (CNXXXXXXXXX.X 或 10位数字)
            match = re.search(r'(CN\d{9}\.\d|[0-9X]{10})', full_text)
            if match:
                patent_number = match.group(1)

            # 生成标题
            title = f"专利无效复审决定 {decision_number or file_path.stem}"

            # 创建文档对象
            decision_doc = InvalidDecisionDocument(
                doc_id="",
                decision_number=decision_number,
                patent_number=patent_number,
                title=title,
                content=full_text.strip(),
                source_file=str(file_path),
                file_type="doc",
                metadata={
                    "extraction_method": "docx2python" if HAS_DOCX2PYTHON else "textutil",
                    "filename": file_path.name
                }
            )

            logger.info(f"   ✅ 解析完成: {len(full_text)}字符")
            return decision_doc

        except Exception as e:
            logger.error(f"   ❌ DOC解析失败 {file_path.name}: {e}")
            return None

    def split_document_into_chunks(self, doc: InvalidDecisionDocument) -> list[dict[str, Any]]:
        """将文档切分成块"""
        chunks = []
        content = doc.content
        total_chars = len(content)

        # 使用doc_id来生成chunk_id
        doc_id_for_chunk = doc.doc_id if doc.doc_id else short_hash(doc.source_file.encode())[:16]

        for i in range(0, total_chars, self.chunk_size):
            chunk_text = content[i:i + self.chunk_size]

            chunk_id = f"decision_{doc_id_for_chunk}_chunk_{i // self.chunk_size:04d}"

            chunk = {
                "chunk_id": chunk_id,
                "doc_id": doc.doc_id,
                "decision_number": doc.decision_number,
                "patent_number": doc.patent_number,
                "title": doc.title,
                "content": chunk_text,
                "chunk_index": i // self.chunk_size,
                "source_file": doc.source_file,
                "metadata": {
                    **doc.metadata,
                    "char_count": len(chunk_text)
                }
            }

            chunks.append(chunk)

        logger.debug(f"   ✅ 切分成 {len(chunks)} 个块")
        return chunks

    async def save_to_postgresql(self, doc: InvalidDecisionDocument):
        """保存到PostgreSQL patent_invalid_decisions表"""
        logger.info(f"💾 保存到PostgreSQL: {doc.title}")

        try:
            cursor = self.pg_conn.cursor()

            # 检查是否已存在
            cursor.execute(
                "SELECT id FROM patent_invalid_decisions WHERE file_path = %s",
                (doc.source_file,)
            )

            existing = cursor.fetchone()
            if existing:
                doc.doc_id = str(existing[0])
                logger.info(f"   ⚠️ 文档已存在，使用现有ID: {doc.doc_id}")
                cursor.close()
                return

            # 插入文档
            cursor.execute(
                """
                INSERT INTO patent_invalid_decisions (
                    decision_number, patent_number, file_path, file_name,
                    document_type, content, metadata, text_length, created_at
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id
                """,
                (
                    doc.decision_number,
                    doc.patent_number,
                    doc.source_file,
                    Path(doc.source_file).name,
                    "invalid_decision",
                    doc.content,
                    Json(doc.metadata),
                    doc.char_count,
                    datetime.now()
                )
            )

            # 获取自动生成的UUID
            doc_id = cursor.fetchone()[0]
            doc.doc_id = str(doc_id)

            self.pg_conn.commit()
            cursor.close()

            logger.info(f"   ✅ 保存成功，ID: {doc.doc_id}")

        except Exception as e:
            logger.error(f"   ❌ 保存失败: {e}")
            self.pg_conn.rollback()
            doc.doc_id = None

    async def generate_and_save_vectors(self, chunks: list[dict[str, Any]]):
        """生成向量并保存到Qdrant"""
        logger.info(f"📊 生成并向量化 {len(chunks)} 个块")

        if not self.use_bge or not self.bge_model:
            logger.info("   ⏭️ BGE模型未加载，跳过向量化")
            return

        # 批量生成向量
        batch_size = 32
        vectors_saved = 0

        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            batch_texts = [chunk["content"] for chunk in batch]
            batch_ids = [chunk["chunk_id"] for chunk in batch]

            try:
                # 生成向量
                logger.info(f"   🔄 生成向量 [{i+1}-{min(i+batch_size, len(chunks))}/{len(chunks)}]")
                embeddings = self.bge_model.encode(
                    batch_texts,
                    normalize_embeddings=True,
                    show_progress_bar=False
                )

                # 准备Qdrant点
                points = []
                for chunk_id, embedding, chunk in zip(batch_ids, embeddings, batch, strict=False):
                    # 使用chunk_id的哈希作为点ID
                    point_id = hash(chunk_id) & 0xFFFFFFFFFFFFFFFF

                    points.append(PointStruct(
                        id=point_id,
                        vector=embedding.tolist(),
                        payload={
                            "chunk_id": chunk_id,
                            "doc_id": chunk.get("doc_id", ""),
                            "decision_number": chunk.get("decision_number", ""),
                            "patent_number": chunk.get("patent_number", ""),
                            "title": chunk["title"],
                            "content": chunk["content"],
                            "chunk_index": chunk["chunk_index"],
                            "source_file": chunk["source_file"],
                            **chunk["metadata"]
                        }
                    ))

                # 保存到Qdrant
                logger.info(f"   💾 保存到Qdrant: {len(points)} 个向量")
                self.qdrant_client.upsert(
                    collection_name=self.collection_name,
                    points=points
                )
                vectors_saved += len(points)

            except Exception as e:
                logger.error(f"   ❌ 批次处理失败: {e}")
                continue

        logger.info(f"   ✅ 完成 {vectors_saved} 个向量保存")

    async def process_file(self, file_path: Path):
        """处理单个文件"""
        logger.info("=" * 60)
        logger.info(f"📄 处理文件: {file_path.name}")
        logger.info("=" * 60)

        # 1. 解析文档
        if file_path.suffix.lower() == '.docx':
            doc = self.parse_docx_document(file_path)
        elif file_path.suffix.lower() == '.doc':
            doc = self.parse_doc_document(file_path)
        else:
            logger.warning(f"   ⚠️ 不支持的文件类型: {file_path.suffix}")
            return

        if not doc:
            logger.error("   ❌ 文档解析失败")
            return

        # 2. 保存到PostgreSQL
        await self.save_to_postgresql(doc)

        if not doc.doc_id:
            logger.error("   ❌ 未获得文档ID，跳过向量化")
            return

        # 3. 切分成块
        chunks = self.split_document_into_chunks(doc)

        # 4. 生成向量并保存到Qdrant
        await self.generate_and_save_vectors(chunks)

        # 5. 保存到知识图谱
        await self.save_to_knowledge_graph(doc)

    async def process_batch(self, max_files: int = None):
        """批量处理文件"""
        logger.info("=" * 60)
        logger.info("🚀 开始批量处理专利无效复审决定")
        logger.info("=" * 60)

        # 获取所有文件（包括.doc和.docx）
        docx_files = list(self.source_dir.glob("*.docx"))
        doc_files = list(self.source_dir.glob("*.doc"))

        docx_files = [f for f in docx_files if f.is_file()]
        doc_files = [f for f in doc_files if f.is_file()]

        logger.info(f"📊 找到 {len(docx_files)} 个DOCX文件, {len(doc_files)} 个DOC文件")

        # 先处理DOCX文件，再处理DOC文件
        files = docx_files + doc_files

        if max_files:
            files = files[:max_files]
            logger.info(f"📋 处理模式: 测试模式（前{max_files}个文件）")
        else:
            logger.info(f"📋 处理模式: 全量模式（全部{len(files)}个文件）")

        for idx, file_path in enumerate(files, 1):
            logger.info(f"\n📄 进度: [{idx}/{len(files)}] {file_path.name}")
            await self.process_file(file_path)

        logger.info("=" * 60)
        logger.info("✅ 批量处理完成")
        logger.info("=" * 60)

    def _escape_nebula_string(self, s: str) -> str:
        """转义NebulaGraph字符串"""
        if not s:
            return ""
        # 转义特殊字符
        s = str(s)
        s = s.replace('\\', '\\\\')  # 反斜杠必须最先转义
        s = s.replace('"', '\\"')     # 双引号
        s = s.replace("'", "\\'")     # 单引号
        s = s.replace('\n', '\\n')    # 换行符
        s = s.replace('\r', '\\r')    # 回车符
        s = s.replace('\t', '\\t')    # 制表符
        return s

    async def save_to_knowledge_graph(self, doc: InvalidDecisionDocument):
        """保存到NebulaGraph patent_rules空间"""
        logger.info(f"🕸️ 保存到知识图谱: {doc.title}")

        if not self.nebula_pool:
            logger.warning("   ⚠️ NebulaGraph未连接，跳过知识图谱保存")
            return

        try:
            session = self.nebula_pool.get_session('root', 'nebula')

            # 确保使用patent_rules空间
            use_space = f"USE {self.space_name};"
            result = session.execute(use_space)
            if not result.is_succeeded():
                logger.warning(f"   ⚠️ 无法使用空间 {self.space_name}: {result.error_msg()}")
                session.release()
                return

            # 提取实体和关系
            entities, relations = self.extract_entities_and_relations(doc)

            # 插入实体顶点
            vertices_inserted = 0
            for entity in entities:
                try:
                    vid = self._escape_nebula_string(entity['vid'])
                    tag = entity['tag']

                    # 正确构建属性值对
                    prop_pairs = []
                    for k, v in entity['props'].items():
                        escaped_key = self._escape_nebula_string(k)
                        escaped_value = self._escape_nebula_string(v)
                        prop_pairs.append(f'"{escaped_key}": "{escaped_value}"')

                    prop_names = ", ".join([f'"{self._escape_nebula_string(k)}"' for k in entity['props'].keys()])
                    props_str = ", ".join(prop_pairs)

                    # 构建INSERT VERTEX语句
                    insert_query = f'INSERT VERTEX {tag}({prop_names}) VALUES "{vid}": {{{props_str}}};'

                    result = session.execute(insert_query)

                    if result.is_succeeded():
                        vertices_inserted += 1
                    else:
                        logger.debug(f"顶点插入失败: {result.error_msg()}")
                except Exception as e:
                    # 顶点可能已存在，跳过
                    logger.debug(f"顶点插入异常: {e}")
                    pass

            logger.info(f"   ✅ 插入 {vertices_inserted}/{len(entities)} 个顶点")

            # 插入关系边
            edges_inserted = 0
            for relation in relations:
                try:
                    src_vid = self._escape_nebula_string(relation['src_vid'])
                    dst_vid = self._escape_nebula_string(relation['dst_vid'])
                    edge_type = relation['edge_type']

                    # 正确构建属性值对
                    prop_pairs = []
                    for k, v in relation['props'].items():
                        escaped_key = self._escape_nebula_string(k)
                        escaped_value = self._escape_nebula_string(v)
                        prop_pairs.append(f'"{escaped_key}": "{escaped_value}"')

                    prop_names = ", ".join([f'"{self._escape_nebula_string(k)}"' for k in relation['props'].keys()])
                    props_str = ", ".join(prop_pairs)

                    # 构建INSERT EDGE语句
                    insert_query = f'INSERT EDGE {edge_type}({prop_names}) VALUES "{src_vid}"->"{dst_vid}": {{{props_str}}};'

                    result = session.execute(insert_query)

                    if result.is_succeeded():
                        edges_inserted += 1
                    else:
                        logger.debug(f"边插入失败: {result.error_msg()}")
                except Exception as e:
                    # 边可能已存在，跳过
                    logger.debug(f"边插入异常: {e}")
                    pass

            logger.info(f"   ✅ 插入 {edges_inserted}/{len(relations)} 条边")

            session.release()
            logger.info("   ✅ 知识图谱保存成功")

        except Exception as e:
            logger.error(f"   ❌ 知识图谱保存失败: {e}")

    def extract_entities_and_relations(self, doc: InvalidDecisionDocument) -> tuple[list[dict], list[dict]]:
        """从文档中提取实体和关系"""
        entities = []
        relations = []
        content = doc.content

        # 1. 创建专利无效决定顶点（主文档）
        doc_vid = f"decision_{doc.doc_id if doc.doc_id else short_hash(doc.source_file.encode())}"
        entities.append({
            'vid': doc_vid,
            'tag': 'patent_decision',
            'props': {
                'decision_id': doc.doc_id if doc.doc_id else '',
                'decision_number': doc.decision_number,
                'patent_number': doc.patent_number,
                'title': doc.title[:100],  # 限制长度
                'content_length': str(doc.char_count),
                'source_file': doc.source_file
            }
        })

        # 2. 提取决定块（将文档切分成逻辑块）
        chunks = content.split('\n\n')
        for idx, chunk in enumerate(chunks[:20]):  # 限制最多20个块
            if len(chunk.strip()) > 50:
                chunk_vid = f"block_{doc_vid}_{idx}"
                entities.append({
                    'vid': chunk_vid,
                    'tag': 'decision_block',
                    'props': {
                        'block_id': chunk_vid,
                        'content': chunk[:500].replace('"', '\\"').replace('\n', ' '),  # 限制长度并转义
                        'block_index': str(idx)
                    }
                })

                # 添加关系：决定包含决定块
                relations.append({
                    'src_vid': doc_vid,
                    'dst_vid': chunk_vid,
                    'edge_type': 'contains',
                    'props': {
                        'sequence': str(idx)
                    }
                })

        # 3. 提取法律依据引用
        # 匹配专利法条款
        law_matches = re.finditer(r'专利法第([一二三四五六七八九十百千]+)条', content)
        for match in law_matches:
            law_ref_id = f"law_patent_{match.group(1)}"
            law_ref_vid = f"ref_{law_ref_id}"

            entities.append({
                'vid': law_ref_vid,
                'tag': 'legal_ref',
                'props': {
                    'ref_id': law_ref_id,
                    'law_name': '专利法',
                    'article': match.group(1),
                    'reference_type': 'statute'
                }
            })

            # 添加关系：决定引用法律依据
            relations.append({
                'src_vid': doc_vid,
                'dst_vid': law_ref_vid,
                'edge_type': 'cites',
                'props': {
                    'citation_type': 'legal_basis'
                }
            })

        # 匹配实施细则条款
        rule_matches = re.finditer(r'专利法实施细则第([一二三四五六七八九十百千]+)条', content)
        for match in rule_matches:
            rule_ref_id = f"rule_implement_{match.group(1)}"
            rule_ref_vid = f"ref_{rule_ref_id}"

            entities.append({
                'vid': rule_ref_vid,
                'tag': 'legal_ref',
                'props': {
                    'ref_id': rule_ref_id,
                    'law_name': '专利法实施细则',
                    'article': match.group(1),
                    'reference_type': 'regulation'
                }
            })

            # 添加关系：决定引用法律依据
            relations.append({
                'src_vid': doc_vid,
                'dst_vid': rule_ref_vid,
                'edge_type': 'cites',
                'props': {
                    'citation_type': 'legal_basis'
                }
            })

        # 4. 提取对比文件（专利号）
        patent_matches = re.finditer(r'CN(\d{9})\.\d', content)
        for match in patent_matches:
            patent_num = match.group(0)
            patent_vid = f"patent_{patent_num.replace('.', '_')}"

            entities.append({
                'vid': patent_vid,
                'tag': 'legal_reference',
                'props': {
                    'ref_id': patent_num,
                    'patent_number': patent_num,
                    'reference_type': 'prior_art'
                }
            })

            # 添加关系：决定引用对比文件
            relations.append({
                'src_vid': doc_vid,
                'dst_vid': patent_vid,
                'edge_type': 'refers_to',
                'props': {
                    'reference_type': 'prior_art'
                }
            })

        logger.debug(f"   📊 提取 {len(entities)} 个实体, {len(relations)} 个关系")
        return entities, relations

    def close(self) -> Any:
        """关闭数据库连接"""
        if self.pg_conn:
            self.pg_conn.close()
            logger.info("✅ PostgreSQL连接已关闭")

        if self.nebula_pool:
            self.nebula_pool.close()
            logger.info("✅ NebulaGraph连接已关闭")


async def main():
    """主函数"""
    importer = InvalidDecisionImporter()

    try:
        # 初始化数据库连接
        await importer.initialize_databases()

        # 批量处理（测试：先处理100个文件）
        await importer.process_batch(max_files=100)

    finally:
        importer.close()


if __name__ == "__main__":
    asyncio.run(main())
