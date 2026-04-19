#!/usr/bin/env python3
"""
专利规则构建系统 - 法律法规数据处理器
Patent Rules Builder - Legal Data Processor

专门处理专利法、实施细则、司法解释等法律法规文档

作者: Athena平台团队
创建时间: 2025-12-20
版本: v2.0.0
"""

from __future__ import annotations
import asyncio
import json
import logging
import re
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any

# 使用安全哈希函数替代不安全的MD5/SHA1
from production.utils.security_helpers import short_hash

# Word处理
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class LegalDocType(Enum):
    """法律文档类型"""
    PATENT_LAW = "专利法"
    IMPLEMENTATION_RULES = "实施细则"
    JUDICIAL_INTERPRETATION = "司法解释"
    GUIDELINE = "审查指南"
    REGULATION = "行政法规"
    CASE_NOTE = "案例笔记"
    OTHER = "其他"

class LegalDocument:
    """法律文档数据结构"""

    def __init__(self):
        self.doc_id: str = ""
        self.doc_type: LegalDocType = LegalDocType.OTHER
        self.title: str = ""
        self.version: str = ""
        self.promulgation_date: str | None = None
        self.effective_date: str | None = None
        self.issuing_authority: str | None = None
        self.content: str = ""
        self.articles: list[dict] = []
        self.chapters: list[dict] = []
        self.modifications_2025: dict | None = None
        self.related_docs: list[str] = []
        self.keywords: list[str] = []
        self.metadata: dict = {}

class LegalDataProcessor:
    """法律法规数据处理器"""

    def __init__(self):
        # 数据源路径
        self.source_dir = Path("/Users/xujian/学习资料/专利/专利法律法规")
        self.output_dir = Path("/Users/xujian/Athena工作平台/production/data/patent_rules")

        # 输出目录结构
        self.legal_dir = self.output_dir / "legal_documents"
        self.laws_dir = self.legal_dir / "laws"
        self.rules_dir = self.legal_dir / "rules"
        self.interpretations_dir = self.legal_dir / "interpretations"
        self.notes_dir = self.legal_dir / "notes"
        self.reports_dir = self.output_dir / "reports"

        # 创建目录
        for dir_path in [self.output_dir, self.legal_dir, self.laws_dir,
                        self.rules_dir, self.interpretations_dir, self.notes_dir, self.reports_dir]:
            dir_path.mkdir(parents=True, exist_ok=True)

        # 法律文档类型映射
        self.doc_type_mapping = {
            "专利法": LegalDocType.PATENT_LAW,
            "实施细则": LegalDocType.IMPLEMENTATION_RULES,
            "司法解释": LegalDocType.JUDICIAL_INTERPRETATION,
            "审查指南": LegalDocType.GUIDELINE,
            "条例": LegalDocType.REGULATION,
            "办法": LegalDocType.REGULATION,
            "规定": LegalDocType.REGULATION,
            "案例": LegalDocType.CASE_NOTE,
            "笔记": LegalDocType.CASE_NOTE,
            "说明": LegalDocType.OTHER
        }

        # 法条模式
        self.article_patterns = [
            r'第[一二三四五六七八九十\d]+条[：:\s]*(.*?)(?=第[一二三四五六七八九十\d]+条|$)',
            r'Article\s+\d+[：:\s]*(.*?)(?=Article\s+\d+|$)',
            r'\d+\.[：:\s]*(.*?)(?=\d+\.[：:\s]|$)'
        ]

        # 章节模式
        self.chapter_patterns = [
            r'第[一二三四五六七八九十\d]+章[：:\s]*(.*?)(?=第[一二三四五六七八九十\d]+章|$)',
            r'Chapter\s+[IVXLCDM\d]+[：:\s]*(.*?)(?=Chapter\s+[IVXLCDM\d]+|$)',
            r'[一二三四五六七八九十\d]+、[：:\s]*(.*?)(?=[一二三四五六七八九十\d]+、|$)'
        ]

        # 2025年修改关键词
        self.modification_2025_keywords = [
            "2025", "二零二五", "AI", "人工智能", "大数据", "比特流",
            "基因编辑", "CRISPR", "生物技术", "算法", "深度学习",
            "神经网络", "机器学习", "区块链"
        ]

        # 处理统计
        self.stats = {
            "start_time": None,
            "documents_processed": 0,
            "articles_extracted": 0,
            "chapters_extracted": 0,
            "modifications_found": 0,
            "doc_types_count": {},
            "errors": []
        }

    def identify_document_type(self, filename: str, content: str = "") -> LegalDocType:
        """识别文档类型"""
        # 首先从文件名识别
        for keyword, doc_type in self.doc_type_mapping.items():
            if keyword in filename:
                # 特殊处理：如果文件名包含"专利法实施细则"，优先识别为实施细则
                if "实施细则" in filename and doc_type == LegalDocType.PATENT_LAW:
                    return LegalDocType.IMPLEMENTATION_RULES
                return doc_type

        # 从内容识别
        if content:
            content_lower = content.lower()
            if "专利法" in content_lower or "patent law" in content_lower:
                return LegalDocType.PATENT_LAW
            elif "实施细则" in content_lower:
                return LegalDocType.IMPLEMENTATION_RULES
            elif "司法解释" in content_lower or "最高法院" in content_lower:
                return LegalDocType.JUDICIAL_INTERPRETATION
            elif "审查指南" in content_lower:
                return LegalDocType.GUIDELINE
            elif "案例" in content_lower or "案件" in content_lower:
                return LegalDocType.CASE_NOTE

        return LegalDocType.OTHER

    def extract_document_info(self, content: str, doc_type: LegalDocType) -> dict:
        """提取文档基本信息"""
        info = {
            "title": "",
            "version": "",
            "promulgation_date": None,
            "effective_date": None,
            "issuing_authority": None
        }

        # 提取标题
        title_patterns = [
            r'^([^。\n]{1,50})(?:法|规定|办法|条例|解释|指南)',
            r'^【([^】]+)】',
            r'^(.+?)(?:\n|$)'
        ]

        for pattern in title_patterns:
            match = re.search(pattern, content, re.MULTILINE)
            if match:
                info["title"] = match.group(1).strip()
                break

        # 提取日期信息
        date_patterns = [
            r'(\d{4}年\d{1,2}月\d{1,2}日)',
            r'(\d{4}-\d{1,2}-\d{1,2})',
            r'(\d{4}/\d{1,2}/\d{1,2})'
        ]

        dates_found = re.findall(r'|'.join(date_patterns), content)
        if dates_found:
            # 第一个日期通常是发布日期
            info["promulgation_date"] = dates_found[0]
            # 最后一个日期可能是生效日期
            if len(dates_found) > 1:
                info["effective_date"] = dates_found[-1]

        # 提取发布机构
        authority_patterns = [
            r'(全国人民代表大会)',
            r'(全国人大常委会)',
            r'(国务院)',
            r'(最高人民法院)',
            r'(国家知识产权局)',
            r'(专利局)'
        ]

        for pattern in authority_patterns:
            if pattern in content:
                info["issuing_authority"] = pattern.strip('()')
                break

        # 版本识别
        if "2020" in content or "修订" in content:
            info["version"] = "2020修订版"
        elif "2008" in content:
            info["version"] = "2008版"
        else:
            info["version"] = "最新版"

        return info

    def extract_articles(self, content: str) -> list[dict]:
        """提取法条"""
        articles = []

        # 使用所有法条模式进行提取
        for pattern in self.article_patterns:
            matches = re.finditer(pattern, content, re.DOTALL | re.IGNORECASE)

            for match in matches:
                article_text = match.group(0).strip()
                article_content = match.group(1).strip() if len(match.groups()) > 0 else ""

                # 提取法条编号
                article_num_match = re.match(r'(第[一二三四五六七八九十\d]+条|Article\s+\d+|\d+\.)', article_text)
                article_num = article_num_match.group(1) if article_num_match else ""

                # 清理内容
                article_content = re.sub(r'\s+', ' ', article_content).strip()

                if article_num and article_content:
                    article = {
                        "article_id": article_num,
                        "content": article_content,
                        "raw_text": article_text,
                        "word_count": len(article_content),
                        "has_2025_modification": self.check_2025_modification(article_content)
                    }
                    articles.append(article)

        return articles

    def extract_chapters(self, content: str) -> list[dict]:
        """提取章节结构"""
        chapters = []

        # 使用所有章节模式进行提取
        for pattern in self.chapter_patterns:
            matches = re.finditer(pattern, content, re.DOTALL | re.IGNORECASE)

            for match in matches:
                chapter_text = match.group(0).strip()
                chapter_title = match.group(1).strip() if len(match.groups()) > 0 else ""

                # 提取章节编号
                chapter_num_match = re.match(r'(第[一二三四五六七八九十\d]+章|Chapter\s+[IVXLCDM\d]+|[一二三四五六七八九十\d]+、)', chapter_text)
                chapter_num = chapter_num_match.group(1) if chapter_num_match else ""

                # 清理标题
                chapter_title = re.sub(r'\s+', ' ', chapter_title).strip()

                if chapter_num and chapter_title:
                    chapter = {
                        "chapter_id": chapter_num,
                        "title": chapter_title,
                        "raw_text": chapter_text,
                        "has_2025_modification": self.check_2025_modification(chapter_title)
                    }
                    chapters.append(chapter)

        return chapters

    def check_2025_modification(self, text: str) -> bool:
        """检查是否包含2025年修改相关内容"""
        text_lower = text.lower()
        return any(keyword.lower() in text_lower for keyword in self.modification_2025_keywords)

    def find_2025_modifications(self, content: str) -> dict:
        """查找2025年修改内容"""
        modifications = {
            "has_modifications": False,
            "keywords_found": [],
            "relevant_sections": []
        }

        content_lower = content.lower()

        # 查找关键词
        for keyword in self.modification_2025_keywords:
            if keyword.lower() in content_lower:
                modifications["keywords_found"].append(keyword)

        # 如果找到关键词，标记有修改
        if modifications["keywords_found"]:
            modifications["has_modifications"] = True

            # 提取相关段落
            sentences = re.split(r'[。\n]', content)
            for sentence in sentences:
                if any(keyword.lower() in sentence.lower() for keyword in self.modification_2025_keywords):
                    modifications["relevant_sections"].append(sentence.strip())

        return modifications

    async def process_word_document(self, file_path: Path) -> LegalDocument | None:
        """处理Word文档"""
        if not DOCX_AVAILABLE:
            logger.warning(f"python-docx未安装，无法处理Word文档: {file_path}")
            return None

        try:
            doc = Document(file_path)

            # 提取文本
            content_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())

            # 处理表格
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_text.append(" | ".join(row_text))

                if table_text:
                    content_parts.append("[表格]")
                    content_parts.extend(table_text)
                    content_parts.append("[/表格]")

            content = "\n".join(content_parts)

            # 创建文档对象
            legal_doc = LegalDocument()
            legal_doc.doc_id = self.generate_doc_id(file_path)
            legal_doc.doc_type = self.identify_document_type(file_path.name, content)
            legal_doc.content = content

            # 提取文档信息
            doc_info = self.extract_document_info(content, legal_doc.doc_type)
            legal_doc.title = doc_info.get("title", file_path.stem)
            legal_doc.version = doc_info.get("version", "最新版")
            legal_doc.promulgation_date = doc_info.get("promulgation_date")
            legal_doc.effective_date = doc_info.get("effective_date")
            legal_doc.issuing_authority = doc_info.get("issuing_authority")

            # 提取法条和章节
            legal_doc.articles = self.extract_articles(content)
            legal_doc.chapters = self.extract_chapters(content)

            # 检查2025年修改
            legal_doc.modifications_2025 = self.find_2025_modifications(content)

            # 提取关键词
            legal_doc.keywords = self.extract_keywords(content, legal_doc.doc_type)

            # 设置元数据
            legal_doc.metadata = {
                "file_path": str(file_path),
                "file_size": file_path.stat().st_size,
                "processed_at": datetime.now().isoformat(),
                "word_count": len(content),
                "articles_count": len(legal_doc.articles),
                "chapters_count": len(legal_doc.chapters)
            }

            return legal_doc

        except Exception as e:
            logger.error(f"处理Word文档失败 {file_path}: {e}")
            return None

    def generate_doc_id(self, file_path: Path) -> str:
        """生成文档ID"""
        # 使用文件名和哈希值生成唯一ID
        file_hash = short_hash(str(file_path).encode())[:8]
        clean_name = re.sub(r'[^\w]', '_', file_path.stem)[:20]
        return f"{clean_name}_{file_hash}"

    def extract_keywords(self, content: str, doc_type: LegalDocType) -> list[str]:
        """提取关键词"""
        keywords = []

        # 基础法律关键词
        base_keywords = [
            "专利", "发明", "实用新型", "外观设计", "申请", "审查",
            "授权", "无效", "侵权", "权利要求", "说明书",
            "优先权", "新颖性", "创造性", "实用性"
        ]

        # 根据文档类型添加特定关键词
        type_keywords = {
            LegalDocType.PATENT_LAW: ["法律", "条款", "责任", "处罚"],
            LegalDocType.IMPLEMENTATION_RULES: ["规则", "程序", "期限", "费用"],
            LegalDocType.JUDICIAL_INTERPRETATION: ["解释", "适用", "裁判", "案例"],
            LegalDocType.GUIDELINE: ["指南", "标准", "操作", "流程"]
        }

        keywords.extend(base_keywords)
        keywords.extend(type_keywords.get(doc_type, []))

        # 从内容中提取实际出现的关键词
        found_keywords = []
        content_lower = content.lower()

        for keyword in keywords:
            if keyword.lower() in content_lower:
                found_keywords.append(keyword)

        return list(set(found_keywords))  # 去重

    async def process_all_documents(self) -> dict:
        """处理所有法律文档"""
        logger.info("="*80)
        logger.info("🚀 开始处理法律法规文档")
        logger.info("="*80)

        self.stats["start_time"] = datetime.now()

        # 查找所有文档
        all_files = []
        for ext in ["*.docx", "*.doc", "*.txt", "*.rtf"]:
            all_files.extend(self.source_dir.glob(ext))

        logger.info(f"找到 {len(all_files)} 个文档文件")

        # 处理每个文档
        for file_path in all_files:
            logger.info(f"\n处理文档: {file_path.name}")

            try:
                # 处理Word文档
                if file_path.suffix.lower() in ['.docx', '.doc']:
                    legal_doc = await self.process_word_document(file_path)
                else:
                    # 其他格式文档的处理（可以扩展）
                    logger.warning(f"暂不支持的文档格式: {file_path.suffix}")
                    continue

                if legal_doc:
                    # 保存文档
                    await self.save_legal_document(legal_doc)

                    # 更新统计
                    self.stats["documents_processed"] += 1
                    self.stats["articles_extracted"] += len(legal_doc.articles)
                    self.stats["chapters_extracted"] += len(legal_doc.chapters)

                    if legal_doc.modifications_2025["has_modifications"]:
                        self.stats["modifications_found"] += 1

                    # 统计文档类型
                    doc_type_name = legal_doc.doc_type.value
                    self.stats["doc_types_count"][doc_type_name] = \
                        self.stats["doc_types_count"].get(doc_type_name, 0) + 1

                    logger.info("  ✅ 处理完成")
                    logger.info(f"     类型: {doc_type_name}")
                    logger.info(f"     法条: {len(legal_doc.articles)} 条")
                    logger.info(f"     章节: {len(legal_doc.chapters)} 个")
                    if legal_doc.modifications_2025["has_modifications"]:
                        logger.info("     ⚠️ 包含2025年修改内容")

            except Exception as e:
                logger.error(f"  ❌ 处理失败: {e}")
                self.stats["errors"].append(f"{file_path.name}: {str(e)}")

        # 生成处理报告
        await self.generate_report()

        self.stats["end_time"] = datetime.now()

        logger.info("\n✅ 所有法律文档处理完成！")
        self.print_summary()

        return self.stats

    async def save_legal_document(self, legal_doc: LegalDocument):
        """保存法律文档"""
        # 根据文档类型选择保存目录
        output_dir = {
            LegalDocType.PATENT_LAW: self.laws_dir,
            LegalDocType.IMPLEMENTATION_RULES: self.rules_dir,
            LegalDocType.JUDICIAL_INTERPRETATION: self.interpretations_dir,
            LegalDocType.CASE_NOTE: self.notes_dir,
            LegalDocType.GUIDELINE: self.legal_dir,
            LegalDocType.REGULATION: self.rules_dir,
            LegalDocType.OTHER: self.legal_dir
        }.get(legal_doc.doc_type, self.legal_dir)

        # 转换为字典格式
        doc_dict = {
            "doc_id": legal_doc.doc_id,
            "doc_type": legal_doc.doc_type.value,
            "title": legal_doc.title,
            "version": legal_doc.version,
            "promulgation_date": legal_doc.promulgation_date,
            "effective_date": legal_doc.effective_date,
            "issuing_authority": legal_doc.issuing_authority,
            "content": legal_doc.content,
            "articles": legal_doc.articles,
            "chapters": legal_doc.chapters,
            "modifications_2025": legal_doc.modifications_2025,
            "related_docs": legal_doc.related_docs,
            "keywords": legal_doc.keywords,
            "metadata": legal_doc.metadata
        }

        # 保存完整文档
        doc_file = output_dir / f"{legal_doc.doc_id}.json"
        with open(doc_file, 'w', encoding='utf-8') as f:
            json.dump(doc_dict, f, ensure_ascii=False, indent=2)

        # 保存法条列表（单独文件，便于检索）
        if legal_doc.articles:
            articles_file = output_dir / f"{legal_doc.doc_id}_articles.json"
            with open(articles_file, 'w', encoding='utf-8') as f:
                json.dump(legal_doc.articles, f, ensure_ascii=False, indent=2)

        # 保存章节列表
        if legal_doc.chapters:
            chapters_file = output_dir / f"{legal_doc.doc_id}_chapters.json"
            with open(chapters_file, 'w', encoding='utf-8') as f:
                json.dump(legal_doc.chapters, f, ensure_ascii=False, indent=2)

    async def generate_report(self):
        """生成处理报告"""
        report = {
            "processing_summary": {
                "start_time": self.stats["start_time"].isoformat(),
                "end_time": self.stats.get("end_time", datetime.now()).isoformat(),
                "duration": (
                    self.stats.get("end_time", datetime.now()) - self.stats["start_time"]
                ).total_seconds()
            },
            "statistics": {
                "documents_processed": self.stats["documents_processed"],
                "articles_extracted": self.stats["articles_extracted"],
                "chapters_extracted": self.stats["chapters_extracted"],
                "modifications_found": self.stats["modifications_found"],
                "doc_types_count": self.stats["doc_types_count"],
                "error_count": len(self.stats["errors"])
            },
            "output_structure": {
                "legal_documents": str(self.legal_dir),
                "laws": str(self.laws_dir),
                "rules": str(self.rules_dir),
                "interpretations": str(self.interpretations_dir),
                "notes": str(self.notes_dir)
            }
        }

        # 保存JSON报告
        report_file = self.reports_dir / f"legal_processing_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(report_file, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)

        # 生成可读报告
        readable_file = self.reports_dir / f"legal_processing_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        with open(readable_file, 'w', encoding='utf-8') as f:
            f.write("法律法规数据处理报告\n")
            f.write("="*50 + "\n\n")
            f.write(f"处理时间: {report['processing_summary']['start_time']}\n")
            f.write(f"完成时间: {report['processing_summary']['end_time']}\n")
            f.write(f"处理耗时: {report['processing_summary']['duration']:.2f} 秒\n\n")

            f.write("处理统计:\n")
            stats = report['statistics']
            f.write(f"  处理文档数: {stats['documents_processed']}\n")
            f.write(f"  提取法条数: {stats['articles_extracted']}\n")
            f.write(f"  提取章节数: {stats['chapters_extracted']}\n")
            f.write(f"  2025年修改文档: {stats['modifications_found']}\n\n")

            f.write("文档类型分布:\n")
            for doc_type, count in stats['doc_types_count'].items():
                f.write(f"  {doc_type}: {count} 个\n")

            if self.stats["errors"]:
                f.write("\n错误列表:\n")
                for error in self.stats["errors"]:
                    f.write(f"  - {error}\n")

        logger.info(f"  📋 处理报告: {readable_file}")

    def print_summary(self) -> Any:
        """打印处理摘要"""
        logger.info("\n📊 处理摘要:")
        logger.info(f"  处理文档: {self.stats['documents_processed']} 个")
        logger.info(f"  提取法条: {self.stats['articles_extracted']} 条")
        logger.info(f"  提取章节: {self.stats['chapters_extracted']} 个")
        logger.info(f"  2025年修改: {self.stats['modifications_found']} 个文档")

        if self.stats.get("start_time") and self.stats.get("end_time"):
            duration = (self.stats["end_time"] - self.stats["start_time"]).total_seconds()
            logger.info(f"  处理耗时: {duration:.2f} 秒")

        if self.stats["doc_types_count"]:
            logger.info("\n📁 文档类型分布:")
            for doc_type, count in self.stats["doc_types_count"].items():
                logger.info(f"  {doc_type}: {count} 个")

# 使用示例
async def main():
    """主函数示例"""
    processor = LegalDataProcessor()
    await processor.process_all_documents()

if __name__ == "__main__":
    asyncio.run(main())
