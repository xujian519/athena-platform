#!/usr/bin/env python3
"""
专利规则构建系统 - 数据处理器
Patent Rules Builder - Data Processor

处理PDF、Word等多模态文档，特别支持2025年修改集成

作者: Athena平台团队
创建时间: 2025-12-20
版本: v2.0.0
"""

from __future__ import annotations
import asyncio
import io
import json
import logging
import os
import re
from concurrent.futures import ThreadPoolExecutor
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

# OCR处理
import cv2
import fitz  # PyMuPDF
import numpy as np

# PDF处理
import pdfplumber
import pytesseract

# 多模态模型
import torch

# Word处理
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter
from transformers import AutoImageProcessor, AutoModel

# 使用安全哈希函数替代不安全的MD5/SHA1
from production.utils.security_helpers import short_hash

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """文档元数据"""
    title: str
    version: str
    source_type: str
    file_type: str
    file_size: int
    hash_md5: str
    created_date: str | None = None
    modified_date: str | None = None
    page_count: int = 0
    sections_count: int = 0
    has_images: bool = False
    modification_2025: bool = False

@dataclass
class SectionData:
    """章节数据"""
    section_id: str
    level: int
    title: str
    content: str
    parent_id: str | None
    full_path: str
    page_range: str
    modification_2025: dict | None = None
    chunk_type: str = "section"  # parent/child
    entities: list[dict] = None
    images: list[dict] = None

@dataclass
class ImageData:
    """图像数据"""
    image_id: str
    page_number: int
    bbox: list[float]  # [x1, y1, x2, y2]
    text_content: str
    ocr_confidence: float
    image_type: str  # cover/toc/content/table/figure
    processed_by: str  # pytesseract/easyocr/multimodal
    image_format: str  # png/jpg
    image_size: tuple[int, int]  # (width, height)
    preprocessing_applied: list[str]  # 应用的预处理
    extraction_method: str  # 提取方法
    quality_score: float  # 图像质量评分

@dataclass
class Modification2025:
    """2025年修改信息"""
    section_id: str
    change_type: str  # added/modified/deleted
    old_content: str
    new_content: str
    reason: str
    application_date: str = "2026-01-01"
    priority: int = 1  # 1=高, 2=中, 3=低

class ImageProcessor:
    """增强的图像处理器"""

    def __init__(self):
        self.preprocessing_steps = {
            'denoise': True,
            'enhance_contrast': True,
            'sharpen': True,
            'binarize': False,
            'deskew': True
        }

    def preprocess_image(self, image: Image.Image) -> tuple[Image.Image, list[str]]:
        """预处理图像以提高OCR识别率"""
        applied_steps = []

        # 1. 降噪
        if self.preprocessing_steps['denoise']:
            image = image.filter(ImageFilter.MedianFilter(size=3))
            applied_steps.append('denoise')

        # 2. 增强对比度
        if self.preprocessing_steps['enhance_contrast']:
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)
            applied_steps.append('enhance_contrast')

        # 3. 锐化
        if self.preprocessing_steps['sharpen']:
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(1.5)
            applied_steps.append('sharpen')

        # 4. 二值化（对扫描文档效果好）
        if self.preprocessing_steps['binarize']:
            gray = image.convert('L')
            # 使用自适应阈值
            img_array = np.array(gray)
            _, binary = cv2.threshold(img_array, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            image = Image.fromarray(binary)
            applied_steps.append('binarize')

        # 5. 倾斜校正
        if self.preprocessing_steps['deskew']:
            image = self._deskew_image(image)
            applied_steps.append('deskew')

        return image, applied_steps

    def _deskew_image(self, image: Image.Image) -> Image.Image:
        """校正图像倾斜"""
        try:
            # 转换为灰度图
            gray = cv2.cvt_color(np.array(image), cv2.COLOR_RGB2GRAY)

            # 边缘检测
            edges = cv2.Canny(gray, 50, 150, aperture_size=3)

            # 霍夫变换检测直线
            lines = cv2.HoughLinesP(edges, 1, np.pi/180, 100, min_line_length=100, max_line_gap=10)

            if lines is not None:
                # 计算倾斜角度
                angles = []
                for line in lines:
                    x1, y1, x2, y2 = line[0]
                    angle = np.arctan2(y2 - y1, x2 - x1) * 180 / np.pi
                    if abs(angle) < 45:  # 只考虑小角度
                        angles.append(angle)

                if angles:
                    # 使用中位数角度
                    median_angle = np.median(angles)

                    # 旋转图像
                    if abs(median_angle) > 0.5:  # 只校正明显的倾斜
                        image = image.rotate(-median_angle, expand=True, fillcolor='white')

            return image
        except Exception as e:
            logger.debug(f"倾斜校正失败: {e}")
            return image

    def calculate_quality_score(self, image: Image.Image) -> float:
        """计算图像质量评分"""
        try:
            # 转换为灰度图
            gray = image.convert('L')
            img_array = np.array(gray)

            # 1. 计算清晰度（拉普拉斯方差）
            laplacian_var = cv2.Laplacian(img_array, cv2.CV_64F).var()

            # 2. 计算对比度
            contrast = img_array.std()

            # 3. 计算亮度分布
            brightness = img_array.mean()

            # 综合评分（0-100）
            clarity_score = min(laplacian_var / 100, 100)  # 清晰度
            contrast_score = min(contrast / 127 * 100, 100)  # 对比度
            brightness_score = 100 - abs(brightness - 127) / 127 * 100  # 亮度

            quality_score = (clarity_score * 0.5 + contrast_score * 0.3 + brightness_score * 0.2)

            return min(quality_score, 100)
        except Exception as e:
            logger.debug(f"质量评分失败: {e}")
            return 50.0  # 默认分数

    def classify_image_type(self, image: Image.Image, text_content: str = "") -> str:
        """分类图像类型"""
        # 简单的启发式分类
        width, height = image.size
        aspect_ratio = width / height

        # 如果有大量文字且结构化，可能是表格
        if text_content:
            lines = text_content.split('\n')
            non_empty_lines = [l for l in lines if l.strip()]

            # 检查是否有表格特征
            has_table_pattern = any('|' in line or '\t' in line for line in non_empty_lines)
            if has_table_pattern and len(non_empty_lines) > 3:
                return "table"

            # 如果有很多短行，可能是目录
            if len(non_empty_lines) > 10:
                avg_line_length = sum(len(l.strip()) for l in non_empty_lines) / len(non_empty_lines)
                if avg_line_length < 50:
                    return "toc"

        # 根据宽高比判断
        if aspect_ratio < 0.8 or aspect_ratio > 1.2:
            # 可能是封面或特殊图形
            if width > 800 and height > 1000:
                return "cover"

        # 默认为内容图像
        return "content"

class AdvancedOCR:
    """高级OCR处理器"""

    def __init__(self):
        self.processors = {
            'easyocr': None,
            'tesseract': True,
            'multimodal': None
        }
        self.initialize_processors()

    def initialize_processors(self) -> Any:
        """初始化OCR处理器"""
        # 初始化EasyOCR
        try:
            import easyocr
            self.processors['easyocr'] = easyocr.Reader(['ch_en', 'en'], gpu=torch.cuda.is_available())
            logger.info("✅ EasyOCR初始化成功")
        except Exception as e:
            logger.warning(f"⚠️ EasyOCR初始化失败: {e}")
            self.processors['easyocr'] = None

        # 检查Tesseract
        try:
            pytesseract.pytesseract.tesseract_cmd = '/usr/local/bin/tesseract' if os.path.exists('/usr/local/bin/tesseract') else 'tesseract'
            pytesseract.get_tesseract_version()
            logger.info("✅ Tesseract可用")
        except Exception as e:
            logger.warning(f"⚠️ Tesseract不可用: {e}")
            self.processors['tesseract'] = False

        # 初始化多模态模型（可选）
        try:
            self.processors['multimodal'] = {
                'processor': AutoImageProcessor.from_pretrained("openai/clip-vit-base-patch32"),
                'model': AutoModel.from_pretrained("openai/clip-vit-base-patch32")
            }
            logger.info("✅ 多模态模型初始化成功")
        except Exception as e:
            logger.warning(f"⚠️ 多模态模型初始化失败: {e}")
            self.processors['multimodal'] = None

    async def extract_text(self, image: Image.Image, preferred_engine: str = "auto") -> dict[str, Any]:
        """提取图像文字，支持多种OCR引擎"""
        results = {
            'text': '',
            'confidence': 0.0,
            'engine_used': '',
            'processing_time': 0,
            'alternatives': []
        }

        start_time = datetime.now()

        # 决定使用哪个引擎
        engines_to_try = []
        if preferred_engine == "auto":
            # 自动选择：优先EasyOCR，备选Tesseract
            if self.processors['easyocr']:
                engines_to_try.append('easyocr')
            if self.processors['tesseract']:
                engines_to_try.append('tesseract')
        else:
            engines_to_try = [preferred_engine]

        # 尝试每个引擎
        for engine in engines_to_try:
            try:
                if engine == 'easyocr' and self.processors['easyocr']:
                    result = self._ocr_with_easyocr(image)
                elif engine == 'tesseract' and self.processors['tesseract']:
                    result = self._ocr_with_tesseract(image)
                else:
                    continue

                # 如果结果可信，使用它
                if result['confidence'] > 0.6 or result['text'].strip():
                    results['text'] = result['text']
                    results['confidence'] = result['confidence']
                    results['engine_used'] = engine
                    results['alternatives'].append(result)
                    break
                else:
                    results['alternatives'].append(result)

            except Exception as e:
                logger.debug(f"{engine} OCR失败: {e}")

        # 如果所有引擎都失败了，尝试使用多模态模型（如果有）
        if not results['text'] and self.processors['multimodal']:
            try:
                result = await self._ocr_with_multimodal(image)
                if result['text']:
                    results['text'] = result['text']
                    results['confidence'] = result['confidence']
                    results['engine_used'] = 'multimodal'
            except Exception as e:
                logger.debug(f"多模态OCR失败: {e}")

        # 计算处理时间
        results['processing_time'] = (datetime.now() - start_time).total_seconds()

        return results

    def _ocr_with_easyocr(self, image: Image.Image) -> dict[str, Any]:
        """使用EasyOCR提取文字"""
        # 转换PIL图像为numpy数组
        img_array = np.array(image)

        # 执行OCR
        results = self.processors['easyocr'].readtext(img_array, detail=1, paragraph=True)

        if results:
            # 合并所有文本
            texts = [res[1] for res in results]
            confidences = [res[2] for res in results]

            return {
                'text': '\n'.join(texts),
                'confidence': sum(confidences) / len(confidences) if confidences else 0,
                'details': results
            }

        return {'text': '', 'confidence': 0, 'details': []}

    def _ocr_with_tesseract(self, image: Image.Image) -> dict[str, Any]:
        """使用Tesseract提取文字"""
        # 配置Tesseract参数
        custom_config = r'--oem 3 --psm 6 -l chi_sim+eng'

        # 执行OCR
        text = pytesseract.image_to_string(image, config=custom_config)

        # 获取置信度
        try:
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT, config=custom_config)
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 50
        except Exception as e:
            logger.debug(f"空except块已触发: {e}")
            avg_confidence = 50  # 默认置信度

        return {
            'text': text,
            'confidence': avg_confidence / 100,  # 转换为0-1范围
            'details': []
        }

    async def _ocr_with_multimodal(self, image: Image.Image) -> dict[str, Any]:
        """使用多模态模型提取文字（实验性）"""
        # 这是一个实验性功能，实际效果可能有限
        # 这里只是框架，实际实现需要专门的视觉语言模型

        # 预处理图像
        processor = self.processors['multimodal']['processor']
        model = self.processors['multimodal']['model']

        # 这里应该调用专门的OCR或视觉语言模型
        # 暂时返回空结果
        return {
            'text': '',
            'confidence': 0,
            'details': []
        }

class PatentRulesDataProcessor:
    """专利规则数据处理器"""

    def __init__(self):
        # 数据源路径
        self.source_dir = Path("/Users/xujian/学习资料/专利/专利法律法规")
        self.output_dir = Path("/Users/xujian/Athena工作平台/production/data/patent_rules")

        # 输出目录结构
        self.raw_dir = self.output_dir / "raw"
        self.processed_dir = self.output_dir / "processed"
        self.vectors_dir = self.output_dir / "vectors"
        self.images_dir = self.output_dir / "images"
        self.reports_dir = self.output_dir / "reports"

        # 创建目录
        for dir_path in [self.output_dir, self.raw_dir, self.processed_dir,
                           self.vectors_dir, self.images_dir, self.reports_dir]:
            dir_path.mkdir(parents=True, exist_ok=True)

        # 2025年修改对照表（模拟数据，实际应从官网获取）
        self.modification_2025 = self._load_modification_2025()

        # 初始化增强的图像处理器
        self.image_processor = ImageProcessor()

        # 初始化高级OCR处理器
        self.advanced_ocr = AdvancedOCR()

        # 线程池用于并行处理
        self.executor = ThreadPoolExecutor(max_workers=4)

        # 处理统计
        self.stats = {
            "start_time": None,
            "documents_processed": 0,
            "sections_extracted": 0,
            "images_extracted": 0,
            "modifications_applied": 0,
            "ocr_processing_time": 0,
            "image_preprocessing_time": 0,
            "errors": []
        }


    def _load_modification_2025(self) -> dict:
        """加载2025年修改对照表"""
        # 这里使用模拟数据，实际应从官方获取
        return {
            "AI相关新增": {
                "新增章节": [
                    "第二部分第十章 特殊情形审查",
                    "第五部分 第六章 人工智能相关申请"
                ],
                "修改内容": {
                    "创造性判断": {
                        "2023": "整体性考虑",
                        "2025": "特别考虑AI和大数据领域的技术进步"
                    },
                    "现有技术": {
                        "2023": "公知技术标准",
                        "2025": "包括网络公开的技术内容"
                    }
                }
            },
            "比特流相关": {
                "新增章节": [
                    "第二部分第五章 计算机程序相关发明"
                ],
                "重要修改": {
                    "保护范围": {
                        "2023": "软件著作权",
                        "2025": "单纯比特流不保护，保护技术方案"
                    }
                }
            },
            "生物技术": {
                "新增内容": [
                    "基因编辑技术的审查指南"
                ],
                "应用案例": [
                    "CRISPR-Cas9专利申请案例",
                    "生物序列专利保护范围"
                ]
            }
        }

    async def process_all_documents(self) -> dict:
        """处理所有文档"""
        logger.info("="*80)
        logger.info("🚀 开始处理专利规则文档")
        logger.info("="*80)

        self.stats["start_time"] = datetime.now()

        # 1. 处理PDF文档
        await self._process_pdf_documents()

        # 2. 处理Word文档
        await self._process_word_documents()

        # 3. 应用2025年修改
        await self._apply_modifications_2025()

        # 4. 数据质量检查
        await self._quality_check()

        # 5. 生成处理报告
        await self._generate_report()

        self.stats["end_time"] = datetime.now()

        logger.info("\n✅ 所有文档处理完成！")
        self._print_summary()

        return self.stats

    async def _process_pdf_documents(self):
        """处理PDF文档"""
        logger.info("\n📚 处理PDF文档...")

        pdf_files = list(self.source_dir.glob("*.pdf"))
        logger.info(f"  找到 {len(pdf_files)} 个PDF文件")

        for pdf_file in pdf_files:
            if pdf_file.name.startswith("专利审查指南"):
                await self._process_patent_guideline_pdf(pdf_file)
            else:
                await self._process_general_pdf(pdf_file)

    async def _process_patent_guideline_pdf(self, pdf_file: Path):
        """处理专利审查指南PDF（特殊处理）"""
        logger.info(f"  处理专利指南: {pdf_file.name}")

        try:
            # 提取PDF基本信息
            metadata = self._extract_pdf_metadata(pdf_file)

            # 提取文本内容
            content, sections = self._extract_guideline_content(pdf_file)

            # 提取图像
            images = await self._extract_pdf_images(pdf_file)

            # 保存原始数据
            doc_data = {
                "metadata": asdict(metadata),
                "content": content,
                "sections": [asdict(s) for s in sections],
                "images": [asdict(img) for img in images],
                "source_file": str(pdf_file)
            }

            # 保存原始数据
            raw_file = self.raw_dir / f"{pdf_file.stem}_raw.json"
            with open(raw_file, 'w', encoding='utf-8') as f:
                json.dump(doc_data, f, ensure_ascii=False, indent=2)

            # 保存处理后的数据
            processed_data = {
                "metadata": asdict(metadata),
                "sections": [asdict(s) for s in sections],
                "entities": [],  # 稍后填充
                "relations": [],  # 稍后填充
                "statistics": {
                    "sections_count": len(sections),
                    "images_count": len(images),
                    "total_chars": len(content)
                }
            }

            processed_file = self.processed_dir / f"{pdf_file.stem}_processed.json"
            with open(processed_file, 'w', encoding='utf-8') as f:
                json.dump(processed_data, f, ensure_ascii=False, indent=2)

            # 保存图像
            if images:
                self._save_images(images, pdf_file.stem)

            # 更新统计
            self.stats["documents_processed"] += 1
            self.stats["sections_extracted"] += len(sections)
            self.stats["images_extracted"] += len(images)

            logger.info(f"    章节: {len(sections)}, 图像: {len(images)}")

        except Exception as e:
            logger.error(f"    处理失败: {e}")
            self.stats["errors"].append(f"PDF处理失败: {pdf_file.name}")

    async def _process_word_documents(self):
        """处理Word文档"""
        logger.info("\n📄 处理Word文档...")

        word_files = list(self.source_dir.glob("*.doc*"))
        logger.info(f"  找到 {len(word_files)} 个Word文档")

        for word_file in word_files:
            await self._process_word_document(word_file)

    async def _process_word_document(self, word_file: Path):
        """处理单个Word文档"""
        logger.info(f"  处理Word文档: {word_file.name}")

        try:
            metadata = self._extract_word_metadata(word_file)
            content = self._extract_word_content(word_file)

            # 简单章节提取
            sections = self._extract_word_sections(content, word_file.name)

            # 检查是否包含2025年修改相关内容
            modification_2025 = self._check_modification_2025(content, word_file.name)

            doc_data = {
                "metadata": asdict(metadata),
                "content": content,
                "sections": [asdict(s) for s in sections],
                "modification_2025": modification_2025,
                "statistics": {
                    "sections_count": len(sections),
                    "total_chars": len(content)
                }
            }

            # 保存数据
            processed_file = self.processed_dir / f"{word_file.stem}_processed.json"
            with open(processed_file, 'w', encoding='utf-8') as f:
                json.dump(doc_data, f, ensure_ascii=False, indent=2)

            self.stats["documents_processed"] += 1
            self.stats["sections_extracted"] += len(sections)

            if modification_2025:
                self.stats["modifications_applied"] += 1
                logger.info("    发现2025年修改内容")

            logger.info(f"    章节: {len(sections)}")

        except Exception as e:
            logger.error(f"    处理失败: {e}")
            self.stats["errors"].append(f"Word处理失败: {word_file.name}")

    def _extract_pdf_metadata(self, pdf_file: Path) -> DocumentMetadata:
        """提取PDF元数据"""
        metadata = DocumentMetadata(
            title=pdf_file.stem,
            version="2023",  # 默认版本，后续会更新
            source_type="PDF",
            file_type="pdf",
            file_size=pdf_file.stat().st_size
        )

        try:
            # 使用PyMuPDF提取元数据
            doc = fitz.open(str(pdf_file))
            metadata.page_count = doc.page_count

            # 尝试提取PDF信息
            if doc.metadata:
                info = doc.metadata
                metadata.title = info.get('title', metadata.title)
                metadata.created_date = str(info.get('creation_date', ''))
                metadata.modified_date = str(info.get('mod_date', ''))

            doc.close()

        except Exception as e:
            logger.warning(f"PDF元数据提取失败: {e}")

        # 计算MD5
        with open(pdf_file, 'rb') as f:
            metadata.hash_md5 = short_hash(f.read())

        return metadata

    def _extract_word_metadata(self, word_file: Path) -> DocumentMetadata:
        """提取Word元数据"""
        metadata = DocumentMetadata(
            title=word_file.stem,
            version="2023",  # 根据文件名推断
            source_type="DOCX",
            file_type="docx",
            file_size=word_file.stat().st_size
        )

        try:
            doc = Document(word_file)

            metadata.title = doc.core_properties.title or metadata.title
            metadata.created_date = str(doc.core_properties.created) if doc.core_properties.created else None
            metadata.modified_date = str(doc.core_properties.modified) if doc.core_properties.modified else None

            doc.close()

        except Exception as e:
            logger.warning(f"Word元数据提取失败: {e}")

        # 计算MD5
        with open(word_file, 'rb') as f:
            metadata.hash_md5 = short_hash(f.read())

        return metadata

    def _extract_guideline_content(self, pdf_file: Path) -> tuple[str, list[SectionData]]:
        """提取专利指南内容（特殊结构化处理）"""
        content = ""
        sections = []

        try:
            # 使用pdfplumber提取文本（保留格式）
            with pdfplumber.open(pdf_file) as pdf:
                current_section = None
                current_content = ""

                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()

                    # 处理表格
                    tables = page.extract_tables()
                    for table in tables:
                        table_text = self._table_to_text(table)
                        page_text += f"\n[表格]\n{table_text}\n[/表格]\n"

                    # 章节识别
                    lines = page_text.split('\n')

                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue

                        # 识别章节标题
                        section_info = self._identify_guideline_section(line)

                        if section_info:
                            # 保存前一个章节
                            if current_section:
                                current_section.content = current_section.content.strip()
                                sections.append(current_section)

                            # 创建新章节
                            current_section = SectionData(
                                section_id=section_info['section_id'],
                                level=section_info['level'],
                                title=line,
                                content="",
                                parent_id=section_info.get('parent_id'),
                                full_path=section_info.get('full_path', line),
                                page_range=f"第{page_num}页",
                                entities=[],
                                images=[]
                            )

                        elif current_section:
                            # 添加到当前章节内容
                            if current_section.content:
                                current_section.content += "\n" + line
                            else:
                                current_section.content = line

                    content += page_text + "\n\n"

                # 保存最后一个章节
                if current_section:
                    current_section.content = current_section.content.strip()
                    sections.append(current_section)

        except Exception as e:
            logger.error(f"指南内容提取失败: {e}")

        return content, sections

    def _extract_word_content(self, word_file: Path) -> str:
        """提取Word内容"""
        try:
            doc = Document(word_file)

            content_parts = []

            # 处理段落
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)

            # 处理表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                table_text = self._table_data_to_text(table_data)
                content_parts.append(f"\n[表格]\n{table_text}\n[/表格]\n")

            doc.close()

            return "\n".join(content_parts)

        except Exception as e:
            logger.error(f"Word内容提取失败: {e}")
            return ""

    def _extract_word_sections(self, content: str, file_name: str) -> list[SectionData]:
        """从Word内容提取章节（简单处理）"""
        sections = []

        # 简单的章节识别
        patterns = [
            r'^第[一二三四五六七八九十\d]+章',
            r'^第[一二三四五六七八九十\d]+节',
            r'^\d+\.\d+',
            r'^附件\d+'
        ]

        lines = content.split('\n')
        current_section = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 检查是否是章节标题
            for pattern in patterns:
                if re.match(pattern, line):
                    # 保存前一个章节
                    if current_section:
                        current_section.content = current_section.content.strip()
                        sections.append(current_section)

                    # 创建新章节
                    section_id = f"{file_name}_{'_'.join(pattern.split()[:2])}"
                    level = len(pattern.split())

                    current_section = SectionData(
                        section_id=section_id,
                        level=level,
                        title=line,
                        content="",
                        parent_id=None,
                        full_path=line,
                        page_range="",
                        entities=[],
                        images=[]
                    )
                    break

        # 保存最后一个章节
        if current_section:
            current_section.content = current_section.content.strip()
            sections.append(current_section)

        return sections

    def _identify_guideline_section(self, line: str) -> dict | None:
        """识别指南章节"""
        # 部分标题
        part_match = re.match(r'^(第[一二三四五六七八九十]+部分)\s*(.*)', line)
        if part_match:
            return {
                'section_id': f"P{self._chinese_to_number(part_match.group(1)[1:-2])}",
                'level': 1,
                'title': line,
                'full_path': line
            }

        # 章标题
        chapter_match = re.match(r'^(第[一二三四五六七八九十\d]+章)\s*(.*)', line)
        if chapter_match:
            chapter_num = self._extract_number(chapter_match.group(1)[1:-1])
            return {
                'section_id': f"C{chapter_num}",
                'level': 2,
                'title': line,
                'full_path': line
            }

        # 节标题
        section_match = re.match(r'^(\d+\.\d+)\s*(.*)', line)
        if section_match:
            return {
                'section_id': f"S{section_match.group(1)}",
                'level': 3,
                'title': line,
                'full_path': line
            }

        # 小节标题
        subsection_match = re.match(r'^(\d+\.\d+\.\d+)\s*(.*)', line)
        if subsection_match:
            return {
                'section_id': f"SS{subsection_match.group(1)}",
                'level': 4,
                'title': line,
                'full_path': line
            }

        return None

    async def _extract_pdf_images(self, pdf_file: Path) -> list[ImageData]:
        """提取PDF图像（增强版）"""
        images = []

        try:
            doc = fitz.open(str(pdf_file))
            logger.info(f"  开始提取 {pdf_file.name} 的图像...")

            # 处理每一页
            for page_num, page in enumerate(doc.pages):
                # 方法1: 提取嵌入的图像
                image_list = page.get_images(full=True)

                for img_index, img_info in enumerate(image_list):
                    try:
                        # 获取图像数据
                        xref = img_info[0]
                        pix = fitz.Pixmap(doc, xref)

                        # 跳过CMYK图像
                        if pix.n - pix.alpha < 4:
                            # 转换为PIL图像
                            img_data = pix.tobytes("png")
                            pil_image = Image.open(io.BytesIO(img_data))

                            # 图像预处理
                            preprocess_start = datetime.now()
                            processed_image, preprocess_steps = self.image_processor.preprocess_image(pil_image)
                            preprocess_time = (datetime.now() - preprocess_start).total_seconds()
                            self.stats["image_preprocessing_time"] += preprocess_time

                            # 获取图像位置（如果可用）
                            rect = img_info[1] if len(img_info) > 1 else fitz.Rect(0, 0, page.rect.width, page.rect.height)
                            bbox = [rect.x0, rect.y0, rect.x1, rect.y1]

                            # OCR识别
                            ocr_result = await self.advanced_ocr.extract_text(processed_image)
                            self.stats["ocr_processing_time"] += ocr_result['processing_time']

                            # 分类图像类型
                            image_type = self.image_processor.classify_image_type(processed_image, ocr_result['text'])

                            # 计算图像质量
                            quality_score = self.image_processor.calculate_quality_score(processed_image)

                            # 创建图像数据
                            image = ImageData(
                                image_id=f"{pdf_file.stem}_p{page_num}_i{img_index}",
                                page_number=page_num,
                                bbox=bbox,
                                text_content=ocr_result['text'],
                                ocr_confidence=ocr_result['confidence'],
                                image_type=image_type,
                                processed_by=ocr_result['engine_used'],
                                image_format="png",
                                image_size=processed_image.size,
                                preprocessing_applied=preprocess_steps,
                                extraction_method="embedded",
                                quality_score=quality_score
                            )
                            images.append(image)

                        pix = None  # 释放内存

                    except Exception as e:
                        logger.debug(f"处理嵌入图像失败: {e}")

                # 方法2: 处理整个页面作为图像（用于扫描文档）
                if page_num == 0 or len(images) == 0:  # 第一页或没有找到图像时
                    try:
                        # 将页面渲染为图像
                        mat = fitz.Matrix(2.0, 2.0)  # 2倍分辨率
                        pix = page.get_pixmap(matrix=mat)
                        img_data = pix.tobytes("png")
                        pil_image = Image.open(io.BytesIO(img_data))

                        # 预处理
                        preprocess_start = datetime.now()
                        processed_image, preprocess_steps = self.image_processor.preprocess_image(pil_image)
                        preprocess_time = (datetime.now() - preprocess_start).total_seconds()
                        self.stats["image_preprocessing_time"] += preprocess_time

                        # OCR
                        ocr_result = await self.advanced_ocr.extract_text(processed_image)
                        self.stats["ocr_processing_time"] += ocr_result['processing_time']

                        # 只处理有内容的页面图像
                        if ocr_result['text'] and len(ocr_result['text'].strip()) > 50:
                            image_type = self.image_processor.classify_image_type(processed_image, ocr_result['text'])
                            quality_score = self.image_processor.calculate_quality_score(processed_image)

                            image = ImageData(
                                image_id=f"{pdf_file.stem}_p{page_num}_full",
                                page_number=page_num,
                                bbox=[0, 0, page.rect.width, page.rect.height],
                                text_content=ocr_result['text'],
                                ocr_confidence=ocr_result['confidence'],
                                image_type=image_type,
                                processed_by=ocr_result['engine_used'],
                                image_format="png",
                                image_size=processed_image.size,
                                preprocessing_applied=preprocess_steps,
                                extraction_method="page_render",
                                quality_score=quality_score
                            )
                            images.append(image)

                    except Exception as e:
                        logger.debug(f"处理页面图像失败: {e}")

            doc.close()
            logger.info(f"  提取完成: {len(images)} 个图像")

        except Exception as e:
            logger.warning(f"图像提取失败 {pdf_file.name}: {e}")
            self.stats["errors"].append(f"图像提取失败: {pdf_file.name}")

        return images


    def _save_images(self, images: list[ImageData], doc_name: str) -> Any:
        """保存提取的图像"""
        doc_images_dir = self.images_dir / doc_name
        doc_images_dir.mkdir(parents=True, exist_ok=True)

        for image in images:
            image_file = doc_images_dir / f"{image['image_id']}.jpg"

            # 这里应该保存实际的图像文件
            # 实际实现时需要从PDF提取图像数据
            logger.debug(f"保存图像: {image_file}")

    def _table_to_text(self, table: list[list[str]]) -> str:
        """将表格转换为文本"""
        rows = []
        for row in table:
            if row and any(cell.strip() for cell in row):
                clean_row = [cell.strip() if cell else "" for cell in row]
                rows.append(" | ".join(clean_row))
        return "\n".join(rows)

    def _table_data_to_text(self, table_data: list[list[str]]) -> str:
        """将表格数据转换为文本"""
        return self._table_to_text(table_data)

    def _check_modification_2025(self, content: str, doc_name: str) -> dict | None:
        """检查是否包含2025年修改相关内容"""
        modification = None

        # 检查2025相关关键词
        keywords_2025 = [
            "2025", "二零二五", "AI", "人工智能", "大数据", "比特流",
            "基因编辑", "CRISPR", "生物技术"
        ]

        if any(keyword.lower() in content.lower() for keyword in keywords_2025):
            # 检查修改对照表
            for category, changes in self.modification_2025.items():
                for change_type, change_data in changes.items():
                    if isinstance(change_data, dict):
                        # 检查旧内容
                        if isinstance(change_data.get('old_content'), str):
                            if change_data['old_content'] in content:
                                modification = {
                                    "type": change_type,
                                    "category": category,
                                    "detected_keywords": [kw for kw in keywords_2025 if kw.lower() in content.lower()],
                                    "applied": False
                                }
                                break
                    else:
                        # 检查新增内容关键词
                        if isinstance(change_data, list):
                            for item in change_data:
                                if isinstance(item, str) and item in content:
                                    modification = {
                                        "type": "new_content",
                                        "category": category,
                                        "detected_keywords": [item],
                                        "content": item,
                                        "applied": False
                                    }
                                    break

        return modification

    async def _apply_modifications_2025(self):
        """应用2025年修改"""
        logger.info("\n🔄 应用2025年修改...")

        processed_files = list(self.processed_dir.glob("*_processed.json"))

        for file_path in processed_files:
            if file_path.name.startswith("专利审查指南"):
                await self._apply_guideline_modifications(file_path)
            else:
                await self._apply_general_modifications(file_path)

    async def _apply_guideline_modifications(self, file_path: Path):
        """应用专利指南的2025修改"""
        try:
            with open(file_path, encoding='utf-8') as f:
                data = json.load(f)

            # 检查是否需要修改
            if data.get("modification_2025"):
                # 应用修改逻辑
                updated_data = await self._merge_modifications_2025(data)

                # 保存更新后的数据
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(updated_data, f, ensure_ascii=False, indent=2)

                self.stats["modifications_applied"] += 1
                logger.info(f"  应用修改: {file_path.name}")

        except Exception as e:
            logger.error(f"   修改应用失败: {e}")
            self.stats["errors"].append(f"修改应用失败: {file_path.name}")

    async def _apply_general_modifications(self, file_path: Path):
        """应用一般文档的2025修改"""
        try:
            with open(file_path, encoding='utf-8') as f:
                data = json.load(f)

            # 检查是否需要修改
            if data.get("modification_2025"):
                # 添加修改标记
                data["modification_2025"]["applied"] = True
                data["metadata"]["version"] = "2025"

                # 保存更新
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)

                self.stats["modifications_applied"] += 1
                logger.info(f"  应用修改标记: {file_path.name}")

        except Exception as e:
            logger.error(f"   修改标记失败: {e}")
            self.stats["errors"].append(f"修改标记失败: {file_path.name}")

    async def _merge_modifications_2025(self, data: dict) -> dict:
        """合并2025年修改到数据"""
        updated_data = data.copy()

        # 更新版本
        updated_data["metadata"]["version"] = "2025"

        # 应用具体修改
        modification = data["modification_2025"]
        if modification:
            if modification.get("type") == "new_content":
                # 添加新章节
                new_section = SectionData(
                    section_id=f"{modification['category']}_NEW_{modification.get('content', '')[:20]}",
                    level=4,
                    title=modification.get('content', ''),
                    content=f"（2025年新增）{modification.get('content', '')}",
                    full_path=f"2025新增 > {modification.get('category')}",
                    modification_2025=modification
                )
                updated_data["sections"].append(asdict(new_section))

            elif modification.get("type") == "deleted_section":
                # 删除章节
                updated_data["sections"] = [
                    s for s in updated_data["sections"]
                    if s.get("section_id") != modification.get("target_id", "")
                ]

            elif modification.get("type") == "modified_content":
                # 修改章节内容
                for section in updated_data["sections"]:
                    if modification.get("old_content") in section["content"]:
                        section["content"] = section["content"].replace(
                            modification["old_content"],
                            modification["new_content"]
                        )
                        # 标记修改
                        if "modification_2025" not in section:
                            section["modification_2025"] = {}
                        section["modification_2025"][modification.get("category", "unknown")] = {
                            "change_type": "content_modified",
                            "old_content": modification["old_content"],
                            "new_content": modification["new_content"]
                        }
                        break

        return updated_data

    async def _quality_check(self):
        """数据质量检查"""
        logger.info("\n🔍 执行数据质量检查...")

        # 检查处理完成的文件
        processed_files = list(self.processed_dir.glob("*_processed.json"))

        quality_report = {
            "total_files": len(processed_files),
            "files_with_sections": 0,
            "files_with_images": 0,
            "files_with_modifications": 0,
            "sections_per_file": [],
            "images_per_file": [],
            "issues_found": []
        }

        for file_path in processed_files:
            try:
                with open(file_path, encoding='utf-8') as f:
                    data = json.load(f)

                sections_count = len(data.get("sections", []))
                images_count = len(data.get("images", []))
                has_modifications = data.get("modification_2025") is not None

                quality_report["files_with_sections"] += (1 if sections_count > 0 else 0)
                quality_report["files_with_images"] += (1 if images_count > 0 else 0)
                quality_report["files_with_modifications"] += (1 if has_modifications else 0)
                quality_report["sections_per_file"].append(sections_count)
                quality_report["images_per_file"].append(images_count)

                # 检查数据完整性
                if not data.get("metadata", {}).get("title"):
                    quality_report["issues_found"].append(f"缺少标题: {file_path.name}")

                if sections_count == 0 and file_path.name.startswith("专利审查指南"):
                    quality_report["issues_found"].append(f"无章节结构: {file_path.name}")

            except Exception:
                quality_report["issues_found"].append(f"读取失败: {file_path.name}")

        # 计算平均值
        if quality_report["files_with_sections"] > 0:
            avg_sections = sum(quality_report["sections_per_file"]) / quality_report["files_with_sections"]
            logger.info(f"  平均章节数: {avg_sections:.1f}")

        if quality_report["files_with_images"] > 0:
            avg_images = sum(quality_report["images_per_file"]) / quality_report["files_with_images"]
            logger.info(f"  平均图像数: {avg_images:.1f}")

        # 保存质量报告
        report_file = self.reports_dir / f"quality_check_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(report_file, 'w', encoding='utf-8') as f:
            json.dump(quality_report, f, ensure_ascii=False, indent=2)

        logger.info(f"  📊 质量报告: {report_file}")

    async def _generate_report(self):
        """生成处理报告"""
        logger.info("\n📋 生成处理报告...")

        # 综计信息
        stats_file = self.reports_dir / f"processing_stats_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(stats_file, 'w', encoding='utf-8') as f:
            json.dump(self.stats, f, ensure_ascii=False, indent=2)

        # 可读报告
        readable_file = self.reports_dir / f"processing_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        with open(readable_file, 'w', encoding='utf-8') as f:
            f.write("专利规则构建系统处理报告\n")
            f.write("="*50 + "\n\n")
            f.write(f"处理时间: {self.stats['start_time']}\n")
            f.write(f"完成时间: {self.stats.get('end_time')}\n\n")

            f.write("处理统计:\n")
            f.write(f"  处理文档数: {self.stats['documents_processed']}\n")
            f.write(f"  提取章节数: {self.stats['sections_extracted']}\n")
            f.write(f"  提取图像数: {self.stats['images_extracted']}\n")
            f.write(f"  应用修改数: {self.stats['modifications_applied']}\n")
            f.write(f"  错误数: {len(self.stats['errors'])}\n\n")

            if self.stats['errors']:
                f.write("错误列表:\n")
                for error in self.stats['errors'][:10]:
                    f.write(f"  - {error}\n")
                if len(self.stats['errors']) > 10:
                    f.write(f"  ... 还有 {len(self.stats['errors']) - 10} 个错误\n")

        logger.info(f"  📋 处理报告: {readable_file}")

    def _print_summary(self) -> Any:
        """打印处理摘要"""
        logger.info("\n📊 处理摘要:")
        logger.info(f"  处理文档: {self.stats['documents_processed']} 个")
        logger.info(f"  提取章节: {self.stats['sections_extracted']} 个")
        logger.info(f"  提取图像: {self.stats['images_extracted']} 个")
        logger.info(f"  应用修改: {self.stats['modifications_applied']} 个")
        logger.info(f"  错误数: {len(self.stats['errors'])} 个")

        if self.stats["start_time"] and self.stats["end_time"]:
            duration = (self.stats["end_time"] - self.stats["start_time"]).total_seconds()
            hours = duration / 3600
            minutes = (duration % 3600) / 60
            logger.info(f"  处理耗时: {int(hours)}小时{int(minutes)}分钟")

    def _chinese_to_number(self, chinese: str) -> int:
        """中文数字转换"""
        mapping = {
            '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
            '六': 6, '七': 7, '八': 8, '九': 9, '十': 10
        }
        return mapping.get(chinese, 0)

    def _extract_number(self, text: str) -> int:
        """提取数字"""
        match = re.search(r'\d+', text)
        return int(match.group()) if match else 0

# 使用示例
async def main():
    """主函数示例"""
    processor = PatentRulesDataProcessor()

    # 处理所有文档
    await processor.process_all_documents()

if __name__ == "__main__":
    asyncio.run(main())
