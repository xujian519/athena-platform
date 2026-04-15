#!/usr/bin/env python3
"""
专利规则构建系统 - 增强版
Patent Rules Builder - Enhanced Version

支持PDF和Word文档，使用增强的实体提取器

作者: Athena平台团队
创建时间: 2025-12-20
版本: v2.0.0
"""

from __future__ import annotations
import asyncio
import json
import logging
import re
import sys
from datetime import datetime
from pathlib import Path

# 添加路径
sys.path.insert(0, str(Path(__file__).parent / "patent_rules_system"))

from nebula_graph_builder import NebulaGraphBuilder
from patent_entity_extractor_pro import PatentEntityExtractorPro
from qdrant_vector_store_simple import DocumentType, QdrantVectorStoreSimple, VectorDocument

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class PatentRulesBuilderEnhanced:
    """专利规则构建系统 - 增强版"""

    def __init__(self):
        self.laws_path = Path("/Users/xujian/学习资料/专利/专利法律法规")
        self.output_path = Path("/Users/xujian/Athena工作平台/production/data/patent_rules")
        self.vector_store = QdrantVectorStoreSimple()
        self.extractor = PatentEntityExtractorPro()
        self.graph_builder = NebulaGraphBuilder()

        # 创建输出目录
        self.output_path.mkdir(parents=True, exist_ok=True)

        # 统计信息
        self.stats = {
            "total_files": 0,
            "processed_files": 0,
            "failed_files": [],
            "total_documents": 0,
            "total_entities": 0,
            "total_relations": 0,
            "start_time": None,
            "end_time": None
        }

    async def scan_documents(self) -> list[dict]:
        """扫描专利法律法规文档（支持多种格式）"""
        logger.info(f"🔍 扫描目录: {self.laws_path}")

        documents = []
        # 支持多种文档格式
        file_types = ['.pdf', '.doc', '.docx', '.txt', '.md']

        for file_path in self.laws_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in file_types:
                # 判断文档类型
                file_name = file_path.name.lower()
                if '专利审查指南' in file_name:
                    doc_type = DocumentType.GUIDELINE
                elif '专利法' in file_name and '实施细则' not in file_name:
                    doc_type = DocumentType.PATENT_LAW
                elif '实施细则' in file_name:
                    doc_type = DocumentType.IMPLEMENTATION_RULES
                elif '最高人民法院' in file_name or '解释' in file_name:
                    doc_type = DocumentType.JUDICIAL_INTERPRETATION
                elif '条例' in file_name:
                    doc_type = DocumentType.PATENT_LAW  # 使用专利法类型
                else:
                    doc_type = DocumentType.PATENT_LAW  # 默认使用专利法类型

                documents.append({
                    'path': str(file_path),
                    'name': file_path.name,
                    'type': doc_type,
                    'size': file_path.stat().st_size,
                    'format': file_path.suffix.lower()
                })

        self.stats['total_files'] = len(documents)
        logger.info(f"✅ 找到 {len(documents)} 个文档")

        # 按格式统计
        format_count = {}
        for doc in documents:
            fmt = doc['format']
            format_count[fmt] = format_count.get(fmt, 0) + 1

        logger.info("文档格式分布:")
        for fmt, count in sorted(format_count.items()):
            logger.info(f"  - {fmt}: {count} 个")

        return documents

    async def process_document(self, doc_info: dict) -> list[VectorDocument | None]:
        """处理单个文档（支持PDF和Word）"""
        file_path = Path(doc_info['path'])
        file_format = doc_info['format']

        try:
            logger.info(f"📖 处理文档: {file_path.name} ({file_format})")

            # 根据格式选择处理方法
            if file_format == '.pdf':
                content = await self._process_pdf(file_path)
            elif file_format in ['.doc', '.docx']:
                content = await self._process_word(file_path)
            else:
                content = await self._process_text(file_path)

            if not content:
                logger.warning(f"⚠️ 文档内容为空: {file_path.name}")
                return None

            # 清理内容
            content = self._clean_content(content)

            # 智能分段
            documents = []
            sections = self._split_content(content, file_format)

            for i, section in enumerate(sections):
                # 创建向量文档
                doc = VectorDocument(
                    doc_id=f"{file_path.stem}_{i}",
                    content=section,
                    doc_type=doc_info['type'],
                    metadata={
                        'source_file': str(file_path),
                        'file_format': file_format,
                        'section_number': i + 1,
                        'total_sections': len(sections),
                        'file_size': doc_info['size'],
                        'document_title': file_path.stem
                    }
                )
                documents.append(doc)

            self.stats['processed_files'] += 1
            self.stats['total_documents'] += len(documents)

            logger.info(f"✅ 成功处理 {len(documents)} 个段落")
            return documents

        except Exception as e:
            logger.error(f"❌ 处理文档失败 {file_path.name}: {str(e)}")
            self.stats['failed_files'].append(str(file_path))
            return None

    async def _process_pdf(self, file_path: Path) -> str:
        """处理PDF文档"""
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                # 提取文本
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            return "\n".join(text_parts)

        except ImportError:
            logger.warning("pdfplumber未安装，尝试使用PyMuPDF")
            try:
                import fitz  # PyMuPDF

                doc = fitz.open(str(file_path))
                text_parts = []

                for page in doc:
                    text_parts.append(page.get_text())

                doc.close()
                return "\n".join(text_parts)

            except ImportError:
                raise ImportError("请安装pdfplumber或PyMuPDF来处理PDF文档") from None

    async def _process_word(self, file_path: Path) -> str:
        """处理Word文档"""
        try:
            import docx

            doc = docx.Document(str(file_path))
            text_parts = []

            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text_parts.append(row_text)

            return "\n".join(text_parts)

        except ImportError:
            raise ImportError("请安装python-docx来处理Word文档") from None

    async def _process_text(self, file_path: Path) -> str:
        """处理纯文本文档"""
        with open(file_path, encoding='utf-8') as f:
            return f.read()

    def _clean_content(self, content: str) -> str:
        """清理文档内容"""
        # 移除多余空行
        content = re.sub(r'\n{3,}', '\n\n', content)

        # 移除页码和页眉页脚模式
        content = re.sub(r'第\s*\d+\s*页', '', content)
        content = re.sub(r'-\s*\d+\s*-', '', content)

        # 统一空白符
        content = re.sub(r'\s+', ' ', content)
        content = re.sub(r'\n\s*\n', '\n\n', content)

        return content.strip()

    def _split_content(self, content: str, file_format: str) -> list[str]:
        """智能分割内容"""
        sections = []

        # 如果是法律文档，按条款分割
        if '法' in content or '条例' in content or '指南' in content:
            # 尝试按条款分割
            clause_pattern = r'(第[一二三四五六七八九十百千万\d]+条[^\n]*\n)'
            clauses = re.split(clause_pattern, content)

            if len(clauses) > 1:
                current_section = ""
                for i, part in enumerate(clauses):
                    if i % 2 == 1:  # 条款标题
                        if current_section:
                            sections.append(current_section.strip())
                        current_section = part
                    else:  # 条款内容
                        current_section += part

                if current_section:
                    sections.append(current_section.strip())

        # 如果没有成功分割，使用通用方法
        if not sections:
            # 按段落分割
            paragraphs = content.split('\n\n')
            paragraphs = [p.strip() for p in paragraphs if len(p.strip()) > 50]

            # 合并短段落
            current_section = ""
            for para in paragraphs:
                if len(current_section + para) < 1000:
                    current_section += para + "\n"
                else:
                    if current_section:
                        sections.append(current_section.strip())
                    current_section = para + "\n"

            if current_section:
                sections.append(current_section.strip())

        return sections

    async def extract_entities_and_relations(self, documents: list[VectorDocument]) -> dict:
        """提取实体和关系（使用增强版提取器）"""
        logger.info("🧠 使用增强版实体提取器...")

        all_entities = []
        all_relations = []

        for doc in documents:
            try:
                # 使用增强版提取器提取实体
                entities = self.extractor.extract_patent_entities(doc.content, doc.doc_id)

                # 提取关系
                relations = self.extractor.extract_relations(
                    doc.content,
                    entities,
                    doc.doc_id
                )

                # 添加元数据
                for entity in entities:
                    entity['document_type'] = doc.doc_type.value

                all_entities.extend(entities)
                all_relations.extend(relations)

            except Exception as e:
                logger.warning(f"⚠️ 提取实体关系失败 {doc.doc_id}: {str(e)}")

        self.stats['total_entities'] = len(all_entities)
        self.stats['total_relations'] = len(all_relations)

        logger.info(f"✅ 提取完成: {len(all_entities)} 个实体, {len(all_relations)} 个关系")

        # 统计实体类型分布
        entity_types = {}
        for entity in all_entities:
            etype = entity['type']
            entity_types[etype] = entity_types.get(etype, 0) + 1

        logger.info("实体类型分布:")
        for etype, count in sorted(entity_types.items(), key=lambda x: x[1], reverse=True):
            logger.info(f"  - {etype}: {count} 个")

        return {
            'entities': all_entities,
            'relations': all_relations
        }

    async def build_vector_database(self, all_documents: list[VectorDocument]):
        """构建向量数据库"""
        logger.info("🗄️ 构建向量数据库...")

        try:
            # 初始化向量库
            await self.vector_store.create_collection()

            # 批量索引文档
            batch_size = 50
            total_batches = (len(all_documents) + batch_size - 1) // batch_size

            for i in range(0, len(all_documents), batch_size):
                batch = all_documents[i:i + batch_size]
                batch_num = i // batch_size + 1

                logger.info(f"📦 处理批次 {batch_num}/{total_batches} ({len(batch)} 个文档)")

                # 批量索引
                await self.vector_store.batch_index(batch)

            logger.info(f"✅ 向量数据库构建完成: {len(all_documents)} 个文档")

        except Exception as e:
            logger.error(f"❌ 构建向量数据库失败: {str(e)}")
            raise

    async def build_knowledge_graph(self, entities_and_relations: dict):
        """构建知识图谱"""
        logger.info("🕸️ 构建知识图谱...")

        try:
            # 初始化图数据库
            await self.graph_builder.initialize_space()

            # 添加实体
            entities = entities_and_relations.get('entities', [])
            if entities:
                # 转换实体格式
                graph_entities = []
                for entity in entities:
                    graph_entities.append({
                        'name': entity['text'],
                        'type': entity['type'],
                        'properties': {
                            'confidence': entity.get('confidence', 0.5),
                            'source': entity.get('source', ''),
                            'document_type': entity.get('document_type', '')
                        }
                    })

                entity_ids = await self.graph_builder.add_entities(graph_entities)
                logger.info(f"✅ 添加 {len(entity_ids)} 个实体")

            # 添加关系
            relations = entities_and_relations.get('relations', [])
            if relations and entities:
                # 转换关系格式
                graph_relations = []
                for relation in relations:
                    graph_relations.append({
                        'subject': relation['subject'],
                        'object': relation['object'],
                        'relation': relation['relation'],
                        'properties': {
                            'confidence': relation.get('confidence', 0.5),
                            'source': relation.get('source', ''),
                            'distance': relation.get('distance', 0)
                        }
                    })

                relation_ids = await self.graph_builder.add_relations(graph_relations)
                logger.info(f"✅ 添加 {len(relation_ids)} 个关系")

            logger.info("✅ 知识图谱构建完成")

        except Exception as e:
            logger.error(f"❌ 构建知识图谱失败: {str(e)}")
            raise

    async def generate_report(self):
        """生成构建报告"""
        self.stats['end_time'] = datetime.now()

        # 计算耗时
        duration = (self.stats['end_time'] - self.stats['start_time']).total_seconds()

        # 生成报告
        report = {
            "构建时间": self.stats['start_time'].isoformat(),
            "完成时间": self.stats['end_time'].isoformat(),
            "总耗时": f"{duration:.2f}秒",
            "文档统计": {
                "扫描文档数": self.stats['total_files'],
                "处理成功数": self.stats['processed_files'],
                "处理失败数": len(self.stats['failed_files']),
                "失败文档": self.stats['failed_files'],
                "生成段落数": self.stats['total_documents']
            },
            "知识提取统计": {
                "提取实体数": self.stats['total_entities'],
                "提取关系数": self.stats['total_relations']
            },
            "输出位置": {
                "向量数据库": str(self.output_path / "vector_db"),
                "知识图谱": str(self.output_path / "knowledge_graph")
            },
            "优化特性": [
                "支持PDF、Word、Txt、Markdown多种格式",
                "增强版实体提取器（专利专业词典）",
                "智能文档分段（支持按条款分割）",
                "实体关系自动构建"
            ]
        }

        # 保存报告
        report_file = self.output_path / f"build_report_enhanced_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(report_file, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)

        logger.info(f"📄 构建报告已保存: {report_file}")

        # 打印摘要
        logger.info("\n" + "="*60)
        logger.info("🎯 专利规则构建系统 - 增强版构建完成")
        logger.info("="*60)
        logger.info(f"📊 处理文档: {self.stats['processed_files']}/{self.stats['total_files']}")
        logger.info(f"📝 生成段落: {self.stats['total_documents']}")
        logger.info(f"🏷️ 提取实体: {self.stats['total_entities']}")
        logger.info(f"🔗 提取关系: {self.stats['total_relations']}")
        logger.info(f"⏱️ 总耗时: {duration:.2f}秒")
        logger.info("="*60)

        return report

    async def run(self):
        """运行构建流程"""
        logger.info("🚀 启动专利规则构建系统 - 增强版")
        self.stats['start_time'] = datetime.now()

        try:
            # 1. 扫描文档
            documents = await self.scan_documents()

            # 2. 处理所有文档
            logger.info("\n📖 开始处理文档...")
            all_documents = []
            all_entities_and_relations = {'entities': [], 'relations': []}

            for doc_info in documents:
                docs = await self.process_document(doc_info)
                if docs:
                    all_documents.extend(docs)

            # 3. 提取实体和关系
            if all_documents:
                logger.info("\n🧠 开始提取实体和关系...")
                all_entities_and_relations = await self.extract_entities_and_relations(all_documents)

            # 4. 构建向量数据库
            if all_documents:
                await self.build_vector_database(all_documents)

            # 5. 构建知识图谱
            if all_entities_and_relations['entities']:
                await self.build_knowledge_graph(all_entities_and_relations)

            # 6. 生成报告
            report = await self.generate_report()

            logger.info("🎉 构建完成!")
            return report

        except Exception as e:
            logger.error(f"❌ 构建失败: {str(e)}")
            import traceback
            logger.error(traceback.format_exc())
            raise

async def main():
    """主函数"""
    builder = PatentRulesBuilderEnhanced()
    await builder.run()

if __name__ == "__main__":
    asyncio.run(main())
