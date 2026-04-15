#!/usr/bin/env python3
"""
测试修复后的决定要点提取完整流程
"""

from __future__ import annotations
import asyncio
import sys
from pathlib import Path

# 添加项目路径
sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent))

from production.scripts.review_decision.docx_only_pipeline import DocxOnlyPipeline


async def test_enhanced_extraction():
    """测试增强后的提取逻辑"""

    pipeline = DocxOnlyPipeline()

    # 测试文件
    test_files = [
        Path("/Volumes/AthenaData/语料/专利/专利复审决定原文/2010101258192.docx"),
        Path("/Volumes/AthenaData/语料/专利/专利复审决定原文/202021583902X.docx"),
    ]

    print('='*70)
    print('🧪 测试增强后的决定要点提取')
    print('='*70)
    print()

    for file_path in test_files:
        if not file_path.exists():
            print(f'⚠️ 文件不存在: {file_path.name}')
            continue

        print(f'📄 文件: {file_path.name}')
        print('-'*70)

        # 提取文本
        from docx import Document
        doc = Document(str(file_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        metadata = {
            'doc_id': file_path.stem,
            'filename': file_path.name,
            'file_size': file_path.stat().st_size,
            'paragraph_count': len(paragraphs)
        }

        # 测试分块
        chunks = pipeline.chunk_decision_text(text, metadata)

        print(f'  段落数: {len(paragraphs)}')
        print(f'  生成块数: {len(chunks)}')
        print()

        # 统计块类型
        block_types = {}
        for chunk in chunks:
            bt = chunk.get('block_type', 'unknown')
            block_types[bt] = block_types.get(bt, 0) + 1

        print('  📊 块类型分布:')
        for bt, count in sorted(block_types.items()):
            emoji = '📌' if bt == 'keypoints' else '📦'
            print(f'    {emoji} {bt}: {count}')

        # 显示决定要点
        keypoints_chunks = [c for c in chunks if c.get('block_type') == 'keypoints']
        if keypoints_chunks:
            print()
            print(f'  ✅ 决定要点提取成功 ({len(keypoints_chunks)}个):')
            for i, kp in enumerate(keypoints_chunks, 1):
                text = kp.get('text', '')[:200].replace('\n', ' ')
                law_refs = kp.get('metadata', {}).get('law_references', [])
                source = kp.get('metadata', {}).get('keypoint_source', 'standard')

                print(f'    [{i}] 来源: {source}')
                print(f'        内容: {text}...')
                if law_refs:
                    print(f'        法律引用: {law_refs}')
                print()
        else:
            print()
            print('  ⚠️ 未提取到决定要点')

        print()
        print('='*70)
        print()


async def test_full_pipeline():
    """测试完整管道（包括智能提取）"""

    pipeline = DocxOnlyPipeline()

    # 只处理1个文件测试
    review_dir = Path("/Volumes/AthenaData/语料/专利/专利复审决定原文")
    test_files = list(review_dir.glob("*.docx"))[:1]

    if not test_files:
        print("⚠️ 未找到测试文件")
        return

    print('='*70)
    print('🚀 测试完整管道流程')
    print('='*70)
    print()

    try:
        # 初始化服务（跳过BGE和Qdrant，只测试分块）
        print('✅ 管道已加载（智能提取器已启用）')
        print()

        # 测试处理
        for file_path in test_files:
            from docx import Document
            doc = Document(str(file_path))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)

            metadata = {
                'doc_id': file_path.stem,
                'filename': file_path.name,
            }

            chunks = pipeline.chunk_decision_text(text, metadata)

            print(f'📄 {file_path.name}')
            print(f'  总块数: {len(chunks)}')

            keypoints_count = sum(1 for c in chunks if c.get('block_type') == 'keypoints')
            print(f'  决定要点: {keypoints_count}')

            # 统计有法律引用的块
            with_law_refs = sum(1 for c in chunks
                                if c.get('metadata', {}).get('law_references'))
            print(f'  有法律引用的块: {with_law_refs}')

            print()

    except Exception as e:
        print(f'❌ 错误: {e}')
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    print("\n" + "="*70)
    print("🚀 开始测试增强的决定要点提取")
    print("="*70 + "\n")

    # 测试1: 增强的提取逻辑
    asyncio.run(test_enhanced_extraction())

    # 测试2: 完整流程
    asyncio.run(test_full_pipeline())

    print("\n" + "="*70)
    print("✅ 测试完成")
    print("="*70 + "\n")
