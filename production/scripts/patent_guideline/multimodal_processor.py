#!/usr/bin/env python3
"""
专利指南多模态文档处理器
Patent Guideline Multimodal Document Processor

处理PDF、Word、图片等多种格式的专利指南文档

作者: Athena平台团队
创建时间: 2025-12-20
"""

from __future__ import annotations
import hashlib
import io
import json
import logging
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

# OCR
import easyocr
import fitz  # PyMuPDF

# PDF处理
import pdfplumber
import pytesseract

# Word处理
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

# 图片处理
from PIL import Image
from PyPDF2 import PdfReader

# 使用安全哈希函数替代不安全的MD5/SHA1

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """文档元数据"""
    title: str
    version: str | None = None
    author: str | None = None
    created_date: str | None = None
    modified_date: str | None = None
    page_count: int = 0
    file_type: str = ""
    file_size: int = 0
    hash_md5: str = ""

@dataclass
class SectionInfo:
    """章节信息"""
    section_id: str
    level: int
    title: str
    content: str
    parent_id: str | None = None
    page_range: str | None = None
    full_path: str = ""
    child_sections: list[str] = None

    def __post_init__(self):
        if self.child_sections is None:
            self.child_sections = []

@dataclass
class TableInfo:
    """表格信息"""
    table_id: str
    title: str
    content: list[list[str]]
    headers: list[str] = None
    caption: str = ""
    page_number: int = 0

class MultimodalDocumentProcessor:
    """多模态文档处理器"""

    def __init__(self):
        # OCR阅读器初始化
        try:
            self.ocr_reader = easyocr.Reader(['ch_en', 'en'])
            logger.info("✅ EasyOCR reader initialized")
        except Exception as e:
            logger.warning(f"⚠️ EasyOCR初始化失败: {e}")
            self.ocr_reader = None

        # 支持的文件类型
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.png': self._process_image,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.tiff': self._process_image,
            '.tif': self._process_image
        }

        # 专利指南特有的章节识别模式
        self.section_patterns = {
            'part': r'^第[一二三四五六七八九十]+部分',
            'chapter': r'^第[一二三四五六七八九十\d]+章',
            'section': r'^\d+\.\d+',
            'subsection': r'^\d+\.\d+\.\d+',
            'case': r'^【例\d+】',
            'note': r'^注\d*：',
            'reference': r'^(参见|根据|依据|按照)'
        }

    def process_document(self, file_path: str) -> dict[str, Any]:
        """处理文档主入口"""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 获取文件信息
        file_ext = file_path.suffix.lower()
        if file_ext not in self.supported_types:
            raise ValueError(f"不支持的文件类型: {file_ext}")

        # 计算文件哈希
        file_hash = self._calculate_file_hash(file_path)
        file_size = file_path.stat().st_size

        logger.info(f"开始处理文档: {file_path.name} ({file_ext})")

        # 提取文档元数据
        metadata = self._extract_metadata(file_path)
        metadata.file_type = file_ext
        metadata.file_size = file_size
        metadata.hash_md5 = file_hash

        # 根据文件类型选择处理方法
        processor_func = self.supported_types[file_ext]
        content = processor_func(file_path)

        # 结构化处理
        sections = self._extract_sections(content, metadata)

        # 提取表格
        tables = self._extract_tables(content)

        # 返回处理结果
        result = {
            "metadata": asdict(metadata),
            "sections": [asdict(s) for s in sections],
            "tables": [asdict(t) for t in tables],
            "raw_content": content if isinstance(content, str) else "",
            "processing_stats": {
                "section_count": len(sections),
                "table_count": len(tables),
                "processing_time": datetime.now().isoformat()
            }
        }

        logger.info(f"✅ 文档处理完成: {len(sections)}个章节, {len(tables)}个表格")
        return result

    def _process_pdf(self, file_path: Path) -> str:
        """处理PDF文件"""
        content_parts = []

        # 方法1: 使用pdfplumber（保留格式）
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        content_parts.append(f"\n--- Page {page_num + 1} ---\n")
                        content_parts.append(page_text)

                    # 处理表格
                    tables = page.extract_tables()
                    for table in tables:
                        table_text = self._table_to_text(table)
                        content_parts.append(f"\n[表格]\n{table_text}\n[/表格]\n")

            logger.info("使用pdfplumber处理完成")
        except Exception as e:
            logger.warning(f"pdfplumber处理失败: {e}")

        # 方法2: 使用PyMuPDF处理图片
        if len(content_parts) == 0:
            try:
                doc = fitz.open(str(file_path))
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    text = page.get_text()
                    content_parts.append(f"\n--- Page {page_num + 1} ---\n")
                    content_parts.append(text)

                    # 如果文字提取效果不好，尝试OCR
                    if len(text.strip()) < 100:
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        ocr_text = self._ocr_image(img_data)
                        if ocr_text:
                            content_parts.append(f"[OCR]\n{ocr_text}\n[/OCR]")

                doc.close()
                logger.info("使用PyMuPDF处理完成")
            except Exception as e:
                logger.warning(f"PyMuPDF处理失败: {e}")

        return "\n".join(content_parts)

    def _process_docx(self, file_path: Path) -> str:
        """处理Word文档"""
        try:
            doc = Document(file_path)
            content_parts = []

            # 处理段落
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)

            # 处理表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                table_text = self._table_data_to_text(table_data)
                content_parts.append(f"\n[表格]\n{table_text}\n[/表格]\n")

            return "\n".join(content_parts)

        except PackageNotFoundError:
            # 如果docx处理失败，尝试使用其他方法
            logger.warning("python-docx处理失败，尝试其他方法")
            # 可以添加antiword或其他工具
            return ""
        except Exception as e:
            logger.error(f"Word文档处理失败: {e}")
            return ""

    def _process_image(self, file_path: Path) -> str:
        """处理图片文件"""
        try:
            # 读取图片
            with open(file_path, 'rb') as f:
                img_data = f.read()

            # OCR识别
            ocr_text = self._ocr_image(img_data)

            if ocr_text:
                return f"[图片文件: {file_path.name}]\n{ocr_text}"
            else:
                logger.warning(f"OCR识别失败: {file_path.name}")
                return f"[图片文件: {file_path.name}]\nOCR识别失败"

        except Exception as e:
            logger.error(f"图片处理失败: {e}")
            return f"[图片文件: {file_path.name}]\n处理错误: {str(e)}"

    def _ocr_image(self, img_data: bytes) -> str:
        """OCR识别图片文字"""
        if not self.ocr_reader:
            # 使用Tesseract作为后备
            try:
                image = Image.open(io.BytesIO(img_data))
                text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                return text
            except Exception as e:
                logger.debug(f"空except块已触发: {e}")
                return ""

        try:
            # 使用EasyOCR
            result = self.ocr_reader.readtext(img_data)
            texts = [item[1] for item in result]
            return "\n".join(texts)
        except Exception as e:
            logger.debug(f"EasyOCR失败: {e}")
            return ""

    def _extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """提取文档元数据"""
        metadata = DocumentMetadata(title=file_path.stem)

        try:
            # PDF元数据
            if file_path.suffix.lower() == '.pdf':
                with open(file_path, 'rb') as f:
                    pdf_reader = PdfReader(f)

                    # 基本信息
                    metadata.page_count = len(pdf_reader.pages)

                    # PDF信息
                    if pdf_reader.metadata:
                        info = pdf_reader.metadata
                        metadata.title = info.get('/Title', metadata.title)
                        metadata.author = info.get('/Author', None)
                        metadata.created_date = str(info.get('/CreationDate', ''))
                        metadata.modified_date = str(info.get('/ModDate', ''))

            # Word元数据
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                doc = Document(file_path)
                metadata.title = doc.core_properties.title or metadata.title
                metadata.author = doc.core_properties.author or None
                metadata.created_date = str(doc.core_properties.created) if doc.core_properties.created else ""
                metadata.modified_date = str(doc.core_properties.modified) if doc.core_properties.modified else ""

        except Exception as e:
            logger.warning(f"元数据提取失败: {e}")

        return metadata

    def _extract_sections(self, content: str, metadata: DocumentMetadata) -> list[SectionInfo]:
        """提取文档章节结构"""
        sections = []
        lines = content.split('\n')

        current_section = None
        section_stack = []  # 用于跟踪层级关系

        for _i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            # 检查是否是章节标题
            section_info = self._identify_section(line)

            if section_info:
                # 处理层级关系
                while section_stack and section_stack[-1]['level'] >= section_info['level']:
                    section_stack.pop()

                parent_id = section_stack[-1]['section_id'] if section_stack else None

                # 构建完整路径
                full_path = section_info['title']
                if parent_id:
                    parent_section = next((s for s in sections if s.section_id == parent_id), None)
                    if parent_section:
                        full_path = parent_section.full_path + " > " + section_info['title']

                # 创建章节对象
                section = SectionInfo(
                    section_id=section_info['section_id'],
                    level=section_info['level'],
                    title=section_info['title'],
                    content="",  # 稍后填充
                    parent_id=parent_id,
                    full_path=full_path
                )

                sections.append(section)
                section_stack.append({
                    'section_id': section.section_id,
                    'level': section.level
                })

                current_section = section
            elif current_section:
                # 添加到当前章节内容
                if current_section.content:
                    current_section.content += "\n" + line
                else:
                    current_section.content = line

        # 为没有内容的章节提供默认内容
        for section in sections:
            if not section.content:
                # 查找后续内容
                section_idx = sections.index(section)
                if section_idx + 1 < len(sections):
                    next_section = sections[section_idx + 1]
                    section.content = f"[章节标题: {section.title}]\n详见下一章节"
                else:
                    section.content = f"[章节标题: {section.title}]"

        return sections

    def _identify_section(self, line: str) -> dict | None:
        """识别章节标题"""
        import re

        # 部分标题
        part_match = re.match(r'^(第[一二三四五六七八九十]+部分)\s*(.*)', line)
        if part_match:
            return {
                'section_id': f"P{self._chinese_to_number(part_match.group(1)[1:-2])}",
                'level': 1,
                'title': line,
                'type': 'part'
            }

        # 章标题
        chapter_match = re.match(r'^(第[一二三四五六七八九十\d]+章)\s*(.*)', line)
        if chapter_match:
            chapter_num = self._extract_number(chapter_match.group(1)[1:-1])
            return {
                'section_id': f"C{chapter_num}",
                'level': 2,
                'title': line,
                'type': 'chapter'
            }

        # 节标题
        section_match = re.match(r'^(\d+\.\d+)\s*(.*)', line)
        if section_match:
            return {
                'section_id': f"S{section_match.group(1).replace('.', '-')}",
                'level': 3,
                'title': line,
                'type': 'section'
            }

        # 小节标题
        subsection_match = re.match(r'^(\d+\.\d+\.\d+)\s*(.*)', line)
        if subsection_match:
            return {
                'section_id': f"SS{subsection_match.group(1).replace('.', '-')}",
                'level': 4,
                'title': line,
                'type': 'subsection'
            }

        # 案例
        case_match = re.match(r'^【例(\d+)】\s*(.*)', line)
        if case_match:
            return {
                'section_id': f"EX{case_match.group(1)}",
                'level': 5,
                'title': line,
                'type': 'case'
            }

        return None

    def _extract_tables(self, content: str) -> list[TableInfo]:
        """提取表格信息"""
        tables = []

        # 查找表格标记
        import re
        table_pattern = r'\[表格\](.*?)\[/表格\]'
        matches = re.findall(table_pattern, content, re.DOTALL)

        for i, table_text in enumerate(matches):
            # 简单解析表格
            rows = table_text.strip().split('\n')
            table_data = []

            for row in rows:
                if row.strip():
                    # 简单的列分割
                    cols = [col.strip() for col in row.split('\t') if col.strip()]
                    if cols:
                        table_data.append(cols)

            if table_data:
                table = TableInfo(
                    table_id=f"T{i+1}",
                    title=f"表格{i+1}",
                    content=table_data,
                    headers=table_data[0] if table_data else []
                )
                tables.append(table)

        return tables

    def _table_to_text(self, table: list[list[str]]) -> str:
        """将表格转换为文本"""
        rows = []
        for row in table:
            if row:
                row_text = " | ".join([str(cell) if cell else "" for cell in row])
                rows.append(row_text)
        return "\n".join(rows)

    def _table_data_to_text(self, table_data: list[list[str]]) -> str:
        """将表格数据转换为文本"""
        return self._table_to_text(table_data)

    def _calculate_file_hash(self, file_path: Path) -> str:
        """计算文件MD5哈希"""
        hash_md5 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

    def _chinese_to_number(self, chinese: str) -> int:
        """中文数字转换"""
        mapping = {
            '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
            '六': 6, '七': 7, '八': 8, '九': 9, '十': 10
        }
        return mapping.get(chinese, 0)

    def _extract_number(self, text: str) -> int:
        """提取数字"""
        import re
        match = re.search(r'\d+', text)
        return int(match.group()) if match else 0

# 使用示例
def main() -> None:
    """主函数示例"""
    processor = MultimodalDocumentProcessor()

    # 处理示例文档
    test_file = "/Users/xujian/Athena工作平台/dev/tools/patent-guideline-system/专利审查指南.pdf"
    if Path(test_file).exists():
        result = processor.process_document(test_file)

        # 打印处理结果
        print("\n=== 文档元数据 ===")
        print(json.dumps(result['metadata'], ensure_ascii=False, indent=2))

        print("\n=== 章节结构 ===")
        for section in result['sections'][:5]:
            print(f"- {section['section_id']}: {section['title']}")
            print(f"  Level: {section['level']}, Parent: {section['parent_id']}")

        print("\n=== 表格信息 ===")
        for table in result['tables']:
            print(f"- {table['table_id']}: {table['title']} ({len(table['content'])} rows)")

if __name__ == "__main__":
    main()
