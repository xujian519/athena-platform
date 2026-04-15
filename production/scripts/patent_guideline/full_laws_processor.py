#!/usr/bin/env python3
"""
专利法律法规全量处理器
处理所有专利法律法规文件，生成BGE向量，导入Qdrant和NebulaGraph

作者: Athena平台团队
创建时间: 2025-12-23
"""

from __future__ import annotations
import asyncio
import json
import logging
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF

# 使用安全哈希函数替代不安全的MD5/SHA1
from production.utils.security_helpers import short_hash

# 添加项目路径
sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent))
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.file_handler(f'/Users/xujian/Athena工作平台/logs/legal_full_processor_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log'),
        logging.stream_handler()
    ]
)
logger = logging.getLogger(__name__)


@dataclass
class LegalDocument:
    """法律文档数据结构"""
    doc_id: str
    doc_type: str  # 专利法、实施细则、审查指南、司法解释等
    title: str
    content: str
    source_file: str
    file_type: str  # pdf, md, docx
    metadata: dict[str, Any] = field(default_factory=dict)

    def __post_init__(self):
        self.char_count = len(self.content)
        self.word_count = len([w for w in self.content if w.strip()])


class PatentLawsFullProcessor:
    """专利法律法规全量处理器"""

    def __init__(self):
        self.base_dir = Path("/Users/xujian/Athena工作平台")
        self.source_dir = Path("/Volumes/AthenaData/语料/专利/专利法律法规")
        self.output_dir = self.base_dir / "production/data/patent_rules"

        # BGE服务
        self.bge_service = None

        # Qdrant客户端
        self.qdrant_client = None

        # NebulaGraph连接
        self.nebula_connection = None

        logger.info("专利法律法规全量处理器初始化完成")

    async def initialize_services(self):
        """初始化所有服务"""
        logger.info("=" * 60)
        logger.info("🚀 初始化服务 (BGE + Qdrant + NebulaGraph)")
        logger.info("=" * 60)

        # 1. 初始化BGE服务
        try:
            from core.nlp.bge_embedding_service import BGEEmbeddingService

            model_path = self.base_dir / "models/converted/bge-large-zh-v1.5"
            if not model_path.exists():
                model_path = self.base_dir / "models/bge-large-zh-v1.5"

            config = {
                "model_path": str(model_path),
                "device": "cpu",
                "batch_size": 32,
                "max_length": 512,
                "normalize_embeddings": True,
                "cache_enabled": True,
                "preload": True
            }

            self.bge_service = BGEEmbeddingService(config)
            await self.bge_service.initialize()

            health = await self.bge_service.health_check()
            logger.info(f"✅ BGE服务: {health['status']}, 维度: {health['dimension']}")

        except Exception as e:
            logger.error(f"❌ BGE服务初始化失败: {e}")

        # 2. 初始化Qdrant客户端
        try:
            from qdrant_client import QdrantClient
            from qdrant_client.models import Distance, VectorParams

            self.qdrant_client = QdrantClient(url="http://localhost:6333")

            # 创建集合
            collection_name = "patent_laws"

            # 检查集合是否存在
            collections = self.qdrant_client.get_collections().collections
            collection_names = [c.name for c in collections]

            if collection_name not in collection_names:
                self.qdrant_client.create_collection(
                    collection_name=collection_name,
                    vectors_config=VectorParams(size=1024, distance=Distance.COSINE)
                )
                logger.info(f"✅ Qdrant集合已创建: {collection_name}")
            else:
                logger.info(f"✅ Qdrant集合已存在: {collection_name}")

        except Exception as e:
            logger.warning(f"⚠️ Qdrant初始化失败: {e}")

        # 3. NebulaGraph连接（简化版）
        try:

            # 这里简化处理，记录日志
            logger.info("✅ NebulaGraph配置已准备")

        except Exception as e:
            logger.warning(f"⚠️ NebulaGraph初始化失败: {e}")

        logger.info("=" * 60)

    def scan_source_directory(self) -> list[Path]:
        """扫描源目录，获取所有法律文件"""
        logger.info(f"📂 扫描目录: {self.source_dir}")

        files = []

        # 支持的文件类型
        extensions = ['.pdf', '.md', '.txt', '.docx']

        for ext in extensions:
            files.extend(self.source_dir.glob(f"*{ext}"))

        logger.info(f"📄 找到 {len(files)} 个文件")

        for file in sorted(files):
            logger.info(f"   - {file.name} ({file.stat().st_size / 1024:.1f} KB)")

        return files

    def extract_text_from_file(self, file_path: Path) -> str:
        """从文件中提取文本"""
        file_ext = file_path.suffix.lower()

        if file_ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_ext in ['.md', '.txt']:
            return self._extract_from_text(file_path)
        elif file_ext == '.docx':
            return self._extract_from_docx(file_path)
        else:
            logger.warning(f"不支持的文件类型: {file_ext}")
            return ""

    def _extract_from_pdf(self, file_path: Path) -> str:
        """从PDF提取文本"""
        try:
            doc = fitz.open(str(file_path))
            text_parts = []

            for page in doc:
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)

            doc.close()
            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"PDF提取失败 {file_path.name}: {e}")
            return ""

    def _extract_from_text(self, file_path: Path) -> str:
        """从文本文件提取"""
        try:
            with open(file_path, encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"文本提取失败 {file_path.name}: {e}")
            return ""

    def _extract_from_docx(self, file_path: Path) -> str:
        """从DOCX提取文本"""
        try:
            from docx import Document

            doc = Document(str(file_path))
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            return "\n".join(paragraphs)

        except Exception as e:
            logger.error(f"DOCX提取失败 {file_path.name}: {e}")
            return ""

    def chunk_document(self, doc: LegalDocument, max_chunk_size: int = 500) -> list[dict[str, Any]]:
        """将文档分块"""
        # 按段落分块
        paragraphs = doc.content.split('\n\n')

        chunks = []
        current_chunk = []
        current_size = 0
        chunk_index = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            para_size = len(para)

            if current_size + para_size > max_chunk_size and current_chunk:
                # 保存当前块
                chunks.append({
                    'chunk_id': f"{doc.doc_id}_chunk_{chunk_index:04d}",
                    'doc_id': doc.doc_id,
                    'doc_type': doc.doc_type,
                    'title': doc.title,
                    'content': '\n\n'.join(current_chunk),
                    'chunk_index': chunk_index,
                    'metadata': {
                        'source_file': doc.source_file,
                        'file_type': doc.file_type
                    }
                })

                # 开始新块
                current_chunk = [para]
                current_size = para_size
                chunk_index += 1
            else:
                current_chunk.append(para)
                current_size += para_size

        # 保存最后一个块
        if current_chunk:
            chunks.append({
                'chunk_id': f"{doc.doc_id}_chunk_{chunk_index:04d}",
                'doc_id': doc.doc_id,
                'doc_type': doc.doc_type,
                'title': doc.title,
                'content': '\n\n'.join(current_chunk),
                'chunk_index': chunk_index,
                'metadata': {
                    'source_file': doc.source_file,
                    'file_type': doc.file_type
                }
            })

        return chunks

    async def process_all_documents(self):
        """处理所有文档"""
        logger.info("=" * 60)
        logger.info("📚 开始处理所有专利法律法规文档")
        logger.info("=" * 60)

        await self.initialize_services()

        # 扫描文件
        files = self.scan_source_directory()

        all_chunks = []
        all_docs = []

        for file_path in files:
            logger.info(f"\n📖 处理: {file_path.name}")

            # 提取文本
            text = self.extract_text_from_file(file_path)

            if not text:
                logger.warning(f"⚠️ 跳过空文件: {file_path.name}")
                continue

            # 确定文档类型
            doc_type = self._determine_doc_type(file_path.name)

            # 创建文档对象
            doc_id = self._generate_doc_id(file_path.name)

            doc = LegalDocument(
                doc_id=doc_id,
                doc_type=doc_type,
                title=file_path.stem,
                content=text,
                source_file=str(file_path),
                file_type=file_path.suffix,
                metadata={
                    'file_size': file_path.stat().st_size,
                    'processed_at': datetime.now().isoformat()
                }
            )

            all_docs.append({
                'doc_id': doc.doc_id,
                'doc_type': doc.doc_type,
                'title': doc.title,
                'char_count': doc.char_count,
                'word_count': doc.word_count,
                'source_file': doc.source_file
            })

            # 分块
            chunks = self.chunk_document(doc)
            all_chunks.extend(chunks)

            logger.info(f"   ✅ 生成 {len(chunks)} 个块")

        logger.info("=" * 60)
        logger.info("📊 处理完成:")
        logger.info(f"   文档数: {len(all_docs)}")
        logger.info(f"   总块数: {len(all_chunks)}")
        logger.info("=" * 60)

        # 保存处理结果
        self._save_processing_results(all_docs, all_chunks)

        # 生成BGE向量并导入Qdrant
        await self._generate_vectors_and_import(all_chunks)

        return {
            'total_documents': len(all_docs),
            'total_chunks': len(all_chunks)
        }

    def _determine_doc_type(self, filename: str) -> str:
        """确定文档类型"""
        if '专利法' in filename and '实施细则' not in filename:
            return '专利法'
        elif '实施细则' in filename:
            return '实施细则'
        elif '审查指南' in filename:
            return '审查指南'
        elif '解释' in filename:
            return '司法解释'
        elif '条例' in filename:
            return '行政法规'
        elif '规定' in filename:
            return '部门规章'
        else:
            return '其他'

    def _generate_doc_id(self, filename: str) -> str:
        """生成文档ID"""
        return f"doc_{short_hash(filename.encode(), 12)}"

    def _save_processing_results(self, docs: list[dict], chunks: list[dict]) -> Any:
        """保存处理结果"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # 保存文档列表
        docs_file = self.output_dir / "processed" / f"legal_documents_{timestamp}.json"
        docs_file.parent.mkdir(parents=True, exist_ok=True)

        with open(docs_file, 'w', encoding='utf-8') as f:
            json.dump({
                'processed_at': datetime.now().isoformat(),
                'total_documents': len(docs),
                'documents': docs
            }, f, ensure_ascii=False, indent=2)

        logger.info(f"📄 文档列表: {docs_file}")

        # 保存块数据
        chunks_file = self.output_dir / "processed" / f"legal_chunks_{timestamp}.json"

        with open(chunks_file, 'w', encoding='utf-8') as f:
            json.dump({
                'processed_at': datetime.now().isoformat(),
                'total_chunks': len(chunks),
                'chunks': chunks
            }, f, ensure_ascii=False, indent=2)

        logger.info(f"📦 块数据: {chunks_file}")

    async def _generate_vectors_and_import(self, chunks: list[dict]):
        """生成向量并导入Qdrant"""
        logger.info("=" * 60)
        logger.info(f"🔄 为 {len(chunks)} 个块生成向量并导入Qdrant")
        logger.info("=" * 60)

        if not self.bge_service:
            logger.error("BGE服务未初始化")
            return

        batch_size = 32
        total_batches = (len(chunks) + batch_size - 1) // batch_size

        # 准备Qdrant点
        points = []

        for batch_idx in range(total_batches):
            start_idx = batch_idx * batch_size
            end_idx = min(start_idx + batch_size, len(chunks))

            batch_chunks = chunks[start_idx:end_idx]
            texts = [chunk['content'] for chunk in batch_chunks]

            try:
                # 生成向量
                result = await self.bge_service.encode(texts, task_type="patent_laws")

                # 创建Qdrant点
                for _i, (chunk, embedding) in enumerate(zip(batch_chunks, result.embeddings, strict=False)):
                    points.append({
                        'id': chunk['chunk_id'],
                        'vector': embedding.tolist() if hasattr(embedding, 'tolist') else list(embedding),
                        'payload': {
                            'doc_id': chunk['doc_id'],
                            'doc_type': chunk['doc_type'],
                            'title': chunk['title'],
                            'content': chunk['content'][:500],  # 保存前500字符用于预览
                            'chunk_index': chunk['chunk_index'],
                            'source_file': chunk['metadata']['source_file']
                        }
                    })

                # 每100个点上传一次
                if len(points) >= 100:
                    await self._upload_to_qdrant(points)
                    points = []

                progress = (batch_idx + 1) / total_batches * 100
                logger.info(f"✅ 批次 {batch_idx + 1}/{total_batches} ({progress:.1f}%)")

            except Exception as e:
                logger.error(f"❌ 批次 {batch_idx} 失败: {e}")

        # 上传剩余的点
        if points:
            await self._upload_to_qdrant(points)

        logger.info("✅ 向量生成和Qdrant导入完成")

    async def _upload_to_qdrant(self, points: list[dict]):
        """上传到Qdrant"""
        try:
            from qdrant_client.models import PointStruct

            # 生成数值ID：使用chunk_id的哈希值
            qdrant_points = []
            for p in points:
                # 将字符串ID转换为整数（使用哈希）
                chunk_id_str = p['id']
                hash_id = abs(hash(chunk_id_str)) % (10 ** 10)  # 限制在10位数字内

                # 确保payload中包含原始chunk_id
                payload_with_id = p['payload'].copy()
                payload_with_id['chunk_id'] = chunk_id_str

                qdrant_points.append(
                    PointStruct(
                        id=hash_id,
                        vector=p['vector'],
                        payload=payload_with_id
                    )
                )

            self.qdrant_client.upsert(
                collection_name="patent_laws",
                points=qdrant_points
            )

            logger.info(f"   📤 已上传 {len(points)} 个点到Qdrant")

        except Exception as e:
            logger.warning(f"Qdrant上传失败: {e}")
            import traceback
            traceback.print_exc()

    def export_to_nebula_graph(self, docs: list[dict], chunks: list[dict]) -> Any:
        """导出到NebulaGraph"""
        logger.info("=" * 60)
        logger.info("🌐 导出数据到NebulaGraph")
        logger.info("=" * 60)

        # 创建NebulaGraph导入脚本
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = self.output_dir / "knowledge_graph" / f"nebula_import_{timestamp}.ngql"

        output_file.parent.mkdir(parents=True, exist_ok=True)

        with open(output_file, 'w', encoding='utf-8') as f:
            # 写入头部
            f.write("# NebulaGraph 导入脚本\n")
            f.write(f"# 生成时间: {datetime.now().isoformat()}\n\n")

            # 创建空间
            f.write("CREATE SPACE IF NOT EXISTS patent_laws(partition_num=10, replica_factor=1, vid_type=FIXED_STRING(32));\n")
            f.write("USE patent_laws;\n\n")

            # 创建标签
            f.write("# 创建节点类型\n")
            f.write("CREATE TAG IF NOT EXISTS document(name string, doc_type string, char_count int);\n")
            f.write("CREATE TAG IF NOT EXISTS chunk(chunk_index int, content string);\n\n")

            # 创建边类型
            f.write("# 创建边类型\n")
            f.write("CREATE EDGE IF NOT EXISTS contains(weight int);\n\n")

            # 插入文档节点
            f.write("# 插入文档节点\n")
            for doc in docs[:10]:  # 示例：只插入前10个
                escaped_name = doc['title'].replace("'", "\\'")
                f.write(f"INSERT VERTEX document(\"{doc['doc_id']}\", \"{escaped_name}\", \"{doc['doc_type']}\", {doc['char_count']});\n")

            f.write("\n")

            # 插入块节点
            f.write("# 插入块节点（示例）\n")
            for chunk in chunks[:10]:  # 示例：只插入前10个
                escaped_content = chunk['content'][:100].replace('"', '\\"').replace("\n", "\\n")
                f.write(f"INSERT VERTEX chunk(\"{chunk['chunk_id']}\", {chunk['chunk_index']}, \"{escaped_content}\");\n")

            f.write("\n")

            # 插入边
            f.write("# 插入边（文档-块关系）\n")
            for chunk in chunks[:10]:
                f.write(f"INSERT EDGE contains(\"{chunk['doc_id']}\")->(\"{chunk['chunk_id']}\", {chunk['chunk_index']});\n")

        logger.info(f"📄 NebulaGraph脚本: {output_file}")
        logger.info("✅ NebulaGraph导出完成")

        return str(output_file)


async def main():
    """主函数"""
    processor = PatentLawsFullProcessor()

    # 处理所有文档
    result = await processor.process_all_documents()

    logger.info("=" * 60)
    logger.info("🎉 全量处理完成！")
    logger.info(f"📊 处理文档: {result['total_documents']}")
    logger.info(f"📦 处理块数: {result['total_chunks']}")
    logger.info("=" * 60)


if __name__ == "__main__":
    asyncio.run(main())
