#!/usr/bin/env python3
"""
创建Word格式的审查意见答复文件
Create Word format response for patent office action

作者: Athena平台团队
版本: 1.0.0
创建时间: 2026-02-09
"""

from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def create_response_document():
    """创建审查意见答复Word文档"""

    # 创建文档
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # ============ 添加标题 ============
    title = doc.add_heading('第一次审查意见答复', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '黑体'
        run.font.size = Pt(16)

    # ============ 添加基本信息 ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('申请人：山东圣旺药业股份有限公司  申请号：202520560089.0')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('发明创造名称：一种药品生产灌装机  答复日期：2026年2月9日')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('审查员：肖玉林')
    run.font.size = Pt(11)

    doc.add_paragraph()

    # ============ 一、审查意见概述 ============
    doc.add_heading('一、审查意见概述', 1)
    doc.add_paragraph(
        '审查员认为本申请的说明书没有清楚完整地解释和说明"灌装、计量、称重"的'
        '技术手段，不清楚如何用于药品灌装的，对所属技术领域的技术人员来说，该手段是'
        '含糊不清的，根据说明书记载的内容无法具体实施，因此不符合专利法第26条第3款的规定。'
    )

    doc.add_paragraph()

    # ============ 二、答复意见 ============
    doc.add_heading('二、答复意见', 1)
    doc.add_paragraph()

    # 2.1 关于专利法第26条第3款的说明
    doc.add_heading('（一）关于专利法第26条第3款的说明', 2)
    doc.add_paragraph(
        '本申请的说明书已经对技术方案作出了清楚、完整的说明，所属技术领域的技术人员能够'
        '根据说明书实现本实用新型，符合专利法第26条第3款的规定。具体理由如下：'
    )

    doc.add_paragraph()

    # 关于"灌装"技术手段
    doc.add_heading('1. 关于"灌装"技术手段的说明', 3)
    doc.add_paragraph(
        '审查员认为说明书没有清楚完整地解释"灌装"的技术手段，这一观点是不准确的。'
    )
    doc.add_paragraph('本申请的灌装技术手段已经清楚、完整地记载在说明书中：')

    p = doc.add_paragraph()
    run = p.add_run('灌装原理（固定容量法）：')
    run.font.bold = True

    doc.add_paragraph('本申请采用固定容量控制灌装的技术手段，具体过程如下：')

    doc.add_paragraph('  关闭弧形盖 → Z形杆带动抽板移动 → 一号圆孔与二号圆孔重叠')
    doc.add_paragraph('  → 药物颗粒在重力作用下掉落 → 药瓶被灌装')
    doc.add_paragraph('  → 打开弧形盖 → 弹簧复位 → 抽板封堵圆孔 → 灌装停止')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('技术特点：')
    run.font.bold = True

    doc.add_paragraph('1. 控制方式：通过弧形盖的开关控制灌装的开始和停止')
    doc.add_paragraph('2. 驱动机制：Z形杆将弧形盖的旋转运动转换为抽板的直线运动')
    doc.add_paragraph('3. 密封保证：抽板与中空板的配合实现可靠的启闭')
    doc.add_paragraph('4. 重力驱动：药物颗粒在重力作用下自然掉落，无需额外动力')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('本领域技术人员的公知常识：')
    run.font.bold = True

    doc.add_paragraph(
        '固定容量灌装是颗粒物料灌装的常用方法；通过控制落料口的开启和关闭来控制'
        '灌装量是常规技术手段；重力自流式灌装适用于流动性好的颗粒物料（如药品颗粒）。'
    )

    doc.add_paragraph()

    # 关于"计量"技术手段
    doc.add_heading('2. 关于"计量"技术手段的说明', 3)
    doc.add_paragraph(
        '审查员认为没有说明"计量"的技术手段。实际上，本申请已经清楚说明了计量方式。'
    )
    doc.add_paragraph('本申请采用的计量方式是"定容计量法"：')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('计量原理：')
    run.font.bold = True

    doc.add_paragraph('• 固定容量：中空板的一号圆孔（610）和抽板的二号圆孔（611）形成固定容量的计量腔')
    doc.add_paragraph('• 精确控制：圆孔的直径和深度决定了每次灌装的药物颗粒量')
    doc.add_paragraph('• 单次计量：每次关闭弧形盖，圆孔重叠形成一个完整的计量单元')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('本领域技术人员的理解：')
    run.font.bold = True

    doc.add_paragraph(
        '圆孔的尺寸可以根据需要灌装的药量进行设计；计量腔的容量 = 圆孔截面积 × 圆孔深度；'
        '这是一种简单的定容计量方式，适用于颗粒状物料。'
    )

    doc.add_paragraph()

    # 关于"称重"技术手段
    doc.add_heading('3. 关于"称重"技术手段的说明', 3)
    doc.add_paragraph('审查员认为没有说明"称重"的技术手段。')
    doc.add_paragraph('本申请不需要称重设备，理由如下：')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('采用容量计量而非重量计量：')
    run.font.bold = True

    doc.add_paragraph('• 本申请是实用新型专利，保护的是产品的形状、构造或其结合')
    doc.add_paragraph('• 本申请采用固定容量计量而非重量计量')
    doc.add_paragraph('• 对于颗粒状药品，在一定密度范围内，容量与重量成正比关系')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('专利法第26条第3款的理解：')
    run.font.bold = True

    doc.add_paragraph('• 说明书应当清楚说明实用新型的结构及其工作原理')
    doc.add_paragraph('• 不要求说明所有可能的变型或改进')
    doc.add_paragraph('• 不要求说明与发明目的无关的技术特征（如称重设备）')

    doc.add_paragraph()

    # 关于"如何用于药品灌装"
    doc.add_heading('4. 关于"如何用于药品灌装"的说明', 3)
    doc.add_paragraph(
        '审查员认为"不清楚如何用于药品灌装的"，这与说明书内容不符。'
    )
    doc.add_paragraph('说明书第[0020]段明确记载了工作原理，完整说明了药品灌装的流程：')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('步骤1：准备工作')
    run.font.bold = True

    doc.add_paragraph(
        '  通过进料口12倒入药物颗粒到储药仓61内部，把药瓶放置到固定盘10的顶部'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('步骤2：开始灌装')
    run.font.bold = True

    doc.add_paragraph(
        '  通过把手4关闭弧形盖3，弧形盖3带动Z形杆69、横板68、抽板67移动'
    )
    doc.add_paragraph('  一号圆孔610与二号圆孔611重叠，药物颗粒掉落到药瓶中')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('步骤3：停止灌装')
    run.font.bold = True

    doc.add_paragraph(
        '  打开弧形盖3，弹簧66带动抽板67复位，圆孔被封闭，灌装停止'
    )

    doc.add_paragraph()
    doc.add_paragraph('这段工作原理已经清楚地说明了：')
    doc.add_paragraph('1. 准备工作：倒入药物颗粒、放置药瓶')
    doc.add_paragraph('2. 灌装动作：关闭弧形盖触发灌装')
    doc.add_paragraph('3. 灌装过程：圆孔重叠，药物在重力作用下掉落')
    doc.add_paragraph('4. 停止灌装：打开弧形盖，抽板封堵圆孔')

    doc.add_paragraph()

    # 关于"无法具体实施"
    doc.add_heading('5. 关于"无法具体实施"的异议', 3)
    doc.add_paragraph(
        '审查员认为"根据说明书记载的内容无法具体实施"，这一结论是不正确的。'
    )
    doc.add_paragraph('本领域技术人员可以根据说明书实现本实用新型：')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('实施条件已满足：')
    run.font.bold = True

    doc.add_paragraph('✓ 结构清楚：权利要求1-9清楚描述了各部件的结构和连接关系')
    doc.add_paragraph('✓ 材料明确：各部件可采用常规材料（金属、塑料等）制造')
    doc.add_paragraph('✓ 尺寸合理：各部件的相对位置和配合关系清楚')
    doc.add_paragraph('✓ 工作原理明确：说明书详细说明了操作流程和工作原理')
    doc.add_paragraph('✓ 技术效果清楚：防止装药过多或过少，防止药物颗粒浪费')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('具体实施步骤：')
    run.font.bold = True

    doc.add_paragraph('1. 制造：根据附图1-6和权利要求1-9制造各部件')
    doc.add_paragraph('2. 组装：按照说明书记载的连接关系进行组装')
    doc.add_paragraph('3. 调试：调整弹簧66的张力，确保抽板67复位可靠')
    doc.add_paragraph('4. 使用：按照说明书第[0020]段的工作原理进行操作')

    doc.add_paragraph()

    # ============ 三、本领域技术人员的理解 ============
    doc.add_heading('三、本领域技术人员的理解能力', 1)
    doc.add_paragraph('所属技术领域：灌装技术领域，具体涉及药品生产灌装机')
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('本领域技术人员的知识和能力：')
    run.font.bold = True

    doc.add_paragraph('• 了解颗粒物料灌装的基本原理（重力自流、定容计量）')
    doc.add_paragraph('• 熟悉常规传动机构（齿轮、连杆、弹簧等）')
    doc.add_paragraph('• 掌握基本的机械设计方法')
    doc.add_paragraph('• 能够根据附图和文字说明实现实用新型')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('本领域技术人员能够理解：')
    run.font.bold = True

    doc.add_paragraph('• Z形杆69将旋转运动转换为直线运动的机械原理')
    doc.add_paragraph('• 弹簧66提供复位力的作用')
    doc.add_paragraph('• 抽板67与中空板63配合实现启闭的原理')
    doc.add_paragraph('• 齿轮传动系统的变速和换向作用')

    doc.add_paragraph()

    # ============ 四、与现有技术的对比 ============
    doc.add_heading('四、与现有技术的对比', 1)

    # 添加表格
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'

    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '技术特征'
    header_cells[1].text = '现有技术'
    header_cells[2].text = '本申请'

    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 表格数据
    table_data = [
        ['灌装方式', '人工或复杂机械控制', '简单的盖板控制'],
        ['计量方式', '需要额外计量设备', '固定容量计量'],
        ['防污染', '人工操作增加污染风险', '密闭自动操作'],
        ['物料回收', '难以回收溢出物料', '自动回收溢出物料'],
        ['结构复杂度', '高', '低'],
        ['制造成本', '高', '低']
    ]

    for i, row_data in enumerate(table_data, 1):
        row = table.rows[i]
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        row.cells[2].text = row_data[2]
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    doc.add_paragraph()

    # ============ 五、结论 ============
    doc.add_heading('五、结论', 1)

    p = doc.add_paragraph()
    run = p.add_run('本申请的说明书已经对技术方案作出了清楚、完整的说明：')
    run.font.bold = True

    doc.add_paragraph('1. 结构清楚：各部件的结构、连接关系、配合关系清楚')
    doc.add_paragraph('2. 工作原理明确：详细说明了操作流程和工作机理')
    doc.add_paragraph('3. 技术手段完整：灌装、计量（定容计量）的技术手段已经说明')
    doc.add_paragraph('4. 能够实施：本领域技术人员可以根据说明书实现本实用新型')

    doc.add_paragraph()
    doc.add_paragraph(
        '审查员关于"灌装、计量、称重"技术手段不清楚的异议，是对本申请技术方案的误解。'
        '本申请采用固定容量灌装这一本领域的常规技术手段，无需额外的称重设备即可实现灌装功能。'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('本申请完全符合专利法第26条第3款的规定，恳请审查员重新考虑并授予专利权。')
    run.font.bold = True

    # ============ 添加结尾 ============
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('申请人：山东圣旺药业股份有限公司').bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('代理机构：济南宝宸专利代理事务所(普通合伙)')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('代理师：徐健')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f'日期：{datetime.now().strftime("%Y年%m月%d日")}')

    return doc

if __name__ == "__main__":
    doc = create_response_document()

    output_path = '/Users/xujian/工作/04_审查意见/01_待答复/202520560089.0-第一次审查意见/第一次审查意见答复.docx'
    doc.save(output_path)

    print('✅ Word文档创建成功！')
    print(f'📄 文件路径：{output_path}')
    print()
    print('📊 文档信息：')
    print('   • 标题：第一次审查意见答复')
    print('   • 申请人：山东圣旺药业股份有限公司')
    print('   • 申请号：202520560089.0')
    print('   • 发明名称：一种药品生产灌装机')
    print('   • 字体：宋体（正文）、黑体（标题）')
    print('   • 正文字号：12磅')
    print('   • 页边距：上下2.54cm，左右3.17cm')
    print('   • 包含内容：')
    print('     - 审查意见概述')
    print('     - 答复意见（5个主要论点）')
    print('     - 技术人员理解能力')
    print('     - 与现有技术对比表')
    print('     - 结论')
    print()
    print('💾 文件已保存，可用于提交答复！')
