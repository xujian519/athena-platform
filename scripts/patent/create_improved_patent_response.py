#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建改进版Word格式的审查意见答复文件
Create Improved Word format response for patent office action

根据AI辩论结果改进答复内容，增加校准方法和精度保证机制说明

作者: Athena平台团队
版本: 2.0.0
创建时间: 2026-02-09
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_improved_response_document():
    """创建改进版审查意见答复Word文档"""

    # 创建文档
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # ============ 添加标题 ============
    title = doc.add_heading('第一次审查意见答复', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '黑体'
        run.font.size = Pt(16)

    # ============ 添加基本信息 ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('申请人：山东圣旺药业股份有限公司  申请号：202520560089.0')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('发明创造名称：一种药品生产灌装机  答复日期：2026年2月9日')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('审查员：肖玉林')
    run.font.size = Pt(11)

    doc.add_paragraph()

    # ============ 一、审查意见概述 ============
    doc.add_heading('一、审查意见概述', 1)
    doc.add_paragraph(
        '审查员认为本申请的说明书没有清楚完整地解释和说明"灌装、计量、称重"的'
        '技术手段，不清楚如何用于药品灌装的，对所属技术领域的技术人员来说，该手段是'
        '含糊不清的，根据说明书记载的内容无法具体实施，因此不符合专利法第26条第3款的规定。'
    )

    doc.add_paragraph()

    # ============ 二、答复意见 ============
    doc.add_heading('二、答复意见', 1)
    doc.add_paragraph()

    # 2.1 关于专利法第26条第3款的说明
    doc.add_heading('（一）关于专利法第26条第3款的总体说明', 2)
    doc.add_paragraph(
        '本申请的说明书已经对技术方案作出了清楚、完整的说明，所属技术领域的技术人员能够'
        '根据说明书实现本实用新型，符合专利法第26条第3款的规定。具体理由如下：'
    )

    doc.add_paragraph()

    # 关于"灌装"技术手段
    doc.add_heading('1. 关于"灌装"技术手段的说明', 3)
    doc.add_paragraph(
        '审查员认为说明书没有清楚完整地解释"灌装"的技术手段，这一观点是不准确的。'
    )
    doc.add_paragraph('本申请的灌装技术手段已经清楚、完整地记载在说明书中：')

    p = doc.add_paragraph()
    run = p.add_run('灌装原理（固定容量法）：')
    run.font.bold = True

    doc.add_paragraph('本申请采用固定容量控制灌装的技术手段，具体过程如下：')

    doc.add_paragraph('  关闭弧形盖 → Z形杆带动抽板移动 → 一号圆孔与二号圆孔重叠')
    doc.add_paragraph('  → 药物颗粒在重力作用下掉落 → 药瓶被灌装')
    doc.add_paragraph('  → 打开弧形盖 → 弹簧复位 → 抽板封堵圆孔 → 灌装停止')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('技术特点：')
    run.font.bold = True

    doc.add_paragraph('1. 控制方式：通过弧形盖的开关控制灌装的开始和停止')
    doc.add_paragraph('2. 驱动机制：Z形杆将弧形盖的旋转运动转换为抽板的直线运动')
    doc.add_paragraph('3. 密封保证：抽板与中空板的配合实现可靠的启闭')
    doc.add_paragraph('4. 重力驱动：药物颗粒在重力作用下自然掉落，无需额外动力')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('本领域技术人员的公知常识：')
    run.font.bold = True

    doc.add_paragraph(
        '固定容量灌装是颗粒物料灌装的常用方法；通过控制落料口的开启和关闭来控制'
        '灌装量是常规技术手段；重力自流式灌装适用于流动性好的颗粒物料（如药品颗粒）。'
    )

    doc.add_paragraph()

    # 关于"计量"技术手段 - 改进版
    doc.add_heading('2. 关于"计量"技术手段的说明', 3)
    doc.add_paragraph(
        '审查员认为没有说明"计量"的技术手段。实际上，本申请已经清楚说明了计量方式。'
    )
    doc.add_paragraph('本申请采用的计量方式是"定容计量法"：')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('计量原理：')
    run.font.bold = True

    doc.add_paragraph('• 固定容量：中空板的一号圆孔（610）和抽板的二号圆孔（611）形成固定容量的计量腔')
    doc.add_paragraph('• 精确控制：圆孔的直径和深度决定了每次灌装的药物颗粒量')
    doc.add_paragraph('• 单次计量：每次关闭弧形盖，圆孔重叠形成一个完整的计量单元')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('本领域技术人员的理解：')
    run.font.bold = True

    doc.add_paragraph(
        '圆孔的尺寸可以根据需要灌装的药量进行设计；计量腔的容量 = 圆孔截面积 × 圆孔深度；'
        '这是一种简单的定容计量方式，适用于颗粒状物料。'
    )

    doc.add_paragraph()

    # ============ 新增：计量精度与校准方法说明 ============
    doc.add_heading('3. 关于"计量准确"的实现机制与校准方法的补充说明', 3)
    doc.add_paragraph(
        '针对审查员关于计量精度实现机制的关切，申请人特此补充说明如下：'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('（1）"计量准确"在本申请中的具体含义：')
    run.font.bold = True

    doc.add_paragraph(
        '本申请所述的"计量准确"首先是指**容积计量的高重复性精度**，这是由以下技术特征保证的：'
    )
    doc.add_paragraph('① 精密加工：一号圆孔和二号圆孔的尺寸可以精确控制在±0.1mm以内，确保每次灌装的容积一致')
    doc.add_paragraph('② 可靠启闭：抽板与中空板的紧密配合确保每次灌装的容积完全相同')
    doc.add_paragraph('③ 机械重复：弹簧复位的可靠性保证了每次动作的一致性')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('（2）重量精度的实现方法：')
    run.font.bold = True

    doc.add_paragraph(
        '对于特定药品的**重量精度**要求，本领域技术人员采用以下标准且必要的校准流程：'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('第一步：物料特性测定')
    run.font.bold = True

    doc.add_paragraph('• 测定该药品颗粒的堆密度（松装密度）或振实密度')
    doc.add_paragraph('• 方法：按照《中国药典》通则"密度测定"方法进行')
    doc.add_paragraph('• 获得参数：例如某药品颗粒的堆密度为0.6g/ml')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('第二步：容积计算')
    run.font.bold = True

    doc.add_paragraph('• 根据目标装量计算所需计量腔容积')
    doc.add_paragraph('• 公式：V = m / ρ （V为容积，m为目标重量，ρ为堆密度）')
    doc.add_paragraph('• 例如：目标装量10g，堆密度0.6g/ml，则所需容积V = 10/0.6 ≈ 16.7ml')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('第三步：结构参数确定')
    run.font.bold = True

    doc.add_paragraph('• 根据计算容积确定圆孔直径和深度')
    doc.add_paragraph('• 对于圆形孔：V = π×(d/2)²×h，其中d为直径，h为深度（板厚）')
    doc.add_paragraph('• 本领域技术人员可根据常规机械设计方法确定具体参数')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('第四步：实验验证与微调')
    run.font.bold = True

    doc.add_paragraph('• 制造初步确定尺寸的抽板和中空板')
    doc.add_paragraph('• 进行试灌装，测量实际装量')
    doc.add_paragraph('• 如有偏差，通过以下方式微调：')
    doc.add_paragraph('  - 更换不同孔径的抽板（抽板为独立部件，便于更换）')
    doc.add_paragraph('  - 或微调抽板行程以调整圆孔重叠面积')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('（5）上述校准流程的法律性质：')
    run.font.bold = True

    doc.add_paragraph(
        '申请人强调，上述"一物料一标定"的方法是定容式颗粒灌装领域的**公知常识**和**标准操作流程**。'
    )
    doc.add_paragraph('参考依据包括：')
    doc.add_paragraph('• 《制药机械行业标准》中关于定容式灌装机的校准要求')
    doc.add_paragraph('• 《药剂学》教材中关于颗粒物料定容灌装的章节')
    doc.add_paragraph('• 制药设备制造商的操作手册和行业标准做法')

    doc.add_paragraph()
    doc.add_paragraph(
        '专利法第26条第3款要求的"能够实现"，是指所属领域技术人员根据说明书和其掌握的'
        '现有技术（包括公知常识）能够实现该技术方案，并不意味着说明书必须记载所有可能'
        '的实施细节或常规操作流程。上述校准流程属于本领域的公知常识，无需在说明书中赘述。'
    )

    doc.add_paragraph()

    # 关于"称重"技术手段
    doc.add_heading('4. 关于"称重"技术手段的说明', 3)
    doc.add_paragraph('审查员认为没有说明"称重"的技术手段。')
    doc.add_paragraph('本申请不需要称重设备，理由如下：')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('采用容量计量而非重量计量：')
    run.font.bold = True

    doc.add_paragraph('• 本申请是实用新型专利，保护的是产品的形状、构造或其结合')
    doc.add_paragraph('• 本申请采用固定容量计量而非重量计量')
    doc.add_paragraph('• 对于颗粒状药品，在一定密度范围内，容量与重量成正比关系')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('专利法第26条第3款的理解：')
    run.font.bold = True

    doc.add_paragraph('• 说明书应当清楚说明实用新型的结构及其工作原理')
    doc.add_paragraph('• 不要求说明所有可能的变型或改进')
    doc.add_paragraph('• 不要求说明与发明目的无关的技术特征（如称重设备）')

    doc.add_paragraph()

    # 关于"如何用于药品灌装"
    doc.add_heading('5. 关于"如何用于药品灌装"的说明', 3)
    doc.add_paragraph(
        '审查员认为"不清楚如何用于药品灌装的"，这与说明书内容不符。'
    )
    doc.add_paragraph('说明书第[0020]段明确记载了工作原理，完整说明了药品灌装的流程：')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('步骤1：准备工作')
    run.font.bold = True

    doc.add_paragraph(
        '  通过进料口12倒入药物颗粒到储药仓61内部，把药瓶放置到固定盘10的顶部'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('步骤2：开始灌装')
    run.font.bold = True

    doc.add_paragraph(
        '  通过把手4关闭弧形盖3，弧形盖3带动Z形杆69、横板68、抽板67移动'
    )
    doc.add_paragraph('  一号圆孔610与二号圆孔611重叠，药物颗粒掉落到药瓶中')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('步骤3：停止灌装')
    run.font.bold = True

    doc.add_paragraph(
        '  打开弧形盖3，弹簧66带动抽板67复位，圆孔被封闭，灌装停止'
    )

    doc.add_paragraph()
    doc.add_paragraph('这段工作原理已经清楚地说明了：')
    doc.add_paragraph('1. 准备工作：倒入药物颗粒、放置药瓶')
    doc.add_paragraph('2. 灌装动作：关闭弧形盖触发灌装')
    doc.add_paragraph('3. 灌装过程：圆孔重叠，药物在重力作用下掉落')
    doc.add_paragraph('4. 停止灌装：打开弧形盖，抽板封堵圆孔')

    doc.add_paragraph()

    # 关于"无法具体实施"
    doc.add_heading('6. 关于"无法具体实施"的异议', 3)
    doc.add_paragraph(
        '审查员认为"根据说明书记载的内容无法具体实施"，这一结论是不正确的。'
    )
    doc.add_paragraph('本领域技术人员可以根据说明书实现本实用新型：')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('实施条件已满足：')
    run.font.bold = True

    doc.add_paragraph('✓ 结构清楚：权利要求1-9清楚描述了各部件的结构和连接关系')
    doc.add_paragraph('✓ 材料明确：各部件可采用常规材料（金属、塑料等）制造')
    doc.add_paragraph('✓ 尺寸合理：各部件的相对位置和配合关系清楚')
    doc.add_paragraph('✓ 工作原理明确：说明书详细说明了操作流程和工作原理')
    doc.add_paragraph('✓ 技术效果清楚：防止装药过多或过少，防止药物颗粒浪费')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('具体实施步骤：')
    run.font.bold = True

    doc.add_paragraph('1. 制造：根据附图1-6和权利要求1-9制造各部件')
    doc.add_paragraph('2. 组装：按照说明书记载的连接关系进行组装')
    doc.add_paragraph('3. 调试：调整弹簧66的张力，确保抽板67复位可靠')
    doc.add_paragraph('4. 校准：按照本领域的公知常识进行物料标定（如前述）')
    doc.add_paragraph('5. 使用：按照说明书第[0020]段的工作原理进行操作')

    doc.add_paragraph()

    # ============ 新增：可调性说明 ============
    doc.add_heading('7. 关于结构可调性的补充说明', 3)
    doc.add_paragraph(
        '针对审查员关于可调性的关切，申请人特此说明：'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('（1）抽板的可更换性：')
    run.font.bold = True

    doc.add_paragraph(
        '说明书及附图中，抽板（67）是一个独立成型的部件，通过与驱动杆（Z形杆69）的简单连接'
        '实现固定。所属领域技术人员可以理解，为适应不同装量的要求，可以：'
    )
    doc.add_paragraph('① 制备多个具有不同孔径二号圆孔的抽板')
    doc.add_paragraph('② 根据目标装量选择合适的抽板进行更换')
    doc.add_paragraph('③ 抽板的独立结构为这种更换提供了物理基础')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('（2）容积调节的替代方案：')
    run.font.bold = True

    doc.add_paragraph('除更换抽板外，所属领域技术人员还可以采用以下常规调节方式：')
    doc.add_paragraph('① 调整抽板行程，改变两圆孔的重叠面积，实现容积微调')
    doc.add_paragraph('② 在中空板上设置多个不同直径的一号圆孔，选择使用')
    doc.add_paragraph('③ 通过调整弹簧张力来控制抽板位置，实现精细调节')

    doc.add_paragraph()
    doc.add_paragraph(
        '这些调节方式属于本领域的常规技术手段，所属领域技术人员根据说明书教导和其掌握'
        '的现有技术即可直接、毫无疑义地确定并实现，无需说明书逐一列举。'
    )

    doc.add_paragraph()

    # ============ 本领域技术人员的理解能力 ============
    doc.add_heading('8. 本领域技术人员的理解能力', 3)
    doc.add_paragraph('所属技术领域：灌装技术领域，具体涉及药品生产灌装机')
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('本领域技术人员的知识和能力：')
    run.font.bold = True

    doc.add_paragraph('• 了解颗粒物料灌装的基本原理（重力自流、定容计量）')
    doc.add_paragraph('• 熟悉常规传动机构（齿轮、连杆、弹簧等）')
    doc.add_paragraph('• 掌握基本的机械设计方法')
    doc.add_paragraph('• 了解药品灌装的行业标准（如《中国药典》"装量差异"要求）')
    doc.add_paragraph('• 掌握定容式灌装机的校准和标定方法（公知常识）')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('本领域技术人员能够理解：')
    run.font.bold = True

    doc.add_paragraph('• Z形杆69将旋转运动转换为直线运动的机械原理')
    doc.add_paragraph('• 弹簧66提供复位力的作用')
    doc.add_paragraph('• 抽板67与中空板63配合实现启闭的原理')
    doc.add_paragraph('• 齿轮传动系统的变速和换向作用')
    doc.add_paragraph('• 如何根据物料特性进行标定和校准（公知常识）')

    doc.add_paragraph()

    # ============ 三、与现有技术的对比 ============
    doc.add_heading('三、与现有技术的对比', 1)

    # 添加表格
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'

    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '技术特征'
    header_cells[1].text = '现有技术'
    header_cells[2].text = '本申请'

    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 表格数据
    table_data = [
        ['灌装方式', '人工或复杂机械控制', '简单的盖板控制'],
        ['计量方式', '需要额外计量设备', '固定容量计量'],
        ['防污染', '人工操作增加污染风险', '密闭自动操作'],
        ['物料回收', '难以回收溢出物料', '自动回收溢出物料'],
        ['结构复杂度', '高', '低'],
        ['制造成本', '高', '低']
    ]

    for i, row_data in enumerate(table_data, 1):
        row = table.rows[i]
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        row.cells[2].text = row_data[2]
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    doc.add_paragraph()

    # ============ 四、结论 ============
    doc.add_heading('四、结论', 1)

    p = doc.add_paragraph()
    run = p.add_run('本申请的说明书已经对技术方案作出了清楚、完整的说明：')
    run.font.bold = True

    doc.add_paragraph('1. 结构清楚：各部件的结构、连接关系、配合关系清楚')
    doc.add_paragraph('2. 工作原理明确：详细说明了操作流程和工作机理')
    doc.add_paragraph('3. 技术手段完整：灌装、计量（定容计量）的技术手段已经说明')
    doc.add_paragraph('4. 能够实施：本领域技术人员可以根据说明书实现本实用新型')
    doc.add_paragraph('5. 精度可实现：通过本领域的公知常识（校准流程）即可实现重量精度要求')

    doc.add_paragraph()
    doc.add_paragraph(
        '审查员关于"灌装、计量、称重"技术手段不清楚的异议，是对本申请技术方案的误解。'
        '本申请采用固定容量灌装这一本领域的常规技术手段，配合本领域的标准校准流程，'
        '无需额外的称重设备即可实现灌装功能。'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('本申请完全符合专利法第26条第3款的规定，恳请审查员重新考虑并授予专利权。')
    run.font.bold = True

    # ============ 添加结尾 ============
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('申请人：山东圣旺药业股份有限公司').bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('代理机构：济南宝宸专利代理事务所(普通合伙)')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('代理师：徐健')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f'日期：{datetime.now().strftime("%Y年%m月%d日")}')

    return doc

if __name__ == "__main__":
    doc = create_improved_response_document()

    output_path = '/Users/xujian/工作/04_审查意见/01_待答复/202520560089.0-第一次审查意见/第一次审查意见答复_改进版.docx'
    doc.save(output_path)

    print('✅ 改进版Word文档创建成功！')
    print(f'📄 文件路径：{output_path}')
    print()
    print('📊 文档信息：')
    print('   • 标题：第一次审查意见答复（改进版）')
    print('   • 申请人：山东圣旺药业股份有限公司')
    print('   • 申请号：202520560089.0')
    print('   • 发明名称：一种药品生产灌装机')
    print('   • 字体：宋体（正文）、黑体（标题）')
    print('   • 正文字号：12磅')
    print('   • 页边距：上下2.54cm，左右3.17cm')
    print()
    print('🔧 改进内容：')
    print('   ✓ 新增"计量准确"实现机制说明')
    print('   ✓ 详细补充校准方法（四步流程）')
    print('   ✓ 明确容积精度与重量精度的关系')
    print('   ✓ 说明校准流程属于公知常识')
    print('   ✓ 补充结构可调性说明')
    print('   ✓ 强调本领域技术人员能够实施')
    print()
    print('💾 文件已保存，可用于提交答复！')
