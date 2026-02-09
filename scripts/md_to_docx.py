#!/usr/bin/env python3
"""
将Markdown文件转换为DOCX文件，保持格式
"""

import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from pathlib import Path


class MarkdownToDocxConverter:
    """Markdown到DOCX转换器"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """设置文档样式"""
        # 设置默认字体
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.doc.styles['Normal'].font.size = Pt(12)

        # 标题样式
        for i in range(1, 7):
            heading_style = self.doc.styles[f'Heading {i}']
            heading_style.font.name = '黑体'
            heading_style.font.bold = True
            heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def _parse_inline_formatting(self, text):
        """解析内联格式（加粗等）"""
        # 处理加粗 **text**
        bold_pattern = r'\*\*([^*]+)\*\*'

        parts = []
        last_end = 0

        for match in re.finditer(bold_pattern, text):
            # 添加前面的普通文本
            if match.start() > last_end:
                parts.append(('text', text[last_end:match.start()]))

            # 添加加粗文本
            parts.append(('bold', match.group(1)))
            last_end = match.end()

        # 添加剩余文本
        if last_end < len(text):
            parts.append(('text', text[last_end:]))

        return parts

    def _add_paragraph_with_formatting(self, text, style=None):
        """添加带格式的段落"""
        if not text or text.strip() == '':
            return

        para = self.doc.add_paragraph()
        if style:
            para.style = style

        # 解析内联格式
        parts = self._parse_inline_formatting(text)

        for part_type, part_text in parts:
            if part_type == 'bold':
                run = para.add_run(part_text)
                run.bold = True
            else:
                para.add_run(part_text)

    def _parse_table(self, lines, start_idx):
        """解析Markdown表格"""
        table_data = []
        idx = start_idx

        # 解析表格行
        while idx < len(lines):
            line = lines[idx].rstrip()

            # 空行或不是表格行
            if not line or not line.startswith('|'):
                break

            # 跳过分隔行（如 |---|---|）
            if re.search(r'^\s*\|?\s*:?-+:?\|', line):
                idx += 1
                continue

            # 解析表格单元格
            cells = line.split('|')
            # 移除首尾空元素（由|开头和结尾产生的空字符串）
            if len(cells) >= 2:
                if cells[0].strip() == '':
                    cells = cells[1:]
                if cells[-1].strip() == '':
                    cells = cells[:-1]

            # 去除每个单元格的首尾空格
            cells = [cell.strip() for cell in cells]

            if cells:
                table_data.append(cells)
            idx += 1

        if not table_data:
            return idx

        # 确定列数（使用最多的列数）
        max_cols = max(len(row) for row in table_data)

        # 创建表格
        table = self.doc.add_table(rows=len(table_data), cols=max_cols)
        table.style = 'Light Grid Accent 1'

        # 填充表格
        for row_idx, row_data in enumerate(table_data):
            for col_idx in range(max_cols):
                cell = table.rows[row_idx].cells[col_idx]

                # 如果当前行这一列有数据
                if col_idx < len(row_data):
                    cell.text = row_data[col_idx]
                else:
                    cell.text = ''

                # 设置表头加粗
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        return idx

    def convert(self, md_content):
        """转换Markdown内容"""
        lines = md_content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # 空行
            if not line.strip():
                self.doc.add_paragraph()
                i += 1
                continue

            # 一级标题 #
            if line.startswith('# '):
                self.doc.add_heading(line[2:], level=1)
                i += 1
                continue

            # 二级标题 ##
            if line.startswith('## '):
                self.doc.add_heading(line[3:], level=2)
                i += 1
                continue

            # 三级标题 ###
            if line.startswith('### '):
                self.doc.add_heading(line[4:], level=3)
                i += 1
                continue

            # 四级标题 ####
            if line.startswith('#### '):
                self.doc.add_heading(line[5:], level=4)
                i += 1
                continue

            # 分隔线
            if line.strip() == '---':
                self.doc.add_paragraph('_' * 80)
                i += 1
                continue

            # 列表项
            if line.strip().startswith('- '):
                # 处理连续列表
                while i < len(lines) and lines[i].strip().startswith('- '):
                    item_text = lines[i].strip()[2:]
                    self._add_paragraph_with_formatting(f'• {item_text}')
                    i += 1
                continue

            # 表格
            if line.strip().startswith('|'):
                i = self._parse_table(lines, i)
                continue

            # 普通段落
            self._add_paragraph_with_formatting(line)
            i += 1

    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)


def convert_md_to_docx(input_file, output_file=None):
    """转换Markdown文件到DOCX"""
    input_path = Path(input_file)

    if not input_path.exists():
        raise FileNotFoundError(f"文件不存在: {input_file}")

    if output_file is None:
        output_file = input_path.with_suffix('.docx')

    # 读取Markdown文件
    with open(input_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 转换
    converter = MarkdownToDocxConverter()
    converter.convert(md_content)
    converter.save(output_file)

    print(f"✅ 转换完成: {output_file}")
    print(f"📊 包含 {len(converter.doc.tables)} 个表格")
    return output_file


if __name__ == '__main__':
    import sys

    if len(sys.argv) < 2:
        print("用法: python md_to_docx.py <markdown文件> [输出文件]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    try:
        convert_md_to_docx(input_file, output_file)
    except Exception as e:
        print(f"❌ 转换失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
